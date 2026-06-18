# -*- coding: utf-8 -*-
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from src.services.attachment_text_extractor import extract_text_from_attachment


def test_plain_text_attachment_extraction():
    raw = "权利人：王汝冲\n坐落：道县月岩西路\n分摊土地使用权面积：15.20平方米"

    extracted = extract_text_from_attachment("房产证.txt", raw.encode("utf-8"))

    assert extracted.method == "plain_text"
    assert "王汝冲" in extracted.text


def test_docx_attachment_extraction():
    stream = BytesIO()
    doc = Document()
    doc.add_paragraph("规划建筑密度：35%-55%")
    doc.add_paragraph("容积率：0.7-1.50")
    doc.save(stream)

    extracted = extract_text_from_attachment("规划条件.docx", stream.getvalue())

    assert extracted.method == "python_docx"
    assert "容积率" in extracted.text


def test_ocr_file_endpoint_extracts_metrics():
    import src.api as api

    client = TestClient(api.app)
    raw = "土地权利人: 王汝冲\n房屋坐落位置: 道县月岩西路\n分摊土地使用权面积: 15.20平方米\n权利性质: 国有划拨\n用途: 城镇住宅用地"

    response = client.post(
        "/api/ocr-file",
        data={"attachment_type": "property_cert"},
        files={"file": ("房产证.txt", raw.encode("utf-8"), "text/plain")},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["extraction_method"] == "plain_text"
    assert body["data"]["land_user"] == "王汝冲"
    assert body["data"]["land_area"] == 15.2


def test_pdf_below_threshold_falls_back_to_paddleocr(monkeypatch):
    import src.services.attachment_text_extractor as extractor

    monkeypatch.setattr(
        extractor,
        "_extract_pdf_text_layer",
        lambda data: extractor.ExtractedAttachmentText(text="短", method="pymupdf_text_layer", page_count=1),
    )
    monkeypatch.setattr(
        extractor,
        "_ocr_pdf_pages_with_paddleocr",
        lambda data: extractor.ExtractedAttachmentText(text="PaddleOCR识别正文", method="paddleocr_pdf_pages", page_count=1),
    )

    extracted = extractor.extract_text_from_attachment("扫描件.pdf", b"fake-pdf", min_text_chars=10)

    assert extracted.method == "paddleocr_pdf_pages"
    assert "低于阈值" in extracted.warnings[0]


def test_image_uses_paddleocr(monkeypatch):
    import src.services.attachment_text_extractor as extractor

    class FakePaddleOCR:
        def ocr(self, path, cls=True):
            return [[[[0, 0], [1, 0], [1, 1], [0, 1]], ("图片识别文本", 0.99)]]

    monkeypatch.setattr(extractor, "_PADDLE_OCR", FakePaddleOCR())
    monkeypatch.setattr(extractor, "_PADDLE_INIT_ERROR", None)

    extracted = extractor.extract_text_from_attachment("附件.png", b"not-real-image")

    assert extracted.method == "paddleocr_image"
    assert "图片识别文本" in extracted.text


def test_paddleocr_incompatible_runtime_guard():
    import src.services.attachment_text_extractor as extractor

    class FakePaddle:
        __version__ = "3.3.1"

    class FakePaddleOCRModule:
        __version__ = "3.5.0"

    try:
        extractor._check_paddleocr_runtime_compatibility(FakePaddle, FakePaddleOCRModule)
    except extractor.AttachmentTextExtractionError as exc:
        assert "ConvertPirAttribute2RuntimeAttribute" in str(exc)
    else:
        raise AssertionError("expected incompatible Paddle runtime guard to raise")
