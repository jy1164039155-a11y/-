# -*- coding: utf-8 -*-
import shutil
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm

from src.services.income_capitalization import calculate_income_capitalization
from src.services.income_capitalization_docx import _fill_image_cell, apply_income_capitalization_to_docx
from src.services.valuation_process_builder import build_valuation_process_draft


PNG_1X1 = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="


def _baozhen_income_payload():
    data = {
        "use_income_cap": True,
        "valuation_date": "2026-06-01",
        "land_location": "道县西洲街道宝珍街六栋九号",
        "land_usage": "住宅用地",
        "land_use_term": "70年",
        "building_area": "205.45",
        "land_area": "50",
    }
    seeded = calculate_income_capitalization(data)
    for item, rent, name, location in zip(
        seeded["rent_instances"],
        ("11.4", "11.4", "11.5"),
        ("金欧雅陶瓷(道州中路店)楼上住宅", "道州锁城楼上住宅", "倩倩便利店(道州中路店)楼上住宅"),
        ("道州中路466号爱莲商业广场", "道州中路466号爱莲商业广场", "道县道州中路397号"),
    ):
        item.update(
            {
                "name": name,
                "location": location,
                "monthly_rent": rent,
                "transaction_date": "2026年6月",
                "usage": "住宅",
                "photo_data": PNG_1X1,
                "location_image_data": PNG_1X1,
                "confirmed": True,
            }
        )
    for factor in seeded["rent_factor_items"]:
        factor["subject_index"] = "100"
        for slot in ("A", "B", "C"):
            factor["cases"][slot].update({"value": "同估价对象", "index": "100", "confirmed": True})
    case_values = {
        "usage": "住宅",
        "transaction_time": "2026.06",
        "transaction_condition": "正常",
        "commercial_prosperity": "较优",
        "bus_convenience": "较优",
        "road_accessibility": "较通畅",
        "infrastructure_guarantee": "95%",
        "environment_quality": "好",
        "public_facilities": "齐全",
        "road_type": "临主干道",
        "ventilation_lighting": "好",
        "newness": "七成新",
        "building_structure": "砖混结构",
        "internal_layout": "合理",
        "decoration": "简单装修",
        "parking": "基本满足停车需求",
        "property_management": "无物业管理",
    }
    for factor in seeded["rent_factor_items"]:
        for slot in ("A", "B", "C"):
            factor["cases"][slot]["value"] = case_values[factor["key"]]
    usage = next(item for item in seeded["rent_factor_items"] if item["key"] == "usage")
    usage["cases"]["A"]["index"] = "101.968"
    usage["cases"]["B"]["index"] = "101.968"
    usage["cases"]["C"]["index"] = "102.041"
    seeded["building_profile"] = {
        "building_location": "道江镇濂溪山庄宝珍街南六栋九号",
        "building_form": "独栋建筑物",
        "built_year": "1998",
        "floor_desc": "4层",
        "owner_floor_desc": "1-4层",
        "structure": "混合结构",
        "exterior": "双飞粉",
        "entrance_door": "不锈钢及木质入户门",
        "windows": "塑钢窗户",
        "security_facilities": "铁艺防盗网",
        "floor_finish": "地贴瓷砖",
        "ceiling_finish": "木质装饰吊顶",
        "interior": "简单装修",
        "maintenance": "维护保养状况一般",
        "newness_rate": "70",
        "newness_desc": "七成新",
        "economic_life_years": "50",
        "used_years": "28",
        "remaining_years": "22",
        "building_area": "205.45",
        "land_area": "50",
        "actual_use": "住宅",
        "current_use_condition": "现状使用条件",
        "current_use_basis": "《评估委托书》",
        "building_area_basis": "《不动产测绘报告书》",
        "land_area_basis": "《宗地图》",
        "plot_ratio": "4.11",
        "set_plot_ratio": "4.11",
    }
    seeded["income_parameters"] = {"vacancy_rate_range": "3%至5%", "vacancy_rate": "4", "rentable_area_ratio": "96"}
    seeded["expense_parameters"] = {
        "management_rate": "2",
        "repair_rate": "2",
        "replacement_base_unit_cost": "1300",
        "regional_adjustment_coefficient": "0.9",
        "cost_growth_rate": "2",
        "cost_base_date": "2023-01-01",
        "valuation_date": "2026-06-01",
        "residual_rate": "2",
        "insurance_rate_permille": "2",
        "property_tax_rate": "4",
        "property_tax_reduction_rate": "50",
    }
    seeded["cap_rate_parameters"] = {
        "land_usage": "住宅用地",
        "income_land_cap_rate": "6.5",
        "income_building_cap_rate": "8.5",
        "use_term_years": "70",
        "confirmed": True,
    }
    data["income_cap_analysis"] = seeded
    return data


def test_baozhen_income_capitalization_gold_values():
    result = calculate_income_capitalization(_baozhen_income_payload())
    values = result["income_results"]

    assert values["average_monthly_rent"] == "11.21"
    assert values["annual_gross_income"] == "2.5470"
    assert values["building_replacement_price"] == "25.7223"
    assert values["total_expense"] == "0.6394"
    assert values["land_net_income"] == "0.9212"
    assert result["income_cap_price"] == "2799.9"
    assert result["complete"] is True


def test_income_capitalization_insurance_uses_permille_not_percent():
    result = calculate_income_capitalization(_baozhen_income_payload())
    text = result["generated_narratives"]["income_cap_expense_narrative"]

    assert "2‰" in text
    assert "房屋年保险费＝房屋现值×2‰\n＝11.6047×2‰\n＝0.0232" in text
    assert "＝11.6047×2%" not in text


def test_income_gross_income_structure_matches_baozhen_template_order():
    result = calculate_income_capitalization(_baozhen_income_payload())
    narratives = result["generated_narratives"]
    text = narratives["income_cap_gross_income_intro"]

    # 验证新语序中的关键段落
    assert "根据受托估价人员现场勘察" in text
    assert "业主房屋所在层数为" in text
    assert "房屋整体经济耐用年限为" in text
    assert "分摊土地使用权面积" in text
    assert "现状容积率为" in text
    assert "4层层" not in text
    assert "1-4层层" not in text
    assert "入户门入户门" not in text
    assert "窗户窗户" not in text
    assert "防盗网防盗网" not in text
    assert "屋内地贴瓷砖" in text
    assert "屋内地贴地贴瓷砖" not in text
    assert "维护保养状况维护保养状况" not in text
    assert "《《" not in text
    assert "》》" not in text
    assert "现状使用条件使用条件" not in text


def test_income_gross_income_hotspots_follow_rendered_physical_description():
    payload = _baozhen_income_payload()

    class EmptyComparableLibrary:
        pass

    draft = build_valuation_process_draft(payload, EmptyComparableLibrary())
    method = next(item for item in draft["methods"] if item["method_key"] == "income_cap")
    section = next(item for item in method["narratives"] if item["key"] == "income_cap_gross_income_intro")
    segments = [item for item in section["segments"] if item.get("field") or item.get("fields")]
    fields = {item.get("field") for item in segments if item.get("field")}

    assert {
        "income_cap_analysis.building_profile.building_form",
        "income_cap_analysis.building_profile.floor_desc",
        "income_cap_analysis.building_profile.owner_floor_desc",
        "income_cap_analysis.building_profile.built_year",
        "income_cap_analysis.building_profile.structure",
        "income_cap_analysis.building_profile.entrance_door",
        "income_cap_analysis.building_profile.windows",
        "income_cap_analysis.building_profile.security_facilities",
        "income_cap_analysis.building_profile.exterior",
        "income_cap_analysis.building_profile.floor_finish",
        "income_cap_analysis.building_profile.ceiling_finish",
        "income_cap_analysis.building_profile.interior",
        "income_cap_analysis.building_profile.maintenance",
        "income_cap_analysis.building_profile.newness_desc",
        "income_cap_analysis.building_profile.economic_life_years",
        "income_cap_analysis.building_profile.used_years",
        "income_cap_analysis.building_profile.remaining_years",
        "income_cap_analysis.building_profile.building_area_basis",
        "income_cap_analysis.building_profile.land_area_basis",
        "income_cap_analysis.building_profile.building_area",
        "income_cap_analysis.building_profile.land_area",
        "income_cap_analysis.building_profile.current_use_basis",
        "valuation_condition_type",
        "set_plot_ratio",
    } <= fields
    assert any(item["text"] == "4" and item.get("field") == "income_cap_analysis.building_profile.floor_desc" for item in segments)
    assert any(item["text"] == "塑钢" and item.get("field") == "income_cap_analysis.building_profile.windows" for item in segments)
    assert any(item["text"] == "《评估委托书》" and item.get("field") == "income_cap_analysis.building_profile.current_use_basis" for item in segments)


def test_income_gross_income_missing_prompts_keep_field_hotspots():
    payload = _baozhen_income_payload()
    profile = payload["income_cap_analysis"]["building_profile"]
    for key in (
        "building_form",
        "floor_desc",
        "owner_floor_desc",
        "built_year",
        "structure",
        "entrance_door",
        "windows",
        "security_facilities",
        "exterior",
        "floor_finish",
        "ceiling_finish",
        "newness_desc",
        "building_area_basis",
        "land_area_basis",
        "current_use_basis",
        "current_use_condition",
        "plot_ratio",
        "set_plot_ratio",
    ):
        profile[key] = ""

    class EmptyComparableLibrary:
        pass

    draft = build_valuation_process_draft(payload, EmptyComparableLibrary())
    method = next(item for item in draft["methods"] if item["method_key"] == "income_cap")
    section = next(item for item in method["narratives"] if item["key"] == "income_cap_gross_income_intro")
    prompts = [item for item in section["segments"] if "【请" in item.get("text", "")]

    assert prompts
    assert all(item.get("field") or item.get("fields") for item in prompts)


def test_income_factor_adjustment_narrative_is_index_driven():
    result = calculate_income_capitalization(_baozhen_income_payload())
    text = result["generated_narratives"]["income_cap_rent_factor_basis"]

    assert "用途修正：以估价对象条件指数100为基准，案例A为101.968" in text
    assert "案例A修正系数为100/101.968=0.9807" in text
    assert "用途相同，因此不用做修正" not in text


def test_income_uses_part_two_plot_ratio_and_condition_and_factor_scales():
    payload = _baozhen_income_payload()
    payload["set_plot_ratio"] = "3.50"
    payload["set_plot_ratio_display"] = "3.50"
    payload["valuation_condition_type"] = "规划"
    analysis = payload["income_cap_analysis"]
    analysis["building_profile"]["plot_ratio"] = "9.99"
    analysis["building_profile"]["set_plot_ratio"] = "9.99"
    analysis["building_profile"]["current_use_condition"] = "旧值"
    road = next(item for item in analysis["rent_factor_items"] if item["key"] == "road_type")
    road["subject_value"] = "临主干道"
    road["subject_level_label"] = "临主干道"
    road["cases"]["A"].update({"value": "临次干道", "level_label": "临次干道", "index": "", "source": "manual_override"})

    result = calculate_income_capitalization(payload)

    assert result["building_profile"]["plot_ratio"] == "3.50"
    assert result["building_profile"]["set_plot_ratio"] == "3.50"
    assert result["building_profile"]["current_use_condition"] == "规划利用条件"
    road = next(item for item in result["rent_factor_items"] if item["key"] == "road_type")
    assert [item["label"] for item in road["levels"]] == ["临主干道", "临次干道", "临支路", "不临路"]
    assert road["cases"]["A"]["index"] == "99.00"


def test_income_migrates_legacy_generic_factor_levels_and_recalculates_indexes():
    payload = _baozhen_income_payload()
    analysis = payload["income_cap_analysis"]
    road = next(item for item in analysis["rent_factor_items"] if item["key"] == "road_accessibility")
    road["subject_value"] = "较优"
    road["subject_level_label"] = "较优"
    road["levels"] = [
        {"label": label, "value": label, "quality_score": 5 - index}
        for index, label in enumerate(["优", "较优", "一般", "较劣", "劣"])
    ]
    road["cases"]["A"].update({"value": "优", "level_label": "优", "index": "", "source": "manual_override"})
    environment = next(item for item in analysis["rent_factor_items"] if item["key"] == "environment_quality")
    environment["subject_value"] = "较好"
    environment["subject_level_label"] = "较好"
    environment["cases"]["A"].update({"value": "好", "level_label": "好", "index": "", "source": "manual_override"})

    result = calculate_income_capitalization(payload)

    road = next(item for item in result["rent_factor_items"] if item["key"] == "road_accessibility")
    assert road["subject_value"] == "较通畅"
    assert [item["label"] for item in road["levels"]] == ["通畅", "较通畅", "一般通畅", "较不通畅", "不通畅"]
    assert road["cases"]["A"]["value"] == "通畅"
    assert road["cases"]["A"]["index"] == "102.00"
    environment = next(item for item in result["rent_factor_items"] if item["key"] == "environment_quality")
    assert environment["subject_value"] == "较好"
    assert environment["cases"]["A"]["index"] == "102.00"


def test_income_usage_is_normalized_and_transaction_condition_defaults_to_normal_100():
    payload = {
        "valuation_date": "2026-06-01",
        "land_usage": "居住用地",
        "income_cap_analysis": {
            "rent_instances": [
                {"slot": "A", "usage": "住宅"},
                {"slot": "B", "usage": "住宅用地"},
                {"slot": "C", "usage": "居住用地"},
            ]
        },
    }

    result = calculate_income_capitalization(payload)

    usage = next(item for item in result["rent_factor_items"] if item["key"] == "usage")
    assert usage["subject_value"] == "住宅"
    assert [usage["cases"][slot]["value"] for slot in ("A", "B", "C")] == ["住宅", "住宅", "住宅"]
    assert [usage["cases"][slot]["index"] for slot in ("A", "B", "C")] == ["100.00", "100.00", "100.00"]
    transaction = next(item for item in result["rent_factor_items"] if item["key"] == "transaction_condition")
    assert [item["label"] for item in transaction["levels"]] == ["正常交易", "非正常交易"]
    assert [transaction["cases"][slot]["value"] for slot in ("A", "B", "C")] == ["正常交易", "正常交易", "正常交易"]
    assert [transaction["cases"][slot]["index"] for slot in ("A", "B", "C")] == ["100.00", "100.00", "100.00"]
    assert [transaction["cases"][slot]["source"] for slot in ("A", "B", "C")] == ["system_default", "system_default", "system_default"]


def test_income_non_normal_transaction_preserves_manual_index_and_reason():
    payload = {"valuation_date": "2026-06-01", "land_usage": "住宅用地"}
    seeded = calculate_income_capitalization(payload)
    transaction = next(item for item in seeded["rent_factor_items"] if item["key"] == "transaction_condition")
    transaction["cases"]["A"].update(
        {
            "value": "非正常交易",
            "level_label": "非正常交易",
            "index": "95",
            "override_reason": "关联方交易，经调查按市场水平上调修正",
            "source": "manual_override",
            "confirmed": True,
        }
    )
    payload["income_cap_analysis"] = seeded

    result = calculate_income_capitalization(payload)

    transaction = next(item for item in result["rent_factor_items"] if item["key"] == "transaction_condition")
    assert transaction["subject_value"] == "正常交易"
    assert transaction["cases"]["A"]["index"] == "95"
    assert transaction["cases"]["A"]["override_reason"] == "关联方交易，经调查按市场水平上调修正"
    assert "非正常交易调整依据：关联方交易，经调查按市场水平上调修正" in result["generated_narratives"]["income_cap_rent_factor_basis"]


def test_income_missing_vacancy_range_requires_review():
    payload = _baozhen_income_payload()
    payload["income_cap_analysis"]["income_parameters"]["vacancy_rate_range"] = ""

    result = calculate_income_capitalization(payload)

    assert result["complete"] is False
    assert any("区域平均空置率区间" in item["message"] for item in result["warnings"])
    assert "平均空置率为【请填写空置率区间】" in result["generated_narratives"]["income_cap_annual_gross_narrative"]


def test_income_replacement_cost_range_defaults_to_max_and_preserves_manual_adoption():
    payload = _baozhen_income_payload()
    analysis = payload["income_cap_analysis"]
    analysis["building_profile"]["structure"] = "砖混结构"
    analysis["building_profile"]["floor_desc"] = "4层"
    analysis["expense_parameters"].pop("replacement_base_unit_cost", None)

    result = calculate_income_capitalization(payload)
    expense = result["expense_parameters"]

    assert expense["replacement_cost_standard_key"] == "brick_concrete_1"
    assert expense["replacement_cost_range_min"] == "1300"
    assert expense["replacement_cost_range_max"] == "1400"
    assert expense["replacement_base_unit_cost"] == "1400"
    assert expense["replacement_cost_adopted_source"] == "range_max_default"
    assert "1300-1400元/平方米" in result["generated_narratives"]["income_cap_expense_narrative"]

    payload["income_cap_analysis"] = result
    payload["income_cap_analysis"]["expense_parameters"]["replacement_base_unit_cost"] = "1350"
    payload["income_cap_analysis"]["expense_parameters"]["replacement_cost_adopted_source"] = "manual_override"
    manual = calculate_income_capitalization(payload)

    assert manual["expense_parameters"]["replacement_base_unit_cost"] == "1350"
    assert manual["expense_parameters"]["replacement_cost_adopted_source"] == "manual_override"


def test_income_process_numeric_hotspots_use_explicit_source_refs_without_unit_noise():
    payload = _baozhen_income_payload()
    payload["building_area"] = "107"
    payload["income_cap_analysis"]["building_profile"]["building_area"] = "107"
    payload["income_cap_analysis"]["income_parameters"]["vacancy_rate"] = "4"
    payload["income_cap_analysis"]["income_parameters"]["rentable_area_ratio"] = "96"

    class EmptyComparableLibrary:
        pass

    draft = build_valuation_process_draft(payload, EmptyComparableLibrary())
    method = next(item for item in draft["methods"] if item["method_key"] == "income_cap")
    segments = {
        section["key"]: [
            segment for segment in section.get("segments", [])
            if segment.get("field") or segment.get("fields")
        ]
        for section in method["narratives"]
    }

    annual_fields = {segment.get("field") for segment in segments["income_cap_annual_gross_narrative"]}
    expense_fields = {segment.get("field") for segment in segments["income_cap_expense_narrative"]}
    total_fields = {segment.get("field") for segment in segments["income_cap_total_price_narrative"]}

    assert "income_cap_analysis.income_parameters.vacancy_rate" in annual_fields
    assert "income_cap_analysis.income_parameters.rentable_area_ratio" in annual_fields
    assert "income_cap_analysis.building_profile.building_area" in annual_fields
    assert any(segment["text"] == "107" and segment.get("field") == "income_cap_analysis.building_profile.building_area" for segment in segments["income_cap_annual_gross_narrative"])
    assert "income_cap_analysis.expense_parameters.repair_rate" in expense_fields
    assert "income_cap_analysis.expense_parameters.cost_growth_rate" in expense_fields
    assert "income_cap_analysis.expense_parameters.insurance_rate_permille" in expense_fields
    assert "income_cap_analysis.building_profile.economic_life_years" in expense_fields
    assert "income_cap_analysis.building_profile.used_years" in expense_fields
    assert "income_cap_analysis.building_profile.remaining_years" in expense_fields
    assert "valuation_date" in expense_fields
    assert "income_cap_analysis.income_results.total_land_price" in total_fields
    assert not any(segment["text"] in {"合理", "元/平方米"} for values in segments.values() for segment in values)


def test_income_formula_variable_explanations_are_on_separate_lines():
    result = calculate_income_capitalization(_baozhen_income_payload())
    narratives = result["generated_narratives"]
    text = narratives["income_cap_method_intro"]

    assert "式中：\nP——土地价格；\nA——土地年纯收益；\nr——土地还原利率；\nn——使用土地的年期或有土地收益的年期。" in text
    assert "（1＋r）^n" in narratives["income_cap_total_price_narrative"]
    assert "（1+6.5%）^70" in narratives["income_cap_total_price_narrative"]


def test_income_factor_tables_do_not_emit_spacer_columns():
    result = calculate_income_capitalization(_baozhen_income_payload())
    tables = {table["key"]: table for table in result["tables"]}

    for key in ("income_rent_condition_rows", "income_rent_index_rows"):
        table = tables[key]
        assert table["columns"] == ["比较因素", "估价对象", "案例A", "案例B", "案例C"]
        assert table["header_rows"][0][0]["label"] == "估价对象及比较实例/比较因素"
        assert all(len(row["cells"]) == 5 for row in table["rows"])
        assert all(len(row.get("cell_refs", [])) == 5 for row in table["rows"])
        assert all(not row["cell_refs"][0] for row in table["rows"])

    correction = tables["income_rent_correction_rows"]
    assert correction["columns"] == ["比较因素", "案例A", "案例B", "案例C"]
    assert correction["header_rows"][0][0]["label"] == "估价对象及比较实例/比较因素"
    assert all(len(row["cells"]) == 4 for row in correction["rows"])
    assert all(len(row.get("cell_refs", [])) == 4 for row in correction["rows"])
    assert all(not row["cell_refs"][0] for row in correction["rows"])
    assert not any(row["cells"][0] == "综合修正系数" for row in correction["rows"])
    assert correction["rows"][0]["cells"][0] == "位置"
    assert correction["rows"][-1]["cells"][0] == "算术平均值(元/m2·月)"

    cap_rate = tables["income_cap_rate_rows"]
    assert len(cap_rate["columns"]) == 2
    assert [row["cells"][0] for row in cap_rate["rows"]] == ["土地还原利率", "房屋还原利率"]
    assert [row["cell_refs"][0] for row in cap_rate["rows"]] == ["", ""]


def test_income_docx_replaces_last_method_block_in_runtime_template(tmp_path):
    template = Path(__file__).resolve().parents[1] / "01_Source" / "02_Word_Templates" / "自动生成的评估报告模板.docx"
    output = tmp_path / "income_dynamic.docx"
    shutil.copy(template, output)

    apply_income_capitalization_to_docx(str(output), _baozhen_income_payload())

    rendered = Document(output)
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "住宅用房月租金＝（11.32+11.52+11.54）/3＝11.46" not in text
    assert "住宅用房月租金＝（11.18+11.18+11.27）/3＝11.21" in text
    assert "三、地价的确定" in text
    method = next(paragraph for paragraph in rendered.paragraphs if paragraph.text.strip() == "★收益还原法")
    assert method.paragraph_format.first_line_indent.cm == pytest.approx(0.99, abs=0.01)
    assert method.runs[0].font.size.pt == 14
    table_title = next(paragraph for paragraph in rendered.paragraphs if paragraph.text.strip() == "表3-13 比较因素条件说明表")
    assert table_title.alignment == 1
    assert table_title.runs[0].bold is False
    condition_table = next(
        table
        for table in rendered.tables
        if table.rows and "估价对象及比较实例/比较因素" in " ".join(cell.text for cell in table.rows[0].cells)
    )
    first_cell = condition_table.rows[0].cells[0]
    assert first_cell.vertical_alignment == 1
    assert first_cell.paragraphs[0].alignment == 1
    assert first_cell.paragraphs[0].runs[0].font.size.pt == 10.5


def test_income_evidence_images_fit_bounded_box_and_keep_caption(tmp_path):
    output = tmp_path / "bounded_image.docx"
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    _fill_image_cell(cell, "实例A位置图", PNG_1X1)
    doc.save(output)

    rendered = Document(output)
    assert "实例A位置图" in rendered.tables[0].cell(0, 0).text
    assert len(rendered.inline_shapes) == 1
    shape = rendered.inline_shapes[0]
    assert shape.width <= Cm(7.0)
    assert shape.height <= Cm(4.6)


def test_income_docx_uses_template_order_for_gross_income_blocks(tmp_path):
    output = tmp_path / "income_order.docx"
    doc = Document()
    doc.add_paragraph("★收益还原法")
    doc.add_paragraph("旧内容")
    doc.add_paragraph("★基准地价系数修正法")
    doc.save(output)

    apply_income_capitalization_to_docx(str(output), _baozhen_income_payload())

    rendered = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert body_text.index("①月租金的确定") < body_text.index("②编制租金案例比较因素表")
    assert body_text.index("租金的比较因素修正说明") < body_text.index("③房地年总收益的确定")
    assert body_text.index("③房地年总收益的确定") < body_text.index("2、确定房地年总费用")
