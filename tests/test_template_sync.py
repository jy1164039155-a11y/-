# -*- coding: utf-8 -*-
import os
import re
import zipfile
from docxtpl import DocxTemplate
from src.schemas.land import LandValuationContract


TEMPLATE_DIR = r"D:\评估报告工具\01_Source\02_Word_Templates"


def runtime_template_path():
    cleaned = os.path.join(TEMPLATE_DIR, "自动生成的评估报告模板_清理后.docx")
    default = os.path.join(TEMPLATE_DIR, "自动生成的评估报告模板.docx")
    return cleaned if os.path.exists(cleaned) else default

def run_linter_for_template(template_path):
    """通用 Word Linter 审计算法"""
    assert os.path.exists(template_path), f"模板文件不存在: {template_path}"
    
    # 1. 自动提取 Word 模板中所有未声明的 {{ 占位变量 }} 集合
    tpl = DocxTemplate(template_path)
    word_vars = tpl.get_undeclared_template_variables()
    
    # 2. 提取 Python 契约中的字段集
    contract_fields = set(LandValuationContract.model_fields.keys())
    
    # 3. 注入地理、权重、及住宅/工业特有条件智能派生白名单字段集
    derived_fields = {
        # 地理与机构智能派生
        "planning_approval_authority", "local_county", "local_county_full", "local_gov", "client_name", "county_name",
        
        # 测算方法与权重合成派生
        "use_cost_approx", "use_cost_approach", "use_market_comp", "use_income_cap", 
        "use_benchmark_corr", "use_residual", "show_price_in_text",
        "cost_approx_price", "market_comp_price", "income_cap_price", "benchmark_corr_price", 
        "residual_price", "weight_logic_type", "dominant_method_name", 
        "formula_display_text", "weight_rationale_text", "adopted_methods_summary",
        
        # 二级内存微渲染长文本方案字段
        "valuation_method_applicability", "final_price_determination", "infrastructure_detail",
        
        # 住宅/新增用地历史权源骨架特有条件派生
        "land_status_type", "gov_approval_name", "gov_approval_no", "gov_approval_date",
        "original_land_owner_desc", "approval_transfer_date", "approval_authority",
        "house_cert_name", "house_cert_no", "area_docs_desc_name", "room_detail_location",
        "proof_doc_name", "proof_doc_date", "buy_year", "buy_location_desc",
        "registered_house_area", "land_cert_name", "land_cert_no", "land_use_type",
        "is_base_price_expired", "base_land_price_expire_limit",
        "final_total_price_cny_upper", "final_total_price_wan", "final_unit_price"
    }
    
    system_allowed_outputs = contract_fields.union(derived_fields)
    
    # 4. 抓捕 Word 里的幽灵变量
    ghost_vars = word_vars - system_allowed_outputs
    
    # 排除部分可能属于 excel 直接填报的特殊临时白名单（如 local_county_full_name 等，若有）
    ghost_vars = {v for v in ghost_vars if not v.startswith("var_")}
    
    if ghost_vars:
        print(f"\n[ERROR] [审计警告] 在模板 {os.path.basename(template_path)} 中发现未在契约中定义的幽灵变量:")
        for var in ghost_vars:
            print(f"  - {var}")
            
    assert not ghost_vars, f"[WARNING] Word 模板与 Python 代码契约脱节！模板中包含了系统未定义的幽灵变量: {ghost_vars}"
    print(f"[SUCCESS] [审计通过] 模板 {os.path.basename(template_path)} 契约对齐完美，无任何幽灵变量！")

def test_industrial_template_sync():
    """审计：工业合并报告模板契约对齐状况"""
    template_path = r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx"
    run_linter_for_template(template_path)

def test_residential_template_sync():
    """审计：住宅合并报告模板契约对齐状况"""
    template_path = r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx"
    run_linter_for_template(template_path)


def test_runtime_template_sync():
    """审计：当前实际渲染母版契约对齐状况"""
    run_linter_for_template(runtime_template_path())


def test_runtime_template_uses_generated_cost_land_class_intro():
    """审计：成本逼近法征收地类说明由正式字段生成，运行母版不得写死具体地类。"""
    from docx import Document

    document = Document(runtime_template_path())
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    assert paragraphs.count("{{ cost_approx_land_class_intro }}") == 1
    assert not any("评估宗地征收地类为林地" in paragraph for paragraph in paragraphs)


def test_template_basis_docs_are_not_project_hardcoded():
    """审计：泛型依据文件字段不得再混入单一项目批复硬编码。"""
    from docx import Document

    hardcoded_doc = "关于道县2025年第七批次建设用地农用地转用和土地征收的批复"
    template_paths = [
        r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx",
        r"D:\评估报告工具\01_Source\02_Word_Templates\自动生成的评估报告模板.docx",
    ]
    for template_path in template_paths:
        doc = Document(template_path)
        all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        assert hardcoded_doc not in all_text


def test_active_templates_use_generated_method_text_blocks():
    """审计：活动模板不得再残留估价方法 use_* 条件块，确价正文由前端/后端生成的大段落接管。"""
    template_paths = [
        r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx",
        r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx",
        runtime_template_path(),
    ]
    for template_path in template_paths:
        with zipfile.ZipFile(template_path) as archive:
            xml = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            )
        assert not re.search(r"{%\s*if\s+(?:not\s+)?use_", xml)
        assert "valuation_method_reasons" in xml
        assert "valuation_method_applicability" in xml


def test_templates_have_no_hardcoded_base_price_policies():
    """自动化审计：全模板库所有活动及合并模板中，均不得再混入写死的基准地价文号、日期或地名等特定项目硬编码。"""
    from docx import Document
    
    template_paths = [
        r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx",
        r"D:\评估报告工具\01_Source\02_Word_Templates\[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx",
        r"D:\评估报告工具\01_Source\02_Word_Templates\自动生成的评估报告模板.docx",
    ]
    
    blacklisted = [
        "道政函〔2019〕85号",
        "2019年11月22日",
        "2019年5月1日",
        "通道侗族自治县自然资源局"
    ]
    
    for template_path in template_paths:
        if not os.path.exists(template_path):
            continue
        doc = Document(template_path)
        all_text = []
        for p in doc.paragraphs:
            all_text.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
                    
        doc_content = "\n".join(all_text)
        for item in blacklisted:
            assert item not in doc_content, f"在 Word 模板 {os.path.basename(template_path)} 中审计到未参数化的特定地区政策硬编码: '{item}'，请将其修改为 Jinja2 变量！"


def test_runtime_market_comparison_block_has_no_sample_cases_or_prices():
    """运行母版的市场比较法必须完全由结构化分析生成。"""
    from docx import Document

    doc = Document(runtime_template_path())
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    for forbidden in (
        "4311242024B000055",
        "4311242024B000016",
        "4311242024B000071",
        "湖南雁晟农业科技有限公司",
        "127.7",
        "128.3",
        "0.645",
    ):
        assert forbidden not in text
    assert "{{ market_comp_step1_instances }}" in text
    assert "{{ market_comp_step4_solve }}" in text
