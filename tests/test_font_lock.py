# -*- coding: utf-8 -*-
import shutil

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from src.api import (
    apply_site_photos_to_docx,
    lock_docx_cjk_fonts,
    normalize_generated_body_paragraph_fonts,
    normalize_generated_text_blocks,
    normalize_report_typography,
    normalize_result_summary_tables,
)


def _east_asia_font(run):
    r_pr = run._r.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:eastAsia"))


def test_lock_docx_cjk_fonts_fills_missing_chinese_font(tmp_path):
    path = tmp_path / "missing_font.docx"
    doc = Document()
    doc.add_paragraph().add_run("变量替换文字")
    doc.save(path)

    lock_docx_cjk_fonts(str(path))

    rendered = Document(path)
    run = rendered.paragraphs[0].runs[0]
    assert _east_asia_font(run) == "仿宋_GB2312"
    assert run.font.size is None


def test_lock_docx_cjk_fonts_preserves_existing_cjk_font(tmp_path):
    path = tmp_path / "existing_font.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("标题文字")
    run.font.name = "Calibri"
    r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "SimSun")
    doc.save(path)

    lock_docx_cjk_fonts(str(path))

    rendered = Document(path)
    run = rendered.paragraphs[0].runs[0]
    assert _east_asia_font(run) == "SimSun"


def test_normalize_generated_text_blocks_splits_numbered_lines(tmp_path):
    path = tmp_path / "numbered_lines.docx"
    doc = Document()
    doc.add_paragraph(
        "1. 土地权利类型：国有建设用地出让土地使用权；\n"
        "2. 土地权利性质：国有储备用地，现拟办理出让手续；"
    )
    doc.save(path)

    normalize_generated_text_blocks(str(path))

    rendered = Document(path)
    texts = [paragraph.text for paragraph in rendered.paragraphs]
    assert texts == [
        "1. 土地权利类型：国有建设用地出让土地使用权；",
        "2. 土地权利性质：国有储备用地，现拟办理出让手续；",
    ]
    for paragraph in rendered.paragraphs:
        assert paragraph.paragraph_format.left_indent is not None
        assert paragraph.paragraph_format.first_line_indent.pt == 0
        assert paragraph.runs[0].font.size.pt == 14


def test_normalize_generated_text_blocks_inherits_explicit_run_size(tmp_path):
    path = tmp_path / "explicit_size.docx"
    doc = Document()
    run = doc.add_paragraph().add_run("1. 第一行\n2. 第二行")
    run.font.size = Pt(11)
    doc.save(path)

    normalize_generated_text_blocks(str(path))

    rendered = Document(path)
    assert [p.runs[0].font.size.pt for p in rendered.paragraphs] == [11, 11]


def test_normalize_generated_text_blocks_inherits_style_size(tmp_path):
    path = tmp_path / "style_size.docx"
    doc = Document()
    doc.styles["Normal"].font.size = Pt(12)
    doc.add_paragraph("1. 第一行\n2. 第二行")
    doc.save(path)

    normalize_generated_text_blocks(str(path))

    rendered = Document(path)
    assert [p.runs[0].font.size.pt for p in rendered.paragraphs] == [12, 12]


def test_normalize_generated_text_blocks_skips_table_paragraphs(tmp_path):
    path = tmp_path / "table_text.docx"
    doc = Document()
    cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run("1. 表内第一行\n2. 表内第二行")
    doc.save(path)

    normalize_generated_text_blocks(str(path))

    rendered = Document(path)
    cell_paragraphs = rendered.tables[0].cell(0, 0).paragraphs
    assert len(cell_paragraphs) == 1
    assert cell_paragraphs[0].text == "1. 表内第一行\n2. 表内第二行"


def test_normalize_generated_body_paragraph_fonts_targets_generated_lines_only(tmp_path):
    path = tmp_path / "generated_body_font.docx"
    doc = Document()
    title_run = doc.add_paragraph().add_run("Report Title")
    title_run.font.size = Pt(18)
    body_run = doc.add_paragraph().add_run("Generated valuation body line")
    body_run.font.size = Pt(9)
    doc.save(path)

    normalize_generated_body_paragraph_fonts(
        str(path),
        {"valuation_method_reasons": "Generated valuation body line"},
    )

    rendered = Document(path)
    assert rendered.paragraphs[0].runs[0].font.size.pt == 18
    assert rendered.paragraphs[1].runs[0].font.size.pt == 14
    assert _east_asia_font(rendered.paragraphs[1].runs[0]) == "仿宋_GB2312"


def test_normalize_report_typography_fixes_template_inherited_body_and_titles(tmp_path):
    path = tmp_path / "template_inherited_typography.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10.5)

    title_run = doc.add_paragraph().add_run("\u571f \u5730 \u4f30 \u4ef7 \u62a5 \u544a")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(10.5)
    doc.add_paragraph("\u7b2c\u4e00\u90e8\u5206  \u6458\u8981")
    doc.add_paragraph("\u6839\u636e\u59d4\u6258\u65b9\u63d0\u4f9b\u8d44\u6599\uff0c\u672c\u6b21\u4f30\u4ef7\u9075\u5faa\u5ba2\u89c2\u539f\u5219\u3002")
    doc.add_paragraph("1. \u571f\u5730\u6743\u5229\u7c7b\u578b\u4e3a\u56fd\u6709\u5efa\u8bbe\u7528\u5730\u4f7f\u7528\u6743\u3002")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "\u4f30\u4ef7\u671f\u65e5\u571f\u5730\u4f7f\u7528\u8005"
    doc.save(path)

    normalize_report_typography(str(path))

    rendered = Document(path)
    title = rendered.paragraphs[0]
    part_heading = rendered.paragraphs[1]
    body = rendered.paragraphs[2]
    numbered_body = rendered.paragraphs[3]
    table_run = rendered.tables[0].cell(0, 0).paragraphs[0].runs[0]

    assert title.runs[0].font.size.pt == 22
    assert title.runs[0].bold is True
    assert _east_asia_font(title.runs[0]) == "\u5b8b\u4f53"
    assert part_heading.runs[0].font.size.pt == 16
    assert _east_asia_font(part_heading.runs[0]) == "\u5b8b\u4f53"
    assert body.runs[0].font.size.pt == 14
    assert body.paragraph_format.line_spacing.pt == 24
    assert _east_asia_font(body.runs[0]) == "\u4eff\u5b8b_GB2312"
    assert numbered_body.runs[0].bold is None
    assert numbered_body.runs[0].font.size.pt == 14
    assert _east_asia_font(numbered_body.runs[0]) == "\u4eff\u5b8b_GB2312"
    assert table_run.font.size is None
    assert _east_asia_font(table_run) is None


def test_normalize_report_typography_aligns_method_headings_and_table_titles(tmp_path):
    path = tmp_path / "report_roles.docx"
    doc = Document()
    doc.add_paragraph("★收益还原法")
    doc.add_paragraph("表3-13 比较因素条件说明表")
    doc.save(path)

    normalize_report_typography(str(path))

    rendered = Document(path)
    method_heading, table_title = rendered.paragraphs
    assert method_heading.paragraph_format.first_line_indent.cm == pytest.approx(0.99, abs=0.01)
    assert method_heading.paragraph_format.line_spacing.pt == 24
    assert method_heading.runs[0].bold is True
    assert table_title.alignment == 1
    assert table_title.paragraph_format.first_line_indent.pt == 0
    assert table_title.paragraph_format.space_before.pt == pytest.approx(7.85, abs=0.05)
    assert table_title.paragraph_format.space_after.pt == pytest.approx(0.5, abs=0.05)
    assert table_title.runs[0].font.size.pt == 14
    assert table_title.runs[0].bold is False


def test_apply_site_photos_removes_intro_when_no_photos(tmp_path):
    path = tmp_path / "no_photos.docx"
    doc = Document()
    doc.add_paragraph("估价对象利用现状如下：")
    doc.add_paragraph("2.规划利用条件")
    doc.save(path)

    apply_site_photos_to_docx(str(path), {})

    rendered = Document(path)
    assert "估价对象利用现状如下：" not in [p.text for p in rendered.paragraphs]
    assert "2.规划利用条件" in [p.text for p in rendered.paragraphs]


def test_apply_site_photos_inserts_picture_and_caption(tmp_path):
    path = tmp_path / "with_photos.docx"
    doc = Document()
    doc.add_paragraph("估价对象利用现状如下：")
    doc.add_paragraph("2.规划利用条件")
    doc.save(path)

    tiny_png = (
        "data:image/png;base64,"
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lWv1LgAAAABJRU5ErkJggg=="
    )
    apply_site_photos_to_docx(
        str(path),
        {"site_photo_data_urls": [tiny_png], "site_photo_captions": ["图1 现场照片"]},
    )

    rendered = Document(path)
    assert len(rendered.inline_shapes) == 1
    assert "图1 现场照片" in [p.text for p in rendered.paragraphs]


def test_normalize_result_summary_tables_sets_small_font_on_result_table(tmp_path):
    template = "D:/评估报告工具/01_Source/02_Word_Templates/自动生成的评估报告模板.docx"
    path = tmp_path / "template_copy.docx"
    shutil.copy(template, path)
    normalize_result_summary_tables(str(path))
    doc = Document(path)
    result_tables = [
        table for table in doc.tables
        if "估价期日土地使用者" in "\n".join(cell.text for row in table.rows[:3] for cell in row.cells[:18])
    ]
    assert result_tables
    first_run = result_tables[0].rows[0].cells[0].paragraphs[0].runs[0]
    assert _east_asia_font(first_run) == "仿宋_GB2312"
    assert first_run.font.size.pt == 8
