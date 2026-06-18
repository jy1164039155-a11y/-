# -*- coding: utf-8 -*-
from pathlib import Path

import pytest
from docx import Document

from src.services.benchmark_correction import calculate_benchmark_correction
from src.services.benchmark_correction_docx import apply_benchmark_correction_to_docx
from src.services.tongdao_benchmark_source import (
    _cap_rates_from_other_table,
    _extract_cap_rate_tables,
)
from src.services.valuation_process_builder import build_valuation_process_draft

BASE_DIR = Path(__file__).resolve().parents[1]


def _tongdao_payload():
    """通道县范例金标准。样本因素必须显式提供，正式配置不得回填具体项目条件。"""
    return {
        "county_name": "通道县",
        "use_benchmark_corr": True,
        "land_usage": "住宅用地",
        "land_level": "二级",
        "plot_ratio": "4.58",
        "land_use_term": "63.97",
        "valuation_date": "2026-06-03",
        "land_area": "19.02",
        "benchmark_corr_analysis": {
            "area_grade": "劣",
            "shape_grade": "较优",
            "landscape_grade": "优",
            "orientation": "南",
            "development_adjustment": "0",
            "development_note": "待估宗地设定开发程度（宗地红线外“五通”、红线内场地平整）与基准地价内涵开发程度一致，故不需开发程度修正，Kf=0。",
            "region_factor_selections": [
                {"sub_factor": "距小区级商服中心距离", "grade": "较优", "description": "距寨上路商贸广场380米"},
                {"sub_factor": "距汽车站距离", "grade": "较劣", "description": "距通道汽车站1200米"},
                {"sub_factor": "临近道路状况", "grade": "较劣", "description": "临次干道育才路"},
                {"sub_factor": "距公交站点距离", "grade": "一般", "description": "距农商行公交站270米"},
                {"sub_factor": "供水保证率", "grade": "优", "description": "≥98"},
                {"sub_factor": "供电保证率", "grade": "优", "description": "≥98"},
                {"sub_factor": "排水状况", "grade": "优", "description": "雨水大时无洪涝积水，排水畅通"},
                {"sub_factor": "公用设施完善度", "grade": "优", "description": "完善"},
                {"sub_factor": "人口密度", "grade": "较优", "description": "4000-4500"},
                {"sub_factor": "噪声污染状况", "grade": "优", "description": "低"},
                {"sub_factor": "水污染状况", "grade": "优", "description": "低"},
                {"sub_factor": "道路规划", "grade": "一般", "description": "生活型次干道"},
            ],
        },
    }


def _result_value(results, key):
    return next(item["value"] for item in results if item["key"] == key)


def test_benchmark_correction_reproduces_specimen_factors():
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    params = result["parameters"]
    # 范例金标准：各修正系数
    assert params["base_land_price"] == "710"
    assert params["kv"] == "1.4369"
    assert params["ky"] == "0.9950"
    assert params["kt"] == "1.0285"
    assert params["sum_ki"] == "2.4020"
    assert params["ks"] == "0.98"
    assert params["ka"] == "1.02"
    assert params["ke"] == "1.05"
    assert params["kto"] == "1.02"
    assert params["kf"] == "0"


def test_benchmark_correction_reproduces_final_price():
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    assert result["benchmark_corr_price"] == "1144.6"
    assert _result_value(result["results"], "benchmark_corr_price") == "1144.6"
    assert result["complete"] is True


def test_benchmark_correction_base_price_narrative_includes_connotation_lead_sentence():
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    narrative = result["generated_narratives"]["benchmark_corr_base_price"]
    assert "地类按商业服务业、居住、公共管理与公共服务、工矿、仓储、公用设施用地区分" in narrative
    assert "各类用地的使用年期均为法定最高使用年限" in narrative
    assert "通道县城区基准地价及其内涵具体情况见下表" in narrative


def test_benchmark_correction_other_cap_rates_are_read_from_table_values():
    parsed = _cap_rates_from_other_table({
        "rows": [
            {
                "土地(房屋)类型": "土地还原率",
                "工矿、仓储用地": "5.81%",
                "公共管理与公共服务用地": "6.26%",
                "公用设施用地": "5.96%",
            }
        ]
    })
    assert parsed == {
        "工矿用地": "5.81",
        "仓储用地": "5.81",
        "公共管理与公共服务用地": "6.26",
        "公用设施用地": "5.96",
    }


def test_extract_cap_rate_tables_uses_parsed_other_table_values():
    """四类还原率以技术报告解析值为准，覆盖内置默认值。"""
    elements = [
        {
            "type": "table",
            "title": "表5-8 工矿、仓储、公共管理与公共服务、公用设施用地还原率表",
            "rows": [
                ["土地(房屋)类型", "工矿、仓储用地", "公共管理与公共服务用地", "公用设施用地"],
                ["土地还原率", "5.83%", "6.28%", "5.98%"],
            ],
        }
    ]
    extracted = _extract_cap_rate_tables(elements)
    assert extracted["cap_rate"]["工矿用地"] == "5.83"
    assert extracted["cap_rate"]["仓储用地"] == "5.83"
    assert extracted["cap_rate"]["公共管理与公共服务用地"] == "6.28"
    assert extracted["cap_rate"]["公用设施用地"] == "5.98"


def test_extract_cap_rate_tables_warns_when_other_table_missing():
    """解析失败不得静默：缺表时显式告警，提醒核对报告原值。"""
    with pytest.warns(UserWarning, match="还原率表"):
        _extract_cap_rate_tables([])


def test_benchmark_correction_docx_missing_start_marker_does_not_silently_skip(tmp_path):
    output = tmp_path / "missing_marker.docx"
    doc = Document()
    doc.add_paragraph("没有基准地价方法标记的模板")
    doc.save(output)

    with pytest.raises(RuntimeError, match="起始标记"):
        apply_benchmark_correction_to_docx(str(output), _tongdao_payload(), str(BASE_DIR))


def test_benchmark_correction_docx_missing_start_marker_when_unused_warns_not_silent(tmp_path):
    """未启用该方法且无起始标记时不得静默跳过：应显式告警，便于核对样本残留。"""
    output = tmp_path / "missing_marker_unused.docx"
    doc = Document()
    doc.add_paragraph("没有基准地价方法标记的模板")
    doc.save(output)

    payload = _tongdao_payload()
    payload["use_benchmark_corr"] = False
    with pytest.warns(UserWarning, match="起始标记"):
        apply_benchmark_correction_to_docx(str(output), payload, str(BASE_DIR))


def test_benchmark_correction_docx_missing_stop_marker_cleans_sample_and_warns(tmp_path):
    """缺少下一方法停止标记时：仍清理样本整段并告警，不得残留模板样本。"""
    output = tmp_path / "missing_stop.docx"
    doc = Document()
    doc.add_paragraph("★基准地价系数修正法")
    doc.add_paragraph("样本残留内容AAA")
    doc.add_paragraph("样本残留内容BBB")
    doc.save(output)

    payload = _tongdao_payload()
    payload["use_benchmark_corr"] = False
    with pytest.warns(UserWarning, match="停止标记"):
        apply_benchmark_correction_to_docx(str(output), payload, str(BASE_DIR))

    reopened = Document(str(output))
    texts = "\n".join(p.text for p in reopened.paragraphs)
    assert "样本残留内容AAA" not in texts
    assert "样本残留内容BBB" not in texts
    assert "★基准地价系数修正法" not in texts


def test_benchmark_correction_docx_start_marker_ignores_body_sentence(tmp_path):
    """防起始标记误匹配正文句子：含句号的正文不得被当作方法起始标记。"""
    from src.services.benchmark_correction_docx import _marker

    doc = Document()
    doc.add_paragraph("基准地价系数修正法估价过程是估价的重要方法之一。")
    real_heading = doc.add_paragraph("★基准地价系数修正法")

    matched = _marker(
        doc,
        ["★基准地价系数修正法", "★公示地价系数修正法", "（一）基准地价系数修正法", "基准地价系数修正法估价过程"],
    )
    assert matched is not None
    assert matched._p is real_heading._p


def test_benchmark_correction_solve_narrative_substitution():
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    solve = result["generated_narratives"]["benchmark_corr_solve"]
    # 公式代入行严格对齐范例 =710×（1+2.4020%）×0.9950×1.4369×1.0285×0.98×1.05×1.02×1.02
    assert "710×（1+2.4020%）×0.9950×1.4369×1.0285×0.98×1.02×1.05×1.02" in solve
    assert "1144.6" in solve


def test_benchmark_correction_date_adjustment_self_consistent_no_typo_residue():
    """G1 回归守卫：期日段叙述与表3-5 统一为 14.07月/0.2%/Kt=1.0285，互相自洽，
    且不得残留范例原件笔误 12.73/1.0258/0.3%。"""
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    params = result["parameters"]
    assert params["months_elapsed"] == "14.07"
    assert params["monthly_growth_rate"] == "0.2"
    assert params["kt"] == "1.0285"
    kt_narrative = result["generated_narratives"]["benchmark_corr_kt"]
    assert "14.07" in kt_narrative
    assert "0.2%" in kt_narrative
    assert "Kt=(1+0.2%)^14.07=1.0285" in kt_narrative
    date_table = next(t for t in result["tables"] if t["key"] == "benchmark_date_adjustment_table")
    table_blob = "".join(str(cell) for row in date_table["rows"] for cell in row.values())
    full_blob = kt_narrative + table_blob
    for typo in ("12.73", "1.0258", "0.3%"):
        assert typo not in full_blob


def test_benchmark_correction_segment_sources_field_convention():
    # 与 CH-3 前端 benchmarkFocusId 约定：方法内部参数前缀 benchmark_corr_analysis.parameters.，
    # 顶层 valuation_date 保持裸键（沿用 i4 规范）。
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    sources = result["narrative_segment_sources"]
    kt_fields = {s["field"] for s in sources["benchmark_corr_kt"]}
    assert "valuation_date" in kt_fields
    assert "benchmark_corr_analysis.parameters.kt" in kt_fields
    solve_fields = {s["field"] for s in sources["benchmark_corr_solve"]}
    expected = {
        "benchmark_corr_analysis.parameters.base_land_price",
        "benchmark_corr_analysis.parameters.sum_ki",
        "benchmark_corr_analysis.parameters.ky",
        "benchmark_corr_analysis.parameters.kv",
        "benchmark_corr_analysis.parameters.kt",
        "benchmark_corr_price",
    }
    assert expected <= solve_fields


def test_benchmark_correction_region_factor_options_auto_coefficient():
    # 优劣度选择即自动取对应修正率：每个区域因素都带 options，且当前 grade 的系数与选中项一致。
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    selections = result["region_factor_selections"]
    assert selections
    for row in selections:
        assert row["options"]
        chosen = next((o for o in row["options"] if o["grade"] == row["grade"]), None)
        assert chosen is not None
        assert chosen["coefficient"] == row["coefficient"]
    # 个别因素同样暴露 options 供优劣度联动。
    individual = result["individual_factors"]
    assert individual["area"]["options"] and individual["shape"]["options"]
    assert individual["orientation"]["options"]


def test_benchmark_correction_region_factor_table_sum():
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    region_table = next(t for t in result["tables"] if t["key"] == "benchmark_region_factor_table")
    total_row = region_table["rows"][-1]
    assert total_row["因素"] == "合计"
    assert total_row["修正系数（%）"] == "2.4020"


def test_benchmark_correction_prefers_current_regional_factors_over_stale_compat_field():
    payload = _tongdao_payload()
    seed = calculate_benchmark_correction(payload, BASE_DIR)
    payload["benchmark_corr_analysis"]["region_factor_selections"] = [
        {"sub_factor": row["sub_factor"], "grade": ""}
        for row in seed["regional_factors"]
    ]
    payload["benchmark_corr_analysis"]["regional_factors"] = [
        {"sub_factor": row["sub_factor"], "level": "一般", "grade": "一般"}
        for row in seed["regional_factors"]
    ]
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert {row["grade"] for row in result["regional_factors"]} == {"一般"}
    assert result["parameters"]["sum_ki"] == "0.0000"


def test_benchmark_correction_blocks_non_tongdao_county():
    payload = _tongdao_payload()
    payload["county_name"] = "道县"
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert result["support_status"] == "unsupported"
    assert result["benchmark_corr_price"] == ""
    assert "仅适配通道县" in result["warnings"][0]["message"]


def test_benchmark_correction_statutory_term_matches_known_usage_rule():
    payload = _tongdao_payload()
    payload["land_usage"] = "交通运输用地"
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["parameters"]["legal_term_years"] == "50"


def test_benchmark_correction_unknown_usage_requires_statutory_term_check():
    payload = _tongdao_payload()
    payload["land_usage"] = "其他用途"
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["parameters"]["legal_term_years"] == ""
    assert any("法定最高出让年期" in item for item in result["support_missing_items"])


def test_benchmark_correction_non_residential_formula_does_not_reuse_residential_factors():
    payload = _tongdao_payload()
    payload["land_usage"] = "工矿用地"
    payload["land_use_term"] = "50"
    result = calculate_benchmark_correction(payload, BASE_DIR)
    formula = result["generated_narratives"]["benchmark_corr_formula"]
    assert "P工=" in formula
    assert "Kv" not in formula
    assert "Ke" not in formula
    assert "Kto" not in formula
    block_keys = {block["key"] for block in result["content_blocks"]}
    assert "benchmark_corr_kv" not in block_keys
    assert "benchmark_corr_ke" not in block_keys
    assert "benchmark_corr_kto" not in block_keys


def _commercial_payload(frontage_mode="non_street"):
    payload = _tongdao_payload()
    payload["land_usage"] = "商业服务业用地"
    payload["land_use_term"] = "40"
    payload["plot_ratio"] = "2"
    payload["land_area"] = "1000"
    payload["benchmark_corr_analysis"] = {
        "frontage_mode": frontage_mode,
        "area_grade": "一般",
        "shape_grade": "一般",
        "ku_grade": "一般",
        "development_adjustment": "0",
    }
    return payload


def _fill_commercial_region_factors(payload):
    seed = calculate_benchmark_correction(payload, BASE_DIR)
    payload["benchmark_corr_analysis"]["region_factor_selections"] = [
        {"sub_factor": row.get("sub_factor") or row.get("indicator"), "grade": "一般"}
        for row in seed.get("regional_factors") or []
    ]
    return payload


def test_benchmark_correction_commercial_requires_ku_confirmation():
    payload = _commercial_payload()
    payload["benchmark_corr_analysis"]["ku_grade"] = ""
    result = calculate_benchmark_correction(_fill_commercial_region_factors(payload), BASE_DIR)
    assert result["complete"] is False
    assert result["benchmark_corr_price"] == ""
    assert any("Ku" in item for item in result["support_missing_items"])
    assert "待补齐后生成完整求价过程" in result["generated_narratives"]["benchmark_corr_solve"]


def test_benchmark_correction_commercial_non_street_uses_level_price_and_region_factors():
    payload = _fill_commercial_region_factors(_commercial_payload("non_street"))
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert result["benchmark_corr_price"]
    assert result["parameters"]["formula_text"] == "2、商业服务业用地（不临街）\nP商=级别基准地价×（1+∑Ki）×Ky×Kv×Kt×Ks×Ka×Ku＋Kf"
    assert "route_price" not in result["parameters"] or result["parameters"]["route_price"] in {"", "0"}
    assert result["parameters"]["ku"] == "1.0000"


def test_benchmark_correction_commercial_route_uses_route_price_without_region_factors():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "80"
    payload["benchmark_corr_analysis"].update({
        "route_segment_id": "B25",
        "route_road_type": "支路",
        "frontage_depth_m": "999",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert result["benchmark_corr_price"]
    assert result["parameters"]["formula_text"] == "1、商业服务业用地（临街）\nP商=路线价×Ky×Kv×Kt×Ks×Ka×Kc×Kk×Kd×Ku＋Kf"
    assert "∑Ki" not in result["generated_narratives"]["benchmark_corr_formula"]
    assert result["parameters"]["route_price"] == "2150"
    assert result["parameters"]["route_road_type"] != "支路"
    assert result["parameters"]["frontage_depth_m"] == "20"
    assert result["parameters"]["kd"] == "1.0000"
    assert result["parameters"]["kk"] == "1.0000"
    narratives = result["generated_narratives"]
    assert "根据估价对象实际情况，本次评估采用基准地价系数修正法评估商业服务业用地土地价格的基本计算公式为" in narratives["benchmark_corr_formula"]
    assert "长征南路（双江大桥至内环路）路线价段" in narratives["benchmark_corr_route_price"]
    assert "商服用地临街宗地特别因素主要包括修正面积分配、街角地、临街深度、临街宽度" in narratives["benchmark_corr_frontage_special_intro"]
    assert "当宗地实际临街深度与路线价设定标准深度不一致时，应进行临街深度修正" in narratives["benchmark_corr_kd"]
    assert "故不需进行街角地修正" in narratives["benchmark_corr_kc"]
    assert narratives["benchmark_corr_solve"].startswith("临街商业服务业用地P商=")
    order = [item["key"] for item in result["content_blocks"] if item["type"] == "narrative"]
    assert order.index("benchmark_corr_formula") < order.index("benchmark_corr_po")
    assert "benchmark_corr_route_price" not in order
    assert order.index("benchmark_corr_kt") < order.index("benchmark_corr_frontage_special_intro") < order.index("benchmark_corr_ks")
    assert order.index("benchmark_corr_frontage_special_intro") < order.index("benchmark_corr_kc") < order.index("benchmark_corr_kd") < order.index("benchmark_corr_kk") < order.index("benchmark_corr_ks") < order.index("benchmark_corr_ka")


def test_benchmark_correction_corner_route_prices_derive_kc_inputs():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "80"
    payload["benchmark_corr_analysis"].update({
        "route_segment_id": "B25",
        "frontage_width_m": "4",
        "is_corner": True,
        "corner_route_price_a": "2150",
        "corner_route_price_b": "1000",
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert result["parameters"]["corner_main_route_price"] == "2150"
    assert result["parameters"]["corner_side_route_price"] == "1000"
    assert result["parameters"]["corner_price_ratio"] == "2.15"
    assert result["parameters"]["kc"]
    assert "相邻两条路线价分别为2150元/平方米、1000元/平方米" in result["generated_narratives"]["benchmark_corr_kc"]


def test_benchmark_correction_area_table_uses_body_cell_spans():
    result = calculate_benchmark_correction(_tongdao_payload(), BASE_DIR)
    area_table = next(item for item in result["tables"] if item["key"] == "benchmark_area_table")
    assert area_table["columns"][:2] == ["用地类型", "项目"]
    assert area_table["rows"][1]["cells"][0]["rowspan"] == 2
    assert area_table["rows"][2]["cells"][0]["hidden"] is True


def test_benchmark_correction_commercial_single_route_rejects_depth_beyond_standard():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "100"
    payload["benchmark_corr_analysis"].update({
        "route_segment_id": "B25",
        "frontage_width_m": "4",
        "frontage_relation": "adjacent",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert result["benchmark_corr_price"] == ""
    assert result["parameters"]["frontage_depth_m"] == "25"
    assert any("切换为临街+不临街拆分路径" in item for item in result["support_missing_items"])


def test_benchmark_correction_commercial_manual_road_grade_matches_road_type():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "72"
    payload["benchmark_corr_analysis"].update({
        "route_price": "2000",
        "route_price_source_note": "按路线价图人工判读。",
        "route_road_grade": "生活型主干道",
        "frontage_relation": "setback",
        "frontage_depth_m": "18",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert result["parameters"]["route_road_type"] == "主干道"
    assert result["parameters"]["route_road_type_source"] == "规划道路等级匹配"
    assert result["parameters"]["frontage_depth_source"] == "manual"
    assert result["parameters"]["frontage_depth_m"] == "18"


def test_benchmark_correction_commercial_setback_depth_beyond_standard_rejects_route_price():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "100"
    payload["benchmark_corr_analysis"].update({
        "route_price": "2000",
        "route_price_source_note": "按路线价图人工判读。",
        "route_road_grade": "生活型主干道",
        "frontage_relation": "setback",
        "frontage_depth_m": "25",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert result["benchmark_corr_price"] == ""
    assert any("不宜采用路线价测算" in item for item in result["support_missing_items"])


def test_benchmark_correction_commercial_split_route_and_non_street_weights_components():
    payload = _fill_commercial_region_factors(_commercial_payload("route_price_plus_non_street"))
    payload["land_area"] = "100"
    payload["benchmark_corr_analysis"].update({
        "route_segment_id": "B25",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert result["frontage_mode"] == "route_price_plus_non_street"
    assert result["parameters"]["formula_text"] == "1、商业服务业用地（临街）\nP商=路线价×Ky×Kv×Kt×Ks×Ka×Kc×Kk×Kd×Ku＋Kf\n2、商业服务业用地（不临街）\nP商=级别基准地价×（1+∑Ki）×Ky×Kv×Kt×Ks×Ka×Ku＋Kf"
    assert result["parameters"]["frontage_standard_depth_m"] == "20"
    assert result["parameters"]["frontage_area_m2"] == "80"
    assert result["parameters"]["non_frontage_area_m2"] == "20"
    assert result["parameters"]["frontage_depth_m"] == "25"
    assert result["result_values"]["route_component_price"]
    assert result["result_values"]["non_street_component_price"]
    assert result["result_values"]["commercial_weighted_price"] == result["benchmark_corr_price"]
    narratives = result["generated_narratives"]
    assert "临街商服用地P商" in narratives["benchmark_corr_solve"]
    assert "不临街商服用地P商" in narratives["benchmark_corr_solve"]
    assert "商业用地综合单价" in narratives["benchmark_corr_solve"]
    assert "临街商服用地面积=标准临街深度×临街宽度" in narratives["benchmark_corr_frontage_special_intro"]


def test_benchmark_correction_commercial_split_requires_non_street_region_factors():
    payload = _commercial_payload("route_price_plus_non_street")
    payload["land_area"] = "100"
    payload["benchmark_corr_analysis"].update({
        "route_segment_id": "B25",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert result["benchmark_corr_price"] == ""
    assert any("区域因素" in item for item in result["support_missing_items"])


def test_benchmark_correction_commercial_split_rejects_negative_residual_area():
    payload = _fill_commercial_region_factors(_commercial_payload("route_price_plus_non_street"))
    payload["land_area"] = "50"
    payload["benchmark_corr_analysis"].update({
        "route_segment_id": "B25",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert result["benchmark_corr_price"] == ""
    assert any("大于商服总面积" in item for item in result["support_missing_items"])


def test_benchmark_correction_commercial_frontage_mode_prefers_current_top_level_value():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "80"
    payload["benchmark_corr_analysis"].update({
        "parameters": {"frontage_mode": "non_street"},
        "route_segment_id": "B25",
        "frontage_depth_m": "999",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["frontage_mode"] == "street_route_price"
    assert result["parameters"]["formula_text"] == "1、商业服务业用地（临街）\nP商=路线价×Ky×Kv×Kt×Ks×Ka×Kc×Kk×Kd×Ku＋Kf"


def test_benchmark_correction_commercial_manual_route_requires_source_note():
    payload = _commercial_payload("street_route_price")
    payload["land_area"] = "80"
    payload["benchmark_corr_analysis"].update({
        "route_price": "2000",
        "route_road_type": "主干道",
        "frontage_depth_m": "999",
        "frontage_width_m": "4",
        "is_corner": False,
    })
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert any("来源说明" in item for item in result["support_missing_items"])
    payload["benchmark_corr_analysis"]["route_price_source_note"] = "估价师按路线价图人工判读。"
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is True
    assert result["parameters"]["route_price"] == "2000"


def test_benchmark_correction_no_sample_default_leakage():
    payload = {
        "county_name": "通道县",
        "use_benchmark_corr": True,
        "land_usage": "住宅用地",
        "land_level": "二级",
        "plot_ratio": "4.58",
        "land_use_term": "63.97",
        "valuation_date": "2026-06-03",
    }
    result = calculate_benchmark_correction(payload, BASE_DIR)
    assert result["complete"] is False
    assert result["benchmark_corr_price"] == ""
    assert any("区域因素" in item for item in result["support_missing_items"])


def test_benchmark_corr_process_in_valuation_draft():
    draft = build_valuation_process_draft(_tongdao_payload(), _FakeLibrary())
    methods = {m["method_key"]: m for m in draft["methods"]}
    assert "benchmark_corr" in methods
    method = methods["benchmark_corr"]
    assert method["status"] == "complete"
    price_result = next(r for r in method["results"] if r["key"] == "benchmark_corr_price")
    assert price_result["value"] == "1144.6"


class _FakeLibrary:
    base_dir = BASE_DIR

    def calculate_market_comparison(self, analysis):
        return analysis

    def build_render_fields(self, analysis):
        return analysis
