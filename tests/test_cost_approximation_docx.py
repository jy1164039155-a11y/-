# -*- coding: utf-8 -*-
import shutil
from pathlib import Path

from docx import Document

from src.services.cost_approximation import calculate_cost_approximation
from src.services.cost_approximation_docx import apply_cost_approximation_to_docx


BASE_DIR = Path(__file__).resolve().parents[1]
TEMPLATE = BASE_DIR / "01_Source" / "02_Word_Templates" / "自动生成的评估报告模板.docx"


def test_cost_postprocessor_replaces_runtime_hardcoding_with_dynamic_content(tmp_path):
    output = tmp_path / "cost_dynamic.docx"
    shutil.copy(TEMPLATE, output)
    data = {
        "use_cost_approx": True,
        "county_name": "道县",
        "valuation_date": "2025-12-04",
        "acquisition_land_class": "水田",
        "land_usage": "商业服务业用地、居住用地",
        "land_development_set": "五通一平",
        "cost_approx_land_class_intro": "本次按水田测算土地取得成本。",
        "cost_approx_process_intro": "根据估价期日有效的征收补偿文件，各项费用如下。",
    }
    analysis = calculate_cost_approximation(data, BASE_DIR)
    for item in (
        analysis["policy_documents"]
        + analysis["acquisition_items"]
        + analysis["tax_items"]
        + analysis["development_items"]
        + analysis["usage_scenarios"]
        + analysis["building_compensation_rows"]
        + analysis["resettlement_population_cases"]
        + analysis["risk_items"]
    ):
        item["confirmed"] = True
    analysis["location_factors"] = [
        {"usage_key": "commercial", "label": "商业综合区位修正", "correction_rate": "11.31", "confirmed": True},
        {"usage_key": "residential", "label": "住宅综合区位修正", "correction_rate": "12.88", "confirmed": True},
    ]
    data["cost_approx_analysis"] = analysis

    apply_cost_approximation_to_docx(str(output), data, str(BASE_DIR))

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "商业服务业用地：\nP=" in text
    assert "居住用地：\nP=" in text
    assert "573.9元/平方米" in text
    assert "600.0元/平方米" in text
    assert "评估宗地征收地类为林地" not in text
    assert "policy_recommendation" not in text
    assert "土地还原率风险因素测算表" not in text


def test_cost_postprocessor_uses_baozhen_nine_step_structure(tmp_path):
    output = tmp_path / "cost_baozhen_structure.docx"
    shutil.copy(TEMPLATE, output)
    data = {
        "use_cost_approx": True,
        "county_name": "道县",
        "valuation_date": "2026-06-01",
        "transfer_purpose_mode": "办理土地使用权出让手续",
        "acquisition_land_class": "水田",
        "land_usage": "居住用地",
        "land_development_set": "五通一平",
    }
    analysis = calculate_cost_approximation(data, BASE_DIR)
    for item in (
        analysis["policy_documents"]
        + analysis["acquisition_items"]
        + analysis["tax_items"]
        + analysis["development_items"]
        + analysis["usage_scenarios"]
        + analysis["location_factors"]
        + analysis["building_compensation_rows"]
        + analysis["resettlement_population_cases"]
        + analysis["risk_items"]
    ):
        item["confirmed"] = True
    data["cost_approx_analysis"] = analysis

    apply_cost_approximation_to_docx(str(output), data, str(BASE_DIR))

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    headings = [
        "1、土地取得费", "2、相关税费（T）", "3、土地开发费", "4、投资利息",
        "5、投资利润", "6、土地增值收益", "7、土地使用年期修正系数",
        "8、区位修正系数", "9、成本逼近法价格计算",
    ]
    positions = [text.index(item) for item in headings]
    assert positions == sorted(positions)
    table_titles = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    for title in (
        "表3-1 建筑物补偿标准",
        "表3-2 平均安置农业人口数确定表",
        "表3-3 估价对象土地取得费及相关税费一览表",
        "表3-4道县城区土地开发费用分项一览表",
        "表3-9 道县城区城镇住宅用地宗地地价区域因素修正系数表（%）",
        "表3-10 宗地地价影响因素条件说明、修正系数表",
    ):
        assert title in table_titles
