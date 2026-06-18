# -*- coding: utf-8 -*-
import base64
import subprocess
import tempfile
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from src.services.comparable_library import ComparableLibrary
from src.services.market_comparison_docx import apply_market_comparison_to_docx


TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/"
    "lWv1LgAAAABJRU5ErkJggg=="
)


def _official_case(guid: str, supervision_no: str, area_sqm: str, unit_price: str):
    total_wan = Decimal(area_sqm) * Decimal(unit_price) / Decimal("10000")
    return {
        "gdGuid": guid,
        "dzBaBh": supervision_no,
        "xmMc": f"测试项目{guid}",
        "tdZl": f"测试坐落{guid}",
        "xzqDm": "431124",
        "xzqFullName": "湖南省永州市道县",
        "tdYt": "工业用地",
        "gyFs": "挂牌出让",
        "gyMj": str(Decimal(area_sqm) / Decimal("10000")),
        "je": str(total_wan),
        "qdRq": "2025-01-01",
        "crNx": "30",
        "tdJb": "城区Ⅳ级",
        "srr": f"受让人{guid}",
    }


def test_market_comparison_docx_uses_effective_semantic_narratives(tmp_path):
    path = tmp_path / "process.docx"
    doc = Document()
    for text in (
        "2.比较实例选择",
        "模板固定实例选择",
        "表3-10  比较实例基本情况表",
        "3.建立价格可比基础",
        "模板固定价格基础",
        "4.比较因素的选择",
        "模板固定因素选择",
        "5.比较因素条件说明",
        "模板固定条件说明",
        "表3-11 比较因素条件说明表",
        "6.编制比较因素条件指数表",
        "模板固定指数依据",
        "表3-13 估价对象及比较因素条件指数表",
        "表3-14 估价对象比较因素修正系数表",
        "模板固定求价",
        "★收益还原法",
    ):
        doc.add_paragraph(text)
    doc.save(path)

    effective = {
        "market_comp_step1_instances": "有效实例选择",
        "market_comp_comparable_basis": "有效价格可比基础",
        "market_comp_factor_selection": "有效因素选择",
        "market_comp_condition_intro": "有效条件说明",
        "market_comp_index_basis": "有效指数依据",
        "market_comp_step4_solve": "有效求价正文",
        "market_comp_verification": "有效核验结论",
    }
    apply_market_comparison_to_docx(
        str(path),
        {
            "use_market_comp": True,
            "market_comp_step1_instances": effective["market_comp_step1_instances"],
            "market_comp_analysis": {"effective_narratives": effective},
        },
    )

    text = "\n".join(item.text for item in Document(path).paragraphs)
    for key, value in effective.items():
        if key == "market_comp_verification":
            continue
        assert value in text
    assert "有效核验结论" not in text
    assert "模板固定价格基础" not in text
    assert "模板固定因素选择" not in text
    assert "模板固定条件说明" not in text
    assert "模板固定指数依据" not in text
    assert "模板固定求价" not in text


def test_manual_market_evidence_images_insert_and_convert_with_soffice(tmp_path):
    soffice = Path("tools/LibreOffice-25.8.7/program/soffice.com").resolve()
    if not soffice.exists():
        pytest.skip("soffice.com not bundled")

    library = ComparableLibrary(str(tmp_path))
    for slot in ("A", "B", "C"):
        library.upsert_official_case(_official_case(slot, f"4311242024B00005{slot}", "10000", "120"))

    ids_by_kind = {"announcement": [], "location": [], "site": []}
    for kind in ids_by_kind:
        for slot in ("A", "B", "C"):
            snapshot = library.create_manual_snapshot(
                slot,
                kind,
                [(f"{slot}_{kind}.png", TINY_PNG)],
            )
            ids_by_kind[kind].append(snapshot["snapshot_id"])

    path = tmp_path / "manual_evidence.docx"
    doc = Document()
    for slot in ("A", "B", "C"):
        doc.add_paragraph(f"比较实例{slot}成交公告（来源：中国土地市场网）")
        doc.add_paragraph(f"比较实例{slot}位置图")
        doc.add_paragraph(f"比较实例{slot}现状图")
        next_slot = chr(ord(slot) + 1)
        doc.add_paragraph(f"比较实例{next_slot}详细情况" if slot != "C" else "3.建立价格可比基础")
    doc.save(path)

    apply_market_comparison_to_docx(
        str(path),
        {
            "use_market_comp": True,
            "market_comp_evidence_snapshot_ids": ids_by_kind["announcement"],
            "market_comp_location_snapshot_ids": ids_by_kind["location"],
            "market_comp_site_snapshot_ids": ids_by_kind["site"],
        },
        library,
    )
    updated_doc = Document(path)
    assert len(updated_doc.inline_shapes) == 9
    for slot in ("A", "B", "C"):
        for marker in ("成交公告", "位置图", "现状图"):
            paragraph = next(
                item
                for item in updated_doc.paragraphs
                if item.text.strip().startswith(f"比较实例{slot}{marker}")
            )
            assert paragraph.paragraph_format.keep_with_next is True

    with tempfile.TemporaryDirectory(prefix="soffice_profile_", dir=tmp_path, ignore_cleanup_errors=True) as profile:
        profile_uri = "file:///" + str(profile).replace("\\", "/")
        result = subprocess.run(
            [
                str(soffice),
                f"-env:UserInstallation={profile_uri}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_path),
                str(path),
            ],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
    assert result.returncode == 0, result.stderr or result.stdout
    pdf_path = path.with_suffix(".pdf")
    assert pdf_path.exists()
    assert pdf_path.stat().st_size > 0


def test_market_comparison_docx_builds_shared_report_tables_with_borders(tmp_path):
    path = tmp_path / "shared_tables.docx"
    doc = Document()
    doc.add_paragraph("2.比较实例选择")
    doc.add_paragraph("模板实例")
    doc.add_paragraph("表3-10  比较实例基本情况表")
    table = doc.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ["估价对象与比较实例", "比较实例A", "比较实例B", "比较实例C"]):
        cell.text = text
    doc.add_paragraph("3.建立价格可比基础")
    doc.add_paragraph("模板基础")
    doc.add_paragraph("4.比较因素的选择")
    doc.add_paragraph("模板选择")
    doc.add_paragraph("5.比较因素条件说明")
    doc.add_paragraph("模板说明")
    doc.add_paragraph("表3-11 比较因素条件说明表")
    table = doc.add_table(rows=1, cols=7)
    table.rows[0].cells[0].text = "待估宗地及本级实例"
    doc.add_paragraph("6.编制比较因素条件指数表")
    doc.add_paragraph("模板指数说明")
    table = doc.add_table(rows=1, cols=5)
    table.rows[0].cells[0].text = "估价对象与比较案例"
    table.rows[0].cells[2].text = "实例A"
    doc.add_paragraph("表3-13 估价对象及比较因素条件指数表")
    table = doc.add_table(rows=1, cols=7)
    table.rows[0].cells[0].text = "待估宗地及比较实例"
    table.rows[0].cells[3].text = "估价对象"
    doc.add_paragraph("表3-14 估价对象比较因素修正系数表")
    table = doc.add_table(rows=1, cols=6)
    table.rows[0].cells[0].text = "待估宗地及比较实例"
    table.rows[0].cells[3].text = "案例A"
    doc.add_paragraph("模板求价")
    doc.add_paragraph("★收益还原法")
    doc.save(path)

    factor_row = {"group": "交易因素", "subgroup": "", "label": "土地用途", "subject": "工矿用地", "a": "工业用地", "b": "工业用地", "c": "工业用地"}
    apply_market_comparison_to_docx(
        str(path),
        {
            "use_market_comp": True,
            "market_comp_analysis": {
                "effective_narratives": {
                    "market_comp_step1_instances": "有效实例",
                    "market_comp_comparable_basis": "有效基础",
                    "market_comp_factor_selection": "有效选择",
                    "market_comp_condition_intro": "有效说明",
                    "market_comp_index_basis": "有效指数说明",
                    "market_comp_step4_solve": "有效求价",
                    "market_comp_verification": "不得进入报告",
                }
            },
            "market_comp_basic_rows": [{"label": "案例来源", "a": "来源A", "b": "来源B", "c": "来源C"}],
            "market_comp_factor_condition_rows": [factor_row],
            "market_comp_time_index_rows": [{"label": "成交日期", "subject": "2025年1月1日", "a": "2024年1月1日", "b": "2024年1月1日", "c": "2024年1月1日"}],
            "market_comp_factor_index_rows": [dict(factor_row, subject="100", a="100", b="100", c="100")],
            "market_comp_correction_rows": [{"group": "交易因素", "subgroup": "", "label": "土地用途", "a": "100/100", "b": "100/100", "c": "100/100"}],
        },
    )

    updated = Document(path)
    text = "\n".join(paragraph.text for paragraph in updated.paragraphs)
    for title in (
        "表3-10 比较实例基本情况表",
        "表3-11 比较因素条件说明表",
        "表3-12 估价对象与比较案例交易期日指数表",
        "表3-13 估价对象及比较因素条件指数表",
        "表3-14 估价对象比较因素修正系数表",
    ):
        assert title in text
    assert "不得进入报告" not in text
    assert len(updated.tables) == 5
    for table in updated.tables:
        borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
        assert borders is not None
        assert borders.find(qn("w:insideH")).get(qn("w:val")) == "single"
