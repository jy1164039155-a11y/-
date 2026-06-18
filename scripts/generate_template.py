import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def generate_template():
    mapping_path = r"02_Process/预检清单.xlsx"
    input_docx = r"01_Source/02_Word_Templates/[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx"
    output_docx = r"01_Source/02_Word_Templates/自动生成的评估报告模板.docx"
    
    if not os.path.exists(mapping_path) or not os.path.exists(input_docx):
        print("输入文件不存在，请检查清单和原始报告路径。")
        return

    # 1. 加载映射清单
    df_mapping = pd.read_excel(mapping_path)
    # 过滤掉没有变量名的行
    df_active = df_mapping[df_mapping['建议变量名'].notna()].copy()
    
    # 建立映射列表（由于可能有重复词，我们需要上下文来区分，但此处先简化为文本替换）
    # 为了处理碎裂 Run，我们会在段落级别进行替换
    replacements = []
    for _, row in df_active.iterrows():
        replacements.append({
            "old": str(row['标黄原词']).strip(),
            "new": f"{{{{ {row['建议变量名']} }}}}"
        })
    
    # 按长度降序排列，避免短词误伤长词
    replacements.sort(key=lambda x: len(x['old']), reverse=True)

    # 2. 读取 Word
    doc = Document(input_docx)
    
    def process_text_container(container):
        """处理段落或单元格中的文本，采用原位修改以保留格式"""
        for para in container.paragraphs:
            runs = para.runs
            i = 0
            while i < len(runs):
                try:
                    is_yellow = (runs[i].font.highlight_color == WD_COLOR_INDEX.YELLOW)
                except:
                    is_yellow = False
                
                if is_yellow:
                    # 1. 寻找连续的黄色高亮块
                    start_idx = i
                    combined_text = runs[i].text
                    j = i + 1
                    while j < len(runs):
                        try:
                            if runs[j].font.highlight_color == WD_COLOR_INDEX.YELLOW:
                                combined_text += runs[j].text
                                j += 1
                            else:
                                break
                        except:
                            break
                    
                    # 2. 对合并后的文本尝试执行替换
                    found_match = False
                    for rep in replacements:
                        if rep['old'] in combined_text:
                            # 在第一个 Run 中执行替换，并清除高亮背景
                            runs[start_idx].text = combined_text.replace(rep['old'], rep['new'])
                            runs[start_idx].font.highlight_color = None
                            
                            # 3. 清空后续的碎裂 Runs，但保留对象以维持段落结构
                            for k in range(start_idx + 1, j):
                                runs[k].text = ""
                                runs[k].font.highlight_color = None
                            found_match = True
                            break
                    
                    if not found_match:
                        # 如果没有匹配到清单里的变量，仅清除高亮，保留原词
                        # 或者可以选择保留高亮提醒用户。这里根据您的反馈，先不清除
                        pass
                    
                    i = j # 跳到下一个块
                else:
                    i += 1

    print("正在生成模板，执行安全替换...")
    
    # 处理正文
    process_text_container(doc)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text_container(cell)

    # 3. 确保输出目录存在
    output_dir = os.path.dirname(output_docx)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc.save(output_docx)
    print(f"模板生成成功！保存路径: {output_docx}")

if __name__ == "__main__":
    generate_template()
