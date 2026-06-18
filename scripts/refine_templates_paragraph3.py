import docx
import os
import shutil
import psutil
from docx.oxml.ns import qn

def kill_file_holder(file_path):
    """
    清理独占文件的进程，杜绝 PermissionError。
    """
    abs_path = os.path.abspath(file_path).lower()
    curr_pid = os.getpid()
    killed_any = False
    TARGET_PROCESS_NAMES = {'winword.exe', 'wps.exe', 'excel.exe', 'python.exe', 'soffice.bin'}
    
    for proc in psutil.process_iter(['pid', 'name']):
        if proc.info['pid'] == curr_pid:
            continue
        try:
            p_name = (proc.info['name'] or "").lower()
            if p_name in TARGET_PROCESS_NAMES:
                open_files = proc.open_files()
                for f in open_files:
                    if f.path and os.path.abspath(f.path).lower() == abs_path:
                        print(f"  [进程清理] 发现占用进程: PID: {proc.info['pid']}, Name: {proc.info['name']}。强行释放锁...")
                        proc.kill()
                        killed_any = True
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    if killed_any:
        import time
        time.sleep(0.5)

def rebuild_paragraph_with_text(p, new_text, font_size=None, font_bold=None):
    """
    高保真重建段落并锁死仿宋_GB2312
    """
    # 物理移除所有旧 runs
    for run in list(p.runs):
        p._element.remove(run._r)
        
    new_run = p.add_run(new_text)
    
    # 设定字号（如果有）
    if font_size:
        new_run.font.size = font_size
    if font_bold is not None:
        new_run.font.bold = font_bold
        
    # 锁死仿宋_GB2312
    new_run.font.name = '仿宋_GB2312'
    rPr = new_run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    rFonts.set(qn('w:ascii'), '仿宋_GB2312')
    rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
    rFonts.set(qn('w:cs'), '仿宋_GB2312')

def transform_paragraph_three(path, doc_name):
    print(f"\n=================== 开始重构 {doc_name} 段落三 ===================")
    kill_file_holder(path)
    
    doc = docx.Document(path)
    paragraphs = list(doc.paragraphs)
    
    start_idx = -1
    end_idx = -1
    
    # 1. 扫描定位
    for idx, p in enumerate(paragraphs):
        text = p.text.strip()
        # 匹配标题行，如：1.地价确定的方法 或 1、地价确定的方法 等
        if ("地价确定的方法" in text or "地价的确定" in text or "地价确定方法" in text) and ("1." in text or "1、" in text):
            start_idx = idx
            
        # 匹配原来的方案三占位符
        if "final_price_determination" in text and "{" in text:
            end_idx = idx
            break
            
    if start_idx == -1 or end_idx == -1:
        print(f"  [Error] {doc_name} 未能成功定位段落三范围！(start: {start_idx}, end: {end_idx})")
        return False
        
    print(f"  [定位成功] 段落三核心区间: 段落 [{start_idx}]({paragraphs[start_idx].text.strip()}) -> 段落 [{end_idx}]({paragraphs[end_idx].text.strip()})")
    
    # 获取标题段落之后第一个正文段落的字号，作为生成模板代码时的字号模板，以确保完美一致
    ref_font_size = None
    ref_font_bold = None
    if start_idx + 1 < len(paragraphs):
        ref_para = paragraphs[start_idx + 1]
        if ref_para.runs:
            ref_font_size = ref_para.runs[0].font.size
            ref_font_bold = ref_para.runs[0].font.bold
            
    # 2. 物理清除 start_idx + 1 到 end_idx 之间的所有旧正文与占位符段落
    for idx in range(end_idx, start_idx, -1):
        p_to_remove = paragraphs[idx]
        p_element = p_to_remove._element
        p_element.getparent().remove(p_element)
        
    # 定义插入辅助函数
    def insert_paragraph_after(paragraph):
        new_p = docx.oxml.OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
        return new_para

    # 3. 构造全新高保真 Jinja2 模板内容
    template_lines = [
        "根据以上评估过程，本次评估采用了{{ adopted_methods_summary }}进行测算。各估价方法测算情况如下：",
        "{% if use_cost_approx -%}",
        "【成本逼近法】成本逼近法是以开发土地所耗费的各项客观费用之和为主要依据，再加上一定的利润、利息、应缴纳的税金和土地增值收益来确定土地价格的方法。考虑到估价对象所在区域有类似征地案例，能够通过征地案例的调查分析，对各项土地取得成本进行量化得出土地价格。{% if show_price_in_text %}成本逼近法测算的测算结果为{{ cost_approx_price }}元/平方米。{% endif %}",
        "{%- endif %}",
        "{% if use_market_comp -%}",
        "【市场比较法】市场比较法是根据替代原理，将待估宗地与具有替代性的，且在估价期日近期市场上交易的类似地产进行比较，对类似地产成交价格作适当修正，以此得到待估宗地比准价格的一种方法。该方法选择的比较案例是市场搜集的比较客观的真实交易，修正体系科学合理。{% if show_price_in_text %}市场比较法测算的测算结果为{{ market_comp_price }}元/平方米，该方法的测算结果能够真实的反映估价对象的客观合理市场价格。{% endif %}",
        "{%- endif %}",
        "{% if use_income_cap -%}",
        "【收益还原法】收益还原法是在估算估价对象在未来每年预期纯收益（正常年纯收益）的基础上，以一定的土地还原率，将评估对象在未来每年的纯收益折算为估价期日收益总和的一种方法。{% if show_price_in_text %}该方法的测算结果能够比较真实的反映待估宗地的客观合理的市场价格。其测算结果为{{ income_cap_price }}元/平方米。{% endif %}",
        "{%- endif %}",
        "{% if use_benchmark_corr -%}",
        "【基准地价系数修正法】基准地价系数修正法是通过对待估宗地的区域条件和个别条件等与其所处区域的平均条件相比较，通过一定的修正系数得到待估宗地在估价期日价格的方法。基准地价是通过科学方法测算而来，是{{ county_name }}人民政府土地市场宏观管理的有效工具，基准地价的适应性较强，运用基准地价系数修正法评估的结果依据较为充分。{% if show_price_in_text %}其测算结果为{{ benchmark_corr_price }}元/平方米。{% endif %}",
        "{%- endif %}",
        "{% if use_residual -%}",
        "【剩余法】剩余法是预计勘察设计、建设完成后不动产的总开发价值，扣除正常的开发成本、利息、利润、税费等，以剩余部分确定起点土地价格的方法。{% if show_price_in_text %}其测算结果为{{ residual_price }}元/平方米。{% endif %}",
        "{%- endif %}",
        "{# --- 核心合成逻辑判定 --- #}",
        "{% if weight_logic_type == \"simple_average\" -%}",
        "综上所述，通过对当地土地市场和土地价格的分析和各方法的可靠性综合分析，根据项目特点，综合考虑宗地所在区域的地价水平，确定以{{ adopted_methods_summary }}评估结果的简单算术平均值[{{ formula_display_text }}]作为本次评估的最终结果。",
        "{%- elif weight_logic_type == \"weighted_average\" -%}",
        "综合分析，两种方法评估的结果内涵一致，考虑到{{ weight_rationale_text }}，根据项目特点，综合考虑估价对象所在区域的地价水平，确定以{{ adopted_methods_summary }}测算结果的加权平均值（取整）作为最终结果较为合理。",
        "{%- elif weight_logic_type == \"single_dominance\" -%}",
        "通过对当地土地市场和土地价格的分析和各方法的可靠性综合分析，根据项目特点，综合考虑宗地所在区域的地价水平，确定以{{ dominant_method_name }}的评估结果作为本次评估的最终结果。",
        "{%- endif %}",
        "{% if not show_price_in_text -%}",
        "具体地价结果确定及各方法测算数值详见“地价结果确定表”。",
        "{%- endif %}"
    ]
    
    # 4. 依次在标题段落之后追加插入这些模板段落
    current_p = paragraphs[start_idx]
    for text in template_lines:
        current_p = insert_paragraph_after(current_p)
        rebuild_paragraph_with_text(current_p, text, font_size=ref_font_size, font_bold=ref_font_bold)
        
    # 5. 回写全局后置大锁死
    def lock_run(run):
        run.font.name = '仿宋_GB2312'
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        rFonts.set(qn('w:cs'), '仿宋_GB2312')

    for p in doc.paragraphs:
        for run in p.runs:
            lock_run(run)
            
    doc.save(path)
    print(f"  [重构成功] {doc_name} 物理替换、段落Run重建与仿宋_GB2312硬锁定圆满完成！")
    return True

if __name__ == "__main__":
    dir_path = r"D:\评估报告工具\01_Source\02_Word_Templates"
    res_dest = os.path.join(dir_path, "[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx")
    ind_dest = os.path.join(dir_path, "[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx")
    
    transform_paragraph_three(res_dest, "住宅合并报告模板")
    transform_paragraph_three(ind_dest, "工业合并报告模板")
