import docx
import os
from docx.oxml.ns import qn

def rebuild_paragraph_with_text(p, new_text):
    """
    高保真重建段落：保留原段落中首个 run 的基本样式属性，清空所有 runs，在单一 run 里完整写入 new_text
    并指定为仿宋_GB2312字体。
    """
    if len(p.runs) > 0:
        first_run = p.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        font_color = first_run.font.color.rgb if (first_run.font.color and hasattr(first_run.font.color, 'rgb')) else None
        
        # 清除段落文本
        p.text = ""
        
        # 单 run 写入
        new_run = p.add_run(new_text)
        new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size
        new_run.font.bold = font_bold
        new_run.font.italic = font_italic
        if font_color:
            new_run.font.color.rgb = font_color
            
        # 强制指定为 仿宋_GB2312
        new_run.font.name = '仿宋_GB2312'
        rPr = new_run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
    else:
        p.text = ""
        new_run = p.add_run(new_text)
        new_run.font.name = '仿宋_GB2312'
        rPr = new_run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')

def apply_fangsong_to_run(run):
    run.font.name = '仿宋_GB2312'
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    rFonts.set(qn('w:ascii'), '仿宋_GB2312')
    rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')

def apply_fangsong_globally(doc):
    """
    对文档中的所有 Runs 实施仿宋_GB2312字体大锁定，确保格式完美无损。
    """
    for p in doc.paragraphs:
        for run in p.runs:
            apply_fangsong_to_run(run)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        apply_fangsong_to_run(run)

def kill_file_holder(file_path):
    import psutil
    import os
    abs_path = os.path.abspath(file_path).lower()
    curr_pid = os.getpid()
    print(f"  正在全系统检索并清理独占文件的进程: {file_path}")
    killed_any = False
    
    # 仅检测可能占用该 Word/Excel 文件的候选进程类型，排除系统级进程以防止卡死
    TARGET_PROCESS_NAMES = {'winword.exe', 'wps.exe', 'excel.exe', 'python.exe', 'soffice.bin'}
    
    for proc in psutil.process_iter(['pid', 'name']):
        if proc.info['pid'] == curr_pid:
            continue
        try:
            p_name = (proc.info['name'] or "").lower()
            if p_name in TARGET_PROCESS_NAMES:
                open_files = proc.open_files()
                for f in open_files:
                    if f.path and os.path.abspath(f.path).lower() == abs_path:
                        print(f"  [占用查杀] 发现占用进程! PID: {proc.info['pid']}, Name: {proc.info['name']}. 正在强行终止以释放锁...")
                        proc.kill()
                        killed_any = True
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    if killed_any:
        import time
        time.sleep(0.5)

def purify_template(docx_path, name):
    print(f"\n--- 正在净化 {name} Word 模板中的长文本段落... ---")
    doc = docx.Document(docx_path)
    
    purified_count = 0
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # 1. 净化方案二：基准地价系数修正法适用性说明段落
        if "道县城镇基准地价的颁布实施时间为2019年11月22日" in text and "城区基准地价估价期日为2019年5月1日" in text:
            print(f"  [方案二] 净化段落 {i} -> '{{{{ valuation_method_applicability }}}}'")
            rebuild_paragraph_with_text(p, "{{ valuation_method_applicability }}")
            purified_count += 1
            
        # 2. 净化方案三：地价加权调和段落
        # 住宅模板
        elif "确定以剩余法的评估结果作为本次评估的最终结果" in text:
            print(f"  [方案三] 净化段落 {i} -> '{{{{ final_price_determination }}}}'")
            rebuild_paragraph_with_text(p, "{{ final_price_determination }}")
            purified_count += 1
        # 工业模板
        elif "确定以成本逼近法 and 市场比较法" in text or ("确定以成本逼近法和市场比较法" in text and ("简单算术" in text or "简单算数" in text) and "最终结果" in text):
            print(f"  [方案三] 净化段落 {i} -> '{{{{ final_price_determination }}}}'")
            rebuild_paragraph_with_text(p, "{{ final_price_determination }}")
            purified_count += 1
            
        # 3. 净化方案四：基础设施与开发程度详细说明段落（限定条件中）
        elif "2.基础设施条件：估价对象土地实际开发程度为宗地红线外" in text and "场地平整" in text and "详见下表" in text:
            print(f"  [方案四] 净化段落 {i} -> '2.{{{{ infrastructure_detail }}}}'")
            rebuild_paragraph_with_text(p, "2.{{ infrastructure_detail }}")
            purified_count += 1

    if purified_count > 0:
        apply_fangsong_globally(doc)
        import time
        success = False
        for attempt in range(5):
            try:
                # 每次尝试前先强行清理独占进程和临时 ~$ 锁文件
                kill_file_holder(docx_path)
                try:
                    dir_n = os.path.dirname(docx_path)
                    for f_name in os.listdir(dir_n):
                        if f_name.startswith("~$") and name in f_name:
                            try:
                                os.remove(os.path.join(dir_n, f_name))
                                print(f"  已自动清理残留的临时锁定文件: {f_name}")
                            except:
                                pass
                except:
                    pass
                
                doc.save(docx_path)
                success = True
                print(f"净化成功！共物理挖空了 {purified_count} 处长文本段落并锁定仿宋字体。")
                break
            except PermissionError as pe:
                print(f"  [第{attempt+1}次保存失败] 遭遇独占: {pe}. 正在尝试解除锁定并重试...")
                time.sleep(1)
        
        if not success:
            print("  警告：直接保存失败，尝试最后的临时中转强制覆盖机制...")
            temp_path = docx_path + ".tmp"
            try:
                doc.save(temp_path)
                kill_file_holder(docx_path)
                if os.path.exists(docx_path):
                    try:
                        os.remove(docx_path)
                    except:
                        os.system(f'del /f /q "{docx_path}"')
                os.rename(temp_path, docx_path)
                print(f"净化成功！采用中转强制覆盖成功挖空了 {purified_count} 处长文本段落。")
            except Exception as final_ex:
                print(f"  [致命异常] 强制覆盖机制亦告失败: {final_ex}")
                raise final_ex
    else:
        print("警告：未在文档中检索到任何需要净化的目标长文本段落。")

if __name__ == "__main__":
    dir_path = r"D:\评估报告工具\01_Source\02_Word_Templates"
    
    res_tmpl = os.path.join(dir_path, "[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx")
    ind_tmpl = os.path.join(dir_path, "[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx")
    
    if os.path.exists(res_tmpl):
        purify_template(res_tmpl, "住宅")
    else:
        print(f"未找到住宅模板: {res_tmpl}")
        
    if os.path.exists(ind_tmpl):
        purify_template(ind_tmpl, "工业")
    else:
        print(f"未找到工业模板: {ind_tmpl}")
        
    print("\n净化模板长文本工作大功告成！")
