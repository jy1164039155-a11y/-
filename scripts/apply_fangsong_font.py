import docx
import os
from docx.oxml.ns import qn

def apply_fangsong_font_to_doc(doc_path):
    print(f"正在处理文档: {os.path.basename(doc_path)}")
    doc = docx.Document(doc_path)
    
    # 1. 设置 Normal 样式的默认字体为 仿宋_GB2312
    try:
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '仿宋_GB2312'
        
        # 强制指定 Normal 样式的中文字体为 仿宋_GB2312
        rPr = normal_font._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        print("  成功修改 Normal 样式默认字体为 仿宋_GB2312。")
    except Exception as e:
        print(f"  修改 Normal 样式失败: {e}")

    # 计数器
    p_runs_count = 0
    t_runs_count = 0

    # 2. 遍历所有段落的 Runs
    for p in doc.paragraphs:
        # 如果段落样式是标题 (Heading)，我们可以选择保留或者也统一设置中文字体
        # 通常标题的中文也使用 仿宋_GB2312（或仿宋加粗），因此我们统一将所有 Runs 的中文字体设为 仿宋_GB2312
        for run in p.runs:
            p_runs_count += 1
            # 设定西文/常规字体名称
            run.font.name = '仿宋_GB2312'
            # 强制指定东亚中文字体
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            rFonts.set(qn('w:ascii'), '仿宋_GB2312')
            rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')

    # 3. 遍历所有表格单元格中的段落和 Runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        t_runs_count += 1
                        run.font.name = '仿宋_GB2312'
                        rPr = run._r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')

    print(f"  段落中处理了 {p_runs_count} 个 Runs。")
    print(f"  表格中处理了 {t_runs_count} 个 Runs。")
    
    try:
        doc.save(doc_path)
        print(f"  成功保存修改后的文档！\n")
    except PermissionError:
        print(f"  [警告] 无法保存 {os.path.basename(doc_path)}，文件可能已被其他程序（如 Word）打开而锁定。请关闭后重试。\n")
    except Exception as e:
        print(f"  保存文档时出错: {e}\n")

if __name__ == "__main__":
    dir_path = r"D:\评估报告工具\01_Source\02_Word_Templates"
    
    file1 = os.path.join(dir_path, "[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx")
    file2 = os.path.join(dir_path, "[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx")
    file3 = os.path.join(dir_path, "自动生成的评估报告模板.docx")
    
    for f in [file1, file2, file3]:
        if os.path.exists(f):
            try:
                apply_fangsong_font_to_doc(f)
            except Exception as e:
                print(f"处理 {os.path.basename(f)} 时出错: {e}")
        else:
            print(f"找不到文件: {f}")
