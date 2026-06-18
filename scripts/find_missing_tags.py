import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def find_missing_tags_precise():
    excel_path = r"01_Source/01_Excel_Sheets/HP_测算表_20260521.xlsx"
    mapping_path = r"02_Process/预检清单.xlsx"
    word_path = r"01_Source/02_Word_Templates/合并报告_原始_20260521.docx"
    output_path = r"02_Process/漏网之鱼清单.xlsx"

    if not os.path.exists(excel_path) or not os.path.exists(word_path):
        print("文件不存在。")
        return

    # 1. 精准提取关键词库
    keywords = set()
    
    # 来源 A: 黄璞测算表 -> 已开发建设土地（划拨办出让）
    print("正在从《已开发建设土地》工作表提取精准关键词...")
    try:
        # 这里使用乱码匹配或索引加载，由于之前 sheet_names 乱码，我们尝试按顺序或模糊匹配
        xl = pd.ExcelFile(excel_path)
        target_sheet = ""
        for name in xl.sheet_names:
            # 根据用户描述，“已开发建设土地”通常是第一个或名称包含相关字眼
            # 之前的 list_dir 显示第一个是 'ѿأã'
            if "已开发建设" in name or name.startswith("ѿ") or name == xl.sheet_names[0]:
                target_sheet = name
                break
        
        if target_sheet:
            df_core = pd.read_excel(excel_path, sheet_name=target_sheet, header=None)
            for _, row in df_core.iterrows():
                for val in row:
                    if pd.notna(val):
                        s = str(val).strip()
                        if len(s) > 1: # 过滤掉单字
                            keywords.add(s)
    except Exception as e:
        print(f"读取 Excel 出错: {e}")

    # 来源 B: 预检清单中的标黄原词
    print("正在从《预检清单》提取已确认关键词...")
    try:
        df_mapping = pd.read_excel(mapping_path)
        existing_words = set(df_mapping['标黄原词'].astype(str).str.strip())
        keywords.update(existing_words)
    except:
        pass

    # 过滤掉一些纯数字或极短词，除非它们在预检清单中
    final_keywords = [k for k in keywords if len(k) > 1]
    # 按长度降序排列
    final_keywords.sort(key=len, reverse=True)

    # 2. 扫描 Word
    print(f"开始精准扫描 Word... 关键词库大小: {len(final_keywords)}")
    doc = Document(word_path)
    missing_items = []

    def scan_container(container, desc):
        for idx, para in enumerate(container.paragraphs):
            text = para.text.strip()
            if not text: continue
            
            # 检查是否已标黄
            is_all_yellow = True
            has_yellow = False
            for run in para.runs:
                try:
                    if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                        has_yellow = True
                    else:
                        is_all_yellow = False
                except:
                    is_all_yellow = False
            
            # 如果整段已经标黄，跳过
            if has_yellow and len(text) < 100: # 短段落标黄视为已处理
                 continue

            # 在未标黄的文本中寻找精准匹配
            for kw in final_keywords:
                # 为了解决用户提到的“项目名称”漏标，我们检查关键词是否在段落中
                if kw in text:
                    # 只有当这个 kw 在段落中对应的位置没有高亮时，才算漏标
                    # 这里简化为：只要出现在没被标黄的段落里，就报出来
                    missing_items.append({
                        "位置": f"{desc}_段落{idx+1}",
                        "疑似漏标词": kw,
                        "上下文": text[:100] + "..." if len(text) > 100 else text
                    })
                    break

    scan_container(doc, "正文")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                scan_container(cell, f"表格{t_idx+1}_行{r_idx+1}_列{c_idx+1}")

    # 3. 结果去重与保存
    if missing_items:
        df_result = pd.DataFrame(missing_items)
        # 按照“疑似漏标词”去重，保留第一个发现的位置
        df_result = df_result.drop_duplicates(subset=['疑似漏标词'])
        df_result.to_excel(output_path, index=False)
        print(f"精准扫描完成！发现 {len(df_result)} 个漏标项。保存至: {output_path}")
    else:
        print("未发现符合条件的漏标项。")

if __name__ == "__main__":
    find_missing_tags_precise()
