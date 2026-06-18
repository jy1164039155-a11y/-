import docx
import os
import shutil
import psutil
from docx.oxml.ns import qn
from docx.shared import Pt

def kill_file_holder(file_path):
    """
    全系统检索并清理独占文件的进程，并物理移除临时锁定文件，杜绝保存与覆盖时的 PermissionError。
    在 Windows 上优化检测，避免对所有系统进程调用 open_files() 导致卡死。
    """
    abs_path = os.path.abspath(file_path).lower()
    curr_pid = os.getpid()
    killed_any = False
    
    # 仅检测可能占用该 Word/Excel 文件的进程类型，排除系统级进程以防止卡死
    TARGET_PROCESS_NAMES = {'winword.exe', 'wps.exe', 'excel.exe', 'python.exe', 'soffice.bin'}
    
    for proc in psutil.process_iter(['pid', 'name']):
        if proc.info['pid'] == curr_pid:
            continue
        try:
            p_name = (proc.info['name'] or "").lower()
            if p_name in TARGET_PROCESS_NAMES:
                open_files = proc.open_files()
                for f in open_files:
                    if f.path and os.path.abspath(f.path).lower() == abs_path:
                        print(f"  [进程清理] 发现占用进程: PID: {proc.info['pid']}, Name: {proc.info['name']}。正在强行终止释放锁...")
                        proc.kill()
                        killed_any = True
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
            
    # 物理清除 Word 产生的临时锁定 ~$ 文件
    try:
        dir_n = os.path.dirname(file_path)
        base_n = os.path.basename(file_path)
        for f_name in os.listdir(dir_n):
            if f_name.startswith("~$") and (base_n[2:] in f_name or f_name[2:] in base_n):
                try:
                    os.remove(os.path.join(dir_n, f_name))
                    print(f"  [锁定清理] 已自动清理临时锁定文件: {f_name}")
                except:
                    pass
    except Exception:
        pass

    if killed_any:
        import time
        time.sleep(0.5)

def rebuild_paragraph_with_text(p, new_text, force_font_size=None):
    """
    高保真重建段落：
    1. 物理切除段落底层的 runs 数组（彻底隔离并剔除空 Run，斩断 XML 字体 None 污染）；
    2. 创建唯一的纯净 Run 写入文本；
    3. 精准恢复首个 run 的字号、粗斜体、颜色等基本样式；
    4. 强力锁死中英文字体为 仿宋_GB2312 (eastAsia, ascii, hAnsi, cs)；
    5. 【字号终极防线】：若抓出的原 Runs 字号为 None（或被默认继承为五号），一律强制并智能死锁在标准的四号字 (Pt(14)) 上，彻底清除五号字跳变硬伤！
    """
    # 提取原段落中首个 run 的基本样式属性
    font_size = None
    font_bold = None
    font_italic = None
    font_color = None
    
    if len(p.runs) > 0:
        first_run = p.runs[0]
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        if first_run.font.color and hasattr(first_run.font.color, 'rgb'):
            font_color = first_run.font.color.rgb

    # 底层 XML 物理移除所有旧 runs，不留任何死空 Run
    for run in list(p.runs):
        p._element.remove(run._r)
        
    # 添加全新唯一的 Run 写入文本
    new_run = p.add_run(new_text)
    
    # 强制锁死或智能兜底字号为四号字 (Pt(14))
    final_sz = force_font_size if force_font_size else font_size
    if not final_sz:
        final_sz = Pt(14) # 标准四号字 (14pt)
        
    new_run.font.size = final_sz
    
    if font_bold is not None:
        new_run.font.bold = font_bold
    if font_italic is not None:
        new_run.font.italic = font_italic
    if font_color:
        new_run.font.color.rgb = font_color
        
    # 强力锁死仿宋_GB2312
    new_run.font.name = '仿宋_GB2312'
    rPr = new_run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    rFonts.set(qn('w:ascii'), '仿宋_GB2312')
    rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
    rFonts.set(qn('w:cs'), '仿宋_GB2312') # 锁定复杂脚本

def insert_paragraph_after_with_format(paragraph, ref_paragraph):
    """
    高保真段落级格式继承器：
    在新创建并插入新段落时，完美复制参考段落的首行缩进、行距、对齐等核心段落格式，
    物理解决新插If-Endif段落没有首行缩进的排版硬伤。
    """
    new_p = docx.oxml.OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    
    if ref_paragraph and ref_paragraph.paragraph_format:
        try:
            new_para.paragraph_format.first_line_indent = ref_paragraph.paragraph_format.first_line_indent
            new_para.paragraph_format.line_spacing = ref_paragraph.paragraph_format.line_spacing
            new_para.paragraph_format.space_before = ref_paragraph.paragraph_format.space_before
            new_para.paragraph_format.space_after = ref_paragraph.paragraph_format.space_after
            new_para.paragraph_format.alignment = ref_paragraph.paragraph_format.alignment
        except Exception:
            pass
    return new_para

def apply_fangsong_globally(doc):
    """
    全局后置高保真大锁死：强制锁死正文与表格中所有 Runs 的 eastAsia、ascii、hAnsi 及 cs 属性为 仿宋_GB2312。
    """
    def lock_run(run):
        run.font.name = '仿宋_GB2312'
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        rFonts.set(qn('w:cs'), '仿宋_GB2312')

    for p in doc.paragraphs:
        for run in p.runs:
            lock_run(run)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        lock_run(run)

def rebuild_method_selection(doc):
    """
    段落二：估价方法选择矩阵 If-Endif 条件骨架物理重构
    """
    rebuilt_count = 0
    while True:
        paragraphs = list(doc.paragraphs)
        start_p_idx = -1
        end_p_idx = -1
        
        # 寻找第一处未重构的方法矩阵区间
        for idx, p in enumerate(paragraphs):
            text = p.text.strip()
            if ("采用的估价方法" in text or "采用的方法" in text or "采用方法" in text) and ("（1）" in text or "(1)" in text):
                if "{%" in text or "{#" in text:
                    continue
                start_p_idx = idx
                
                # 寻找结束标志
                for end_idx in range(start_p_idx + 1, len(paragraphs)):
                    end_text = paragraphs[end_idx].text.strip()
                    if (("估价结果" in end_text and ("三、" in end_text or "四、" in end_text or "3." in end_text or "4." in end_text or end_text == "估价结果" or end_text.startswith("三、") or end_text.startswith("四、"))) or
                        "地价确定的方法" in end_text or "估价测算过程" in end_text or 
                        end_text.startswith("四、") or end_text.startswith("五、") or end_text.startswith("三、") or 
                        end_text.startswith("4.") or end_text.startswith("5.") or end_text.startswith("3.")):
                        end_p_idx = end_idx
                        break
                break
                
        if start_p_idx == -1 or end_p_idx == -1:
            break
            
        print(f"  [物理重构] 成功定位方法选择矩阵区间: 段落 [{start_p_idx}] -> [{end_p_idx}]")
        
        # 提取参考段落
        ref_para = None
        if start_p_idx + 1 < len(paragraphs):
            ref_para = paragraphs[start_p_idx + 1]

        # 反向物理删除旧段落
        for idx in range(end_p_idx - 1, start_p_idx, -1):
            p_to_remove = paragraphs[idx]
            p_element = p_to_remove._element
            p_element.getparent().remove(p_element)
            
        # 终审修正后的完美无损 Jinja2 替换列表（强行锁死四号字 Pt(14)）
        new_paras_texts = [
            '{% if use_market_comp %}市场比较法：估价对象位于{{ land_location }}，属于开发程度较好的区域，区域内同类土地及房地产交易市场较活跃，类似交易案例易于搜集，适宜采用市场比较法进行评估。{% endif %}',
            '{% if use_income_cap %}收益还原法：估价对象位于{{ land_location }}，属于具备获取稳定纯收益条件的区域，其同类物业租赁或经营收益易于测算，可以通过测算未来纯收益并还原求取地价。{% endif %}',
            '{% if use_residual %}剩余法：估价对象为具备开发潜力的土地或在建工程，其设定用途为{{ land_usage_short }}，适宜通过测算开发完成后的房地产价值并扣除各项正常建设开发成本费用、税费、利息及利润后，求取土地使用权价值。{% endif %}',
            '{% if use_cost_approx or use_cost_approach %}成本逼近法：估价对象为{{ land_usage_short }}，其土地开发成本、征地安置补偿费用及相关税费等各项历史及重置成本容易搜集并测算，适宜通过累加重置成本、土地开发费用、利息、利润及土地增值等求取土地使用权价值。{% endif %}',
            '{% if use_benchmark_corr %}基准地价系数修正法：估价对象位于城区基准地价覆盖红线范围内，其土地用途符合基准地价适用体系，且基准地价处于有效期内，适宜利用当地最新的基准地价及其修正系数体系，进行时点、区域、个别因素修正后求取土地使用权价值。{% endif %}',
            '（2）不采用的估价方法及理由：',
            '{% if not use_market_comp %}市场比较法：估价对象所在区域同类用途土地或房地产的实际交易案例极其匮乏，且交易市场不够活跃，难以收集到符合比较要求的类似交易案例，故不适宜采用市场比较法进行评估。{% endif %}',
            '{% if not use_income_cap %}收益还原法：{% if land_usage_short == "工业用地" or land_usage_short == "工业" %}估价对象为工业用地，通常由企业自用，区域内同类工业厂房的租赁市场极不活跃，难以收集到客观的租金收益数据，故不采用收益还原法进行评估。{% else %}估价对象所在区域同类住宅及房地产租赁市场不活跃，难以收集到客观、准确的租金及经营性纯收益数据，故不采用收益还原法进行评估。{% endif %}{% endif %}',
            '{% if not use_residual %}剩余法：{% if land_usage_short == "工业用地" or land_usage_short == "工业" %}估价对象为工业用地，区域内工业厂房整体开发及再次转售案例极少，难以客观准确预测开发完成后的工业房地产价值及各项开发费用，故不采用剩余法评估。{% else %}估价对象所在区域的房地产开发市场趋于饱和或不活跃，估价对象为已建成并正常使用的房产，不具备重新投资开发建设的实际条件，开发完成后的房地产价值及各项开发费用难以准确预测，故不采用剩余法评估。{% endif %}{% endif %}',
            '{% if not use_cost_approx and not use_cost_approach %}成本逼近法：{% if land_usage_short == "工业用地" or land_usage_short == "工业" %}由于估价对象所在区域的土地征收及开发历史较为久远，相关的征地补偿标准、基础设施开发配套费用等重置成本难以客观、准确地测算和还原，故不采用成本逼近法进行评估。{% else %}估价对象设定用途为{{ land_usage_short }}，属于典型的经营性房地产，其土地价值更多地取决于未来的增值与收益空间，成本与公开市场价值的关联度较低，故本次评估不采用成本逼近法。{% endif %}{% endif %}',
            '{% if not use_benchmark_corr %}基准地价系数修正法（公示地价系数修正法）：{% if is_base_price_expired %}根据自然资源部及湖南省自然资源厅关于基准地价更新管理规定，当本地基准地价体系成果自发布实施至估价期日已超过{{ base_land_price_expire_limit }}年未全面更新时，该基准地价系数修正法强制停用，为确保评估结果客观合理与审计合规，故本次不采用基准地价系数修正法。{% else %}估价对象宗地虽处于当地最新基准地价成果有效期内，但由于其物理位置超出了城区最新基准地价成果覆盖的红线范围，或者其土地用途极其特殊、缺少对应的修正体系，难以客观进行系数修正，故本次不采用基准地价系数修正法。{% endif %}{% endif %}',
            '综上所述，经对上述估价方法适用性的科学分析与甄选，本次评估最终采用{{ adopted_methods_summary }}最适用于估价对象，确保了评估结果的科学、合理与准确。'
        ]

        # 还原标题
        rebuild_paragraph_with_text(paragraphs[start_p_idx], "（1）采用的估价方法及理由：{# REBUILT_METHOD_MATRIX #}")
        
        current_p = paragraphs[start_p_idx]
        for text in new_paras_texts:
            # 引入首行缩进等格式继承，并强制指定字号为 Pt(14) 四号字，防止任何五号字残留！
            current_p = insert_paragraph_after_with_format(current_p, ref_para)
            rebuild_paragraph_with_text(current_p, text, force_font_size=Pt(14))
            
        rebuilt_count += 1
    return rebuilt_count > 0

def rebuild_paragraph_three(doc):
    """
    大段落三（确价与加权逻辑）：高保真 8 段式包裹重构方案
    本版本彻底去除了所有外层包裹的破折号吞换行控制，将段落换行权完全交回给 Word 桌面端，彻底封杀合并行乌龙！
    并且通过 force_font_size=Pt(14) 强制大段落三的字号在生成时 100% 锁定为四号字！
    支持 while True 循环重构所有匹配到的确价大段落。
    """
    rebuilt_count = 0
    while True:
        paragraphs = list(doc.paragraphs)
        start_idx = -1
        end_idx = -1
        
        # 扫描定位标题与原来包含最终确定价格等字眼的旧段落
        for idx, p in enumerate(paragraphs):
            text = p.text.strip()
            if ("地价确定的方法" in text or "地价的确定" in text or "地价确定方法" in text) and ("1." in text or "1、" in text):
                if idx + 1 < len(paragraphs) and ("adopted_methods_summary" in paragraphs[idx + 1].text or "{%" in paragraphs[idx + 1].text):
                    continue
                start_idx = idx
                
                # 寻找结束标志
                for end_idx in range(start_idx + 1, len(paragraphs)):
                    end_text = paragraphs[end_idx].text.strip()
                    if ("final_price_determination" in end_text or "确定以剩余法" in end_text or 
                        ("确定以成本逼近" in end_text and "简单算术平均" in end_text) or 
                        ("确定以成本逼近" in end_text and "简单算数平均" in end_text) or 
                        ("确定以" in end_text and "作为最终结果" in end_text) or
                        ("确定以" in end_text and "作为本次评估的最终结果" in end_text) or
                        "采用的估价方法" in end_text or "采用的方法" in end_text or
                        end_text.startswith("二、") or end_text.startswith("三、") or end_text.startswith("四、") or
                        end_text.startswith("2.") or end_text.startswith("3.") or end_text.startswith("4.")):
                        end_idx = end_idx
                        break
                break
                
        if start_idx == -1 or end_idx == -1:
            break
            
        print(f"  [物理重构] 定位段落三区间: 段落 [{start_idx}]({paragraphs[start_idx].text.strip()}) -> 段落 [{end_idx}]({paragraphs[end_idx].text.strip()})")
        
        # 提取格式参考段落
        ref_para = None
        if start_idx + 1 < len(paragraphs):
            ref_para = paragraphs[start_idx + 1]
                
        # 物理删除中间旧文字及原占位符段落
        for idx in range(end_idx, start_idx, -1):
            p_to_remove = paragraphs[idx]
            p_element = p_to_remove._element
            p_element.getparent().remove(p_element)
            
        # 8 段式完美包裹方案（去除破折号控制符，无损保留换行）
        template_lines = [
            "根据以上评估过程，本次评估采用了{{ adopted_methods_summary }}进行测算。各估价方法测算情况如下：",
            "{% if use_cost_approx %}【成本逼近法】成本逼近法是以开发土地所耗费的各项客观费用之和为主要依据，再加上一定的利润、利息、应缴纳的税金和土地增值收益来确定土地价格的方法。考虑到估价对象所在区域有类似征地案例，能够通过征地案例的调查分析，对各项土地取得成本进行量化得出土地价格。{% if show_price_in_text %}成本逼近法测算的测算结果为{{ cost_approx_price }}元/平方米。{% endif %}{% endif %}",
            "{% if use_market_comp %}【市场比较法】市场比较法是根据替代原理，将待估宗地与具有替代性的，且在估价期日近期市场上交易的类似地产进行比较，对类似地产成交价格作适当修正，以此得到待估宗地比准价格的一种方法。该方法选择的比较案例是市场搜集的比较客观的真实交易，修正体系科学合理.{% if show_price_in_text %}市场比较法测算的测算结果为{{ market_comp_price }}元/平方米，该方法的测算结果能够真实的反映估价对象的客观合理市场价格。{% endif %}{% endif %}",
            "{% if use_income_cap %}【收益还原法】收益还原法是在估算估价对象在未来每年预期纯收益（正常年纯收益）的基础上，以一定的土地还原率，将评估对象在未来每年的纯收益折算为估价期日收益总和的一种方法。{% if show_price_in_text %}该方法的测算结果能够比较真实的反映待估宗地的客观合理的市场价格。其测算结果为{{ income_cap_price }}元/平方米。{% endif %}{% endif %}",
            "{% if use_benchmark_corr %}【基准地价系数修正法】基准地价系数修正法是通过对待估宗地的区域条件和个别条件等与其所处区域的平均条件相比较，通过一定的修正系数得到待估宗地在估价期日价格的方法。基准地价是通过科学方法测算而来，是{{ county_name }}人民政府土地市场宏观管理的有效工具，基准地价的适应性较强，运用基准地价系数修正法评估的结果依据较为充分。{% if show_price_in_text %}其测算结果为{{ benchmark_corr_price }}元/平方米。{% endif %}{% endif %}",
            "{% if use_residual %}【剩余法】剩余法是预计勘察设计、建设完成后不动产的总开发价值，扣除正常的开发成本、利息、利润、税费等，以剩余部分确定起点土地价格的方法。{% if show_price_in_text %}其测算结果为{{ residual_price }}元/平方米。{% endif %}{% endif %}",
            "{% if weight_logic_type == \"simple_average\" %}综上所述，通过对当地土地市场和土地价格的分析和各方法的可靠性综合分析，根据项目特点，综合考虑宗地所在区域的地价水平，确定以{{ adopted_methods_summary }}评估结果的简单算术平均值[{{ formula_display_text }}]作为本次评估的最终结果。{% elif weight_logic_type == \"weighted_average\" %}综合分析，两种方法评估的结果内涵一致，考虑到{{ weight_rationale_text }}，根据项目特点，综合考虑估价对象所在区域的地价水平，确定以{{ adopted_methods_summary }}测算结果的加权平均值（取整）作为最终结果较为合理。{% elif weight_logic_type == \"single_dominance\" %}通过对当地土地市场和土地价格的分析和各方法的可靠性综合分析，根据项目特点，综合考虑宗地所在区域的地价水平，确定以{{ dominant_method_name }}的评估结果作为本次评估的最终结果。{% endif %}",
            "{% if not show_price_in_text %}具体地价结果确定及各方法测算数值详见“地价结果确定表”。{% endif %}"
        ]
        
        current_p = paragraphs[start_idx]
        for text in template_lines:
            current_p = insert_paragraph_after_with_format(current_p, ref_para)
            rebuild_paragraph_with_text(current_p, text, force_font_size=Pt(14))
            
        print(f"  [物理重构] 段落三 8段式完美包裹物理注入、首行缩进继承与字号硬死锁顺利完成。")
        rebuilt_count += 1
        
    return rebuilt_count > 0

def transform_residential(src_path, dest_path):
    print(f"\n[住宅模板] 正在基于干净原始件重构: {os.path.basename(src_path)}")
    kill_file_holder(dest_path)
    shutil.copy(src_path, dest_path)
    doc = docx.Document(dest_path)
    
    # 骨架定义
    history_skeleton = (
        "{%- if land_status_type == \"new_grant\" -%}"
        "根据《{{ gov_approval_name }}》（{{ gov_approval_no }}{%- if gov_approval_date -%}，{{ gov_approval_date }}{%- endif -%}），"
        "估价对象原为{{ original_land_owner_desc }}，于{{ approval_transfer_date }}经{{ approval_authority }}办理了农用地转用手续后，转变为国有新增建设用地。"
        "宗地范围内土地征收安置补偿已全部落实到位，土地权利清晰，现为国有储备土地。"
        "{%- elif land_status_type == \"historical_unregistered\" -%}"
        "根据{{ house_cert_name }}（{{ house_cert_no }}），估价对象位于{{ land_location }}，根据{{ area_docs_desc_name }}，"
        "估价对象位于{{ room_detail_location }}，建筑面积{{ building_area }}平方米，分摊土地使用权面积{{ land_area }}平方米，"
        "委托方拟按现状利用条件为其办理土地使用权出让手续。"
        "根据{{ proof_doc_name }}（{{ proof_doc_date }}），估价对象土地使用者为{{ land_user }}，{{ user_identity }}，"
        "根据房改房政策于{{ buy_year }}购得位于{{ buy_location_desc }}住房一套，房产证号：{{ house_cert_no }}，"
        "房屋登记面积{{ registered_house_area }}平方米。其房屋分摊土地使用权类型为国有划拨，因历史遗留问题，土地暂未办理土地权属登记。"
        "{%- elif land_status_type == \"registered\" -%}"
        "根据《{{ land_cert_name }}》（{{ land_cert_no }}），估价对象已取得土地权利证书，权利人为{{ land_user }}，"
        "土地使用权类型为{{ land_use_type }}，使用权面积为{{ land_area }}平方米，建筑面积为{{ building_area }}平方米。"
        "{%- endif -%}"
    )
    
    investigation_skeleton = (
        "{%- if land_status_type == \"new_grant\" -%}"
        "估价对象为拟出让用地，未办理登记手续，故暂无土地使用权人、土地权利证书编号、登记时间、地籍图号、宗地号、终止日期等相关信息。"
        "{%- elif land_status_type == \"historical_unregistered\" -%}"
        "估价对象为划拨用地，未办理登记手续，故暂无土地权利证书编号、登记时间、地籍图号、宗地号、终止日期等相关信息。"
        "{%- elif land_status_type == \"registered\" -%}"
        "估价对象已办理权属登记，相关权利证书编号、登记时间等登记状况详见“土地登记状况表”。"
        "{%- endif -%}"
    )
    
    def process_paras(paras):
        for p in paras:
            text = p.text.strip()
            if not text:
                continue
                
            new_text = text
            
            # (1) 历史权源骨架
            if "根据《房屋所有权证》" in text and "皮革厂家属楼" in text and "因历史遗留问题" in text:
                new_text = history_skeleton
            # (2) 调查附记骨架
            elif text.startswith("估价对象为划拨用地，未办理登记手续，故暂无土地权利证书编号"):
                new_text = investigation_skeleton
            # (5) 方案四净化
            elif "2.基础设施条件：估价对象土地实际开发程度为宗地红线外" in text and "场地平整" in text and "详见下表" in text:
                new_text = "2.{{ infrastructure_detail }}"
                
            # (6) 土地用途精细净化 (10处)
            elif "估价对象设定用途为住宅用地" in text:
                new_text = new_text.replace("估价对象设定用途为住宅用地", "估价对象设定用途为{{ land_usage_short }}")
            elif "出让最高年限为住宅用地70年" in text:
                new_text = new_text.replace("出让最高年限为住宅用地70年", "出让最高年限为{{ land_usage_short }}70年")
            elif "出让最高年限为住宅用地" in text and "70年" in text and "设定" in text:
                new_text = new_text.replace("住宅用地", "{{ land_usage_short }}")
            elif "故设定估价对象为住宅用地" in text and "容积率" in text:
                new_text = new_text.replace("故设定估价对象为住宅用地", "故设定估价对象为{{ land_usage_short }}")
            elif "土地使用年期为" in text and "住宅用地" in text:
                new_text = new_text.replace("住宅用地", "{{ land_usage_short }}")
            
            # 物理重建段落
            if new_text != text:
                rebuild_paragraph_with_text(p, new_text)
                continue
                
            # 常规强类型语义占位符替换
            new_text = text
            if "现场勘查日期为" in text or "设定估价期日" in text or "估价期日的设定" in text or "于估价期日" in text or "估价期日设定为" in text:
                new_text = new_text.replace("2026年5月18日", "{{ valuation_date }}")
                new_text = new_text.replace("道县自然资源局", "{{ client_name }}")
                new_text = new_text.replace("道县月岩西路", "{{ land_location }}")
                
            if "委托方" in text or "自然资源局" in text:
                new_text = new_text.replace("道县自然资源局", "{{ client_name }}")
                new_text = new_text.replace("道县月岩西路", "{{ land_location }}")
                new_text = new_text.replace("王汝冲", "{{ land_user }}")
                
            if "坐落" in text or "位置" in text or "位于道县月岩西路" in text:
                new_text = new_text.replace("道县月岩西row" if "row" in new_text else "道县月岩西路", "{{ land_location }}")
                
            if "土地用途的设定" in text or "土地用途" in text or "设定土地用途" in text:
                new_text = new_text.replace("城镇住宅用地", "{{ land_usage }}")
                new_text = new_text.replace("住宅用地", "{{ land_usage }}")
                
            if "容积率" in text or "容积率为" in text or "设定容积率" in text:
                new_text = new_text.replace("4.83", "{{ plot_ratio }}")
                
            if "土地面积" in text or "用地面积" in text or "分摊土地使用权面积" in text or "分摊土地面积" in text:
                new_text = new_text.replace("15.2", "{{ land_area }}")
                
            if "建筑总面积" in text or "建筑面积为" in text or "建筑面积" in text:
                new_text = new_text.replace("73.415", "{{ building_area }}")
                new_text = new_text.replace("74.55", "{{ registered_house_area }}")
                
            if "土地使用年限" in text or "土地使用年期" in text or "出让年期为" in text or "使用年限为" in text:
                new_text = new_text.replace("70年", "{{ land_use_term }}")
                
            if "二〇二" in text and "五月" in text:
                new_text = new_text.replace("二〇二六年五月二十日", "{{ report_date }}")
                
            if new_text != text:
                rebuild_paragraph_with_text(p, new_text)

    # 1. 处理正文段落
    process_paras(doc.paragraphs)
    
    # 2. 物理重构方法矩阵（段落二）
    rebuild_method_selection(doc)
    
    # 3. 物理重构确价段落（段落三）8段式完美包裹与首行缩进复制注入
    rebuild_paragraph_three(doc)
    
    # 4. 表格替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    if not text.strip():
                        continue
                    new_text = text
                    new_text = new_text.replace("道县自然资源局", "{{ client_name }}")
                    new_text = new_text.replace("道县月岩西路", "{{ land_location }}")
                    new_text = new_text.replace("2026年5月18日", "{{ valuation_date }}")
                    new_text = new_text.replace("住宅用地", "{{ land_usage }}")
                    new_text = new_text.replace("城镇住宅用地", "{{ land_usage }}")
                    new_text = new_text.replace("二〇二六年五月二十日", "{{ report_date }}")
                    new_text = new_text.replace("王汝冲", "{{ land_user }}")
                    new_text = new_text.replace("70年", "{{ land_use_term }}")
                    new_text = new_text.replace("4.83", "{{ plot_ratio }}")
                    
                    if "分摊土地" in new_text or "土地面积" in new_text:
                        new_text = new_text.replace("15.2", "{{ land_area }}")
                    if "建筑面积" in new_text:
                        new_text = new_text.replace("73.415", "{{ building_area }}")
                        new_text = new_text.replace("74.55", "{{ registered_house_area }}")
                        
                    new_text = new_text.replace("（湖南）湘兆财土地[2026]（估）字第ZC0005号", "{{ report_no }}")
                    
                    if new_text != text:
                        rebuild_paragraph_with_text(p, new_text)
                        
    # 5. 全局仿宋硬锁定
    apply_fangsong_globally(doc)
    doc.save(dest_path)
    print("  [住宅模板] 全局8段包裹式物理重构与字体、段落缩进锁定大功告成！")

def transform_industrial(src_path, dest_path):
    print(f"\n[工业模板] 正在基于干净原始件重构: {os.path.basename(src_path)}")
    kill_file_holder(dest_path)
    shutil.copy(src_path, dest_path)
    doc = docx.Document(dest_path)
    
    # 骨架定义
    history_skeleton = (
        "{%- if land_status_type == \"new_grant\" -%}"
        "根据《{{ gov_approval_name }}》（{{ gov_approval_no }}{%- if gov_approval_date -%}，{{ gov_approval_date }}{%- endif -%}），"
        "估价对象原为{{ original_land_owner_desc }}，于{{ approval_transfer_date }}经{{ approval_authority }}办理了农用地转用手续后，转变为国有新增建设用地。"
        "宗地范围内土地征收安置补偿已全部落实到位，土地权利清晰，现为国有储备土地。"
        "{%- elif land_status_type == \"historical_unregistered\" -%}"
        "根据{{ house_cert_name }}（{{ house_cert_no }}），估价对象位于{{ land_location }}，根据{{ area_docs_desc_name }}，"
        "估价对象位于{{ room_detail_location }}，建筑面积{{ building_area }}平方米，分摊土地使用权面积{{ land_area }}平方米，"
        "委托方拟按现状利用条件为其办理土地使用权出让手续。"
        "根据{{ proof_doc_name }}（{{ proof_doc_date }}），估价对象土地使用者为{{ land_user }}，{{ user_identity }}，"
        "根据房改房政策于{{ buy_year }}购得位于{{ buy_location_desc }}住房一套，房产证号：{{ house_cert_no }}，"
        "房屋登记面积{{ registered_house_area }}平方米。其房屋分摊土地使用权类型为国有划拨，因历史遗留问题，土地暂未办理土地权属登记。"
        "{%- elif land_status_type == \"registered\" -%}"
        "根据《{{ land_cert_name }}》（{{ land_cert_no }}），估价对象已取得土地权利证书，权利人为{{ land_user }}，"
        "土地使用权类型为{{ land_use_type }}，使用权面积为{{ land_area }}平方米，建筑面积为{{ building_area }}平方米。"
        "{%- endif -%}"
    )
    
    investigation_skeleton = (
        "{%- if land_status_type == \"new_grant\" -%}"
        "估价对象为拟出让用地，未办理登记手续，故暂无土地使用权人、土地权利证书编号、登记时间、地籍图号、宗地号、终止日期等相关信息。"
        "{%- elif land_status_type == \"historical_unregistered\" -%}"
        "估价对象为划拨用地，未办理登记手续，故暂无土地权利证书编号、登记时间、地籍图号、宗地号、终止日期等相关信息。"
        "{%- elif land_status_type == \"registered\" -%}"
        "估价对象已办理权属登记，相关权利证书编号、登记时间等登记状况详见“土地登记状况表”。"
        "{%- endif -%}"
    )
    
    def process_paras(paras):
        for p in paras:
            text = p.text.strip()
            if not text:
                continue
                
            new_text = text
            
            # (1) 历史权源骨架
            if "根据《关于道县2025年第七批次建设用地" in text and "农用地转用 and 土地征收的批复" in text or ("根据《关于道县2025年第七批次建设用地" in text and "农用地转用和土地征收的批复" in text):
                new_text = history_skeleton
            # (2) 调查附记骨架
            elif text.startswith("估价对象为拟出让用地，未办理登记手续，故暂无"):
                new_text = investigation_skeleton
            # (5) 方案四净化
            elif "2.基础设施条件：估价对象土地实际开发程度为宗地红线外" in text and "场地平整" in text and "详见下表" in text:
                new_text = "2.{{ infrastructure_detail }}"
                
            # (6) 土地用途精细净化 (11处)
            elif "一宗工业用地国有建设用地使用权市场价格评估" in text:
                new_text = new_text.replace("一宗工业用地国有建设用地使用权", "一宗{{ land_usage_short }}国有建设用地使用权")
            elif "一宗工业用地国有建设用地使用权实施挂牌" in text:
                new_text = new_text.replace("一宗工业用地国有建设用地使用权", "一宗{{ land_usage_short }}国有建设用地使用权")
            elif "规划用途为工业用地" in text and "土地利用现状为" in text:
                new_text = new_text.replace("规划用途为工业用地", "规划用途为{{ land_usage_short }}")
            elif "规划用途为工业用地" in text:
                new_text = new_text.replace("规划用途为工业用地", "规划用途为{{ land_usage_short }}")
            elif "一宗工业用地国有建设用地使用权" in text and "建设用地规划红线图" in text:
                new_text = new_text.replace("一宗工业用地国有建设用地使用权", "一宗{{ land_usage_short }}国有建设用地使用权")
            elif "出让年限为工业用地" in text and "{{ land_use_term }}" in text:
                new_text = new_text.replace("出让年限为工业用地{{ land_use_term }}", "出让年限为{{ land_usage_short }}{{ land_use_term }}")
            elif "出让年限为工业用地" in text and "30年" in text:
                new_text = new_text.replace("工业用地", "{{ land_usage_short }}")
            elif "故设定估价对象为工业用地" in text and "容积率" in text:
                new_text = new_text.replace("故设定估价对象为工业用地", "故设定估价对象为{{ land_usage_short }}")
            elif "原为代开用地，现为工业用地" in text:
                new_text = new_text.replace("原为代开用地，现为工业用地", "原为代开用地，现为{{ land_usage_short }}")
            elif "设定用途为工业用地" in text and "评估设定开发程度" in text:
                new_text = new_text.replace("设定用途为工业用地", "设定用途为{{ land_usage_short }}")
            elif "一宗工业用地国有建设用地使用权" in text and "拟挂牌出让" in text:
                new_text = new_text.replace("一宗工业用地国有建设用地使用权", "一宗{{ land_usage_short }}国有建设用地使用权")
            elif "一宗工业用地国有建设用地使用权" in text and "（湖南省）" in text:
                new_text = new_text.replace("一宗工业用地国有建设用地使用权", "一宗{{ land_usage_short }}国有建设用地使用权")
                
            # 物理重建段落
            if new_text != text:
                rebuild_paragraph_with_text(p, new_text)
                continue
                
            # 常规强类型语义占位符替换
            new_text = text
            if "现场勘查日期为" in text or "设定估价期日" in text or "估价期日的设定" in text or "于估价期日" in text or "估价期日设定为" in text:
                new_text = new_text.replace("2025年11月26日", "{{ valuation_date }}")
                new_text = new_text.replace("道县自然资源局", "{{ client_name }}")
                new_text = new_text.replace("道县湖南紫金新材料产业集聚区", "{{ land_location }}")
                
            if "委托方" in text or "自然资源局" in text:
                new_text = new_text.replace("道县自然资源局", "{{ client_name }}")
                new_text = new_text.replace("道县湖南紫金新材料产业集聚区", "{{ land_location }}")
                new_text = new_text.replace("待出让后确定", "{{ land_user }}")
                
            if "位于道县湖南紫金新材料" in text or "坐落" in text or "位置" in text:
                new_text = new_text.replace("道县湖南紫金新材料产业集聚区", "{{ land_location }}")
                
            if "土地用途的设定" in text or "土地用途" in text or "设定土地用途" in text:
                new_text = new_text.replace("工业用地", "{{ land_usage }}")
                
            if "容积率" in text or "容积率为" in text or "设定容积率" in text:
                new_text = new_text.replace("1.5", "{{ plot_ratio }}")
                new_text = new_text.replace("1.50", "{{ plot_ratio }}")
                
            if "土地面积" in text or "用地面积" in text or "使用权面积" in text:
                new_text = new_text.replace("4829.6", "{{ land_area }}")
                
            if "土地使用年限" in text or "土地使用年期" in text or "出让年期为" in text or "使用年限为" in text:
                new_text = new_text.replace("30年", "{{ land_use_term }}")
                
            if "二〇二" in text and "十一月" in text:
                new_text = new_text.replace("二〇二五年十一月二十八日", "{{ report_date }}")
                
            # 【规划条件限制指标精准替换：物理消除硬编码，全面支持建筑密度、容积率范围、绿地率与建筑限高】
            if "规划条件的函" in new_text or "规划建筑密度" in new_text or "绿地率" in new_text or "建筑限高" in new_text:
                new_text = new_text.replace("35%-55%", "{{ building_density_min }}-{{ building_density_max }}")
                new_text = new_text.replace("0.7-{{ plot_ratio }}", "{{ plot_ratio_min }}-{{ plot_ratio }}")
                new_text = new_text.replace("0.7-1.5", "{{ plot_ratio_min }}-{{ plot_ratio }}")
                new_text = new_text.replace("0.7-1.50", "{{ plot_ratio_min }}-{{ plot_ratio }}")
                new_text = new_text.replace("绿地率≤15%", "绿地率{{ greening_rate }}")
                new_text = new_text.replace("建筑限高24米", "建筑限高{{ building_height_limit }}")

            # 【终审合规修正：规划指标批准机关精准占位符纠正】
            if "规划条件的函" in new_text and "client_name" in new_text:
                # 纠正错误的 {{ client_name }} 为符合技术审查的专属变量 {{ planning_approval_authority }}
                new_text = new_text.replace("{{ client_name }}", "{{ planning_approval_authority }}")
                
            if new_text != text:
                rebuild_paragraph_with_text(p, new_text)

    # 1. 正文替换
    process_paras(doc.paragraphs)
    
    # 2. 物理重构方法矩阵（段落二）
    rebuild_method_selection(doc)
    
    # 3. 物理重构确价段落（段落三）8段式完美包裹与首行缩进复制注入
    rebuild_paragraph_three(doc)
    
    # 4. 表格单元格替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    if not text.strip():
                        continue
                    new_text = text
                    new_text = new_text.replace("道县自然资源局", "{{ client_name }}")
                    new_text = new_text.replace("道县湖南紫金新材料产业集聚区", "{{ land_location }}")
                    new_text = new_text.replace("2025年11月26日", "{{ valuation_date }}")
                    new_text = new_text.replace("工业用地", "{{ land_usage }}")
                    new_text = new_text.replace("二〇二五年十一月二十八日", "{{ report_date }}")
                    new_text = new_text.replace("待出让后确定", "{{ land_user }}")
                    new_text = new_text.replace("30年", "{{ land_use_term }}")
                    new_text = new_text.replace("1.5", "{{ plot_ratio }}")
                    new_text = new_text.replace("1.50", "{{ plot_ratio }}")
                    
                    if "土地面积" in new_text or "使用权面积" in new_text:
                        new_text = new_text.replace("4829.6", "{{ land_area }}")
                        
                    new_text = new_text.replace("（湖南）湘兆财土地[2025]（估）字第ZC0018号", "{{ report_no }}")
                    
                    if new_text != text:
                        rebuild_paragraph_with_text(p, new_text)

    # 5. 全局仿宋锁定
    apply_fangsong_globally(doc)
    doc.save(dest_path)
    print("  [工业模板] 全局8段包裹式物理重构与字体、段落缩进锁定大功告成！")

if __name__ == "__main__":
    dir_path = r"D:\评估报告工具\01_Source\02_Word_Templates"
    
    # 住宅文件定义
    res_src = os.path.join(dir_path, "道县月岩西路皮革厂家属楼3栋4层404房——合并报告_备份_原始.docx")
    res_dest = os.path.join(dir_path, "[模板]道县月岩西路皮革厂家属楼3栋4层404房——合并报告.docx")
    
    # 工业文件定义
    ind_src_org = os.path.join(dir_path, "道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx")
    ind_bak = os.path.join(dir_path, "道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告_备份_原始.docx")
    
    if os.path.exists(ind_src_org) and not os.path.exists(ind_bak):
        shutil.copy(ind_src_org, ind_bak)
        print(f"已为原始工业文件创建物理备份: {ind_bak}")
        
    ind_dest = os.path.join(dir_path, "[模板]道县湖南紫金新材料产业集聚区（2025018号地块）——合并报告（定稿）.docx")
    
    if os.path.exists(res_src):
        transform_residential(res_src, res_dest)
    else:
        print(f"找不到住宅原始备份: {res_src}")
        
    if os.path.exists(ind_bak):
        transform_industrial(ind_bak, ind_dest)
    else:
        print(f"找不到工业原始备份: {ind_bak}")
        
    print("\n[所有模板的 8段包裹大段落三物理重构 与 四重仿宋锁定 全部圆满达成！]")
