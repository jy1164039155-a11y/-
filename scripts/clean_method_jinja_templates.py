# -*- coding: utf-8 -*-
"""Replace old valuation-method Jinja condition blocks with generated text placeholders."""

from __future__ import annotations

from pathlib import Path

from docx import Document


TEMPLATE_DIR = Path("01_Source/02_Word_Templates")
RUNTIME_TEMPLATE_NAME = "自动生成的评估报告模板.docx"
CLEAN_RUNTIME_TEMPLATE_NAME = "自动生成的评估报告模板_清理后.docx"

REPLACEMENT_RANGES = (
    (
        "{% if use_market_comp %}市场比较法",
        "{% if not use_benchmark_corr %}",
        "{{ valuation_method_reasons }}",
    ),
    (
        "{% if use_cost_approx %}【成本逼近法】",
        "{% if use_residual %}【剩余法】",
        "{{ valuation_method_applicability }}",
    ),
)

TEXT_REPLACEMENTS = {
    "2.{{ infrastructure_detail }}": "{{ infrastructure_detail }}",
    "2．{{ infrastructure_detail }}": "{{ infrastructure_detail }}",
}


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def replace_range_once(doc: Document, start_marker: str, end_marker: str, placeholder: str) -> bool:
    paragraphs = list(doc.paragraphs)
    start_index = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip().startswith(start_marker):
            start_index = idx
            break
    if start_index is None:
        return False

    end_index = None
    for idx in range(start_index, len(paragraphs)):
        if paragraphs[idx].text.strip().startswith(end_marker):
            end_index = idx
            break
    if end_index is None:
        return False

    replace_paragraph_text(paragraphs[start_index], placeholder)
    for paragraph in paragraphs[start_index + 1 : end_index + 1]:
        remove_paragraph(paragraph)
    return True


def clean_template(path: Path) -> int:
    doc = Document(path)
    changes = 0
    made_change = True
    while made_change:
        made_change = False
        for start_marker, end_marker, placeholder in REPLACEMENT_RANGES:
            if replace_range_once(doc, start_marker, end_marker, placeholder):
                changes += 1
                made_change = True
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in TEXT_REPLACEMENTS:
            replace_paragraph_text(paragraph, TEXT_REPLACEMENTS[text])
            changes += 1
    if changes:
        try:
            doc.save(path)
        except PermissionError:
            if path.name != RUNTIME_TEMPLATE_NAME:
                raise
            fallback_path = path.with_name(CLEAN_RUNTIME_TEMPLATE_NAME)
            doc.save(fallback_path)
            print(f"{path.name}: locked, wrote cleaned runtime copy -> {fallback_path.name}")
    return changes


def main() -> None:
    for path in sorted(TEMPLATE_DIR.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        changes = clean_template(path)
        if changes:
            print(f"{path.name}: replaced {changes} method block(s)")


if __name__ == "__main__":
    main()
