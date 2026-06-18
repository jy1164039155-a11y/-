import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def scan_docx_highlights(file_path):
    doc = Document(file_path)
    results = []

    def process_runs(runs, location_desc):
        i = 0
        while i < len(runs):
            run = runs[i]
            # 安全读取高亮颜色，防止 python-docx 抛出 ValueError
            try:
                highlight = run.font.highlight_color
            except Exception:
                highlight = None

            if highlight == WD_COLOR_INDEX.YELLOW:
                full_text = run.text
                start_idx = i
                # 合并连续的高亮 Runs
                j = i + 1
                while j < len(runs):
                    try:
                        next_h = runs[j].font.highlight_color
                    except Exception:
                        next_h = None
                    if next_h == WD_COLOR_INDEX.YELLOW:
                        full_text += runs[j].text
                        j += 1
                    else:
                        break
                
                # 提取上下文
                # 前文：取当前高亮块之前所有非高亮文本的末尾
                pre_text = "".join([r.text for r in runs[max(0, start_idx-5):start_idx]])[-10:]
                # 后文：取当前高亮块之后所有非高亮文本的开头
                post_text = "".join([r.text for r in runs[j:j+5]])[:10]
                
                results.append({
                    "位置": location_desc,
                    "标黄原词": full_text.strip(),
                    "前文上下文": pre_text,
                    "后文上下文": post_text,
                    "建议变量名": "", # 留给用户填写
                    "对应Excel字段": "" # 留给用户填写
                })
                i = j # 跳过已合并的
            else:
                i += 1

    # 扫描正文段落
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            process_runs(para.runs, f"段落_{idx+1}")

    # 扫描表格
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    process_runs(para.runs, f"表格{t_idx+1}_行{r_idx+1}_列{c_idx+1}")

    return results

if __name__ == "__main__":
    input_file = r"01_Source/02_Word_Templates/合并报告_原始_20260521.docx"
    output_dir = "02_Process"
    output_file = os.path.join(output_dir, "预检清单.xlsx")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    print(f"正在扫描文档: {input_file}")
    if os.path.exists(input_file):
        data = scan_docx_highlights(input_file)
        if data:
            df = pd.DataFrame(data)
            # 去重：相同的词在相同的上下文中可能出现多次，但预检清单可以先保留位置信息
            df.to_excel(output_file, index=False)
            print(f"扫描完成！共发现 {len(df)} 处标黄内容。清单已保存至: {output_file}")
        else:
            print("未在文档中发现黄色高亮内容。")
    else:
        print(f"错误：找不到输入文件 {input_file}")
