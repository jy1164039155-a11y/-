# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import warnings
from typing import Any, Dict, List, Optional

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.services.benchmark_correction import calculate_benchmark_correction
from src.services.report_docx_styles import format_report_paragraph, format_report_table


NARRATIVE_HEADINGS = {
    "benchmark_corr_method_intro": "基准地价系数修正法估价过程",
    "benchmark_corr_base_price": "1、基准地价评估依据",
    "benchmark_corr_formula": "2、基准地价系数修正法的基本公式",
    "benchmark_corr_po": "3、确定基准地价（Po）",
    "benchmark_corr_route_price": "路线价（Po）的确定",
    "benchmark_corr_kv": "4、容积率修正系数（Kv）的确定",
    "benchmark_corr_ky": "5、土地使用年期修正系数（Ky）的确定",
    "benchmark_corr_kt": "6、期日修正系数（Kt）的确定",
    "benchmark_corr_ki": "7、区域因素修正系数（∑Ki）的确定",
    "benchmark_corr_frontage_special_intro": "商服用地临街宗地特别因素修正",
    "benchmark_corr_kd": "临街深度修正系数（Kd）的确定",
    "benchmark_corr_kk": "临街宽度修正系数（Kk）的确定",
    "benchmark_corr_kc": "街角地修正系数（Kc）的确定",
    "benchmark_corr_ks": "8、宗地面积修正系数（Ks）的确定",
    "benchmark_corr_ka": "9、宗地形状修正系数（Ka）的确定",
    "benchmark_corr_ku": "周边土地利用类型修正系数（Ku）的确定",
    "benchmark_corr_ke": "10、景观环境修正系数（Ke）的确定",
    "benchmark_corr_kto": "11、建筑物朝向修正系数（Kto）的确定",
    "benchmark_corr_kf": "12、开发程度修正额（Kf）的确定",
    "benchmark_corr_solve": "13、基准地价系数修正法地价计算",
}


def _remove_node(node: Any) -> None:
    parent = node.getparent()
    if parent is not None:
        parent.remove(node)


def _is_internal_subheading(text: str) -> bool:
    """编号小标题（如“1、基准地价评估依据”）或正文句子不应被当作方法起始标记。"""
    if re.match(r"^\d+\s*[、.．]", text):
        return True
    # 正文句子通常含句末标点；方法/章节标题不含句号。
    if "。" in text:
        return True
    return False


def _marker(document: Document, marker_texts: List[str]) -> Optional[Paragraph]:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or _is_internal_subheading(text):
            continue
        if any(text.startswith(marker) for marker in marker_texts):
            return paragraph
    return None


def _next_method_marker(document: Document, start: Paragraph) -> Optional[Paragraph]:
    seen = False
    for paragraph in document.paragraphs:
        if paragraph._p is start._p:
            seen = True
            continue
        if not seen:
            continue
        text = paragraph.text.strip()
        if (
            text.startswith("★")
            or text.startswith("三、地价的确定")
            or (re.match(r"^（[一二三四五六七八九十]+）", text) and "法" in text and "基准地价" not in text)
        ):
            return paragraph
    return None


def _remove_existing_block(start: Paragraph, stop: Paragraph) -> None:
    node = start._p
    stop_node = stop._p
    while node is not None and node is not stop_node:
        next_node = node.getnext()
        _remove_node(node)
        node = next_node


def _append_stop_marker(document: Document) -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.add_run("")
    return paragraph


def _insert_paragraph_before(document: Document, stop: Paragraph, text: str, *, role: str = "body") -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    format_report_paragraph(paragraph, role=role)
    paragraph._p.getparent().remove(paragraph._p)
    stop._p.addprevious(paragraph._p)
    return paragraph


def _narrative_line_role(text: str) -> str:
    stripped = str(text or "").strip()
    if re.match(r"^[A-Za-z\u4e00-\u9fff]*[A-Za-z0-9住]*\s*[=＝]", stripped):
        return "formula"
    return "body"


def _row_values(row: Dict[str, Any], columns: List[str]) -> List[str]:
    if row.get("cells"):
        values = []
        for value in row.get("cells") or []:
            if isinstance(value, dict):
                values.append(str(value.get("label") or value.get("value") or value.get("text") or ""))
            else:
                values.append(str(value or ""))
        return values
    return [str(row.get(column) or "") for column in columns]


def _row_cells(row: Dict[str, Any], columns: List[str]) -> List[Any]:
    if row.get("cells"):
        return list(row.get("cells") or [])
    return [{"label": row.get(column) or ""} for column in columns]


def _cell_label(cell: Any) -> str:
    if isinstance(cell, dict):
        return str(cell.get("label") or cell.get("value") or cell.get("text") or "")
    return str(cell or "")


def _cell_span(cell: Any, key: str) -> int:
    if isinstance(cell, dict):
        try:
            return max(1, int(cell.get(key) or 1))
        except (TypeError, ValueError):
            return 1
    return 1


def _insert_table_before(document: Document, stop: Paragraph, section: Dict[str, Any]) -> None:
    columns = list(section.get("columns") or [])
    rows = list(section.get("rows") or [])
    header_rows = list(section.get("header_rows") or [])
    if not columns:
        return
    _insert_paragraph_before(document, stop, str(section.get("report_title") or section.get("title") or ""), role="table_title")
    table = document.add_table(rows=max(1, len(header_rows) or 1), cols=len(columns))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    if header_rows:
        for row_index, header in enumerate(header_rows):
            col_index = 0
            for item in header:
                while col_index < len(columns) and table.cell(row_index, col_index).text:
                    col_index += 1
                if col_index >= len(columns):
                    break
                target = table.cell(row_index, col_index)
                target.text = _cell_label(item)
                colspan = _cell_span(item, "colspan")
                rowspan = _cell_span(item, "rowspan")
                end_col = min(len(columns) - 1, col_index + colspan - 1)
                end_row = min(len(header_rows) - 1, row_index + rowspan - 1)
                if end_col != col_index or end_row != row_index:
                    target.merge(table.cell(end_row, end_col))
                col_index += colspan
    else:
        for cell, label in zip(table.rows[0].cells, columns):
            cell.text = str(label or "")
    body_start = len(table.rows)
    for _ in rows:
        table.add_row()
    for offset, row in enumerate(rows):
        row_index = body_start + offset
        col_index = 0
        for item in _row_cells(row, columns):
            if isinstance(item, dict) and item.get("hidden"):
                col_index += _cell_span(item, "colspan")
                continue
            while col_index < len(columns) and table.cell(row_index, col_index).text:
                col_index += 1
            if col_index >= len(columns):
                break
            target = table.cell(row_index, col_index)
            target.text = _cell_label(item)
            colspan = _cell_span(item, "colspan")
            rowspan = _cell_span(item, "rowspan")
            end_col = min(len(columns) - 1, col_index + colspan - 1)
            end_row = min(len(table.rows) - 1, row_index + rowspan - 1)
            if end_col != col_index or end_row != row_index:
                target.merge(table.cell(end_row, end_col))
            col_index += colspan
    table._tbl.getparent().remove(table._tbl)
    stop._p.addprevious(table._tbl)
    _format_table(table, header_rows=max(1, len(header_rows) or 1))


def _format_table(table: Table, *, header_rows: int = 1) -> None:
    column_count = max(1, len(table.columns))
    font_size = 9.0 if column_count >= 6 else 10.5
    widths = tuple(16.1 / column_count for _ in range(column_count))
    format_report_table(table, widths_cm=widths, font_size_pt=font_size, header_rows=header_rows)


def apply_benchmark_correction_to_docx(docx_path: str, data: Dict[str, Any], base_dir: str) -> None:
    document = Document(docx_path)
    start = _marker(document, ["★基准地价系数修正法", "★公示地价系数修正法", "（一）基准地价系数修正法", "基准地价系数修正法估价过程"])
    if start is None:
        if data.get("use_benchmark_corr"):
            raise RuntimeError("未找到基准地价系数修正法模板起始标记，无法写入动态正文。")
        # 未启用该方法且无起始标记：无可定位的样本块，显式告警而非静默跳过。
        warnings.warn(
            "未找到基准地价系数修正法模板起始标记，无法定位样本块；"
            "若模板内仍存在该方法样本段落，请人工核对是否残留。",
            stacklevel=2,
        )
        return
    stop = _next_method_marker(document, start)
    if stop is None:
        # 缺少“下一方法/章节”停止标记：在文末追加锚点以便清理样本块，避免整段残留；显式告警。
        warnings.warn(
            "未找到基准地价系数修正法之后的下一方法/章节停止标记，"
            "将在文末追加停止锚点后清理样本块，请确认其后无需保留的内容。",
            stacklevel=2,
        )
        stop = _append_stop_marker(document)

    analysis = calculate_benchmark_correction(data, base_dir) if data.get("use_benchmark_corr") else {}
    support_status = str(analysis.get("support_status") or "")
    if not data.get("use_benchmark_corr") or support_status == "unsupported":
        _remove_existing_block(start, stop)
        document.save(docx_path)
        return

    effective = analysis.get("effective_narratives") or analysis.get("generated_narratives") or {}
    table_map = {item.get("key"): item for item in analysis.get("tables") or []}
    _remove_existing_block(start, stop)
    _insert_paragraph_before(document, stop, "★基准地价系数修正法", role="method_heading")
    if analysis.get("warnings"):
        warning_text = "；".join(str(item.get("message") or "") for item in analysis.get("warnings") or [] if item.get("message"))
        _insert_paragraph_before(document, stop, f"【待校核：{warning_text}】")

    for block in analysis.get("content_blocks") or []:
        block_type = str(block.get("type") or "")
        key = str(block.get("key") or "")
        if block_type == "narrative":
            heading = NARRATIVE_HEADINGS.get(key, "")
            if heading:
                heading_paragraph = _insert_paragraph_before(document, stop, heading)
                heading_paragraph.paragraph_format.keep_with_next = True
            text = str(effective.get(key) or "").strip()
            for line in text.splitlines() or ["【待校核】"]:
                role = _narrative_line_role(line)
                paragraph = _insert_paragraph_before(document, stop, line, role=role)
                if role == "formula":
                    paragraph.paragraph_format.keep_together = True
        elif block_type == "table" and key in table_map:
            _insert_table_before(document, stop, table_map[key])

    document.save(docx_path)
