# -*- coding: utf-8 -*-
from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import threading
import time
import urllib.error
import urllib.request
import uuid
from datetime import date, datetime, timedelta
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from src.schemas.comparable import FactorScheme, MarketComparisonAnalysis
from src.services.land_usage import (
    LAND_USAGE_OPTION_BY_KEY,
    infer_land_usage_key,
    land_usage_first_level_label,
    official_land_usage_code,
)
from src.services.valuation_builder import MARKET_COMP_INTRO


LANDCHINA_API = "https://api.landchina.com"
LANDCHINA_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/142.0.0.0 Safari/537.36"
)
LANDCHINA_RESULT_URL = (
    "https://www.landchina.com/#/landSupplyDetail?id={gd_guid}"
    "&type=%E4%BE%9B%E5%9C%B0%E7%BB%93%E6%9E%9C&path=0"
)
CASE_SLOTS = ("A", "B", "C")
COMPARABLE_BASIS_TEMPLATE = "\n".join(
    [
        "选取了比较实例后，一般应先对这些比较实例的成交价格进行换算处理，对价格内涵和形式进行“标准化”，使它们之间的口径一致、相互可比，为后续对比较实例成交价格进行修正和调整建立一个共同的基础。",
        "建立可比基础一般要做以下工作：①统一付款方式；②统一币种和货币单位；③统一面积内涵和面积单位等。",
        "根据对比较实例成交价格内涵和形式的调查分析与待估宗地价格比较，两者相同，具体情况为：付款方式为在交易日期一次性付清，税费负担为各付各税，计价单位为人民币元/平方米.土地面积，面积内涵不包含非土地成分，土地实物范围没有特殊情况，土地面积单位为平方米等。",
        "可比实例与估价对象在付款方式、币种、货币单位、面积内涵和面积单位是统一的，价格可比基础相同，故无需进行调整。",
    ]
)
LANDCHINA_CRAWL_PAGE_SIZE = 40
LANDCHINA_CRAWL_MAX_TOTAL = 6000
LANDCHINA_DIRECT_REQUEST_INTERVAL = 0.35
LANDCHINA_DIRECT_DETAIL_INTERVAL = 1.1
LANDCHINA_PROXY_REQUEST_INTERVAL = 2.0
LANDCHINA_PROXY_DETAIL_INTERVAL = 2.8
LANDCHINA_REGION_ALIASES = {
    "431124": {"道县", "湖南省永州市道县"},
    "431230": {"通道县", "通道侗族自治县", "湖南省怀化市通道县", "湖南省怀化市通道侗族自治县"},
}


def _libreoffice_ascii_temp_root() -> Path:
    """Return a writable ASCII-only temp root for LibreOffice conversion."""
    candidates = [Path(r"C:\tmp"), Path(tempfile.gettempdir())]
    for root in candidates:
        try:
            root.mkdir(parents=True, exist_ok=True)
            with tempfile.TemporaryDirectory(prefix="lo_probe_", dir=root):
                pass
            return root
        except Exception:
            continue
    return Path(tempfile.gettempdir())


class LandChinaRateLimitError(RuntimeError):
    pass


class LandChinaAccessGate:
    def __init__(self, db_path: Path):
        self.db_path = Path(db_path)
        self._init_db()

    def _connect(self) -> sqlite3.Connection:
        connection = sqlite3.connect(self.db_path, timeout=30)
        connection.row_factory = sqlite3.Row
        return connection

    def _init_db(self) -> None:
        with self._connect() as conn:
            conn.executescript(
                """
                CREATE TABLE IF NOT EXISTS landchina_access_state (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    next_request_at REAL NOT NULL DEFAULT 0,
                    blocked_until REAL NOT NULL DEFAULT 0,
                    last_http_status TEXT NOT NULL DEFAULT '',
                    last_reason TEXT NOT NULL DEFAULT '',
                    updated_at TEXT NOT NULL
                );
                INSERT OR IGNORE INTO landchina_access_state(id, updated_at) VALUES (1, '');
                CREATE TABLE IF NOT EXISTS landchina_crawl_lease (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    job_id TEXT NOT NULL DEFAULT '',
                    expires_at REAL NOT NULL DEFAULT 0,
                    updated_at TEXT NOT NULL
                );
                INSERT OR IGNORE INTO landchina_crawl_lease(id, updated_at) VALUES (1, '');
                """
            )
            now = time.time()
            conn.execute(
                """
                UPDATE landchina_access_state
                SET blocked_until=0,
                    next_request_at=0,
                    last_reason=CASE
                        WHEN last_reason = '' THEN '已清除旧版重定向冷却状态。'
                        ELSE last_reason || '；已清除旧版重定向冷却状态。'
                    END,
                    updated_at=?
                WHERE last_http_status IN ('301', '302')
                  AND (blocked_until > ? OR next_request_at > ?)
                """,
                (_now(), now, now),
            )

    def status(self) -> Dict[str, Any]:
        now = time.time()
        with self._connect() as conn:
            access = conn.execute("SELECT * FROM landchina_access_state WHERE id=1").fetchone()
            lease = conn.execute("SELECT * FROM landchina_crawl_lease WHERE id=1").fetchone()
        blocked_until = float(access["blocked_until"] or 0)
        lease_expires_at = float(lease["expires_at"] or 0)
        return {
            "blocked": blocked_until > now,
            "blocked_until": datetime.fromtimestamp(blocked_until).isoformat(timespec="seconds") if blocked_until > now else "",
            "remaining_seconds": max(int(blocked_until - now), 0),
            "last_http_status": access["last_http_status"] or "",
            "last_reason": access["last_reason"] or "",
            "active_job_id": lease["job_id"] if lease_expires_at > now else "",
            "active_job_expires_at": (
                datetime.fromtimestamp(lease_expires_at).isoformat(timespec="seconds")
                if lease_expires_at > now
                else ""
            ),
        }

    def reserve_request(self, interval: float) -> None:
        now = time.time()
        with self._connect() as conn:
            conn.execute("BEGIN IMMEDIATE")
            row = conn.execute("SELECT * FROM landchina_access_state WHERE id=1").fetchone()
            blocked_until = float(row["blocked_until"] or 0)
            if blocked_until > now:
                retry_at = datetime.fromtimestamp(blocked_until).strftime("%Y-%m-%d %H:%M:%S")
                raise LandChinaRateLimitError(f"官网访问冷却中，请于 {retry_at} 后重试。")
            scheduled_at = max(now, float(row["next_request_at"] or 0))
            conn.execute(
                "UPDATE landchina_access_state SET next_request_at=?, updated_at=? WHERE id=1",
                (scheduled_at + max(interval, 0), _now()),
            )
        delay = scheduled_at - now
        if delay > 0:
            time.sleep(delay)

    def block(self, seconds: float, http_status: Any, reason: str) -> None:
        now = time.time()
        with self._connect() as conn:
            conn.execute("BEGIN IMMEDIATE")
            row = conn.execute("SELECT blocked_until FROM landchina_access_state WHERE id=1").fetchone()
            blocked_until = max(float(row["blocked_until"] or 0), now + max(seconds, 0))
            conn.execute(
                """
                UPDATE landchina_access_state
                SET blocked_until=?, next_request_at=?, last_http_status=?, last_reason=?, updated_at=?
                WHERE id=1
                """,
                (blocked_until, blocked_until, str(http_status or ""), reason, _now()),
            )

    def record_failure(self, http_status: Any, reason: str) -> None:
        with self._connect() as conn:
            conn.execute("BEGIN IMMEDIATE")
            conn.execute(
                """
                UPDATE landchina_access_state
                SET last_http_status=?, last_reason=?, updated_at=?
                WHERE id=1
                """,
                (str(http_status or ""), reason, _now()),
            )

    def clear_cooldown(self, reason: str = "") -> None:
        with self._connect() as conn:
            conn.execute("BEGIN IMMEDIATE")
            conn.execute(
                """
                UPDATE landchina_access_state
                SET blocked_until=0,
                    next_request_at=0,
                    last_http_status='',
                    last_reason=?,
                    updated_at=?
                WHERE id=1
                """,
                (reason, _now()),
            )

    def acquire_crawl(self, job_id: str, ttl_seconds: int = 21600) -> bool:
        now = time.time()
        with self._connect() as conn:
            conn.execute("BEGIN IMMEDIATE")
            row = conn.execute("SELECT job_id, expires_at FROM landchina_crawl_lease WHERE id=1").fetchone()
            if row["job_id"] and float(row["expires_at"] or 0) > now and row["job_id"] != job_id:
                return False
            conn.execute(
                "UPDATE landchina_crawl_lease SET job_id=?, expires_at=?, updated_at=? WHERE id=1",
                (job_id, now + ttl_seconds, _now()),
            )
        return True

    def release_crawl(self, job_id: str) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                UPDATE landchina_crawl_lease SET job_id='', expires_at=0, updated_at=?
                WHERE id=1 AND job_id=?
                """,
                (_now(), job_id),
            )


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _decimal(value: Any, default: str = "0") -> Decimal:
    try:
        return Decimal(str(value if value not in (None, "") else default))
    except (InvalidOperation, ValueError):
        return Decimal(default)


def _q(value: Decimal, places: str) -> str:
    return str(value.quantize(Decimal(places), rounding=ROUND_HALF_UP))


def _date_text(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, (int, float)):
        return datetime.fromtimestamp(float(value) / 1000).strftime("%Y-%m-%d")
    text = str(value).strip()
    if "T" in text:
        return text.split("T", 1)[0]
    return text[:10]


def _cn_date(value: Any) -> str:
    text = _date_text(value)
    try:
        parsed = datetime.strptime(text, "%Y-%m-%d")
        return f"{parsed.year}年{parsed.month}月{parsed.day}日"
    except ValueError:
        return text


def _report_text(value: Any) -> str:
    text = str(value if value is not None else "").strip()
    if not text or text == "______" or text.startswith(("【请", "【待")):
        return ""
    return text


def _narrative_help_text(value: Any) -> str:
    text = _report_text(value)
    if text.startswith("请"):
        text = text[1:].lstrip()
    return text


def _content_hash(data: Any) -> str:
    raw = json.dumps(data, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _repair_mojibake(value: Any) -> Any:
    if isinstance(value, dict):
        return {key: _repair_mojibake(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_repair_mojibake(item) for item in value]
    if not isinstance(value, str) or not value:
        return value
    try:
        repaired = value.encode("latin-1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return value
    original_cjk = sum("\u4e00" <= char <= "\u9fff" for char in value)
    repaired_cjk = sum("\u4e00" <= char <= "\u9fff" for char in repaired)
    return repaired if repaired_cjk > original_cjk else value


COMMON_FACTORS = [
    {"key": "land_usage", "label": "土地用途", "group": "交易因素", "source": "official"},
    {"key": "transaction_time", "label": "交易时间", "group": "交易因素", "source": "calculated"},
    {"key": "transaction_condition", "label": "交易情况", "group": "交易因素", "source": "manual"},
    {"key": "transaction_method", "label": "交易方式", "group": "交易因素", "source": "official"},
    {"key": "use_term", "label": "土地使用年限", "group": "交易因素", "source": "official"},
    {"key": "development_level", "label": "开发程度", "group": "个别因素", "source": "manual"},
    {"key": "parcel_scale", "label": "宗地规模", "group": "个别因素", "source": "official"},
    {"key": "parcel_shape", "label": "宗地形状", "group": "个别因素", "source": "manual"},
    {"key": "terrain", "label": "地势", "group": "个别因素", "source": "manual"},
    {"key": "geology", "label": "地质条件", "group": "个别因素", "source": "manual"},
    {"key": "planning_restriction", "label": "规划限制", "group": "个别因素", "source": "manual"},
]

USAGE_FACTORS = {
    "residential": ["商服繁华度", "公共交通便捷度", "教育医疗配套", "居住环境质量", "未来发展潜力"],
    "commercial": ["商服繁华度", "人流集聚度", "临路状况", "公共交通便捷度", "未来发展潜力"],
    "industrial": ["距高速公路出入口距离", "距货运火车站距离", "临近道路状况", "产业集聚程度", "公共配套设施完善度", "环境质量优劣度"],
    "warehouse": ["距高速公路出入口距离", "距货运场站距离", "临近道路状况", "物流产业集聚程度", "基础设施完善度"],
    "public": ["公共交通便捷度", "服务人口覆盖度", "公共配套完善度", "环境质量优劣度"],
    "transportation": ["路网衔接程度", "交通流量条件", "区域可达性", "基础设施完善度"],
    "utility": ["服务覆盖条件", "管网衔接条件", "区域可达性", "环境影响程度"],
    "green": ["区域可达性", "景观环境质量", "公共服务辐射度", "周边开发成熟度"],
    "special": ["区域可达性", "用途适配程度", "基础设施完善度", "环境限制程度"],
    "other": ["区域可达性", "基础设施完善度", "环境质量优劣度", "用途适配程度"],
}

FACTOR_LEVEL_PRESETS = {
    "商服繁华度": (("优", "108"), ("较优", "104"), ("一般", "100"), ("较劣", "96"), ("劣", "92")),
    "公共交通便捷度": (("优", "108"), ("较优", "104"), ("一般", "100"), ("较劣", "96"), ("劣", "92")),
    "临近道路状况": (("通畅", "104"), ("较通畅", "102"), ("一般通畅", "100"), ("较不通畅", "98"), ("不通畅", "96")),
    "环境质量优劣度": (("好", "104"), ("较好", "102"), ("一般", "100"), ("较差", "98"), ("差", "96")),
    "居住环境质量": (("好", "104"), ("较好", "102"), ("一般", "100"), ("较差", "98"), ("差", "96")),
    "景观环境质量": (("好", "104"), ("较好", "102"), ("一般", "100"), ("较差", "98"), ("差", "96")),
    "parcel_shape": (("规则", "102"), ("较规则", "101"), ("一般", "100"), ("较不规则", "99"), ("不规则", "98")),
    "terrain": (("平坦", "102"), ("较平坦", "101"), ("一般", "100"), ("略有起伏", "99"), ("起伏较大", "98")),
    "geology": (("良好", "102"), ("较好", "101"), ("一般", "100"), ("较差", "99"), ("差", "98")),
    "planning_restriction": (("限制少", "102"), ("限制较少", "101"), ("一般", "100"), ("限制较多", "99"), ("限制多", "98")),
    "transaction_condition": (("正常", "100"), ("有利", "102"), ("较有利", "101"), ("较不利", "99"), ("不利", "98")),
    "development_level": (("七通一平", "104"), ("五通一平", "102"), ("三通一平", "100"), ("基础设施较弱", "98")),
}


def _factor_levels(key: str, label: str) -> List[Dict[str, str]]:
    preset = FACTOR_LEVEL_PRESETS.get(key) or FACTOR_LEVEL_PRESETS.get(label)
    if not preset:
        return []
    return [
        {
            "label": level_label,
            "index": level_index,
            "description": f"选择“{level_label}”时建议采用指数 {level_index}，采用前须由估价师结合实例资料确认。",
        }
        for level_label, level_index in preset
    ]


def _normalize_factor_definition(definition: Dict[str, Any], order: int) -> Dict[str, Any]:
    item = dict(definition)
    item.setdefault("required", True)
    item.setdefault("source", "manual")
    item.setdefault("note", "")
    item.setdefault("help_text", item.get("note") or "请结合实例资料、现场调查和估价经验判定。")
    item.setdefault("levels", _factor_levels(str(item.get("key") or ""), str(item.get("label") or "")))
    item.setdefault("review_status", "needs_review")
    item.setdefault("order", order)
    item.setdefault("enabled", True)
    return item


def normalize_factor_scheme(scheme: Dict[str, Any], land_usage_key: str) -> Dict[str, Any]:
    payload = dict(scheme or {})
    payload["land_usage_key"] = land_usage_key
    payload.setdefault("name", f"{land_usage_key}市场比较法因子方案")
    payload["factors"] = [
        _normalize_factor_definition(definition, index)
        for index, definition in enumerate(payload.get("factors") or [])
    ]
    return payload


def default_factor_scheme(land_usage_key: str) -> Dict[str, Any]:
    usage_key = land_usage_key if land_usage_key in USAGE_FACTORS else "other"
    factors = [
        _normalize_factor_definition(dict(item, required=True, note=""), index)
        for index, item in enumerate(COMMON_FACTORS)
    ]
    factors.extend(
        _normalize_factor_definition(
            {
                "key": f"regional_{index + 1}",
                "label": label,
                "group": "区域因素",
                "required": True,
                "source": "manual",
                "note": "系统提供建议值，采用前须由估价师确认。",
            },
            len(COMMON_FACTORS) + index,
        )
        for index, label in enumerate(USAGE_FACTORS[usage_key])
    )
    return {"land_usage_key": land_usage_key, "name": f"{land_usage_key}市场比较法因子方案", "factors": factors}


class LandChinaClient:
    def __init__(
        self,
        timeout: float = 20.0,
        request_interval: float = LANDCHINA_DIRECT_REQUEST_INTERVAL,
        detail_interval: float = LANDCHINA_DIRECT_DETAIL_INTERVAL,
        access_gate: Optional[LandChinaAccessGate] = None,
        proxy_url: Optional[str] = None,
        proxy_token: Optional[str] = None,
    ):
        self.timeout = timeout
        self.request_interval = request_interval
        self.detail_interval = detail_interval
        self.access_gate = access_gate
        proxy_url = os.environ.get("LANDCHINA_PROXY_URL", "") if proxy_url is None else proxy_url
        proxy_token = os.environ.get("LANDCHINA_PROXY_TOKEN", "") if proxy_token is None else proxy_token
        self.proxy_url = str(proxy_url or "").strip().rstrip("/")
        self.proxy_token = str(proxy_token or "").strip()
        self.rate_limit_hits = 0
        self._rate_lock = threading.Lock()
        self._next_request_at = 0.0

    def _wait_for_request_slot(self, interval: float) -> None:
        if self.access_gate:
            self.access_gate.reserve_request(interval)
            return
        with self._rate_lock:
            now = time.monotonic()
            scheduled_at = max(now, self._next_request_at)
            self._next_request_at = scheduled_at + max(interval, 0)
        delay = scheduled_at - now
        if delay > 0:
            time.sleep(delay)

    def _cool_down(self, seconds: float, http_status: Any = "", reason: str = "") -> None:
        if self.access_gate:
            self.access_gate.block(seconds, http_status, reason)
            return
        with self._rate_lock:
            self._next_request_at = max(self._next_request_at, time.monotonic() + max(seconds, 0))

    @staticmethod
    def _request_hash(path: str, current: Optional[datetime] = None) -> str:
        action = path.rstrip("/").rsplit("/", 1)[-1]
        day = (current or datetime.now()).day
        raw = f"{LANDCHINA_USER_AGENT}{day}{action}"
        return hashlib.sha256(raw.encode("utf-8")).hexdigest()

    def _record_failure(self, http_status: Any, reason: str) -> None:
        if self.access_gate:
            self.access_gate.record_failure(http_status, reason)

    @staticmethod
    def _read_http_error_body(exc: urllib.error.HTTPError) -> str:
        try:
            raw = exc.read()
        except Exception:
            return ""
        if not raw:
            return ""
        try:
            return raw.decode("utf-8", errors="replace")
        except AttributeError:
            return str(raw)

    def _build_request(self, path: str, payload: Dict[str, Any]) -> urllib.request.Request:
        if self.proxy_url:
            body = json.dumps({"path": path, "payload": payload}, ensure_ascii=False).encode("utf-8")
            headers = {
                "Content-Type": "application/json;charset=UTF-8",
                "Accept": "application/json, text/plain, */*",
                "User-Agent": LANDCHINA_USER_AGENT,
            }
            if self.proxy_token:
                headers["X-LandChina-Proxy-Token"] = self.proxy_token
            return urllib.request.Request(
                self.proxy_url,
                data=body,
                method="POST",
                headers=headers,
            )

        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        return urllib.request.Request(
            LANDCHINA_API + path,
            data=body,
            method="POST",
            headers={
                "Content-Type": "application/json;charset=UTF-8",
                "Accept": "application/json, text/plain, */*",
                "Accept-Language": "zh-CN,zh;q=0.9",
                "Cache-Control": "no-cache",
                "Origin": "https://www.landchina.com",
                "Referer": "https://www.landchina.com/",
                "User-Agent": LANDCHINA_USER_AGENT,
                "Hash": self._request_hash(path),
            },
        )

    def _open_request(self, request: urllib.request.Request):
        # Both access modes must bypass the workstation's system/environment proxy.
        # In direct mode, inheriting a stale HTTP(S)_PROXY makes the UI claim
        # "本机直连" while urllib is still routing through another server.
        opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
        return opener.open(request, timeout=self.timeout)

    def _post(
        self,
        path: str,
        payload: Dict[str, Any],
        retries: int = 2,
        interval: Optional[float] = None,
    ) -> Dict[str, Any]:
        last_error: Optional[Exception] = None
        for attempt in range(retries + 1):
            self._wait_for_request_slot(self.request_interval if interval is None else interval)
            try:
                request = self._build_request(path, payload)
                with self._open_request(request) as response:
                    result = json.loads(response.read().decode("utf-8"))
                result = _repair_mojibake(result)
                if result.get("code") != 200:
                    raise RuntimeError(result.get("msg") or f"LandChina API code={result.get('code')}")
                return result
            except urllib.error.HTTPError as exc:
                last_error = exc
                if exc.code in (301, 302):
                    location = (exc.headers.get("Location") if exc.headers else "") or ""
                    reason = f"官网接口被重定向（HTTP {exc.code}）"
                    if location:
                        reason += f"，跳转至 {location}"
                    if "#/404" in location or "/404" in location:
                        reason += "；可能是官方接口规则变化或请求签名未通过。"
                    self._record_failure(exc.code, reason)
                    last_error = RuntimeError(reason)
                    if self.access_gate or "#/404" in location or "/404" in location:
                        break
                    if attempt < retries:
                        time.sleep(min(2 ** attempt, 8))
                elif self.proxy_url and exc.code == 401:
                    reason = "云服务器中转认证失败，请检查本机 ProxyToken 是否与云服务器 -Token 完全一致。"
                    self._record_failure(exc.code, reason)
                    last_error = RuntimeError(reason)
                    break
                elif exc.code in (500, 502, 503, 504):
                    body = self._read_http_error_body(exc)
                    proxy_detail = ""
                    official_status = None
                    message = ""
                    if body:
                        try:
                            parsed_body = json.loads(body)
                        except json.JSONDecodeError:
                            parsed_body = None
                        if isinstance(parsed_body, dict):
                            official_status = parsed_body.get("official_status")
                            message = parsed_body.get("message") or parsed_body.get("error")
                            location = parsed_body.get("location")
                            detail_parts = []
                            if official_status not in (None, ""):
                                detail_parts.append(f"官网状态 {official_status}")
                            if location:
                                detail_parts.append(f"跳转至 {location}")
                            if message:
                                detail_parts.append(str(message))
                            proxy_detail = "，" + "，".join(detail_parts) if detail_parts else ""
                    reason = (
                        f"云服务器中转访问官网临时失败（HTTP {exc.code}）{proxy_detail}"
                        if self.proxy_url
                        else f"官网接口临时失败（HTTP {exc.code}）"
                    )
                    timeout_like = self.proxy_url and any(
                        marker in str(message).lower() for marker in ("超时", "timeout", "timed out")
                    )
                    cooldown_seconds = 0 if timeout_like else 300 if self.proxy_url else 180
                    self._cool_down(cooldown_seconds, exc.code, reason)
                    last_error = LandChinaRateLimitError(reason)
                    if self.access_gate:
                        break
                    if attempt < retries:
                        time.sleep(min(2 ** attempt, 8))
                elif exc.code in (403, 418, 429):
                    self.rate_limit_hits += 1
                    retry_after = exc.headers.get("Retry-After") if exc.headers else None
                    try:
                        retry_after_seconds = float(retry_after or 0)
                    except ValueError:
                        retry_after_seconds = 0
                    location = exc.headers.get("Location") if exc.headers else ""
                    if exc.code == 403:
                        cooldown_seconds = max(retry_after_seconds, 1800)
                    else:
                        cooldown_seconds = max(retry_after_seconds, 300)
                    reason = f"官网访问受限（HTTP {exc.code}）"
                    if location:
                        reason += f"，跳转至 {location}"
                    self._cool_down(cooldown_seconds, exc.code, reason)
                    last_error = LandChinaRateLimitError(reason)
                    if self.access_gate:
                        break
                elif attempt < retries:
                    time.sleep(min(2 ** attempt, 8))
            except (urllib.error.URLError, TimeoutError, RuntimeError, json.JSONDecodeError) as exc:
                last_error = exc
                if attempt < retries:
                    time.sleep(min(2 ** attempt, 8))
            if attempt >= retries:
                break
        if isinstance(last_error, LandChinaRateLimitError):
            raise last_error
        raise RuntimeError(f"中国土地市场网请求失败：{last_error}") from last_error

    def result_list(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        return self._post("/tGdxm/result/list", payload, retries=3).get("data") or {}

    def result_detail(self, gd_guid: str) -> Dict[str, Any]:
        result = self._post(
            "/tGdxm/result/detail",
            {"gdGuid": gd_guid},
            retries=2,
            interval=self.detail_interval,
        )
        detail = result.get("data") or {}
        if not detail:
            raise RuntimeError("中国土地市场网详情接口返回空数据")
        return detail

    def regions(self, parent_xzq: str = "") -> List[Dict[str, Any]]:
        result = self._post("/bptFieldEnum/xzq", {"parentXzq": parent_xzq})
        return result.get("data") or []

    def land_usage_tree(self) -> List[Dict[str, Any]]:
        result = self._post("/bptFieldEnum/tdytTreeList", {})
        return result.get("data") or []


class ComparableLibrary:
    def __init__(self, base_dir: str):
        self.base_dir = Path(base_dir)
        self.data_dir = self.base_dir / "data"
        self.evidence_dir = self.data_dir / "comparable_evidence"
        self.db_path = self.data_dir / "comparable_library.sqlite3"
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.evidence_dir.mkdir(parents=True, exist_ok=True)
        self.access_gate = LandChinaAccessGate(self.db_path)
        self.client = LandChinaClient(access_gate=self.access_gate)
        self.jobs: Dict[str, Dict[str, Any]] = {}
        self.jobs_lock = threading.Lock()
        self._init_db()
        self._apply_proxy_config()
        self._reclassify_existing_cases()

    def _connect(self) -> sqlite3.Connection:
        connection = sqlite3.connect(self.db_path, timeout=30)
        connection.row_factory = sqlite3.Row
        return connection

    def _init_db(self) -> None:
        with self._connect() as conn:
            conn.executescript(
                """
                CREATE TABLE IF NOT EXISTS comparable_cases (
                    id TEXT PRIMARY KEY,
                    gd_guid TEXT NOT NULL UNIQUE,
                    electronic_supervision_no TEXT,
                    official_json TEXT NOT NULL,
                    manual_json TEXT NOT NULL DEFAULT '{}',
                    manual_draft_json TEXT NOT NULL DEFAULT '{}',
                    official_hash TEXT NOT NULL,
                    fetched_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                );
                CREATE UNIQUE INDEX IF NOT EXISTS idx_case_supervision
                    ON comparable_cases(electronic_supervision_no)
                    WHERE electronic_supervision_no IS NOT NULL AND electronic_supervision_no != '';
                CREATE TABLE IF NOT EXISTS comparable_case_snapshots (
                    snapshot_id TEXT PRIMARY KEY,
                    case_id TEXT NOT NULL,
                    snapshot_hash TEXT NOT NULL,
                    snapshot_json TEXT NOT NULL,
                    docx_path TEXT,
                    pdf_path TEXT,
                    image_paths_json TEXT NOT NULL DEFAULT '[]',
                    created_at TEXT NOT NULL,
                    UNIQUE(case_id, snapshot_hash)
                );
                CREATE TABLE IF NOT EXISTS comparable_case_history (
                    history_id TEXT PRIMARY KEY,
                    case_id TEXT NOT NULL,
                    official_hash TEXT NOT NULL,
                    official_json TEXT NOT NULL,
                    fetched_at TEXT NOT NULL,
                    UNIQUE(case_id, official_hash)
                );
                CREATE TABLE IF NOT EXISTS factor_schemes (
                    land_usage_key TEXT PRIMARY KEY,
                    scheme_json TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS landchina_proxy_config (
                    id INTEGER PRIMARY KEY CHECK (id = 1),
                    enabled INTEGER NOT NULL DEFAULT 0,
                    proxy_url TEXT NOT NULL DEFAULT '',
                    proxy_token TEXT NOT NULL DEFAULT '',
                    updated_at TEXT NOT NULL
                );
                """
            )
            case_columns = {
                str(row["name"])
                for row in conn.execute("PRAGMA table_info(comparable_cases)").fetchall()
            }
            if "manual_draft_json" not in case_columns:
                conn.execute(
                    "ALTER TABLE comparable_cases ADD COLUMN manual_draft_json TEXT NOT NULL DEFAULT '{}'"
                )
            conn.execute(
                """
                INSERT OR IGNORE INTO landchina_proxy_config
                    (id, enabled, proxy_url, proxy_token, updated_at)
                VALUES (1, ?, ?, ?, ?)
                """,
                (
                    1 if self.client.proxy_url else 0,
                    self.client.proxy_url,
                    self.client.proxy_token,
                    _now(),
                ),
            )
        for key in USAGE_FACTORS:
            if not self.get_factor_scheme(key, create=False):
                self.save_factor_scheme(key, default_factor_scheme(key))

    @staticmethod
    def _bool_config(value: Any) -> bool:
        if isinstance(value, str):
            return value.strip().lower() in {"1", "true", "yes", "y", "on", "启用", "是"}
        return bool(value)

    def _read_proxy_config_row(self, conn: sqlite3.Connection) -> sqlite3.Row:
        return conn.execute("SELECT * FROM landchina_proxy_config WHERE id=1").fetchone()

    def _proxy_config_payload(self, row: sqlite3.Row) -> Dict[str, Any]:
        enabled = bool(row["enabled"])
        proxy_url = str(row["proxy_url"] or "").strip()
        token = str(row["proxy_token"] or "")
        return {
            "enabled": enabled,
            "proxy_url": proxy_url,
            "token_set": bool(token),
            "updated_at": row["updated_at"] or "",
            "access_channel": "cloud_proxy" if enabled and proxy_url else "direct",
        }

    def _apply_proxy_config(self) -> Dict[str, Any]:
        with self._connect() as conn:
            row = self._read_proxy_config_row(conn)
        enabled = bool(row["enabled"])
        proxy_url = str(row["proxy_url"] or "").strip()
        proxy_token = str(row["proxy_token"] or "").strip()
        self.client.proxy_url = proxy_url.rstrip("/") if enabled and proxy_url else ""
        self.client.proxy_token = proxy_token if enabled else ""
        self.client.request_interval = (
            LANDCHINA_PROXY_REQUEST_INTERVAL if self.client.proxy_url else LANDCHINA_DIRECT_REQUEST_INTERVAL
        )
        self.client.detail_interval = (
            LANDCHINA_PROXY_DETAIL_INTERVAL if self.client.proxy_url else LANDCHINA_DIRECT_DETAIL_INTERVAL
        )
        return self._proxy_config_payload(row)

    def get_proxy_config(self) -> Dict[str, Any]:
        with self._connect() as conn:
            row = self._read_proxy_config_row(conn)
        payload = self._proxy_config_payload(row)
        payload["active"] = bool(self.client.proxy_url)
        return payload

    def save_proxy_config(self, config: Dict[str, Any]) -> Dict[str, Any]:
        previous_endpoint = self.client.proxy_url or "direct"
        with self._connect() as conn:
            row = self._read_proxy_config_row(conn)
            enabled = bool(row["enabled"])
            proxy_url = str(row["proxy_url"] or "").strip()
            proxy_token = str(row["proxy_token"] or "")
            if "enabled" in config:
                enabled = self._bool_config(config.get("enabled"))
            if "proxy_url" in config:
                proxy_url = str(config.get("proxy_url") or "").strip()
            if config.get("clear_token"):
                proxy_token = ""
            elif str(config.get("proxy_token") or "").strip():
                proxy_token = str(config.get("proxy_token") or "").strip()
            conn.execute(
                """
                UPDATE landchina_proxy_config
                SET enabled=?, proxy_url=?, proxy_token=?, updated_at=?
                WHERE id=1
                """,
                (1 if enabled else 0, proxy_url.rstrip("/"), proxy_token, _now()),
            )
        payload = self._apply_proxy_config()
        current_endpoint = self.client.proxy_url or "direct"
        if current_endpoint != previous_endpoint:
            self.access_gate.clear_cooldown(
                f"访问通道已由 {previous_endpoint} 切换为 {current_endpoint}，已清除旧通道冷却状态。"
            )
        return payload

    def test_proxy_config(self) -> Dict[str, Any]:
        config = self._apply_proxy_config()
        if not config["enabled"]:
            return {"ok": False, "status": "disabled", "message": "云服务器爬取未启用。", "config": config}
        if not config["proxy_url"]:
            return {"ok": False, "status": "missing_url", "message": "请先填写云服务器地址。", "config": config}
        if not config["token_set"]:
            return {"ok": False, "status": "missing_token", "message": "请先填写云服务器 Token。", "config": config}
        try:
            self.client.land_usage_tree()
            return {"ok": True, "status": "ok", "message": "云服务器中转连接正常。", "config": self.get_proxy_config()}
        except Exception as exc:
            return {
                "ok": False,
                "status": "failed",
                "message": str(exc),
                "config": self.get_proxy_config(),
                "access_status": self.access_gate.status(),
            }

    def normalize_official_case(self, detail: Dict[str, Any]) -> Dict[str, Any]:
        gd_guid = str(detail.get("gdGuid") or "").strip()
        area_sqm = _decimal(detail.get("gyMj")) * Decimal("10000")
        total_price_wan = _decimal(detail.get("je"))
        unit_price = total_price_wan * Decimal("10000") / area_sqm if area_sqm else Decimal("0")
        usage_raw = str(detail.get("tdYt") or "").strip()
        usage_key = infer_land_usage_key(usage_raw) or "other"
        record_level = str(detail.get("_record_level") or "").strip()
        if not record_level:
            record_level = "detail" if any(key in detail for key in ("dzBaBh", "xmMc", "je", "srr")) else "list"
        official_fields = {
            key: value for key, value in detail.items() if key not in {"_record_level", "_detail_error"}
        }
        return {
            "id": gd_guid,
            "source": "中国土地市场网",
            "source_url": LANDCHINA_RESULT_URL.format(gd_guid=gd_guid),
            "gd_guid": gd_guid,
            "electronic_supervision_no": detail.get("dzBaBh"),
            "project_name": detail.get("xmMc"),
            "location": detail.get("tdZl"),
            "administrative_region_code": detail.get("xzqDm"),
            "administrative_region_name": detail.get("xzqFullName"),
            "land_usage_raw": usage_raw,
            "land_usage_key": usage_key,
            "land_usage_first_level": land_usage_first_level_label(usage_key),
            "supply_method": detail.get("gyFs"),
            "area_sqm": _q(area_sqm, "0.01") if area_sqm else None,
            "total_price_wan": _q(total_price_wan, "0.01") if total_price_wan else None,
            "unit_price_sqm": _q(unit_price, "0.01") if unit_price else None,
            "transaction_date": _date_text(detail.get("qdRq")),
            "use_term_years": str(detail.get("crNx") or "") or None,
            "land_level": detail.get("tdJb"),
            "recipient": detail.get("srr"),
            "plot_ratio_min": str(detail.get("minRjl") or "") or None,
            "plot_ratio_max": str(detail.get("maxRjl") or "") or None,
            "detail_status": "complete" if record_level == "detail" else "partial",
            "detail_error": detail.get("_detail_error"),
            "detail_fetched_at": _now() if record_level == "detail" else None,
            "official_fields": official_fields,
            "fetched_at": _now(),
        }

    def _reclassify_existing_cases(self) -> None:
        with self._connect() as conn:
            rows = conn.execute("SELECT id, official_json FROM comparable_cases").fetchall()
            for row in rows:
                official = json.loads(row["official_json"])
                usage_key = infer_land_usage_key(official.get("land_usage_raw")) or "other"
                first_level = land_usage_first_level_label(usage_key)
                if (
                    official.get("land_usage_key") == usage_key
                    and official.get("land_usage_first_level") == first_level
                ):
                    continue
                official["land_usage_key"] = usage_key
                official["land_usage_first_level"] = first_level
                official_json = json.dumps(official, ensure_ascii=False, sort_keys=True, default=str)
                conn.execute(
                    "UPDATE comparable_cases SET official_json=?, official_hash=?, updated_at=? WHERE id=?",
                    (
                        official_json,
                        hashlib.sha256(official_json.encode("utf-8")).hexdigest(),
                        _now(),
                        row["id"],
                    ),
                )

    def _merge_official_case(self, existing: Dict[str, Any], incoming: Dict[str, Any]) -> Dict[str, Any]:
        merged = dict(existing)
        for key, value in incoming.items():
            if value not in (None, "", {}, []):
                merged[key] = value
        merged_fields = dict(existing.get("official_fields") or {})
        merged_fields.update(incoming.get("official_fields") or {})
        merged["official_fields"] = merged_fields
        if existing.get("detail_status") == "complete" and incoming.get("detail_status") != "complete":
            merged["detail_status"] = "complete"
            merged["detail_error"] = None
            merged["detail_fetched_at"] = existing.get("detail_fetched_at")
        elif incoming.get("detail_status") == "complete":
            merged["detail_error"] = None
        return merged

    def upsert_official_case(self, detail: Dict[str, Any]) -> Dict[str, Any]:
        normalized = self.normalize_official_case(detail)
        if not normalized["gd_guid"]:
            raise ValueError("供地结果详情缺少 gdGuid")
        now = _now()
        with self._connect() as conn:
            existing = conn.execute(
                "SELECT id, manual_json, official_hash, official_json, fetched_at FROM comparable_cases "
                "WHERE gd_guid=? OR electronic_supervision_no=? LIMIT 1",
                (normalized["gd_guid"], normalized.get("electronic_supervision_no") or ""),
            ).fetchone()
            case_id = existing["id"] if existing else normalized["gd_guid"]
            manual_json = existing["manual_json"] if existing else "{}"
            if existing:
                normalized = self._merge_official_case(json.loads(existing["official_json"]), normalized)
            official_json = json.dumps(normalized, ensure_ascii=False, sort_keys=True, default=str)
            official_hash = hashlib.sha256(official_json.encode("utf-8")).hexdigest()
            if existing and existing["official_hash"] != official_hash:
                conn.execute(
                    """
                    INSERT OR IGNORE INTO comparable_case_history
                        (history_id, case_id, official_hash, official_json, fetched_at)
                    VALUES (?, ?, ?, ?, ?)
                    """,
                    (
                        uuid.uuid4().hex,
                        case_id,
                        existing["official_hash"],
                        existing["official_json"],
                        existing["fetched_at"],
                    ),
                )
            conn.execute(
                """
                INSERT INTO comparable_cases
                    (id, gd_guid, electronic_supervision_no, official_json, manual_json, official_hash, fetched_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(id) DO UPDATE SET
                    gd_guid=excluded.gd_guid,
                    electronic_supervision_no=excluded.electronic_supervision_no,
                    official_json=excluded.official_json,
                    official_hash=excluded.official_hash,
                    fetched_at=excluded.fetched_at,
                    updated_at=excluded.updated_at
                """,
                (
                    case_id,
                    normalized["gd_guid"],
                    normalized.get("electronic_supervision_no"),
                    official_json,
                    manual_json,
                    official_hash,
                    normalized["fetched_at"],
                    now,
                ),
            )
        return self.get_case(case_id) or normalized

    def _row_to_case(self, row: sqlite3.Row) -> Dict[str, Any]:
        official = json.loads(row["official_json"])
        manual = json.loads(row["manual_json"] or "{}")
        manual_draft = json.loads(row["manual_draft_json"] or "{}")
        result = dict(official)
        result.update({key: value for key, value in manual.items() if value not in (None, "")})
        result["manual_fields"] = manual
        result["manual_draft_fields"] = manual_draft
        result["official_fields"] = official.get("official_fields") or {}
        result["updated_at"] = row["updated_at"]
        result["id"] = row["id"]
        return result

    def get_case(self, case_id: str) -> Optional[Dict[str, Any]]:
        with self._connect() as conn:
            row = conn.execute("SELECT * FROM comparable_cases WHERE id=?", (case_id,)).fetchone()
        return self._row_to_case(row) if row else None

    def _case_filter_sql(self, filters: Dict[str, Any]) -> tuple[str, List[Any]]:
        clauses, params = [], []
        mapping = {
            "xzq_dm": "$.administrative_region_code",
            "land_usage_key": "$.land_usage_key",
            "supply_method": "$.supply_method",
            "electronic_supervision_no": "$.electronic_supervision_no",
        }
        for key, json_path in mapping.items():
            value = str(filters.get(key) or "").strip()
            if value:
                clauses.append(f"json_extract(official_json, '{json_path}') LIKE ?")
                params.append(f"%{value}%")
        start_date = str(filters.get("start_date") or "").strip()
        end_date = str(filters.get("end_date") or "").strip()
        location = str(filters.get("location") or "").strip()
        if start_date:
            clauses.append("json_extract(official_json, '$.transaction_date') >= ?")
            params.append(start_date)
        if end_date:
            clauses.append("json_extract(official_json, '$.transaction_date') <= ?")
            params.append(end_date)
        if location:
            clauses.append("json_extract(official_json, '$.location') LIKE ?")
            params.append(f"%{location}%")
        keyword = str(filters.get("keyword") or "").strip()
        if keyword:
            clauses.append("(official_json LIKE ? OR manual_json LIKE ? OR manual_draft_json LIKE ?)")
            params.extend((f"%{keyword}%", f"%{keyword}%", f"%{keyword}%"))
        where = " WHERE " + " AND ".join(clauses) if clauses else ""
        return where, params

    def list_cases(self, filters: Dict[str, Any]) -> Dict[str, Any]:
        where, params = self._case_filter_sql(filters)
        page = max(int(filters.get("page") or 1), 1)
        page_size = min(max(int(filters.get("page_size") or 30), 1), 100)
        with self._connect() as conn:
            total = conn.execute(f"SELECT COUNT(*) FROM comparable_cases{where}", params).fetchone()[0]
            rows = conn.execute(
                f"SELECT * FROM comparable_cases{where} ORDER BY fetched_at DESC LIMIT ? OFFSET ?",
                (*params, page_size, (page - 1) * page_size),
            ).fetchall()
        return {"total": total, "page": page, "page_size": page_size, "items": [self._row_to_case(row) for row in rows]}

    def list_case_regions(self) -> List[Dict[str, Any]]:
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT code, name, COUNT(*) AS count
                FROM (
                    SELECT
                        TRIM(COALESCE(json_extract(official_json, '$.administrative_region_code'), '')) AS code,
                        TRIM(COALESCE(json_extract(official_json, '$.administrative_region_name'), '')) AS name
                    FROM comparable_cases
                )
                WHERE code <> '' OR name <> ''
                GROUP BY code, name
                ORDER BY name COLLATE NOCASE, code
                """
            ).fetchall()
        grouped: Dict[str, Dict[str, Any]] = {}
        for row in rows:
            raw_code = str(row["code"] or "").strip()
            code = raw_code[:6] if len(raw_code) > 6 and raw_code.endswith("000") else raw_code
            name = str(row["name"] or code).strip()
            key = code or name
            item = grouped.setdefault(key, {"code": code, "name": name, "count": 0})
            if len(name) > len(str(item.get("name") or "")):
                item["name"] = name
            item["count"] += int(row["count"] or 0)
        return sorted(
            grouped.values(),
            key=lambda item: (str(item.get("name") or ""), str(item.get("code") or "")),
        )

    @staticmethod
    def _export_cell(value: Any) -> str:
        if value in (None, ""):
            return ""
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
        return str(value)

    def export_cases_csv(self, filters: Dict[str, Any]) -> bytes:
        where, params = self._case_filter_sql(filters)
        columns = [
            ("电子监管号", "electronic_supervision_no"),
            ("项目名称", "project_name"),
            ("坐落", "location"),
            ("行政区代码", "administrative_region_code"),
            ("行政区", "administrative_region_name"),
            ("土地用途", "land_usage_raw"),
            ("用途分类", "land_usage_first_level"),
            ("供地方式", "supply_method"),
            ("土地使用权人", "recipient"),
            ("合同签订日期", "transaction_date"),
            ("面积(平方米)", "area_sqm"),
            ("使用年限", "use_term_years"),
            ("土地级别", "land_level"),
            ("成交价格(万元)", "total_price_wan"),
            ("地面单价(元/平方米)", "unit_price_sqm"),
            ("容积率下限", "plot_ratio_min"),
            ("容积率上限", "plot_ratio_max"),
            ("详情状态", "detail_status"),
            ("来源链接", "source_url"),
            ("人工补充字段", "manual_fields"),
        ]
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT * FROM comparable_cases{where} ORDER BY fetched_at DESC",
                params,
            ).fetchall()
        output = io.StringIO(newline="")
        writer = csv.writer(output)
        writer.writerow([label for label, _key in columns])
        for row in rows:
            item = self._row_to_case(row)
            writer.writerow([self._export_cell(item.get(key)) for _label, key in columns])
        return output.getvalue().encode("utf-8-sig")

    def patch_case(self, case_id: str, manual_fields: Dict[str, Any]) -> Dict[str, Any]:
        existing = self.get_case(case_id)
        if not existing:
            raise KeyError(case_id)
        merged = dict(existing.get("manual_fields") or {})
        for key, value in manual_fields.items():
            if value in (None, ""):
                merged.pop(key, None)
            else:
                merged[key] = value
        with self._connect() as conn:
            conn.execute(
                "UPDATE comparable_cases SET manual_json=?, updated_at=? WHERE id=?",
                (json.dumps(merged, ensure_ascii=False), _now(), case_id),
            )
        return self.get_case(case_id) or {}

    def patch_case_draft(self, case_id: str, manual_draft_fields: Dict[str, Any]) -> Dict[str, Any]:
        existing = self.get_case(case_id)
        if not existing:
            raise KeyError(case_id)
        merged = dict(existing.get("manual_draft_fields") or {})
        for key, value in manual_draft_fields.items():
            if value in (None, ""):
                merged.pop(key, None)
            else:
                merged[key] = value
        with self._connect() as conn:
            conn.execute(
                "UPDATE comparable_cases SET manual_draft_json=?, updated_at=? WHERE id=?",
                (json.dumps(merged, ensure_ascii=False), _now(), case_id),
            )
        return self.get_case(case_id) or {}

    def confirm_case_manual_fields(self, case_id: str, fields: Optional[List[str]] = None) -> Dict[str, Any]:
        existing = self.get_case(case_id)
        if not existing:
            raise KeyError(case_id)
        draft = dict(existing.get("manual_draft_fields") or {})
        selected = set(fields or draft.keys())
        confirmed = dict(existing.get("manual_fields") or {})
        remaining = dict(draft)
        for key in selected:
            if key not in draft:
                continue
            value = draft[key]
            if value in (None, ""):
                confirmed.pop(key, None)
            else:
                confirmed[key] = value
            remaining.pop(key, None)
        with self._connect() as conn:
            conn.execute(
                """
                UPDATE comparable_cases
                SET manual_json=?, manual_draft_json=?, updated_at=?
                WHERE id=?
                """,
                (
                    json.dumps(confirmed, ensure_ascii=False),
                    json.dumps(remaining, ensure_ascii=False),
                    _now(),
                    case_id,
                ),
            )
        return self.get_case(case_id) or {}

    def get_factor_scheme(self, land_usage_key: str, create: bool = True) -> Optional[Dict[str, Any]]:
        with self._connect() as conn:
            row = conn.execute("SELECT scheme_json, updated_at FROM factor_schemes WHERE land_usage_key=?", (land_usage_key,)).fetchone()
        if row:
            result = normalize_factor_scheme(json.loads(row["scheme_json"]), land_usage_key)
            result["updated_at"] = row["updated_at"]
            return result
        if create:
            scheme = default_factor_scheme(land_usage_key)
            self.save_factor_scheme(land_usage_key, scheme)
            return scheme
        return None

    def save_factor_scheme(self, land_usage_key: str, scheme: Dict[str, Any]) -> Dict[str, Any]:
        validated = FactorScheme(**normalize_factor_scheme(scheme, land_usage_key))
        payload = validated.model_dump()
        now = _now()
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO factor_schemes(land_usage_key, scheme_json, updated_at) VALUES (?, ?, ?)
                ON CONFLICT(land_usage_key) DO UPDATE SET scheme_json=excluded.scheme_json, updated_at=excluded.updated_at
                """,
                (land_usage_key, json.dumps(payload, ensure_ascii=False), now),
            )
        return dict(payload, updated_at=now)

    def create_crawl_job(self, filters: Dict[str, Any]) -> str:
        job_id = uuid.uuid4().hex
        with self.jobs_lock:
            self.jobs[job_id] = {
                "job_id": job_id,
                "status": "queued",
                "filters": filters,
                "total": 0,
                "processed": 0,
                "saved": 0,
                "complete": 0,
                "partial": 0,
                "failed_details": 0,
                "rate_limited": 0,
                "phase": "queued",
                "stopped_reason": "",
                "errors": [],
                "cancel_requested": False,
                "created_at": _now(),
            }
        threading.Thread(target=self._run_crawl_job, args=(job_id,), daemon=True).start()
        return job_id

    def get_access_status(self) -> Dict[str, Any]:
        status = self.access_gate.status()
        status["proxy_enabled"] = bool(self.client.proxy_url)
        status["proxy_url"] = self.client.proxy_url
        status["proxy_config"] = self.get_proxy_config()
        status["access_channel"] = "cloud_proxy" if self.client.proxy_url else "direct"
        return status

    def get_crawl_job(self, job_id: str) -> Optional[Dict[str, Any]]:
        with self.jobs_lock:
            job = self.jobs.get(job_id)
            return dict(job) if job else None

    def cancel_crawl_job(self, job_id: str) -> Optional[Dict[str, Any]]:
        with self.jobs_lock:
            job = self.jobs.get(job_id)
            if job:
                job["cancel_requested"] = True
            return dict(job) if job else None

    def _update_job(self, job_id: str, **changes: Any) -> None:
        with self.jobs_lock:
            if job_id in self.jobs:
                self.jobs[job_id].update(changes)

    def _selected_usage_key(self, filters: Dict[str, Any]) -> str:
        key = str(filters.get("land_usage_key") or "").strip()
        if key in LAND_USAGE_OPTION_BY_KEY:
            return key
        legacy_usage = str(filters.get("land_usage") or "").strip()
        return infer_land_usage_key(legacy_usage) if legacy_usage else ""

    @staticmethod
    def _crawl_location_filter(filters: Dict[str, Any]) -> Optional[str]:
        location = str(filters.get("location") or "").strip()
        if not location:
            return None
        xzq_dm = str(filters.get("xzq_dm") or "").strip()
        if xzq_dm and location in LANDCHINA_REGION_ALIASES.get(xzq_dm, set()):
            return None
        return location

    def _crawl_payload(
        self,
        filters: Dict[str, Any],
        page_num: int,
        start_date: str,
        end_date: str,
        usage_code: str = "",
    ) -> Dict[str, Any]:
        return {
            "pageNum": page_num,
            "pageSize": LANDCHINA_CRAWL_PAGE_SIZE,
            "xzqDm": filters.get("xzq_dm") or None,
            "gyFs": filters.get("supply_method") or None,
            "tdYt": usage_code or None,
            "startDate": f"{start_date} 00:00:00" if start_date else "",
            "endDate": f"{end_date} 23:59:59" if end_date else "",
            "dzBaBh": filters.get("electronic_supervision_no") or None,
            "tdZl": self._crawl_location_filter(filters),
        }

    def _date_ranges(self, filters: Dict[str, Any], usage_code: str = "") -> List[tuple[str, str]]:
        start_text = str(filters.get("start_date") or "")
        end_text = str(filters.get("end_date") or "")
        if not start_text or not end_text:
            return [(start_text, end_text)]

        start = date.fromisoformat(start_text)
        end = date.fromisoformat(end_text)
        if start > end:
            return [(start_text, end_text)]
        ranges: List[tuple[str, str]] = []
        current = start
        while current <= end:
            if current.month == 12:
                next_month = date(current.year + 1, 1, 1)
            else:
                next_month = date(current.year, current.month + 1, 1)
            range_end = min(next_month - timedelta(days=1), end)
            ranges.append((current.isoformat(), range_end.isoformat()))
            current = range_end + timedelta(days=1)
        return ranges

    @staticmethod
    def _split_date_range(start_text: str, end_text: str) -> List[tuple[str, str]]:
        start = date.fromisoformat(start_text)
        end = date.fromisoformat(end_text)
        if start >= end:
            return [(start_text, end_text)]
        midpoint = start + timedelta(days=(end - start).days // 2)
        return [
            (start.isoformat(), midpoint.isoformat()),
            ((midpoint + timedelta(days=1)).isoformat(), end.isoformat()),
        ]

    @staticmethod
    def _should_split_listing_error(exc: Exception) -> bool:
        text = str(exc).lower()
        if any(marker in text for marker in ("401", "403", "418", "429", "限流", "认证失败")):
            return False
        return any(marker in text for marker in ("超时", "timeout", "timed out"))

    @staticmethod
    def _crawl_page_count(first_page: Dict[str, Any]) -> int:
        pages = int(first_page.get("pages") or 0)
        total = int(first_page.get("total") or 0)
        if total:
            pages = max(pages, (total + LANDCHINA_CRAWL_PAGE_SIZE - 1) // LANDCHINA_CRAWL_PAGE_SIZE)
        if first_page.get("list"):
            pages = max(pages, 1)
        return pages

    def _run_crawl_job(self, job_id: str) -> None:
        job = self.get_crawl_job(job_id) or {}
        filters = job.get("filters") or {}
        if not self.access_gate.acquire_crawl(job_id):
            access_status = self.access_gate.status()
            self._update_job(
                job_id,
                status="failed",
                phase="blocked",
                access_status=access_status,
                errors=["已有另一项土地市场网抓取任务正在运行，请等待该任务完成。"],
                finished_at=_now(),
            )
            return
        rate_limit_start = self.client.rate_limit_hits
        self._update_job(job_id, status="running", phase="listing", started_at=_now())
        try:
            access_status = self.access_gate.status()
            if access_status["blocked"]:
                raise LandChinaRateLimitError(
                    f"官网访问冷却中，请于 {access_status['blocked_until'].replace('T', ' ')} 后重试。"
                )
            selected_usage_key = self._selected_usage_key(filters)
            usage_code = official_land_usage_code(selected_usage_key)
            self._update_job(
                job_id,
                land_usage_key=selected_usage_key,
                official_land_usage_code=usage_code,
            )
            ranges = self._date_ranges(filters, usage_code)
            self._update_job(job_id, range_count=len(ranges), range_processed=0)
            records_by_guid: Dict[str, Dict[str, Any]] = {}
            errors: List[str] = []
            listing_stopped_reason = ""
            range_index = 0
            while range_index < len(ranges):
                start_date, end_date = ranges[range_index]
                try:
                    first = self.client.result_list(self._crawl_payload(filters, 1, start_date, end_date, usage_code))
                except LandChinaRateLimitError as exc:
                    if start_date != end_date and self._should_split_listing_error(exc):
                        ranges[range_index : range_index + 1] = self._split_date_range(start_date, end_date)
                        self._update_job(
                            job_id,
                            range_count=len(ranges),
                            range_processed=range_index,
                            listed=len(records_by_guid),
                            errors=(errors + [f"{start_date}~{end_date}: 列表请求超时，已自动拆分区间继续抓取。"])[-20:],
                        )
                        continue
                    listing_stopped_reason = f"官网列表访问受限，已保存本次已发现记录；请等待冷却后再次抓取继续补齐。{exc}"
                    errors.append(f"{start_date}~{end_date}: 列表抓取受限：{exc}")
                    break
                total_results = int(first.get("total") or 0)
                if total_results > LANDCHINA_CRAWL_MAX_TOTAL and start_date != end_date:
                    ranges[range_index : range_index + 1] = self._split_date_range(start_date, end_date)
                    self._update_job(job_id, range_count=len(ranges), listed=len(records_by_guid))
                    continue
                pages = self._crawl_page_count(first)

                def collect_items(page: Dict[str, Any]) -> None:
                    for item in page.get("list") or []:
                        if item.get("gdGuid") and (
                            not selected_usage_key
                            or infer_land_usage_key(item.get("tdYt")) == selected_usage_key
                        ):
                            records_by_guid[item["gdGuid"]] = item

                collect_items(first)
                for page_num in range(2, pages + 1):
                    if (self.get_crawl_job(job_id) or {}).get("cancel_requested"):
                        self._update_job(job_id, status="cancelled", finished_at=_now())
                        return
                    try:
                        page = self.client.result_list(
                            self._crawl_payload(filters, page_num, start_date, end_date, usage_code)
                        )
                    except LandChinaRateLimitError as exc:
                        if start_date != end_date and self._should_split_listing_error(exc):
                            ranges[range_index : range_index + 1] = self._split_date_range(start_date, end_date)
                            self._update_job(
                                job_id,
                                range_count=len(ranges),
                                range_processed=range_index,
                                listed=len(records_by_guid),
                                errors=(errors + [f"{start_date}~{end_date} 第 {page_num} 页: 列表请求超时，已自动拆分区间继续抓取。"])[-20:],
                            )
                            break
                        listing_stopped_reason = f"官网列表访问受限，已保存本次已发现记录；请等待冷却后再次抓取继续补齐。{exc}"
                        errors.append(f"{start_date}~{end_date} 第 {page_num} 页: 列表抓取受限：{exc}")
                        break
                    collect_items(page)
                if listing_stopped_reason:
                    break
                if range_index < len(ranges) and ranges[range_index] != (start_date, end_date):
                    continue
                range_index += 1
                self._update_job(
                    job_id,
                    range_processed=range_index,
                    range_count=len(ranges),
                    listed=len(records_by_guid),
                )

            total = len(records_by_guid)
            if listing_stopped_reason and total == 0:
                raise LandChinaRateLimitError(listing_stopped_reason)
            self._update_job(job_id, total=total, phase="saving_list")
            saved, processed, failed_details = 0, 0, 0
            complete_guids = set()
            for guid, item in records_by_guid.items():
                try:
                    saved_case = self.upsert_official_case(dict(item, _record_level="list"))
                    saved += 1
                    if saved_case.get("detail_status") == "complete":
                        complete_guids.add(guid)
                except Exception as exc:
                    errors.append(f"{guid}: 列表记录入库失败：{exc}")
            self._update_job(
                job_id,
                saved=saved,
                complete=len(complete_guids),
                partial=max(saved - len(complete_guids), 0),
                phase="enriching_details",
                errors=errors[-20:],
            )

            stopped_reason = listing_stopped_reason
            for guid, item in records_by_guid.items():
                if (self.get_crawl_job(job_id) or {}).get("cancel_requested"):
                    self._update_job(
                        job_id,
                        status="cancelled",
                        processed=processed,
                        saved=saved,
                        complete=len(complete_guids),
                        partial=max(saved - len(complete_guids), 0),
                        failed_details=failed_details,
                        rate_limited=self.client.rate_limit_hits - rate_limit_start,
                        finished_at=_now(),
                    )
                    return
                processed += 1
                existing_case = self.get_case(guid)
                if (
                    existing_case
                    and existing_case.get("detail_status") == "complete"
                    and not filters.get("refresh_complete_details")
                ):
                    complete_guids.add(guid)
                    self._update_job(
                        job_id,
                        processed=processed,
                        saved=saved,
                        complete=len(complete_guids),
                        partial=max(saved - len(complete_guids), 0),
                        failed_details=failed_details,
                        rate_limited=self.client.rate_limit_hits - rate_limit_start,
                        errors=errors[-20:],
                    )
                    continue
                try:
                    detail = dict(self.client.result_detail(guid), _record_level="detail")
                    saved_case = self.upsert_official_case(detail)
                    if saved_case.get("detail_status") == "complete":
                        complete_guids.add(guid)
                except LandChinaRateLimitError as exc:
                    failed_details += 1
                    errors.append(f"{guid}: 详情请求受官网限流：{exc}")
                    self.upsert_official_case(dict(item, _record_level="list", _detail_error=str(exc)))
                    detail_stopped_reason = "官网持续返回限流状态，已停止详情补全；列表记录均已保留，可稍后重新抓取补全。"
                    stopped_reason = f"{stopped_reason}；{detail_stopped_reason}" if stopped_reason else detail_stopped_reason
                    break
                except Exception as exc:
                    failed_details += 1
                    errors.append(f"{guid}: 详情补全失败：{exc}")
                    self.upsert_official_case(dict(item, _record_level="list", _detail_error=str(exc)))
                self._update_job(
                    job_id,
                    processed=processed,
                    saved=saved,
                    complete=len(complete_guids),
                    partial=max(saved - len(complete_guids), 0),
                    failed_details=failed_details,
                    rate_limited=self.client.rate_limit_hits - rate_limit_start,
                    errors=errors[-20:],
                )
            self._update_job(
                job_id,
                status="completed",
                phase="completed",
                processed=processed,
                saved=saved,
                complete=len(complete_guids),
                partial=max(saved - len(complete_guids), 0),
                failed_details=failed_details,
                rate_limited=self.client.rate_limit_hits - rate_limit_start,
                stopped_reason=stopped_reason,
                errors=errors[-20:],
                finished_at=_now(),
            )
        except Exception as exc:
            self._update_job(
                job_id,
                status="failed",
                phase="blocked" if isinstance(exc, LandChinaRateLimitError) else "failed",
                access_status=self.access_gate.status(),
                rate_limited=self.client.rate_limit_hits - rate_limit_start,
                errors=[str(exc)],
                finished_at=_now(),
            )
        finally:
            self.access_gate.release_crawl(job_id)

    def _suggest_factor_value(self, factor: Dict[str, Any], subject: Dict[str, Any], case: Dict[str, Any]) -> tuple[Any, str, bool]:
        key = factor["key"]
        if key == "land_usage":
            matches = subject.get("land_usage_key") == case.get("land_usage_key")
            return case.get("land_usage_raw"), "100" if matches else "", matches
        if key == "transaction_method":
            return case.get("supply_method"), "100", bool(case.get("supply_method"))
        if key == "use_term":
            matches = str(case.get("use_term_years") or "") == str(subject.get("land_use_term_years") or "")
            return case.get("use_term_years"), "100" if matches else "", matches
        if key == "parcel_scale":
            subject_area = _decimal(subject.get("land_area"))
            case_area = _decimal(case.get("area_sqm"))
            if not subject_area or not case_area:
                return case.get("area_sqm"), "", False
            ratio = max(subject_area, case_area) / min(subject_area, case_area)
            suggested = "100" if ratio <= Decimal("1.5") else "101" if case_area > subject_area else "99"
            return case.get("area_sqm"), suggested, False
        if key == "transaction_time":
            return case.get("transaction_date"), "", False
        return (case.get(key) or case.get("manual_fields", {}).get(key), "", False)

    def prepare_analysis(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        subject = analysis.get("subject") or {}
        case_ids = analysis.get("case_ids") or []
        if len(case_ids) != 3:
            raise ValueError("市场比较法必须固定选择三宗比较实例 A/B/C。")
        supplied_cases = analysis.get("selected_cases") or []
        supplied_by_id = {str(case.get("id") or ""): case for case in supplied_cases}
        cases = [supplied_by_id.get(str(case_id)) or self.get_case(case_id) for case_id in case_ids]
        if any(case is None for case in cases):
            raise ValueError("所选比较实例不存在或已从实例库移除。")
        usage_key = str(subject.get("land_usage_key") or "other")
        scheme = analysis.get("factor_scheme_snapshot") or self.get_factor_scheme(usage_key) or default_factor_scheme(usage_key)
        supplied = {item.get("key"): item for item in analysis.get("factors") or []}
        factors = []
        for definition in scheme.get("factors") or []:
            if not definition.get("enabled", True):
                continue
            existing = supplied.get(definition["key"]) or {}
            case_values = existing.get("cases") or {}
            for slot, case in zip(CASE_SLOTS, cases):
                if slot not in case_values:
                    value, index, confirmed = self._suggest_factor_value(definition, subject, case or {})
                    case_values[slot] = {
                        "value": value,
                        "index": index,
                        "confirmed": confirmed,
                        "source": definition.get("source") or "manual",
                    }
            report_confirmed = all(
                case_values.get(slot, {}).get("confirmed")
                and _decimal(case_values.get(slot, {}).get("index"))
                for slot in CASE_SLOTS
            )
            factors.append(
                {
                    "key": definition["key"],
                    "label": definition["label"],
                    "group": definition["group"],
                    "required": definition.get("required", True),
                    "source": definition.get("source") or "manual",
                    "help_text": definition.get("help_text") or definition.get("note") or "",
                    "levels": definition.get("levels") or [],
                    "review_status": (
                        "confirmed"
                        if report_confirmed
                        else existing.get("review_status") or definition.get("review_status") or "needs_review"
                    ),
                    "subject_value": existing.get("subject_value") or self._subject_factor_value(definition["key"], subject),
                    "subject_index": str(existing.get("subject_index") or "100"),
                    "cases": case_values,
                }
            )
        return dict(
            analysis,
            subject=subject,
            case_ids=case_ids,
            selected_cases=cases,
            factor_scheme_snapshot=scheme,
            factors=factors,
        )

    def _subject_factor_value(self, key: str, subject: Dict[str, Any]) -> Any:
        mapping = {
            "land_usage": "land_usage",
            "transaction_time": "valuation_date",
            "transaction_method": "right_type",
            "use_term": "land_use_term_years",
            "development_level": "land_development_set",
            "parcel_scale": "land_area",
        }
        return subject.get(mapping.get(key, key))

    def calculate_market_comparison(self, analysis_data: Dict[str, Any]) -> Dict[str, Any]:
        analysis = self.prepare_analysis(analysis_data)
        subject = analysis["subject"]
        cases = analysis["selected_cases"]
        monthly_rate = _decimal(analysis.get("monthly_growth_rate"), "0.13") / Decimal("100")
        warnings: List[Dict[str, str]] = []
        factors = analysis["factors"]
        valuation_date_text = _date_text(subject.get("valuation_date"))
        valuation_day = date.fromisoformat(valuation_date_text) if valuation_date_text else None

        rule_review_labels: List[str] = []
        missing_rule_level_labels: List[str] = []
        for factor in factors:
            if factor.get("review_status") != "confirmed":
                rule_review_labels.append(str(factor.get("label") or factor.get("key") or "未命名因素"))
            if (
                factor.get("review_status") != "confirmed"
                and factor.get("source") == "manual"
                and not factor.get("levels")
            ):
                missing_rule_level_labels.append(str(factor.get("label") or factor.get("key") or "未命名因素"))
        if rule_review_labels:
            warnings.append(
                {
                    "level": "info",
                    "message": (
                        f"全局规则管理中有 {len(rule_review_labels)} 项规则仍待校核；"
                        "这只影响后续新建报告的规则模板，当前报告请以因素确认结果为准。"
                    ),
                    "target": "rules",
                }
            )
        if missing_rule_level_labels:
            warnings.append(
                {
                    "level": "warning",
                    "message": (
                        f"全局规则管理中有 {len(missing_rule_level_labels)} 项人工判断规则尚未配置等级与建议指数。"
                    ),
                    "target": "rules",
                }
            )

        for factor in factors:
            if factor["key"] == "transaction_time" and valuation_day:
                for slot, case in zip(CASE_SLOTS, cases):
                    case_text = _date_text(case.get("transaction_date"))
                    if not case_text:
                        continue
                    case_day = date.fromisoformat(case_text)
                    months = Decimal(str((valuation_day - case_day).days)) / Decimal("30")
                    growth = (Decimal("1") + monthly_rate) ** months
                    index = Decimal("100") / growth if growth else Decimal("100")
                    factor["cases"][slot]["index"] = _q(index, "0.01")
                    factor["cases"][slot]["confirmed"] = True
            if factor["key"] == "use_term":
                rate = _decimal(analysis.get("land_reduction_rate"), "5.4") / Decimal("100")
                subject_years = _decimal(subject.get("land_use_term_years"))
                for slot, case in zip(CASE_SLOTS, cases):
                    case_years = _decimal(case.get("use_term_years"))
                    if not rate or not subject_years or not case_years:
                        continue
                    subject_term = Decimal("1") - Decimal("1") / ((Decimal("1") + rate) ** subject_years)
                    case_term = Decimal("1") - Decimal("1") / ((Decimal("1") + rate) ** case_years)
                    coefficient = subject_term / case_term if case_term else Decimal("1")
                    factor["cases"][slot]["index"] = _q(Decimal("100") / coefficient, "0.01")
                    factor["cases"][slot]["confirmed"] = True

        calculations, corrected_prices = [], []
        complete = True
        for slot, case in zip(CASE_SLOTS, cases):
            coefficient = Decimal("1")
            for factor in factors:
                item = factor["cases"].get(slot) or {}
                index = _decimal(item.get("index"))
                subject_index = _decimal(factor.get("subject_index"), "100")
                if factor.get("required") and (not item.get("confirmed") or not index):
                    complete = False
                    warnings.append(
                        {
                            "level": "warning",
                            "message": f"实例{slot}的“{factor['label']}”尚未确认。",
                            "target": "factors",
                            "factor_key": str(factor.get("key") or ""),
                        }
                    )
                    continue
                if index:
                    coefficient *= subject_index / index
            unit_price = _decimal(case.get("unit_price_sqm"))
            corrected = unit_price * coefficient
            calculations.append(
                {
                    "slot": slot,
                    "case_id": case["id"],
                    "unit_price_sqm": _q(unit_price, "0.01"),
                    "correction_coefficient": _q(coefficient, "0.0001"),
                    "corrected_price": _q(corrected, "0.01"),
                }
            )
            corrected_prices.append(corrected)
            if abs(coefficient - Decimal("1")) > Decimal("0.30"):
                warnings.append({"level": "warning", "message": f"实例{slot}综合修正幅度超过30%。", "target": "factors"})
            if valuation_day and case.get("transaction_date"):
                if (valuation_day - date.fromisoformat(_date_text(case["transaction_date"]))).days > 1096:
                    warnings.append({"level": "warning", "message": f"实例{slot}交易时间距估价期日超过三年。", "target": "instances"})
            if subject.get("land_usage_key") and subject.get("land_usage_key") != case.get("land_usage_key"):
                warnings.append({"level": "warning", "message": f"实例{slot}土地用途与估价对象不一致。", "target": "instances"})
        if corrected_prices:
            high, low = max(corrected_prices), min(corrected_prices)
            if low and (high - low) / low > Decimal("0.40"):
                warnings.append({"level": "warning", "message": "三宗比较实例修正后的比准价格差异超过40%。", "target": "factors"})
        market_price = sum(corrected_prices, Decimal("0")) / Decimal("3") if len(corrected_prices) == 3 else Decimal("0")
        result = dict(
            analysis,
            calculations=calculations,
            warnings=list({item["message"]: item for item in warnings}.values()),
            complete=complete and len(cases) == 3,
            market_comp_price=_q(market_price, "0.1") if corrected_prices else None,
        )
        result.update(self.build_render_fields(result))
        return result

    @staticmethod
    def _case_factor_values(slot: str, factors: Iterable[Dict[str, Any]]) -> Dict[str, str]:
        values: Dict[str, str] = {}
        for factor in factors:
            item = (factor.get("cases") or {}).get(slot) or {}
            if not item.get("confirmed"):
                continue
            value = _report_text(item.get("value") or item.get("level_label"))
            if not value:
                continue
            key = str(factor.get("key") or "").strip()
            label = str(factor.get("label") or key).strip()
            if key:
                values[key] = value
            if label:
                values[label] = value
        return values

    def _case_description(self, slot: str, case: Dict[str, Any], factors: Optional[Iterable[Dict[str, Any]]] = None) -> str:
        manual = case.get("manual_fields") or {}
        factor_values = self._case_factor_values(slot, factors or [])

        def pick(*keys: str) -> str:
            for key in keys:
                for source in (manual, factor_values, case):
                    value = _report_text(source.get(key) if isinstance(source, dict) else None)
                    if value:
                        return value
            return ""

        recipient = _report_text(case.get("recipient"))
        supply_method = _report_text(case.get("supply_method"))
        supervision_no = _report_text(case.get("electronic_supervision_no"))
        if recipient and supply_method:
            lead = f"该实例为{recipient}以{supply_method}方式取得"
        elif recipient:
            lead = f"该实例受让人为{recipient}"
        elif supply_method:
            lead = f"该实例以{supply_method}方式成交"
        else:
            lead = "该实例为中国土地市场网收录的可比成交案例"
        if supervision_no:
            lead += f"（电子监管号为{supervision_no}）"

        core_clauses = [lead]
        location = _report_text(case.get("location"))
        if location:
            core_clauses.append(f"位于{location}")
        land_usage = _report_text(case.get("land_usage_raw"))
        if land_usage:
            core_clauses.append(f"土地用途为{land_usage}")
        area_sqm = _report_text(case.get("area_sqm"))
        if area_sqm:
            core_clauses.append(f"土地面积为{area_sqm}平方米")
        use_term = _report_text(case.get("use_term_years"))
        if use_term:
            core_clauses.append(f"土地使用年限为{use_term if use_term.endswith('年') else use_term + '年'}")
        land_level = _report_text(case.get("land_level"))
        if land_level:
            core_clauses.append(f"土地级别为{land_level}")

        condition_clauses = []
        road_network = pick("road_network", "regional_road_network", "surrounding_road_network")
        if road_network:
            condition_clauses.append(f"区域内路网由{road_network}构成")
        traffic = pick("traffic_condition", "traffic", "road_access", "临近道路状况")
        if traffic:
            condition_clauses.append(f"交通{traffic}")
        utilities = pick("infrastructure_guarantee", "water_power_guarantee", "utilities_guarantee")
        if utilities:
            condition_clauses.append(f"水电综合保证率{utilities}")
        adjacent_road = pick("adjacent_road", "road_frontage")
        if adjacent_road:
            condition_clauses.append(f"宗地临{adjacent_road}")
        parcel_shape = pick("parcel_shape", "宗地形状")
        if parcel_shape:
            condition_clauses.append(f"形状{parcel_shape}")
        terrain = pick("terrain", "地势")
        if terrain:
            condition_clauses.append(f"地势{terrain}")
        geology = pick("geology", "地质条件")
        if geology:
            condition_clauses.append(f"地质条件{geology}")
        bearing_capacity = pick("bearing_capacity", "land_bearing_capacity")
        if bearing_capacity:
            condition_clauses.append(f"土地承载力{bearing_capacity}")
        disaster_frequency = pick("disaster_frequency", "natural_disaster_frequency", "natural_disaster")
        if disaster_frequency:
            condition_clauses.append(f"自然灾害发生频率{disaster_frequency}")

        used_condition_labels = {
            "development_level",
            "transaction_condition",
            "road_network",
            "regional_road_network",
            "surrounding_road_network",
            "traffic_condition",
            "traffic",
            "road_access",
            "临近道路状况",
            "infrastructure_guarantee",
            "water_power_guarantee",
            "utilities_guarantee",
            "adjacent_road",
            "road_frontage",
            "parcel_shape",
            "宗地形状",
            "terrain",
            "地势",
            "geology",
            "地质条件",
            "bearing_capacity",
            "land_bearing_capacity",
            "disaster_frequency",
            "natural_disaster_frequency",
            "natural_disaster",
        }
        skipped_factor_keys = {
            "land_usage",
            "transaction_time",
            "transaction_condition",
            "transaction_method",
            "use_term",
            "parcel_scale",
            "development_level",
        }
        for factor in factors or []:
            key = str(factor.get("key") or "").strip()
            label = str(factor.get("label") or key).strip()
            if key in skipped_factor_keys or key in used_condition_labels or label in used_condition_labels:
                continue
            item = (factor.get("cases") or {}).get(slot) or {}
            if not item.get("confirmed"):
                continue
            value = _report_text(item.get("value") or item.get("level_label"))
            if not value or value in {"与估价对象一致", "一致"}:
                continue
            if label and not value.startswith(label):
                condition_clauses.append(f"{label}{value}")
            else:
                condition_clauses.append(value)

        transaction_condition = pick("transaction_condition") or "正常"
        transaction_date = _cn_date(case.get("transaction_date"))
        total_price = _report_text(case.get("total_price_wan"))
        development = pick("development_level")
        unit_price = _report_text(case.get("unit_price_sqm"))
        price_clauses = []
        if transaction_date:
            price_clauses.append(f"实例交易时间为{transaction_date}")
        if transaction_condition:
            price_clauses.append(f"交易{transaction_condition}")
        if total_price:
            price_clauses.append(f"成交总价为人民币{total_price}万元")
        if development:
            price_clauses.append(f"交易时宗地基础设施状况为{development}")
        if unit_price and supply_method:
            price_clauses.append(f"土地使用权{supply_method}价格为{unit_price}元/平方米")
        elif unit_price:
            price_clauses.append(f"成交地面地价为{unit_price}元/平方米")

        first_sentence = "，".join(core_clauses + condition_clauses) + "。"
        second_sentence = "，".join(price_clauses)
        return f"比较实例{slot}：{first_sentence}{second_sentence + '。' if second_sentence else ''}"

    def build_render_fields(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        cases = analysis.get("selected_cases") or []
        factors = analysis.get("factors") or []
        descriptions = [self._case_description(slot, case, factors) for slot, case in zip(CASE_SLOTS, cases)]
        intro = [
            "根据中华人民共和国国家标准《城镇土地估价规程》[GB/T18508-2014，2014年7月24日]规定，选择比较实例应符合下列要求：",
            "①选择比较实例数量应达到3个以上（含3个），且估价期日距比较实例的交易时间原则上不超过3年；",
            "②比较实例与估价对象条件的相似性大于差异性；",
            "③比较实例应为用途与估价对象用途相同，土地条件基本一致；属同一供需圈内相邻地区或类似地区的正常（或可修正为正常）交易案例。",
            "估价人员通过对估价对象类似区域地产比较实例的调查与分析，选择符合以上条件的三个交易案例作为比较实例，确定的各比较实例具体情况如下：",
            *descriptions,
        ]
        basic_labels = [
            ("案例来源", lambda c: c.get("source")),
            ("电子监管号", lambda c: c.get("electronic_supervision_no")),
            ("土地使用权人", lambda c: c.get("recipient")),
            ("土地级别", lambda c: c.get("land_level")),
            ("位置", lambda c: c.get("location")),
            ("成交总价（万元）", lambda c: c.get("total_price_wan")),
            ("成交地面地价（元/㎡）", lambda c: c.get("unit_price_sqm")),
            ("土地用途", lambda c: c.get("land_usage_raw")),
            ("土地面积（㎡）", lambda c: c.get("area_sqm")),
            ("合同签订日期", lambda c: _cn_date(c.get("transaction_date"))),
            ("交易情况", lambda c: (c.get("manual_fields") or {}).get("transaction_condition") or "正常"),
            ("交易方式", lambda c: c.get("supply_method")),
            (
                "土地使用年限",
                lambda c: (
                    _report_text(c.get("use_term_years"))
                    if _report_text(c.get("use_term_years")).endswith("年")
                    else f"{_report_text(c.get('use_term_years'))}年" if _report_text(c.get("use_term_years")) else ""
                ),
            ),
        ]
        basic_field_keys = {
            "案例来源": "source",
            "电子监管号": "electronic_supervision_no",
            "土地使用权人": "recipient",
            "土地级别": "land_level",
            "位置": "location",
            "成交总价（万元）": "total_price_wan",
            "成交地面地价（元/㎡）": "unit_price_sqm",
            "土地用途": "land_usage_raw",
            "土地面积（㎡）": "area_sqm",
            "合同签订日期": "transaction_date",
            "交易情况": "manual_fields.transaction_condition",
            "交易方式": "supply_method",
            "土地使用年限": "use_term_years",
        }
        basic_rows = []
        for label, getter in basic_labels:
            field_key = basic_field_keys.get(label, "")
            basic_rows.append(
                {
                    "label": label,
                    **{slot.lower(): getter(case) for slot, case in zip(CASE_SLOTS, cases)},
                    "cell_refs": [
                        "",
                        *[
                            f"market_comp_analysis.selected_cases.{index}.{field_key}" if field_key else ""
                            for index, _case in enumerate(cases)
                        ],
                    ],
                }
            )
        subject = analysis.get("subject") or {}
        calculations = {item["slot"]: item for item in analysis.get("calculations") or []}

        def subgroup(factor: Dict[str, Any]) -> str:
            label = str(factor.get("label") or "")
            group = str(factor.get("group") or "")
            if group == "区域因素" and label in {"距高速公路出入口距离", "距货运火车站距离", "距货运场站距离", "临近道路状况"}:
                return "交通条件"
            return label if group in {"区域因素", "个别因素"} else ""

        factor_condition_rows = [
            {
                "group": "",
                "subgroup": "",
                "label": "项目名称",
                "subject": subject.get("project_name") or subject.get("land_location") or "待校核",
                **{slot.lower(): case.get("project_name") for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "",
                    "",
                    "project_name",
                    *[f"market_comp_analysis.selected_cases.{index}.project_name" for index, _case in enumerate(cases)],
                ],
            },
            {
                "group": "",
                "subgroup": "",
                "label": "位置",
                "subject": subject.get("land_location") or "待校核",
                **{slot.lower(): case.get("location") for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "",
                    "",
                    "land_location_full",
                    *[f"market_comp_analysis.selected_cases.{index}.location" for index, _case in enumerate(cases)],
                ],
            },
            {
                "group": "",
                "subgroup": "",
                "label": "售价（元/㎡）",
                "subject": "待估",
                **{slot.lower(): case.get("unit_price_sqm") for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "",
                    "",
                    "",
                    *[f"market_comp_analysis.selected_cases.{index}.unit_price_sqm" for index, _case in enumerate(cases)],
                ],
            },
        ]
        factor_index_rows = [
            {
                "group": "",
                "subgroup": "",
                "label": "座落",
                "subject": subject.get("land_location") or "待校核",
                **{slot.lower(): case.get("location") for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "",
                    "",
                    "land_location_full",
                    *[f"market_comp_analysis.selected_cases.{index}.location" for index, _case in enumerate(cases)],
                ],
            },
            {
                "group": "",
                "subgroup": "",
                "label": "交易价格（元/平方米）",
                "subject": "待估",
                **{slot.lower(): case.get("unit_price_sqm") for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "",
                    "",
                    "",
                    *[f"market_comp_analysis.selected_cases.{index}.unit_price_sqm" for index, _case in enumerate(cases)],
                ],
            },
        ]
        correction_rows = [
            {
                "group": "",
                "subgroup": "",
                "label": "交易价格（元/平方米）",
                **{slot.lower(): case.get("unit_price_sqm") for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "",
                    "",
                    *[f"market_comp_analysis.selected_cases.{index}.unit_price_sqm" for index, _case in enumerate(cases)],
                ],
            }
        ]
        for factor in analysis.get("factors") or []:
            factor_key = str(factor.get("key") or "")
            condition = {
                "group": factor["group"],
                "subgroup": subgroup(factor),
                "label": factor["label"],
                "subject": factor.get("subject_value"),
                "factor_key": factor_key,
                "cell_refs": [
                    "",
                    "",
                    f"market_comp_analysis.factors.{factor_key}",
                    f"market_comp_analysis.factors.{factor_key}.subject_value",
                    *[
                        f"market_comp_analysis.factors.{factor_key}.cases.{slot}.value"
                        for slot in CASE_SLOTS
                    ],
                ],
            }
            index_row = {
                "group": factor["group"],
                "subgroup": subgroup(factor),
                "label": factor["label"],
                "subject": factor.get("subject_index") or "100",
                "factor_key": factor_key,
                "cell_refs": [
                    "",
                    "",
                    f"market_comp_analysis.factors.{factor_key}",
                    f"market_comp_analysis.factors.{factor_key}.subject_index",
                    *[
                        f"market_comp_analysis.factors.{factor_key}.cases.{slot}.index"
                        for slot in CASE_SLOTS
                    ],
                ],
            }
            correction = {
                "group": factor["group"],
                "subgroup": subgroup(factor),
                "label": factor["label"],
                "factor_key": factor_key,
                "cell_refs": [
                    "",
                    "",
                    f"market_comp_analysis.factors.{factor_key}",
                    *[
                        f"market_comp_analysis.factors.{factor_key}.cases.{slot}.index"
                        for slot in CASE_SLOTS
                    ],
                ],
            }
            subject_index = _decimal(factor.get("subject_index"), "100")
            for slot in CASE_SLOTS:
                item = factor.get("cases", {}).get(slot) or {}
                condition[slot.lower()] = item.get("value")
                index_row[slot.lower()] = item.get("index")
                case_index = _decimal(item.get("index"))
                correction[slot.lower()] = f"{_q(subject_index, '0.01')}/{_q(case_index, '0.01')}" if case_index else ""
            factor_condition_rows.append(condition)
            factor_index_rows.append(index_row)
            correction_rows.append(correction)
        for label, key in (("综合因素修正系数", "correction_coefficient"), ("比准价格（元/㎡）", "corrected_price")):
            correction_rows.append(
                {
                    "group": "测算结果",
                    "subgroup": "",
                    "label": label,
                    **{slot.lower(): calculations.get(slot, {}).get(key) for slot in CASE_SLOTS},
                    "cell_refs": [
                        "",
                        "",
                        "",
                        *[f"market_comp_analysis.calculations.{slot}.{key}" for slot in CASE_SLOTS],
                    ],
                }
            )
        corrected = [calculations.get(slot, {}).get("corrected_price") for slot in CASE_SLOTS]
        market_comp_price = _report_text(analysis.get("market_comp_price"))
        correction_rows.append(
            {
                "group": "测算结果",
                "subgroup": "",
                "label": "土地单价（元/㎡）",
                **{slot.lower(): market_comp_price for slot in CASE_SLOTS},
                "cell_refs": ["", "", "", *["market_comp_analysis.market_comp_price" for _slot in CASE_SLOTS]],
            }
        )
        time_factor = next((item for item in factors if item.get("key") == "transaction_time"), {})
        time_cases = time_factor.get("cases") or {}
        time_index_rows = [
            {
                "label": "成交日期",
                "subject": _cn_date(subject.get("valuation_date")),
                **{slot.lower(): _cn_date(case.get("transaction_date")) for slot, case in zip(CASE_SLOTS, cases)},
                "cell_refs": [
                    "",
                    "valuation_date",
                    *[
                        f"market_comp_analysis.selected_cases.{index}.transaction_date"
                        for index, _case in enumerate(cases)
                    ],
                ],
            },
            {
                "label": "综合修正系数",
                "subject": "0",
                **{
                    slot.lower(): _q(abs(_decimal(time_cases.get(slot, {}).get("index"), "100") - Decimal("100")), "0.01")
                    for slot in CASE_SLOTS
                },
                "cell_refs": [
                    "",
                    "",
                    *[f"market_comp_analysis.factors.transaction_time.cases.{slot}.index" for slot in CASE_SLOTS],
                ],
            },
            {
                "label": "交易期日指数",
                "subject": time_factor.get("subject_index") or "100",
                **{slot.lower(): time_cases.get(slot, {}).get("index") for slot in CASE_SLOTS},
                "cell_refs": [
                    "",
                    "market_comp_analysis.factors.transaction_time.subject_index",
                    *[f"market_comp_analysis.factors.transaction_time.cases.{slot}.index" for slot in CASE_SLOTS],
                ],
            },
        ]
        if all(corrected) and market_comp_price:
            step4 = (
                "经过比较分析，采用各因素修正系数连乘法，求算各比较实例经因素修正后达到估价对象条件时的比准价格。"
                "经过比较分析，三个比准价格中价格相近。故取三个比准价格的简单算术平均值作为市场比较法的评估结果，"
                f"即：（{corrected[0]}+{corrected[1]}+{corrected[2]}）÷3≈{market_comp_price}元/平方米。"
            )
        else:
            step4 = (
                "经过比较分析，采用各因素修正系数连乘法，求算各比较实例经因素修正后达到估价对象条件时的比准价格。"
                "待完成三宗比较实例的因素确认及测算后，取三个比准价格的简单算术平均值作为市场比较法的评估结果。"
            )
        scheme = analysis.get("factor_scheme_snapshot") or {}
        if not scheme:
            scheme = default_factor_scheme(str(subject.get("land_usage_key") or "other"))
        factor_defs = factors or [item for item in scheme.get("factors") or [] if item.get("enabled", True)]
        factor_by_key = {str(item.get("key") or ""): item for item in factor_defs}
        group_factors: Dict[str, List[str]] = {}
        for item in factor_defs:
            group_factors.setdefault(str(item.get("group") or "其他因素"), []).append(
                str(item.get("label") or item.get("key") or "未命名因素")
            )
        selection_descriptions = {
            "transaction_time": "根据地价指数，确定估价期日修正系数。",
            "transaction_condition": "主要指交易是否正常、公开、公平、自愿的交易。",
            "transaction_method": "主要指拍卖、招标、挂牌、协议出让、转让、企业改制、抵押等。",
            "land_usage": "是指土地的具体用途。",
            "use_term": "指待估宗地和比较实例的土地使用年期。",
        }
        factor_selection_lines = ["根据估价对象的宗地条件，影响估价对象价格的主要因素有："]
        selection_items: List[str] = []
        for key in ("transaction_time", "transaction_condition", "transaction_method", "land_usage", "use_term"):
            factor = factor_by_key.get(key)
            if factor:
                selection_items.append(f"{factor.get('label')}：{selection_descriptions[key]}")
        for group in ("区域因素", "个别因素"):
            labels = group_factors.get(group) or []
            if labels:
                selection_items.append(f"{group}：主要有{'、'.join(labels)}等。")
        numerals = "①②③④⑤⑥⑦⑧⑨⑩"
        factor_selection_lines.extend(
            f"{numerals[index] if index < len(numerals) else str(index + 1) + '、'}{text}"
            for index, text in enumerate(selection_items)
        )
        factor_selection = "\n".join(factor_selection_lines)
        broad_labels = [
            str(item.get("label") or item.get("key"))
            for item in factor_defs
            if str(item.get("group") or "") == "交易因素"
        ]
        broad_labels.extend(group for group in ("区域因素", "个别因素") if group_factors.get(group))
        condition_intro = (
            f"根据估价对象的特点，综合确定{'、'.join(broad_labels)}作为本次估价的比较因素，"
            "估价对象与可比实例的比较因素具体情况见下表："
        )
        monthly_growth_rate = _report_text(analysis.get("monthly_growth_rate")) or "0.13"
        land_reduction_rate = _report_text(analysis.get("land_reduction_rate")) or "5.4"
        index_lines = [
            "为将各比较实例和估价对象间的各种差别量化，然后反映到地价水平的差别上，必须将前述的估价对象和比较实例的各种比较因素条件转化为可比的定量条件指数。",
            "以估价对象的各因素条件为基础，相应指数为100，将比较实例相应因素条件与估价对象相比较，确定相应指数。",
        ]
        letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        subject_years = _report_text(subject.get("land_use_term_years"))
        needs_time_adjustment = any(
            _date_text(case.get("transaction_date")) != _date_text(subject.get("valuation_date"))
            for case in cases
            if case.get("transaction_date") and subject.get("valuation_date")
        )
        needs_term_adjustment = any(
            _report_text(case.get("use_term_years")) != subject_years
            for case in cases
            if case.get("use_term_years") and subject_years
        )
        for index, factor in enumerate(factor_defs):
            label = factor.get("label") or factor.get("key") or "未命名因素"
            key = factor.get("key")
            index_lines.append(f"{letters[index] if index < len(letters) else index + 1}、{label}")
            case_values = [
                _report_text((factor.get("cases") or {}).get(slot, {}).get("value"))
                for slot in CASE_SLOTS
            ]
            case_indexes = [
                _report_text((factor.get("cases") or {}).get(slot, {}).get("index"))
                for slot in CASE_SLOTS
            ]
            subject_value = _report_text(factor.get("subject_value"))
            if key == "transaction_time" and needs_time_adjustment:
                index_lines.append(
                    f"估价对象与比较实例的交易时间存在差异，根据交易时间修正规则，"
                    f"本次采用月平均增长率{monthly_growth_rate}%确定交易期日指数，具体见交易期日指数表。"
                )
            elif key == "use_term" and needs_term_adjustment:
                index_lines.append(
                    f"估价对象与比较实例的土地使用年期存在差异，采用土地还原率{land_reduction_rate}%及剩余年期公式进行修正。"
                )
            elif subject_value and all(value == subject_value for value in case_values if value) and all(case_values):
                index_lines.append(
                    f"估价对象与三个比较实例的{label}均为{subject_value}，条件一致，不需进行修正，因素条件指数均确定为100。"
                )
            else:
                details = "；".join(
                    f"案例{slot}条件为{value or '待校核'}，指数为{case_index or '待校核'}"
                    for slot, value, case_index in zip(CASE_SLOTS, case_values, case_indexes)
                )
                index_lines.append(
                    f"以估价对象{label}条件为基础，相应指数为{factor.get('subject_index') or '100'}；{details}。"
                )
        index_lines.append("根据以上比较因素指数的说明，编制比较因素条件指数表。")
        generated_narratives = {
            "market_comp_method_intro": MARKET_COMP_INTRO,
            "market_comp_step1_instances": "\n".join(intro),
            "market_comp_comparable_basis": COMPARABLE_BASIS_TEMPLATE,
            "market_comp_factor_selection": factor_selection,
            "market_comp_condition_intro": condition_intro,
            "market_comp_index_basis": "\n".join(index_lines),
            "market_comp_step4_solve": step4,
            "market_comp_verification": "",
        }
        overrides = analysis.get("narrative_overrides") or {}
        effective_narratives = {
            key: ("" if key == "market_comp_verification" else str(overrides[key]) if key in overrides else value)
            for key, value in generated_narratives.items()
        }
        def segment_source(
            text: Any,
            field: str,
            prefix: str = "",
            suffix: str = "",
            priority: int = 0,
        ) -> Dict[str, str]:
            return {
                "text": str(text or "").strip(),
                "field": field,
                "prefix": prefix,
                "suffix": suffix,
                "priority": str(priority),
            }

        narrative_segment_sources: Dict[str, List[Dict[str, str]]] = {
            key: [] for key in generated_narratives
        }
        for index, (slot, case) in enumerate(zip(CASE_SLOTS, cases)):
            case_prefix = f"比较实例{slot}："
            narrative_segment_sources["market_comp_step1_instances"].extend(
                [
                    segment_source(
                        case.get("electronic_supervision_no"),
                        f"market_comp_analysis.selected_cases.{index}.electronic_supervision_no",
                        prefix="电子监管号为",
                        suffix="）",
                        priority=10,
                    ),
                    segment_source(
                        case.get("project_name"),
                        f"market_comp_analysis.selected_cases.{index}.project_name",
                        prefix="项目名称为",
                        suffix="，",
                        priority=8,
                    ),
                    segment_source(
                        case.get("location"),
                        f"market_comp_analysis.selected_cases.{index}.location",
                        prefix="位于",
                        suffix="，",
                        priority=8,
                    ),
                    segment_source(
                        case.get("unit_price_sqm"),
                        f"market_comp_analysis.selected_cases.{index}.unit_price_sqm",
                        prefix="价格为",
                        suffix="元/平方米",
                        priority=10,
                    ),
                    segment_source(
                        case.get("transaction_date"),
                        f"market_comp_analysis.selected_cases.{index}.transaction_date",
                        prefix="合同签订日期为",
                        suffix="，",
                        priority=8,
                    ),
                    segment_source(
                        case_prefix,
                        f"market_comp_analysis.selected_cases.{index}",
                        priority=1,
                    ),
                ]
            )
        for factor in factors:
            factor_key = str(factor.get("key") or "")
            label = str(factor.get("label") or factor_key or "").strip()
            if not factor_key or not label:
                continue
            narrative_segment_sources["market_comp_factor_selection"].append(
                segment_source(label, f"market_comp_analysis.factors.{factor_key}", suffix="：", priority=3)
            )
            narrative_segment_sources["market_comp_condition_intro"].append(
                segment_source(label, f"market_comp_analysis.factors.{factor_key}", priority=1)
            )
            narrative_segment_sources["market_comp_index_basis"].append(
                segment_source(label, f"market_comp_analysis.factors.{factor_key}", prefix="、", priority=1)
            )
            subject_value = _report_text(factor.get("subject_value"))
            if subject_value:
                narrative_segment_sources["market_comp_index_basis"].append(
                    segment_source(
                        subject_value,
                        f"market_comp_analysis.factors.{factor_key}.subject_value",
                        prefix=f"估价对象与三个比较实例的{label}均为",
                        suffix="，条件一致",
                        priority=8,
                    )
                )
            for slot in CASE_SLOTS:
                case_item = (factor.get("cases") or {}).get(slot) or {}
                case_value = _report_text(case_item.get("value"))
                case_index = _report_text(case_item.get("index"))
                if case_value:
                    narrative_segment_sources["market_comp_index_basis"].append(
                        segment_source(
                            case_value,
                            f"market_comp_analysis.factors.{factor_key}.cases.{slot}.value",
                            prefix=f"案例{slot}条件为",
                            suffix="，指数为",
                            priority=7,
                        )
                    )
                if case_index:
                    narrative_segment_sources["market_comp_index_basis"].append(
                        segment_source(
                            case_index,
                            f"market_comp_analysis.factors.{factor_key}.cases.{slot}.index",
                            prefix=f"案例{slot}条件为{case_value or '待校核'}，指数为",
                            priority=9,
                        )
                    )
        if needs_time_adjustment:
            narrative_segment_sources["market_comp_index_basis"].append(
                segment_source(
                    monthly_growth_rate,
                    "market_comp_analysis.monthly_growth_rate",
                    prefix="本次采用月平均增长率",
                    suffix="%确定",
                    priority=12,
                )
            )
        if needs_term_adjustment:
            narrative_segment_sources["market_comp_index_basis"].append(
                segment_source(
                    land_reduction_rate,
                    "market_comp_analysis.land_reduction_rate",
                    prefix="采用土地还原率",
                    suffix="%及",
                    priority=12,
                )
            )
        if all(corrected) and market_comp_price:
            narrative_segment_sources["market_comp_step4_solve"].extend(
                [
                    segment_source(
                        corrected[0],
                        "market_comp_analysis.calculations.A.corrected_price",
                        prefix="即：（",
                        suffix="+",
                        priority=12,
                    ),
                    segment_source(
                        corrected[1],
                        "market_comp_analysis.calculations.B.corrected_price",
                        prefix=f"{corrected[0]}+",
                        suffix="+",
                        priority=12,
                    ),
                    segment_source(
                        corrected[2],
                        "market_comp_analysis.calculations.C.corrected_price",
                        prefix=f"{corrected[1]}+",
                        suffix="）÷3",
                        priority=12,
                    ),
                    segment_source(
                        market_comp_price,
                        "market_comp_analysis.market_comp_price",
                        prefix="≈",
                        suffix="元/平方米",
                        priority=12,
                    ),
                ]
            )
        return {
            "generated_narratives": generated_narratives,
            "effective_narratives": effective_narratives,
            "narrative_segment_sources": narrative_segment_sources,
            "market_comp_step1_instances": effective_narratives["market_comp_step1_instances"],
            "market_comp_comparable_basis": effective_narratives["market_comp_comparable_basis"],
            "market_comp_factor_selection": effective_narratives["market_comp_factor_selection"],
            "market_comp_condition_intro": effective_narratives["market_comp_condition_intro"],
            "market_comp_index_basis": effective_narratives["market_comp_index_basis"],
            "market_comp_verification": "",
            "market_comp_basic_rows": basic_rows,
            "market_comp_factor_condition_rows": factor_condition_rows,
            "market_comp_time_index_rows": time_index_rows,
            "market_comp_factor_index_rows": factor_index_rows,
            "market_comp_correction_rows": correction_rows,
            "market_comp_step4_solve": effective_narratives["market_comp_step4_solve"],
            "instance_a_desc": descriptions[0] if len(descriptions) > 0 else "",
            "instance_b_desc": descriptions[1] if len(descriptions) > 1 else "",
            "instance_c_desc": descriptions[2] if len(descriptions) > 2 else "",
            "comparable_case_count": len(cases),
        }

    def create_snapshot(self, case_id: str) -> Dict[str, Any]:
        case = self.get_case(case_id)
        if not case:
            raise KeyError(case_id)
        snapshot_hash = _content_hash(case)
        with self._connect() as conn:
            row = conn.execute(
                "SELECT * FROM comparable_case_snapshots WHERE case_id=? AND snapshot_hash=?",
                (case_id, snapshot_hash),
            ).fetchone()
        if row and row["pdf_path"] and Path(row["pdf_path"]).exists():
            return self._snapshot_row(row)
        snapshot_id = row["snapshot_id"] if row else uuid.uuid4().hex
        folder = self.evidence_dir / snapshot_hash
        folder.mkdir(parents=True, exist_ok=True)
        json_path = folder / "official_snapshot.json"
        json_path.write_text(json.dumps(case, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
        docx_path = folder / "成交公告证据页.docx"
        pdf_path = folder / "成交公告证据页.pdf"
        self._build_evidence_docx(case, docx_path)
        self._convert_docx_to_pdf(docx_path, pdf_path)
        image_paths = self._pdf_to_png(pdf_path, folder)
        with self._connect() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO comparable_case_snapshots
                    (snapshot_id, case_id, snapshot_hash, snapshot_json, docx_path, pdf_path, image_paths_json, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    snapshot_id,
                    case_id,
                    snapshot_hash,
                    json.dumps(case, ensure_ascii=False, default=str),
                    str(docx_path),
                    str(pdf_path),
                    json.dumps(image_paths, ensure_ascii=False),
                    _now(),
                ),
            )
            row = conn.execute("SELECT * FROM comparable_case_snapshots WHERE snapshot_id=?", (snapshot_id,)).fetchone()
        return self._snapshot_row(row)

    def create_manual_snapshot(
        self,
        case_id: str,
        evidence_kind: str,
        files: List[tuple[str, bytes]],
    ) -> Dict[str, Any]:
        case = self.get_case(case_id)
        if not case:
            raise KeyError(case_id)
        kind = str(evidence_kind or "announcement").strip()
        if kind not in {"announcement", "location", "site"}:
            raise ValueError("证据类型仅支持成交公告、位置图或现状图。")
        image_files: List[tuple[str, bytes, str]] = []
        for name, raw in files:
            suffix = Path(name or "").suffix.lower()
            if suffix not in {".png", ".jpg", ".jpeg", ".bmp", ".webp"}:
                raise ValueError("证据截图仅支持 png、jpg、jpeg、bmp、webp 图片。")
            if not raw:
                raise ValueError("上传图片为空。")
            image_files.append((name, raw, suffix))
        digest = hashlib.sha256()
        digest.update(json.dumps(case, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8"))
        digest.update(kind.encode("utf-8"))
        for name, raw, _suffix in image_files:
            digest.update(str(name).encode("utf-8", errors="ignore"))
            digest.update(raw)
        snapshot_hash = f"manual_{kind}_{digest.hexdigest()}"
        snapshot_id = uuid.uuid4().hex
        folder = self.evidence_dir / snapshot_hash[:32]
        folder.mkdir(parents=True, exist_ok=True)
        image_paths: List[str] = []
        for index, (name, raw, suffix) in enumerate(image_files, start=1):
            safe_name = re.sub(r'[\\/:*?"<>|]', "_", Path(name).stem or f"evidence_{index}")
            path = folder / f"{index:02d}_{safe_name}{suffix}"
            path.write_bytes(raw)
            image_paths.append(str(path))
        snapshot_json = dict(
            case,
            evidence_kind=kind,
            evidence_source="manual_upload",
            evidence_note="估价师人工上传的官网截图或现场调查图片。",
        )
        (folder / "manual_snapshot.json").write_text(
            json.dumps(snapshot_json, ensure_ascii=False, indent=2, default=str),
            encoding="utf-8",
        )
        with self._connect() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO comparable_case_snapshots
                    (snapshot_id, case_id, snapshot_hash, snapshot_json, docx_path, pdf_path, image_paths_json, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    snapshot_id,
                    case_id,
                    snapshot_hash,
                    json.dumps(snapshot_json, ensure_ascii=False, default=str),
                    "",
                    "",
                    json.dumps(image_paths, ensure_ascii=False),
                    _now(),
                ),
            )
            row = conn.execute("SELECT * FROM comparable_case_snapshots WHERE snapshot_id=?", (snapshot_id,)).fetchone()
        return self._snapshot_row(row)

    def get_snapshot(self, snapshot_id: str) -> Optional[Dict[str, Any]]:
        with self._connect() as conn:
            row = conn.execute("SELECT * FROM comparable_case_snapshots WHERE snapshot_id=?", (snapshot_id,)).fetchone()
        return self._snapshot_row(row) if row else None

    def _snapshot_row(self, row: sqlite3.Row) -> Dict[str, Any]:
        return {
            "snapshot_id": row["snapshot_id"],
            "case_id": row["case_id"],
            "snapshot_hash": row["snapshot_hash"],
            "data": json.loads(row["snapshot_json"]),
            "docx_path": row["docx_path"],
            "pdf_path": row["pdf_path"],
            "image_paths": json.loads(row["image_paths_json"] or "[]"),
            "created_at": row["created_at"],
        }

    def _build_evidence_docx(self, case: Dict[str, Any], path: Path) -> None:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("中国土地市场网供地结果证据快照")
        run.bold = True
        run.font.size = Pt(16)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        fields = [
            ("行政区", case.get("administrative_region_name")),
            ("电子监管号", case.get("electronic_supervision_no")),
            ("项目名称", case.get("project_name")),
            ("项目位置", case.get("location")),
            ("土地面积（㎡）", case.get("area_sqm")),
            ("土地用途", case.get("land_usage_raw")),
            ("供地方式", case.get("supply_method")),
            ("土地使用年限", case.get("use_term_years")),
            ("土地级别", case.get("land_level")),
            ("成交价格（万元）", case.get("total_price_wan")),
            ("成交地面地价（元/㎡）", case.get("unit_price_sqm")),
            ("土地使用权人", case.get("recipient")),
            ("约定容积率", f"{case.get('plot_ratio_min') or ''} - {case.get('plot_ratio_max') or ''}".strip(" -")),
            ("合同签订日期", _cn_date(case.get("transaction_date"))),
            ("来源链接", case.get("source_url")),
            ("抓取时间", case.get("fetched_at")),
        ]
        for label, value in fields:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value or "")
        note = doc.add_paragraph("*以上数据信息以原始国有建设用地使用权出让合同或国有建设用地划拨决定书等为准。")
        note.runs[0].font.size = Pt(9)
        doc.save(path)

    def _find_soffice(self) -> str:
        candidates = [
            self.base_dir / "tools" / "LibreOffice-25.8.7" / "program" / "soffice.com",
            Path(r"C:\Program Files\LibreOffice\program\soffice.com"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.com"),
        ]
        found = shutil.which("soffice.com") or shutil.which("soffice")
        if found:
            candidates.append(Path(found))
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        raise RuntimeError("未找到 LibreOffice soffice.com，无法生成成交公告证据 PDF。")

    def _convert_docx_to_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        with tempfile.TemporaryDirectory(prefix="lo_evidence_", dir=_libreoffice_ascii_temp_root()) as stage:
            stage_dir = Path(stage)
            staged_docx = stage_dir / "input.docx"
            staged_pdf = stage_dir / "input.pdf"
            profile = stage_dir / "profile"
            profile.mkdir(parents=True, exist_ok=True)
            shutil.copy2(docx_path, staged_docx)
            profile_uri = "file:///" + str(profile).replace("\\", "/")
            try:
                process = subprocess.run(
                    [
                        self._find_soffice(),
                        "--headless",
                        "--invisible",
                        "--nologo",
                        "--nofirststartwizard",
                        "--nolockcheck",
                        "--norestore",
                        f"-env:UserInstallation={profile_uri}",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(stage_dir),
                        str(staged_docx),
                    ],
                    capture_output=True,
                    timeout=120,
                    check=False,
                )
            except subprocess.TimeoutExpired as exc:
                raise RuntimeError("soffice.com 证据 PDF 转换超时。") from exc
            if not staged_pdf.exists():
                pdf_candidates = sorted(
                    stage_dir.glob("*.pdf"),
                    key=lambda path: path.stat().st_mtime,
                    reverse=True,
                )
                if pdf_candidates:
                    staged_pdf = pdf_candidates[0]
            if process.returncode != 0 or not staged_pdf.exists() or staged_pdf.stat().st_size <= 0:
                stderr = process.stderr.decode(errors="ignore") if process.stderr else ""
                stdout = process.stdout.decode(errors="ignore") if process.stdout else ""
                raise RuntimeError(f"soffice.com 证据 PDF 转换失败：{stderr or stdout}")
            shutil.copy2(staged_pdf, pdf_path)

    def _pdf_to_png(self, pdf_path: Path, folder: Path) -> List[str]:
        try:
            import fitz  # type: ignore
        except ImportError as exc:
            raise RuntimeError("缺少 PyMuPDF，请先执行 pip install PyMuPDF。") from exc
        paths: List[str] = []
        with fitz.open(pdf_path) as document:
            for index, page in enumerate(document):
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_path = folder / f"page_{index + 1}.png"
                pix.save(str(image_path))
                paths.append(str(image_path))
        return paths

    def snapshot_images_as_data_urls(self, snapshot_ids: Iterable[str]) -> List[List[str]]:
        result: List[List[str]] = []
        for snapshot_id in snapshot_ids:
            snapshot = self.get_snapshot(snapshot_id)
            urls = []
            for path in (snapshot or {}).get("image_paths") or []:
                raw = Path(path).read_bytes()
                urls.append("data:image/png;base64," + base64.b64encode(raw).decode("ascii"))
            result.append(urls)
        return result
