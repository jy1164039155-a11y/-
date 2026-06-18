# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.services.market_comparison_structure import (
    build_market_comparison_table_sections,
    table_row_values,
)
from src.services.report_docx_styles import format_report_paragraph, format_report_table


CASE_SLOTS = ("A", "B", "C")


def _remove_node(node: Any) -> None:
    parent = node.getparent()
    if parent is not None:
        parent.remove(node)


def _insert_paragraph_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.add_run(text)
    return result


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is not None:
        return
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _set_cell_margins(cell: Any, value: int = 45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = margins.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_table(table: Table, *, font_name: str = "仿宋_GB2312", font_size: float = 10.0) -> None:
    width_map = {
        4: (3.2, 3.8, 3.8, 3.8),
        5: (2.6, 3.0, 3.0, 3.0, 3.0),
        6: (1.6, 2.0, 2.5, 2.8, 2.8, 2.8),
        7: (1.4, 1.8, 2.2, 2.3, 2.3, 2.3, 2.3),
    }
    widths = width_map.get(len(table.columns))
    format_report_table(table, widths_cm=widths, font_size_pt=font_size)


def _replace_table(document: Document, old_table: Table, headers: List[str], rows: Iterable[List[Any]]) -> Table:
    style = old_table.style
    new_table = document.add_table(rows=1, cols=len(headers))
    try:
        new_table.style = style
    except Exception:
        new_table.style = "Table Grid"
    for cell, value in zip(new_table.rows[0].cells, headers):
        cell.text = str(value or "")
    for values in rows:
        cells = new_table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value if value not in (None, "") else "")
    old_table._tbl.addnext(new_table._tbl)
    _remove_node(old_table._tbl)
    _format_table(new_table)
    return new_table


def _replace_structured_table(document: Document, old_table: Table, section: Dict[str, Any]) -> Table:
    columns = section.get("columns") or []
    header_rows = section.get("header_rows") or []
    new_table = document.add_table(rows=max(len(header_rows), 1), cols=len(columns))
    try:
        new_table.style = old_table.style
    except Exception:
        new_table.style = "Table Grid"

    occupied = set()
    for row_index, row_spec in enumerate(header_rows):
        column_index = 0
        for item in row_spec:
            while (row_index, column_index) in occupied:
                column_index += 1
            colspan = int(item.get("colspan") or 1)
            rowspan = int(item.get("rowspan") or 1)
            end_row = row_index + rowspan - 1
            end_column = column_index + colspan - 1
            target = new_table.cell(row_index, column_index)
            if end_row != row_index or end_column != column_index:
                target = target.merge(new_table.cell(end_row, end_column))
            target.text = str(item.get("label") or "")
            for rr in range(row_index, end_row + 1):
                for cc in range(column_index, end_column + 1):
                    occupied.add((rr, cc))
            column_index = end_column + 1

    body_start = len(new_table.rows)
    for row in section.get("rows") or []:
        cells = new_table.add_row().cells
        for cell, value in zip(cells, table_row_values(section, row)):
            cell.text = str(value if value not in (None, "") else "")

    for column_index in section.get("group_columns") or []:
        start = body_start
        while start < len(new_table.rows):
            value = new_table.cell(start, column_index).text.strip()
            end = start
            while value and end + 1 < len(new_table.rows):
                next_value = new_table.cell(end + 1, column_index).text.strip()
                if next_value != value:
                    break
                end += 1
            if value and end > start:
                for row_index in range(start + 1, end + 1):
                    new_table.cell(row_index, column_index).text = ""
                new_table.cell(start, column_index).merge(new_table.cell(end, column_index))
            start = end + 1

    old_table._tbl.addnext(new_table._tbl)
    _remove_node(old_table._tbl)
    table_font_size = 9.0 if section.get("key") == "market_comp_basic_rows" else 10.0
    format_report_table(new_table, widths_cm={
        4: (3.2, 3.8, 3.8, 3.8),
        5: (2.6, 3.0, 3.0, 3.0, 3.0),
        6: (1.6, 2.0, 2.5, 2.8, 2.8, 2.8),
        7: (1.4, 1.8, 2.2, 2.3, 2.3, 2.3, 2.3),
    }.get(len(new_table.columns)), font_size_pt=table_font_size, header_rows=max(len(header_rows), 1))
    return new_table


def _table_key(table: Table) -> str:
    header = " ".join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ""
    if "估价对象与比较实例" in header and "比较实例A" in header:
        return "market_comp_basic_rows"
    if "待估宗地及本级实例" in header:
        return "market_comp_factor_condition_rows"
    if "估价对象与比较案例" in header:
        return "market_comp_time_index_rows"
    if "待估宗地及比较实例" in header and "估价对象" in header:
        return "market_comp_factor_index_rows"
    if "待估宗地及比较实例" in header and "案例A" in header:
        return "market_comp_correction_rows"
    return ""


def _ensure_title_before_table(table: Table, title: str) -> None:
    def format_title(paragraph: Paragraph) -> None:
        format_report_paragraph(paragraph, role="table_title")

    previous = table._tbl.getprevious()
    if previous is not None and previous.tag.endswith("}p"):
        paragraph = Paragraph(previous, table._parent)
        if paragraph.text.strip() == title:
            format_title(paragraph)
            return
        if paragraph.text.strip().startswith(title.split(" ", 1)[0]):
            paragraph.text = title
            format_title(paragraph)
            return
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    paragraph = Paragraph(new_p, table._parent)
    paragraph.add_run(title)
    format_title(paragraph)


def _paragraph_between(document: Document, start_text: str, stop_text: str) -> List[Paragraph]:
    body = list(document.element.body)
    started = False
    result: List[Paragraph] = []
    for node in body:
        if not node.tag.endswith("}p"):
            continue
        paragraph = Paragraph(node, document)
        text = paragraph.text.strip()
        if text == start_text:
            started = True
            continue
        if started and text == stop_text:
            break
        if started:
            result.append(paragraph)
    return result


def _replace_paragraph_block(document: Document, start_text: str, stop_text: str, lines: Iterable[str]) -> None:
    targets = _paragraph_between(document, start_text, stop_text)
    stop = next((p for p in document.paragraphs if p.text.strip() == stop_text), None)
    if stop is None:
        return
    for paragraph in targets:
        _remove_node(paragraph._p)
    for line in lines:
        if str(line or "").strip():
            paragraph = _insert_paragraph_before(stop, str(line).strip())
            format_report_paragraph(paragraph)


def _factor_rows(rows: Iterable[Dict[str, Any]], include_subject: bool) -> List[List[Any]]:
    result = []
    for item in rows:
        values = [item.get("group"), item.get("label")]
        if include_subject:
            values.append(item.get("subject"))
        values.extend(item.get(slot.lower()) for slot in CASE_SLOTS)
        result.append(values)
    return result


def _evidence_paths(data: Dict[str, Any], comparable_library: Any, field_name: str) -> List[List[str]]:
    snapshot_ids = data.get(field_name) or []
    result: List[List[str]] = []
    for snapshot_id in snapshot_ids[:3]:
        snapshot = comparable_library.get_snapshot(snapshot_id) if comparable_library else None
        result.append(list((snapshot or {}).get("image_paths") or []))
    while len(result) < 3:
        result.append([])
    return result


def _find_marker(document: Document, marker_texts: Iterable[str]) -> Optional[Paragraph]:
    markers = tuple(marker_texts)
    return next((p for p in document.paragraphs if p.text.strip().startswith(markers)), None)


def _insert_images_after_marker(
    document: Document,
    marker: Optional[Paragraph],
    image_paths: List[str],
    stop_prefixes: Iterable[str],
) -> None:
        if marker is None:
            return
        existing_paths = [Path(image_path) for image_path in image_paths if Path(image_path).exists()]
        if existing_paths:
            marker.paragraph_format.keep_with_next = True
        for image_node in list(marker._p.xpath(".//w:drawing | .//w:pict")):
            _remove_node(image_node)
        stop_prefix_tuple = tuple(stop_prefixes)
        node = marker._p.getnext()
        while node is not None:
            next_node = node.getnext()
            if node.tag.endswith("}p"):
                paragraph = Paragraph(node, document)
                if paragraph.text.strip().startswith(stop_prefix_tuple):
                    break
                if not paragraph.text.strip() or node.xpath(".//w:drawing") or node.xpath(".//w:pict"):
                    _remove_node(node)
            elif node.tag.endswith("}tbl"):
                table = Table(node, document)
                table_text = "".join(cell.text.strip() for row in table.rows for cell in row.cells)
                if not table_text:
                    _remove_node(node)
                else:
                    break
            node = next_node
        anchor: Paragraph = marker
        for path in existing_paths:
            new_p = OxmlElement("w:p")
            anchor._p.addnext(new_p)
            image_paragraph = Paragraph(new_p, anchor._parent)
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.paragraph_format.first_line_indent = Pt(0)
            image_paragraph.paragraph_format.space_before = Pt(3)
            image_paragraph.paragraph_format.space_after = Pt(3)
            image_paragraph.add_run().add_picture(str(path), width=Cm(15.2))
            anchor = image_paragraph


def _insert_evidence_images(document: Document, data: Dict[str, Any], comparable_library: Any) -> None:
    announcement_paths = _evidence_paths(data, comparable_library, "market_comp_evidence_snapshot_ids")
    location_paths = _evidence_paths(data, comparable_library, "market_comp_location_snapshot_ids")
    site_paths = _evidence_paths(data, comparable_library, "market_comp_site_snapshot_ids")
    for slot, announcement, location, site in zip(CASE_SLOTS, announcement_paths, location_paths, site_paths):
        next_detail = f"比较实例{chr(ord(slot) + 1)}详细情况"
        announcement_marker = _find_marker(document, [f"比较实例{slot}成交公告"])
        _insert_images_after_marker(
            document,
            announcement_marker,
            announcement,
            (f"比较实例{slot}位置图", f"比较实例{slot}现状图", f"比较实例{slot}现场照片", next_detail, "3.建立价格可比基础"),
        )
        location_marker = _find_marker(document, [f"比较实例{slot}位置图"])
        _insert_images_after_marker(
            document,
            location_marker,
            location,
            (f"比较实例{slot}现状图", f"比较实例{slot}现场照片", next_detail, "3.建立价格可比基础"),
        )
        site_marker = _find_marker(document, [f"比较实例{slot}现状图", f"比较实例{slot}现场照片"])
        _insert_images_after_marker(
            document,
            site_marker,
            site,
            (next_detail, "3.建立价格可比基础"),
        )


def apply_market_comparison_to_docx(
    docx_path: str,
    data: Dict[str, Any],
    comparable_library: Optional[Any] = None,
) -> None:
    """Replace the runtime template's sample market-comparison block with report-owned data."""
    if not data.get("use_market_comp"):
        return

    document = Document(docx_path)
    analysis = data.get("market_comp_analysis") or {}
    effective = analysis.get("effective_narratives") or {}
    table_sections = {
        item["key"]: item
        for item in build_market_comparison_table_sections(dict(data, market_comp_analysis=analysis))
    }
    for table in list(document.tables):
        key = _table_key(table)
        if key and key in table_sections:
            _ensure_title_before_table(table, table_sections[key]["report_title"])

    def narrative(key: str, fallback: str = "") -> str:
        return str(effective.get(key) if key in effective else data.get(key) or fallback).strip()

    step1 = narrative("market_comp_step1_instances")
    step1_lines = step1.splitlines() if step1 else ["【请完成市场比较法三宗比较实例选择】"]
    basic_title = (
        "表3-10 比较实例基本情况表"
        if any(p.text.strip() == "表3-10 比较实例基本情况表" for p in document.paragraphs)
        else "表3-10  比较实例基本情况表"
    )
    _replace_paragraph_block(document, "2.比较实例选择", basic_title, step1_lines)

    subject = analysis.get("subject") or {}
    _replace_paragraph_block(
        document,
        "3.建立价格可比基础",
        "4.比较因素的选择",
        narrative("market_comp_comparable_basis", "【请校核价格可比基础】").splitlines(),
    )
    _replace_paragraph_block(
        document,
        "4.比较因素的选择",
        "5.比较因素条件说明",
        narrative("market_comp_factor_selection", "【请校核比较因素选择】").splitlines(),
    )
    _replace_paragraph_block(
        document,
        "5.比较因素条件说明",
        "表3-11 比较因素条件说明表",
        narrative("market_comp_condition_intro", "【请校核比较因素条件说明】").splitlines(),
    )
    time_title = "表3-12 估价对象与比较案例交易期日指数表"
    index_title = "表3-13 估价对象及比较因素条件指数表"
    index_stop = time_title if any(p.text.strip() == time_title for p in document.paragraphs) else index_title
    _replace_paragraph_block(
        document,
        "6.编制比较因素条件指数表",
        index_stop,
        narrative("market_comp_index_basis", "【请校核因素条件指数编制依据】").splitlines(),
    )
    if index_stop == time_title:
        _replace_paragraph_block(document, time_title, index_title, [])

    solve = narrative("market_comp_step4_solve", "【请完成市场比较法因素确认及价格测算】")
    _replace_paragraph_block(document, "表3-14 估价对象比较因素修正系数表", "★收益还原法", [solve])

    for table in list(document.tables):
        key = _table_key(table)
        if key and key in table_sections:
            _replace_structured_table(document, table, table_sections[key])

    _insert_evidence_images(document, data, comparable_library)
    document.save(docx_path)
