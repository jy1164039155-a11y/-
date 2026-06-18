# -*- coding: utf-8 -*-
from __future__ import annotations

import hashlib
import json
import os
import re
from copy import deepcopy
from datetime import date
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document

from src.services.cost_policy_config import (
    default_location_template_key,
    find_policy_rule,
    get_location_template,
    get_risk_scheme,
    load_cost_policy_config,
)


D0 = Decimal("0")
D1 = Decimal("1")
MU_TO_SQM = Decimal("0.0015")
MONEY_QUANT = Decimal("0.01")
HUNAN_PAID_LAND_USE_FEE_STANDARDS: Dict[str, str] = {
    "一等": "140",
    "二等": "120",
    "三等": "100",
    "四等": "80",
    "五等": "64",
    "六等": "56",
    "七等": "48",
    "八等": "42",
    "九等": "34",
    "十等": "28",
    "十一等": "24",
    "十二等": "20",
    "十三等": "16",
    "十四等": "14",
    "十五等": "10",
}
COST_PRICING_SCENARIO_SPECS: List[Dict[str, Any]] = [
    {"key": "current", "label": "当前方案（基线）", "overrides": {}},
    {"key": "paid_land_14", "label": "有偿使用费十四等", "overrides": {"water_conservancy_fund": "十四等"}},
    {"key": "paid_land_12", "label": "有偿使用费十二等", "overrides": {"water_conservancy_fund": "十二等"}},
    {"key": "reclamation_superior", "label": "耕地开垦费优等水田", "overrides": {"farmland_reclamation_fee": "优等水田"}},
    {
        "key": "combo_prudent",
        "label": "税费审慎组合",
        "overrides": {"water_conservancy_fund": "十四等", "farmland_reclamation_fee": "高等水田"},
    },
]
COST_BASIS_ATTACHMENT_DIR = Path("01_Source") / "03_attachment" / "成本逼近法政策与测算依据" / "01_征地补偿与区片"
LOCAL_COMPENSATION_ATTACHMENT_DIR = (
    Path("01_Source") / "03_attachment" / "成本逼近法政策与测算依据" / "02_青苗房屋及附着物补偿"
)
YONGZHOU_LOCAL_COMPENSATION_PDF = "永州市人民政府关于印发《永州市集体土地与房屋征收补偿安置办法》的通知.pdf"
YONGZHOU_LOCAL_COMPENSATION_XLSX = "永州市人民政府关于印发《永州市集体土地与房屋征收补偿安置办法》的通知.xlsx"
YONGZHOU_BUILDING_COMPENSATION_CONFIG = Path("src") / "config" / "yongzhou_building_compensation_rows.json"
HUAIHUA_BUILDING_COMPENSATION_CONFIG = Path("src") / "config" / "huaihua_building_compensation_rows.json"
YONGZHOU_GREEN_SEEDLING_CONFIG = Path("src") / "config" / "yongzhou_green_seedling_standards.json"
HUAIHUA_GREEN_SEEDLING_CONFIG = Path("src") / "config" / "huaihua_green_seedling_standards.json"
COLLECTIVE_LAND_COMPENSATION_DIR = Path("01_Source") / "03_attachment" / "集体土地与房屋征收补偿安置办法"
PROVINCE_COMPENSATION_DOC = "湖南省人民政府关于调整湖南省征地补偿标准的通知湘政发〔2024〕1号.docx"
PROVINCE_SCOPE_PDF = "湖南省自然资源厅关于调整征收农用地补偿区片空间范围的通知 - 湖南省自然资源厅.pdf"
YONGZHOU_ZONE_TABLE = "永州市农用地补偿区片划分表.docx"
PRICE_QUANT = Decimal("0.1")
FACTOR_QUANT = Decimal("0.0001")
PROVINCE_POLICY_NAME = "湖南省人民政府关于调整湖南省征地补偿标准的通知"
PROVINCE_POLICY_NO = "湘政发〔2024〕1号"
PROVINCE_POLICY_EFFECTIVE_DATE = date(2024, 1, 31)
PROVINCE_SCOPE_POLICY_NAME = "湖南省自然资源厅关于调整征收农用地补偿区片空间范围的通知"
PROVINCE_SCOPE_POLICY_NO = "湘自资发〔2024〕7号"
_DOCUMENT_HASH_CACHE: Dict[tuple[str, int, int], str] = {}
_PROVINCE_COMPENSATION_ROW_CACHE: Dict[tuple[str, int, int, str], Dict[str, str]] = {}
_YONGZHOU_XLSX_VALIDATION_CACHE: Dict[tuple[str, int, int, str], Dict[str, Any]] = {}
_JSON_CONFIG_CACHE: Dict[tuple[str, int, int], Dict[str, Any]] = {}
_BUILDING_ROWS_CACHE: Dict[tuple, List[Dict[str, Any]]] = {}


def _path_sig(path: Path) -> tuple[int, int]:
    try:
        stat = path.stat()
        return (stat.st_mtime_ns, stat.st_size)
    except OSError:
        return (0, 0)


def _load_json_config_cached(path: Path) -> Dict[str, Any]:
    """File-aware cached JSON loader. The parsed dict is cached keyed by
    (abspath, mtime_ns, size) so repeated reads during a single interactive
    /calculate request avoid re-reading and re-parsing the policy config from
    disk. Callers must treat the returned dict as read-only (deepcopy before
    mutating); existing callers already deepcopy the slices they edit."""
    try:
        stat = path.stat()
    except OSError:
        return {}
    cache_key = (os.path.normcase(os.path.abspath(path)), stat.st_mtime_ns, stat.st_size)
    cached = _JSON_CONFIG_CACHE.get(cache_key)
    if cached is not None:
        return cached
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        data = {}
    if not isinstance(data, dict):
        data = {}
    _JSON_CONFIG_CACHE[cache_key] = data
    return data

DAO_LOCAL_POLICY = {
    "name": "永州市人民政府关于印发《永州市集体土地与房屋征收补偿安置办法》的通知",
    "document_no": "永政发〔2024〕4号",
    "date": "2024年5月11日",
}
LAND_MANAGEMENT_LAW_REFERENCE = (
    "《中华人民共和国土地管理法》（根据2019年8月26日第十三届全国人民代表大会常务委员会第十二次会议"
    "《关于修改〈中华人民共和国土地管理法〉、〈中华人民共和国城市房地产管理法〉的决定》第三次修正）"
)
STATIC_TAX_POLICY_REFERENCES = (
    "湖南省人民政府关于印发《湖南省水利建设基金筹集和使用管理办法》的通知（湘政发[2011]27号）",
    "《关于调整新增建设用地土地有偿使用费征收管理有关政策的通知》（湘财综[2018]42号）",
    "《湖南省人民代表大会常务委员会关于湖南省耕地占用税适用税额的决定》",
    "《湖南省耕地开垦费征收使用管理办法》",
    "湖南省人力资源社会保障厅《关于进一步做好被征地农民社会保障工作的通知》（湘人社规〔2023〕1号，2023年4月28日）",
    "《湖南省森林植被恢复费征收使用管理实施办法》（湘财税〔2024〕10号）",
)

DEVELOPMENT_COST_RANGES = {
    "通路": ("14", "16"),
    "供水": ("12", "14"),
    "排水": ("14", "16"),
    "通电": ("12", "14"),
    "通讯": ("5", "7"),
    "供气": ("6", "8"),
    "供热": ("6", "8"),
    "场地平整": ("13", "15"),
}

RESIDENTIAL_LOCATION_GRADE_DESCRIPTIONS = {
    "距汽车站的距离": ("近", "较近", "一般", "较远", "远"),
    "距火车站距离": ("近", "较近", "一般", "较远", "远"),
    "距商服中心距离": ("近", "较近", "一般", "较远", "远"),
    "临路类型": ("混合型主干道", "生活型主干道", "交通型主道", "交通次主道", "支路"),
    "路网密度": ("高", "较高", "一般", "较低", "低"),
    "对外交通便利度": ("便利", "较便利", "一般便利", "较不便利", "不便利"),
    "环境质量优劣度": ("优", "较优", "一般", "较劣", "劣"),
    "人口密度": ("密集", "较密集", "一般", "较稀疏", "稀疏"),
    "基础设施完善度": ("完善", "较完善", "一般完善", "较不完善", "不完善"),
    "地形状况": ("平坦", "较平坦", "基本平坦", "较不平坦", "不平坦"),
    "宗地面积": ("适中", "面积较大", "面积一般", "面积较小", "面积小"),
    "宗地形状": ("规则", "较规则", "基本规则", "较不规则", "不规则"),
}

ACQUISITION_LAND_CLASS_TREE = {
    "耕地": ("水田", "水浇地", "旱地"),
    "园地": ("果园", "茶园", "其他园地"),
    "林地": ("乔木林地", "竹林地", "灌木林地", "其他林地"),
    "草地": ("天然牧草地", "人工牧草地", "其他草地"),
    "其他农用地": ("设施农用地", "农村道路", "坑塘水面", "沟渠", "其他"),
}
ACQUISITION_SUBCLASS_TO_CLASS = {
    subclass: land_class
    for land_class, subclasses in ACQUISITION_LAND_CLASS_TREE.items()
    for subclass in subclasses
}

BAOZHEN_RISK_FACTORS = (
    ("政策风险", "产业政策", "A", "0.3234", "0.2957"),
    ("政策风险", "土地使用制度改革", "C", "0.3234", "0.2121"),
    ("政策风险", "住房制度", "B", "0.3234", "0.2734"),
    ("政策风险", "环保变化", "B", "0.3234", "0.2188"),
    ("经济风险", "市场供求", "A", "0.3922", "0.3434"),
    ("经济风险", "财务风险", "B", "0.3922", "0.2119"),
    ("经济风险", "管理风险", "C", "0.3922", "0.2365"),
    ("经济风险", "当地经济发展", "B", "0.3922", "0.2082"),
    ("社会风险", "城市规划", "A", "0.2844", "0.3762"),
    ("社会风险", "区域发展", "A", "0.2844", "0.3113"),
    ("社会风险", "治安", "A", "0.2844", "0.3125"),
)
RISK_LEVEL_ADJUSTMENTS = {"D": "0", "C": "2", "B": "4", "A": "8"}

RESIDENTIAL_LOCATION_FACTORS = (
    ("交通条件", "距汽车站的距离", "2"),
    ("交通条件", "距火车站距离", "2"),
    ("繁华程度", "距商服中心距离", "2"),
    ("交通条件", "临路类型", "2"),
    ("交通条件", "路网密度", "2"),
    ("交通条件", "对外交通便利度", "3"),
    ("环境条件", "环境质量优劣度", "3"),
    ("人口状况", "人口密度", "2"),
    ("基础设施", "基础设施完善度", "3"),
    ("个别因素", "地形状况", "1"),
    ("个别因素", "宗地面积", "1"),
    ("个别因素", "宗地形状", "2"),
)


def _acquisition_land_classes(data: Dict[str, Any]) -> tuple[str, str]:
    raw_class = str(data.get("acquisition_land_class") or "").strip()
    raw_subclass = str(data.get("acquisition_land_subclass") or "").strip()
    if raw_subclass:
        return ACQUISITION_SUBCLASS_TO_CLASS.get(raw_subclass, raw_class or "其他农用地"), raw_subclass
    if raw_class in ACQUISITION_SUBCLASS_TO_CLASS:
        return ACQUISITION_SUBCLASS_TO_CLASS[raw_class], raw_class
    if raw_class in ACQUISITION_LAND_CLASS_TREE:
        default_subclass = ACQUISITION_LAND_CLASS_TREE[raw_class][0]
        return raw_class, default_subclass
    return "耕地", "水田"


def _expected_city(county_name: Any, province_row: Dict[str, Any] | None = None) -> str:
    row_city = str((province_row or {}).get("city") or "").strip()
    if row_city:
        return row_city
    county = str(county_name or "")
    if "通道" in county:
        return "怀化市"
    if "道县" in county:
        return "永州市"
    return ""


def _decimal(value: Any, default: Decimal = D0) -> Decimal:
    if value in (None, ""):
        return default
    text = str(value).strip().replace("%", "").replace(",", "")
    try:
        return Decimal(text)
    except (InvalidOperation, ValueError):
        return default


def _money(value: Decimal) -> str:
    return format(value.quantize(MONEY_QUANT, rounding=ROUND_HALF_UP), "f")


def _price(value: Decimal) -> str:
    return format(value.quantize(PRICE_QUANT, rounding=ROUND_HALF_UP), "f")


def _factor(value: Decimal) -> str:
    return format(value.quantize(FACTOR_QUANT, rounding=ROUND_HALF_UP), "f")


def _trim_decimal_text(value: Decimal) -> str:
    text = format(value.normalize(), "f")
    return text.rstrip("0").rstrip(".") if "." in text else text


def _qmoney(value: Decimal) -> Decimal:
    return value.quantize(MONEY_QUANT, rounding=ROUND_HALF_UP)


def _qprice(value: Decimal) -> Decimal:
    return value.quantize(PRICE_QUANT, rounding=ROUND_HALF_UP)


def _qfactor(value: Decimal) -> Decimal:
    return value.quantize(FACTOR_QUANT, rounding=ROUND_HALF_UP)


def _percent_decimal(value: Any) -> Decimal:
    return _decimal(value) / Decimal("100")


def _bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "yes", "on", "是", "已确认"}


def _date(value: Any) -> date | None:
    numbers = [int(item) for item in re.findall(r"\d+", str(value or ""))[:3]]
    if len(numbers) != 3:
        return None
    try:
        return date(numbers[0], numbers[1], numbers[2])
    except ValueError:
        return None


def _cn_date(value: Any) -> str:
    parsed = _date(value)
    if parsed:
        return f"{parsed.year}年{parsed.month}月{parsed.day}日"
    return str(value or "").strip()


def _policy_reference(policy: Dict[str, Any]) -> str:
    name = str(policy.get("name") or "").strip()
    if not name:
        return ""
    title = name if "《" in name or name.startswith("中华人民共和国") else f"《{name}》"
    document_no = str(policy.get("document_no") or "").strip()
    publish_date = _cn_date(policy.get("publish_date") or policy.get("effective_date"))
    details = "，".join(value for value in (document_no, publish_date) if value)
    return f"{title}（{details}）" if details else title


def _policy_workflow_note(rule: Dict[str, Any]) -> str:
    """Workbench citation + internal audit; must not flow into generated report narratives."""
    parts = [str(rule.get("source_note") or "").strip()]
    internal = str(rule.get("internal_note") or "").strip()
    if internal:
        parts.append(internal)
    return "；".join(part for part in parts if part)


def _local_policy_defaults(county_name: str) -> Dict[str, str]:
    if "通道" in county_name:
        return {"name": "", "document_no": "", "date": ""}
    if "道县" in county_name:
        return dict(DAO_LOCAL_POLICY)
    return {"name": "", "document_no": "", "date": ""}


def _load_yongzhou_green_seedling_config(base_dir: str | Path) -> Dict[str, Any]:
    path = Path(base_dir) / YONGZHOU_GREEN_SEEDLING_CONFIG
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def _load_huaihua_green_seedling_config(base_dir: str | Path) -> Dict[str, Any]:
    path = Path(base_dir) / HUAIHUA_GREEN_SEEDLING_CONFIG
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def _load_green_seedling_config(base_dir: str | Path, county_name: str = "") -> Dict[str, Any]:
    county = str(county_name or "").strip()
    if "通道" in county:
        config = _load_huaihua_green_seedling_config(base_dir)
        if config:
            return config
    config = _load_yongzhou_green_seedling_config(base_dir)
    aliases = [str(item) for item in config.get("county_aliases") or []]
    if config and county and _county_matches_aliases(county, aliases):
        return config
    return {}


def _county_matches_aliases(county_name: str, aliases: Iterable[str]) -> bool:
    county = str(county_name or "").strip()
    if not county:
        return False
    alias_list = [str(item) for item in aliases if item]
    normalized = {_normalize_location_key(item) for item in alias_list}
    if _normalize_location_key(county) in normalized:
        return True
    # Fall back to robust county matching so that names carrying a township,
    # street or city qualifier (e.g. '道县西洲街道', '永州市道县') still resolve
    # to their configured alias instead of silently returning no policy data.
    return any(_county_matches(county, alias) for alias in alias_list)


def _green_seedling_policy_source_path(base_dir: str | Path, config: Dict[str, Any]) -> str:
    source_file = str(config.get("policy_source_file") or "").strip()
    if not source_file:
        return ""
    direct = Path(base_dir) / source_file
    if direct.exists():
        return str(direct)
    collective_dir = _collective_land_compensation_dir(base_dir)
    fallback = collective_dir / Path(source_file).name
    if fallback.exists():
        return str(fallback)
    nested = collective_dir / source_file.replace("\\", "/").split("怀化/", 1)[-1] if "怀化" in source_file else source_file
    if isinstance(nested, str):
        nested_path = collective_dir / "怀化" / Path(nested).name if (collective_dir / "怀化").exists() else Path(nested)
        if nested_path.exists():
            return str(nested_path)
    return str(direct)


def green_seedling_standard_for_county(base_dir: str | Path, county_name: str, land_subclass: str) -> Dict[str, str]:
    config = _load_green_seedling_config(base_dir, county_name)
    policy_no = str(config.get("policy_document_no") or ("待校核文号" if "通道" in str(county_name or "") else "永政发〔2024〕4号"))
    source_path = _green_seedling_policy_source_path(base_dir, config) if config else ""
    aliases = [str(item) for item in config.get("county_aliases") or []]
    subclass = str(land_subclass or "").strip()
    if config and _county_matches_aliases(county_name, aliases):
        for rule in config.get("standards") or []:
            subclasses = [str(item) for item in rule.get("land_subclasses") or []]
            if subclass in subclasses or any(item in subclass or subclass in item for item in subclasses):
                return {
                    "standard_per_mu": str(rule.get("standard_per_mu") or ""),
                    "label": str(rule.get("label") or subclass),
                    "attachment": str(rule.get("attachment") or ""),
                    "source_note": (
                        f"{policy_no}"
                        f" {rule.get('attachment') or '附件3'} {rule.get('label') or subclass}"
                        f" {rule.get('standard_per_mu') or ''}元/亩"
                    ),
                    "source_path": source_path or str(_local_compensation_policy_xlsx_path(base_dir) or _local_compensation_policy_path(base_dir)),
                    "review_status": str(config.get("review_status") or "verified_partial"),
                }
        if "林" in subclass:
            fallback = next((rule for rule in config.get("standards") or [] if rule.get("key") == "forest_patch"), {})
            if fallback:
                return {
                    "standard_per_mu": str(fallback.get("standard_per_mu") or "4000"),
                    "label": str(fallback.get("label") or "成片林地"),
                    "attachment": str(fallback.get("attachment") or "附件3/5"),
                    "source_note": f"{policy_no} 成片林地 {fallback.get('standard_per_mu') or '4000'}元/亩",
                    "source_path": source_path or str(_local_compensation_policy_xlsx_path(base_dir) or _local_compensation_policy_path(base_dir)),
                    "review_status": str(config.get("review_status") or "verified_partial"),
                }
        if any(item in subclass for item in ("水田", "水浇地")):
            fallback = next((rule for rule in config.get("standards") or [] if rule.get("key") == "paddy"), {})
            standard = str(fallback.get("standard_per_mu") or ("1500" if "通道" in str(county_name or "") else "2800"))
            return {
                "standard_per_mu": standard,
                "label": "水田/水浇地",
                "attachment": "附件3",
                "source_note": f"{policy_no} 附件3 水田 {standard}元/亩",
                "source_path": source_path or str(_local_compensation_policy_xlsx_path(base_dir) or _local_compensation_policy_path(base_dir)),
                "review_status": str(config.get("review_status") or "verified_partial"),
            }
        if "旱" in subclass:
            fallback = next((rule for rule in config.get("standards") or [] if rule.get("key") == "dryland"), {})
            standard = str(fallback.get("standard_per_mu") or ("1000" if "通道" in str(county_name or "") else "2000"))
            return {
                "standard_per_mu": standard,
                "label": "旱地",
                "attachment": "附件3",
                "source_note": f"{policy_no} 附件3 旱地 {standard}元/亩",
                "source_path": source_path or str(_local_compensation_policy_xlsx_path(base_dir) or _local_compensation_policy_path(base_dir)),
                "review_status": str(config.get("review_status") or "verified_partial"),
            }
    if "林" in subclass:
        return {"standard_per_mu": "4000", "label": "林地", "attachment": "", "source_note": "默认口径：林地4000元/亩（待匹配县市本地政策）", "source_path": "", "review_status": "pending"}
    if any(item in subclass for item in ("水田", "水浇地")):
        return {"standard_per_mu": "2800", "label": "水田/水浇地", "attachment": "", "source_note": "默认口径：水田2800元/亩（待匹配县市本地政策）", "source_path": "", "review_status": "pending"}
    if "旱" in subclass:
        return {"standard_per_mu": "2000", "label": "旱地", "attachment": "", "source_note": "默认口径：旱地2000元/亩（待匹配县市本地政策）", "source_path": "", "review_status": "pending"}
    return {"standard_per_mu": "", "label": subclass, "attachment": "", "source_note": "未匹配到青苗补偿标准", "source_path": "", "review_status": "pending"}


def cost_basis_attachment_inventory(base_dir: str | Path, county_name: str = "") -> List[Dict[str, Any]]:
    county = str(county_name or "").strip()
    city = _expected_city(county)
    seedling_config = _load_green_seedling_config(base_dir, county)
    building_config = _load_building_compensation_config(base_dir, county)
    seedling_count = len(seedling_config.get("standards") or [])
    building_catalog = building_config.get("grade_catalog") or {}
    building_count = sum(len(rows or []) for rows in building_catalog.values())
    entries: List[Dict[str, Any]] = []
    path_resolvers = {
        "local_compensation_xlsx": _local_compensation_policy_xlsx_path,
        "local_compensation_pdf": _local_compensation_policy_path,
        "province_compensation": _policy_path,
        "zone_table": lambda root: _scope_table_path(root, city) if city else None,
    }
    for source in (seedling_config.get("policy_files") or []) + (building_config.get("policy_files") or [] if building_config.get("policy_files") else []):
        resolver = path_resolvers.get(str(source.get("path_key") or ""))
        path = resolver(base_dir) if resolver else Path(str(source.get("path") or ""))
        if not path or not Path(path).exists():
            continue
        counties = [str(item) for item in source.get("counties") or seedling_config.get("county_aliases") or []]
        if county and counties and not _county_matches_aliases(county, counties):
            continue
        entries.append(
            {
                "category": str(source.get("category") or "政策附件"),
                "label": str(source.get("label") or Path(path).name),
                "source_path": str(path),
                "counties": counties,
                "status": str(source.get("status") or "available"),
            }
        )
    if not entries and seedling_config.get("policy_files"):
        for source in seedling_config.get("policy_files") or []:
            resolver = path_resolvers.get(str(source.get("path_key") or ""))
            path = resolver(base_dir) if resolver else None
            entries.append(
                {
                    "category": str(source.get("category") or "政策附件"),
                    "label": str(source.get("label") or ""),
                    "source_path": str(path) if path and Path(path).exists() else "",
                    "counties": [str(item) for item in source.get("counties") or seedling_config.get("county_aliases") or []],
                    "status": "missing" if not path or not Path(path).exists() else str(source.get("status") or "available"),
                }
            )
    if _policy_path(base_dir).exists():
        entries.append(
            {
                "category": "征地补偿与区片",
                "label": PROVINCE_POLICY_NAME,
                "source_path": str(_policy_path(base_dir)),
                "counties": ["*"],
                "status": "available",
            }
        )
    zone_city = _expected_city(county)
    zone_path = _scope_table_path(base_dir, zone_city) if zone_city else None
    if zone_path and zone_path.exists():
        entries.append(
            {
                "category": "征地补偿与区片",
                "label": f"{zone_city}农用地补偿区片划分表",
                "source_path": str(zone_path),
                "counties": [zone_city] if zone_city else ["*"],
                "status": "available",
            }
        )
    policy_config = load_cost_policy_config()
    policy_root = Path(base_dir) / "01_Source" / "03_attachment" / "成本逼近法政策与测算依据"
    rules_by_file: Dict[str, List[Dict[str, Any]]] = {}
    for rule in policy_config.get("policy_rules") or []:
        source_file = str(rule.get("source_file") or "").strip()
        if not source_file:
            continue
        rules_by_file.setdefault(source_file, []).append(rule)
    fee_labels = {
        "paid_land_use_fee": "新增建设用地土地有偿使用费",
        "water_conservancy_fund": "水利建设基金",
        "farmland_occupation_tax": "耕地占用税",
        "farmland_reclamation_fee": "耕地开垦费",
        "social_security_fund": "社会保障费",
        "forest_restoration_fee": "森林植被恢复费",
    }
    fee_targets = {
        "paid_land_use_fee": "cost_approx_analysis.tax_items.water_conservancy_fund.grade_name",
        "water_conservancy_fund": "cost_approx_analysis.tax_items.water_conservancy_fund.grade_name",
        "farmland_occupation_tax": "cost_approx_analysis.tax_items.farmland_occupation_tax.grade_name",
        "farmland_reclamation_fee": "cost_approx_analysis.tax_items.farmland_reclamation_fee.grade_name",
        "social_security_fund": "cost_approx_analysis.tax_items.social_security_fund.amount_per_sqm",
        "forest_restoration_fee": "cost_approx_analysis.tax_items.forest_restoration_fee.grade_name",
    }
    for source_file, rules in rules_by_file.items():
        source_path = ""
        if policy_root.exists():
            matched_paths = [path for path in policy_root.rglob(Path(source_file).name) if path.is_file()]
            if matched_paths:
                source_path = str(matched_paths[0])
        fee_keys = sorted({str(rule.get("fee_key") or "") for rule in rules if rule.get("fee_key")})
        if not fee_keys:
            continue
        if all(str(rule.get("formula_key") or "") == "manual_default" for rule in rules):
            structured_status = "manual_input"
            next_action = "由估价师填写采用值并确认。"
        elif any(str(rule.get("grade_name") or "").strip() for rule in rules):
            structured_status = "needs_grade_selection"
            next_action = "按县市、估价期日和征收细分类推荐等别；用户可改选后重新确认。"
        else:
            structured_status = "structured"
            next_action = "已进入税费规则配置，可按地区自动匹配。"
        entries.append(
            _apply_inventory_structured_fields(
                {
                    "category": "税费与专项费用",
                    "label": "、".join(fee_labels.get(key, key) for key in fee_keys),
                    "source_path": source_path,
                    "counties": ["*"],
                    "status": "available" if source_path else "missing",
                },
                structured_status=structured_status,
                structured_item_count=len(rules),
                price_fields=[fee_targets.get(key, f"cost_approx_analysis.tax_items.{key}") for key in fee_keys],
                target_ref=fee_targets.get(fee_keys[0], ""),
                next_action=next_action,
            )
        )
    deduped: Dict[str, Dict[str, Any]] = {}
    for item in entries:
        deduped[item.get("source_path") or item.get("label") or ""] = item
    for item in _collective_land_attachment_inventory(base_dir, county):
        deduped[item.get("source_path") or item.get("label") or ""] = item
    result_entries = list(deduped.values())
    for item in result_entries:
        item.setdefault("counties_display", _dedup_counties_display(item.get("counties") or []))
        item.setdefault("status_label", _attachment_status_label(item.get("status")))
        label = str(item.get("label") or "")
        source_name = Path(str(item.get("source_path") or label)).name
        if not item.get("structured_status"):
            if (
                label == PROVINCE_POLICY_NAME
                or "征地补偿标准" in label
                or "区片综合地价" in label
            ):
                _apply_inventory_structured_fields(
                    item,
                    structured_status="needs_grade_selection",
                    structured_item_count=3,
                    price_fields=[
                        "cost_approx_analysis.compensation_zone",
                        "cost_approx_analysis.acquisition_items.land_compensation.standard_value",
                    ],
                    target_ref="cost_approx_analysis.compensation_zone",
                    next_action="按县市和坐落推荐区片，正式采用前需确认区片。",
                )
            elif "区片划分表" in label:
                _apply_inventory_structured_fields(
                    item,
                    structured_status="structured",
                    structured_item_count=1,
                    price_fields=["cost_approx_analysis.compensation_zone_suggestion"],
                    target_ref="cost_approx_analysis.compensation_zone_suggestion",
                    next_action="用于自动推荐征地区片。",
                )
            elif any(token in label or token in source_name for token in ("永政发", "怀化", "房屋", "青苗", "附着物")) and (
                seedling_count or building_count
            ):
                _apply_inventory_structured_fields(
                    item,
                    structured_status="structured",
                    structured_item_count=seedling_count + building_count,
                    price_fields=[
                        "cost_approx_analysis.green_seedling_standard_per_mu",
                        "cost_approx_analysis.building_compensation_rows.0.standard",
                    ],
                    target_ref="cost_approx_analysis.building_compensation_rows.0.standard",
                    next_action="青苗按地类匹配；房屋及附着物可从政策目录添加或选择等级。",
                )
            elif item.get("status") == "scan_pending_structuring":
                _apply_inventory_structured_fields(
                    item,
                    structured_status="pending_structuring",
                    target_ref="cost_approx_analysis.cost_basis_attachment_inventory",
                    next_action="已收集原件，下一步加入配置后才能自动取价。",
                )
            else:
                _apply_inventory_structured_fields(
                    item,
                    structured_status="reference_only",
                    target_ref="cost_approx_analysis.cost_basis_attachment_inventory",
                    next_action="作为正文依据或核对说明，不直接改变费用计算。",
                )
    return result_entries


def _green_seedling_standard(land_subclass: str, county_name: str = "", base_dir: str | Path | None = None) -> str:
    matched = green_seedling_standard_for_county(base_dir or Path("."), county_name, land_subclass)
    return str(matched.get("standard_per_mu") or "")


def _green_seedling_standard_source(county_name: str = "", land_subclass: str = "", base_dir: str | Path | None = None) -> str:
    matched = green_seedling_standard_for_county(base_dir or Path("."), county_name, land_subclass)
    note = str(matched.get("source_note") or "")
    path = str(matched.get("source_path") or "")
    if path:
        return f"{note}（{Path(path).name}）"
    return note or "按征收地类细分类默认口径（待匹配县市本地政策）"


def _compensation_zone_options(row: Dict[str, str]) -> List[str]:
    options: List[str] = []
    if row.get("zone_i"):
        options.append("Ⅰ")
    if row.get("zone_ii"):
        options.append("Ⅱ")
    if row.get("zone_iii"):
        options.append("Ⅲ")
    return options or ["Ⅰ", "Ⅱ", "Ⅲ"]


def _default_location_factors(template: Dict[str, Any], scenarios: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    usage_keys = [str(item.get("key") or "") for item in scenarios]
    default_usage = usage_keys[0] if len(usage_keys) == 1 else ""
    return [
        {
            "key": factor.get("key") or "",
            "usage_key": default_usage,
            "group": factor.get("group") or "",
            "label": factor.get("label") or "",
            "description": "",
            "level": "",
            "levels": list(factor.get("levels") or []),
            "grade_amplitude": str(factor.get("grade_amplitude") or ""),
            "weight": "",
            "correction_rate": "0.00",
            "source": "configured_location_template",
            "enabled": True,
            "confirmed": False,
        }
        for factor in template.get("factors") or []
    ]


def _enabled_location_factors(analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
    return [item for item in analysis.get("location_factors") or [] if item.get("enabled") is not False]


def _grade_correction(level: str, amplitude: Any, levels: Iterable[str] | None = None) -> Decimal:
    configured_levels = [str(item) for item in levels or []]
    if level and level in configured_levels:
        center = Decimal(len(configured_levels) - 1) / Decimal("2")
        return _decimal(amplitude) * (center - Decimal(configured_levels.index(level)))
    multiplier = {"优": D1 * 2, "较优": D1, "一般": D0, "较劣": -D1, "劣": D1 * -2}.get(
        str(level or "").strip(),
        D0,
    )
    return _decimal(amplitude) * multiplier


def _first_existing_path(base_dir: str | Path, *relative_paths: Path) -> Path:
    base = Path(base_dir)
    for relative in relative_paths:
        path = base / relative
        if path.exists():
            return path
    return base / relative_paths[0]


def _policy_path(base_dir: str | Path) -> Path:
    return _first_existing_path(
        base_dir,
        COST_BASIS_ATTACHMENT_DIR / PROVINCE_COMPENSATION_DOC,
        Path("01_Source") / "03_attachment" / PROVINCE_COMPENSATION_DOC,
    )


def _scope_policy_path(base_dir: str | Path) -> Path:
    return _first_existing_path(
        base_dir,
        COST_BASIS_ATTACHMENT_DIR / PROVINCE_SCOPE_PDF,
        Path("01_Source")
        / "03_attachment"
        / "湖南省自然资源厅关于调整征收农用地补偿区片空间范围的通知"
        / PROVINCE_SCOPE_PDF,
        Path("01_Source") / "03_attachment" / PROVINCE_SCOPE_PDF,
    )


def _local_compensation_policy_path(base_dir: str | Path) -> Path:
    return _first_existing_path(
        base_dir,
        LOCAL_COMPENSATION_ATTACHMENT_DIR / YONGZHOU_LOCAL_COMPENSATION_PDF,
        Path("01_Source") / "03_attachment" / YONGZHOU_LOCAL_COMPENSATION_PDF,
        _collective_land_compensation_dir(base_dir) / YONGZHOU_LOCAL_COMPENSATION_PDF,
    )


def _local_compensation_policy_xlsx_path(base_dir: str | Path) -> Path:
    return _first_existing_path(
        base_dir,
        LOCAL_COMPENSATION_ATTACHMENT_DIR / YONGZHOU_LOCAL_COMPENSATION_XLSX,
        Path("01_Source") / "03_attachment" / YONGZHOU_LOCAL_COMPENSATION_XLSX,
        _collective_land_compensation_dir(base_dir) / YONGZHOU_LOCAL_COMPENSATION_XLSX,
    )


def _collective_land_compensation_dir(base_dir: str | Path) -> Path:
    direct = Path(base_dir) / COLLECTIVE_LAND_COMPENSATION_DIR
    if direct.exists():
        return direct
    attachment_root = Path(base_dir) / "01_Source" / "03_attachment"
    if not attachment_root.exists():
        return direct
    for item in attachment_root.iterdir():
        if item.is_dir() and "集体土地" in item.name and "补偿安置" in item.name:
            return item
    return direct


def _collective_land_attachment_inventory(base_dir: str | Path, county_name: str = "") -> List[Dict[str, Any]]:
    county = str(county_name or "").strip()
    folder = _collective_land_compensation_dir(base_dir)
    if not folder.exists():
        return []
    region_filters: List[str]
    if "通道" in county:
        region_filters = ["怀化"]
    elif "道县" in county:
        region_filters = ["永州", "道县"]
    else:
        region_filters = ["*"]
    entries: List[Dict[str, Any]] = []
    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in {".pdf", ".doc", ".docx", ".xlsx"}:
            continue
        rel = path.relative_to(folder).as_posix()
        region = "怀化" if rel.startswith("怀化/") or rel.startswith("怀化\\") else (
            "道县" if "道县" in path.name else "永州"
        )
        if region_filters != ["*"] and region not in region_filters and not (
            region == "永州" and "永州" in region_filters
        ):
            continue
        counties = ["通道县", "通道侗族自治县"] if region == "怀化" else (
            ["道县", "道县城区"] if region in {"永州", "道县"} else ["*"]
        )
        if county and counties != ["*"] and not _county_matches_aliases(county, counties):
            continue
        category = "青苗房屋及附着物"
        if "青苗" in path.name:
            category = "青苗补偿"
        elif any(token in path.name for token in ("房屋", "装修", "征收补偿")):
            category = "房屋及附着物补偿"
        elif "坟墓" in path.name:
            category = "坟墓补偿"
        entries.append(
            {
                "category": category,
                "label": path.name,
                "source_path": str(path),
                "region": region,
                "counties": counties,
                "counties_display": _dedup_counties_display(counties),
                "status": "scan_pending_structuring",
                "status_label": _attachment_status_label("scan_pending_structuring"),
            }
        )
    return entries


def _flatten_sheet_text(sheet: Any) -> str:
    parts: List[str] = []
    for row in sheet.iter_rows(values_only=True):
        for value in row:
            if value is None:
                continue
            text = str(value).strip()
            if text:
                parts.append(text)
    return "\n".join(parts)


def _parse_yongzhou_local_compensation_xlsx(path: Path) -> Dict[str, Any]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    sheet = workbook.active
    flat_text = _flatten_sheet_text(sheet)
    compact = re.sub(r"\s+", "", flat_text)
    attachment1: Dict[str, Any] = {}
    current_category = ""
    for row in sheet.iter_rows(min_row=150, max_row=175, values_only=True):
        values = list(row)
        if values and values[0]:
            text0 = str(values[0]).replace("\n", "")
            if text0 and not text0.startswith("—") and "附件" not in text0:
                current_category = text0
        grade = str(values[1] or "").replace("\n", "") if len(values) > 1 else ""
        if current_category == "砖混结构" and grade == "二等":
            house_rate = values[23] if len(values) > 23 else None
            decor_rate = values[26] if len(values) > 26 else None
            attachment1 = {
                "category": "砖混结构",
                "grade": "二等",
                "house_rate": str(int(house_rate)) if isinstance(house_rate, (int, float)) else "",
                "decoration_rate": str(int(decor_rate)) if isinstance(decor_rate, (int, float)) else "",
                "combined_rate": "",
            }
            if attachment1["house_rate"] and attachment1["decoration_rate"]:
                attachment1["combined_rate"] = str(int(attachment1["house_rate"]) + int(attachment1["decoration_rate"]))
            break
    moving_fee = ""
    transition_fee = ""
    vacating_awards: List[str] = []
    article31 = re.search(r"按(\d+)元/户支付搬家费", compact)
    if article31:
        moving_fee = article31.group(1)
    article31_transition = re.search(r"按(\d+)元/户/月一次性支付安置过渡费", compact)
    if article31_transition:
        transition_fee = article31_transition.group(1)
    vacating_awards = sorted(set(re.findall(r"给予(\d+)元/m[²2]", compact)))
    workbook.close()
    return {
        "moving_fee_per_household": moving_fee,
        "transition_fee_per_household_month": transition_fee,
        "vacating_awards_per_sqm": vacating_awards,
        "attachment1_brick_mixed_grade_ii": attachment1,
    }


def validate_yongzhou_local_compensation_xlsx(base_dir: str | Path, config_rows: Iterable[Dict[str, Any]] | None = None) -> Dict[str, Any]:
    xlsx_path = _local_compensation_policy_xlsx_path(base_dir)
    rows = list(config_rows or [])
    result: Dict[str, Any] = {
        "ok": False,
        "source_path": str(xlsx_path),
        "source_hash": _document_hash(xlsx_path) if xlsx_path.exists() else "",
        "checks": [],
        "warnings": [],
        "parsed": {},
    }
    if not xlsx_path.exists():
        result["warnings"].append("未找到永州市补偿办法 Excel，无法执行转换校验。")
        return result
    stat = xlsx_path.stat()
    config_hash = hashlib.sha256(
        json.dumps(rows, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
    ).hexdigest()
    cache_key = (str(xlsx_path.resolve()), stat.st_mtime_ns, stat.st_size, config_hash)
    cached = _YONGZHOU_XLSX_VALIDATION_CACHE.get(cache_key)
    if cached is not None:
        return deepcopy(cached)
    try:
        parsed = _parse_yongzhou_local_compensation_xlsx(xlsx_path)
    except Exception as exc:
        result["warnings"].append(f"Excel 解析失败：{exc}")
        return result
    result["parsed"] = parsed
    rows_by_key = {str(item.get("key") or ""): item for item in rows}
    checks: List[Dict[str, str]] = []

    def add_check(key: str, label: str, expected: str, actual: str, status: str, detail: str = "") -> None:
        checks.append(
            {
                "key": key,
                "label": label,
                "expected": expected,
                "actual": actual,
                "status": status,
                "detail": detail,
            }
        )

    moving_expected = parsed.get("moving_fee_per_household") or ""
    moving_actual = str((rows_by_key.get("moving_fee") or {}).get("standard") or "")
    add_check(
        "moving_fee",
        "搬家费（第三十一条）",
        moving_expected,
        moving_actual,
        "match" if moving_expected and moving_expected == moving_actual else "mismatch",
        "Excel 正文：按2000元/户支付搬家费",
    )
    transition_expected = parsed.get("transition_fee_per_household_month") or ""
    transition_actual = str((rows_by_key.get("transition_fee") or {}).get("standard") or "")
    add_check(
        "transition_fee",
        "过渡费（第三十一条）",
        transition_expected,
        transition_actual,
        "match" if transition_expected and transition_expected == transition_actual else "mismatch",
        "Excel 正文：1000元/户/月，不超过十二个月",
    )
    vacating_actual = str((rows_by_key.get("vacating_award") or {}).get("standard") or "")
    vacating_candidates = list(parsed.get("vacating_awards_per_sqm") or [])
    vacating_status = "match" if vacating_actual in vacating_candidates else "partial"
    add_check(
        "vacating_award",
        "主动腾地奖（第三十三条）",
        "、".join(vacating_candidates) or "待解析",
        vacating_actual,
        vacating_status,
        "政策含 80/100/100 元/㎡ 三档奖励，表3-1 通常取签约或腾地奖 100",
    )
    attachment1 = parsed.get("attachment1_brick_mixed_grade_ii") or {}
    house_actual = str((rows_by_key.get("house_compensation") or {}).get("standard") or "")
    combined = str(attachment1.get("combined_rate") or "")
    house_status = "match" if house_actual == combined else "mismatch"
    add_check(
        "house_compensation",
        "房屋补偿费（附件1 砖混二等）",
        combined or f"{attachment1.get('house_rate', '')}+{attachment1.get('decoration_rate', '')}",
        house_actual,
        house_status,
        "表3-1 项目口径可与附件1 合计值不同，需估价师确认建筑等级/成新度",
    )
    result["checks"] = checks
    hard_mismatch = [item for item in checks if item["status"] == "mismatch" and item["key"] in {"moving_fee", "transition_fee"}]
    soft_mismatch = [item for item in checks if item["status"] in {"mismatch", "partial"} and item["key"] in {"house_compensation", "vacating_award"}]
    result["ok"] = not hard_mismatch
    if hard_mismatch:
        result["warnings"].extend(f"{item['label']}转换不一致：Excel={item['expected']}，配置={item['actual']}" for item in hard_mismatch)
    if soft_mismatch:
        result["warnings"].extend(f"{item['label']}需人工确认：Excel={item['expected']}，配置={item['actual']}" for item in soft_mismatch)
    _YONGZHOU_XLSX_VALIDATION_CACHE[cache_key] = deepcopy(result)
    return result


def _load_yongzhou_building_compensation_config(base_dir: str | Path) -> Dict[str, Any]:
    return _load_json_config_cached(Path(base_dir) / YONGZHOU_BUILDING_COMPENSATION_CONFIG)


def _load_huaihua_building_compensation_config(base_dir: str | Path) -> Dict[str, Any]:
    return _load_json_config_cached(Path(base_dir) / HUAIHUA_BUILDING_COMPENSATION_CONFIG)


def _load_building_compensation_config(base_dir: str | Path, county_name: str = "") -> Dict[str, Any]:
    county = str(county_name or "").strip()
    if "通道" in county:
        return _load_huaihua_building_compensation_config(base_dir)
    if "道县" in county:
        return _load_yongzhou_building_compensation_config(base_dir)
    return {}


def building_compensation_grade_options(
    base_dir: str | Path,
    county_name: str = "",
    row_key: str = "",
) -> List[Dict[str, Any]]:
    county = str(county_name or "").strip()
    config = _load_building_compensation_config(base_dir, county)
    aliases = [str(item) for item in config.get("county_aliases") or []]
    if not config or not _county_matches_aliases(county, aliases):
        return []
    catalog = config.get("grade_catalog") or {}
    if row_key:
        return deepcopy(catalog.get(row_key) or [])
    return deepcopy(catalog)


def building_compensation_add_catalog(base_dir: str | Path, county_name: str = "") -> List[Dict[str, Any]]:
    county = str(county_name or "").strip()
    config = _load_building_compensation_config(base_dir, county)
    aliases = [str(item) for item in config.get("county_aliases") or []]
    if not config or not _county_matches_aliases(county, aliases):
        return []
    # catalog_templates carries template metadata (label/unit/计算基数) for items
    # offered in 「从政策目录添加」 WITHOUT putting them into the default form rows;
    # this keeps the default building_compensation_rows = 范例 while still letting
    # the add-catalog expose the full structured policy with correct units.
    templates_by_key = {
        str(key): {"key": str(key), **(value or {})}
        for key, value in (config.get("catalog_templates") or {}).items()
    }
    rows_by_key = {**templates_by_key, **{str(row.get("key") or ""): row for row in config.get("rows") or []}}
    catalog_keys = set((config.get("grade_catalog") or {}).keys())
    catalog: List[Dict[str, Any]] = []
    for row_key, options in (config.get("grade_catalog") or {}).items():
        default_option = next((opt for opt in (options or []) if opt.get("is_default")), (options or [{}])[0] if options else {})
        template = deepcopy(rows_by_key.get(row_key) or {})
        if not template:
            template = {
                "key": row_key,
                "label": row_key,
                "calculation_basis": "按合法建筑面积",
                "quantity": "",
                "divisor": "1",
            }
        # Backfill standard / unit from the default grade option so catalog_templates
        # (which intentionally omit a sample standard to keep default rows clean) still
        # expose the policy default value in 「从政策目录添加」.
        if template.get("standard") in (None, ""):
            template["standard"] = str(default_option.get("standard") or "")
        if not template.get("standard_unit"):
            template["standard_unit"] = str(default_option.get("standard_unit") or "元/平方米")
        catalog.append(
            {
                "row_key": row_key,
                "label": str(template.get("label") or row_key),
                "grade_options": deepcopy(options),
                "template": template,
            }
        )
    for row in config.get("rows") or []:
        row_key = str(row.get("key") or "")
        if row_key in catalog_keys:
            continue
        catalog.append(
            {
                "row_key": row_key,
                "label": str(row.get("label") or row_key),
                "grade_options": [],
                "template": deepcopy(row),
            }
        )
    signing_options = [
        option
        for options in (config.get("grade_catalog") or {}).values()
        for option in (options or [])
        if str(option.get("key") or "").startswith("signing_")
    ]
    if signing_options and "signing_award" not in {entry["row_key"] for entry in catalog}:
        parent = deepcopy(rows_by_key.get("vacating_award") or rows_by_key.get("house_compensation") or {})
        default_option = signing_options[0]
        catalog.append(
            {
                "row_key": "signing_award",
                "label": "签约奖",
                "grade_options": deepcopy(signing_options),
                "template": {
                    **parent,
                    "key": "signing_award",
                    "label": "签约奖",
                    "standard": str(default_option.get("standard") or "80"),
                    "standard_unit": "元/平方米",
                    "calculation_basis": str(parent.get("calculation_basis") or "按合法建筑面积"),
                    "note": str(default_option.get("note") or "第三十三条：签约奖"),
                },
            }
        )
    return catalog


def building_compensation_policy_help(base_dir: str | Path, county_name: str = "") -> Dict[str, Any]:
    county = str(county_name or "").strip()
    config = _load_building_compensation_config(base_dir, county)
    aliases = [str(item) for item in config.get("county_aliases") or []]
    if not config or not _county_matches_aliases(county, aliases):
        return {}
    templates = config.get("catalog_templates") or {}
    row_labels = {str(key): str((value or {}).get("label") or "") for key, value in templates.items()}
    row_units = {str(key): str((value or {}).get("standard_unit") or "") for key, value in templates.items()}
    for row in config.get("rows") or []:
        row_key = str(row.get("key") or "")
        if row.get("label"):
            row_labels[row_key] = str(row.get("label"))
        if row.get("standard_unit"):
            row_units[row_key] = str(row.get("standard_unit"))
    entries: List[Dict[str, Any]] = []
    for row_key, options in (config.get("grade_catalog") or {}).items():
        for option in options:
            entries.append(
                {
                    "row_key": row_key,
                    "row_label": row_labels.get(row_key, row_key),
                    "option_key": str(option.get("key") or ""),
                    "label": str(option.get("label") or ""),
                    "standard": str(option.get("standard") or ""),
                    "standard_unit": str(option.get("standard_unit") or row_units.get(row_key) or "元/平方米"),
                    "note": str(option.get("note") or ""),
                    "house_rate": str(option.get("house_rate") or ""),
                    "decoration_rate": str(option.get("decoration_rate") or ""),
                }
            )
    catalog_keys = set((config.get("grade_catalog") or {}).keys())
    for row in config.get("rows") or []:
        row_key = str(row.get("key") or "")
        if row_key in catalog_keys:
            continue
        entries.append(
            {
                "row_key": row_key,
                "row_label": str(row.get("label") or row_key),
                "option_key": "",
                "label": str(row.get("label") or row_key),
                "standard": str(row.get("standard") or ""),
                "standard_unit": str(row.get("standard_unit") or ""),
                "note": str(row.get("note") or ""),
            }
        )
    return {
        "policy_name": str(config.get("policy_name") or ""),
        "policy_document_no": str(config.get("policy_document_no") or ""),
        "source_file": str(config.get("policy_source_file") or ""),
        "entries": entries,
        "paid_land_use_fee_standards": [
            {"label": grade, "standard": value, "standard_unit": "元/平方米", "row_label": "新增建设用地有偿使用费", "note": "湘财综〔2018〕42号附件1"}
            for grade, value in HUNAN_PAID_LAND_USE_FEE_STANDARDS.items()
        ],
    }


def building_compensation_rows_for_county(base_dir: str | Path, county_name: str = "") -> List[Dict[str, Any]]:
    county = str(county_name or "").strip()
    config = _load_building_compensation_config(base_dir, county)
    aliases = [str(item) for item in config.get("county_aliases") or []]
    if not config or not _county_matches_aliases(county, aliases):
        return _default_building_compensation_rows()
    collective_dir = _collective_land_compensation_dir(base_dir)
    xlsx_path = _local_compensation_policy_xlsx_path(base_dir)
    pdf_path = _local_compensation_policy_path(base_dir)
    # File-aware memoization: building the full per-county row set restats and
    # deepcopies the whole policy catalog, and it is called repeatedly within one
    # /calculate request (main calc + each pricing-assistant scenario). Cache the
    # built rows keyed by county + source-file signatures; callers mutate rows so
    # always hand back a deepcopy.
    cache_key = (
        os.path.normcase(os.path.abspath(Path(base_dir))),
        _normalize_location_key(county),
        _path_sig(Path(base_dir) / (YONGZHOU_BUILDING_COMPENSATION_CONFIG if "通道" not in county else HUAIHUA_BUILDING_COMPENSATION_CONFIG)),
        _path_sig(xlsx_path),
        _path_sig(pdf_path),
    )
    cached_rows = _BUILDING_ROWS_CACHE.get(cache_key)
    if cached_rows is not None:
        return deepcopy(cached_rows)
    validation = validate_yongzhou_local_compensation_xlsx(base_dir, config.get("rows") or []) if "通道" not in county else {"checks": [], "warnings": [], "ok": False}
    policy_pdf = Path(base_dir) / str(config.get("policy_source_file") or "")
    if not policy_pdf.exists() and config.get("policy_source_file"):
        policy_pdf = collective_dir / Path(str(config.get("policy_source_file") or "")).name
    if xlsx_path.exists() and "通道" not in county:
        source_path = str(xlsx_path)
        source_hash = validation.get("source_hash") or _document_hash(xlsx_path)
        review_status = "verified_partial" if validation.get("ok") else "conversion_warning"
        review_note = "；".join(validation.get("warnings") or []) or str(config.get("review_note") or "")
    elif policy_pdf.exists():
        source_path = str(policy_pdf)
        source_hash = _document_hash(policy_pdf)
        review_status = str(config.get("review_status") or "scan_pending_structuring")
        review_note = str(config.get("review_note") or "政策扫描件待结构化")
    else:
        source_path = str(pdf_path) if pdf_path.exists() else ""
        source_hash = _document_hash(pdf_path) if pdf_path.exists() else ""
        review_status = str(config.get("review_status") or "pending_ocr")
        review_note = str(config.get("review_note") or "扫描政策原件待 OCR 复核")
    rows: List[Dict[str, Any]] = []
    checks_by_key = {item["key"]: item for item in validation.get("checks") or []}
    for item in config.get("rows") or []:
        row = deepcopy(item)
        row.setdefault("confirmed", False)
        row_source = source_path
        source_doc = str(row.get("source_doc") or "")
        if source_doc and collective_dir.exists():
            doc_path = collective_dir / source_doc.replace("\\", "/")
            if doc_path.exists():
                row_source = str(doc_path)
        row["source_path"] = row_source
        row["source_hash"] = _document_hash(Path(row_source)) if row_source and Path(row_source).exists() else source_hash
        row["review_status"] = review_status
        row["pending_review"] = review_status not in {"verified", "verified_partial"} or checks_by_key.get(row.get("key"), {}).get("status") in {"mismatch", "partial"}
        row["source_note"] = f"{config.get('policy_name') or DAO_LOCAL_POLICY['name']}（{config.get('policy_document_no') or DAO_LOCAL_POLICY['document_no']}）"
        conversion = checks_by_key.get(str(row.get("key") or ""))
        if conversion:
            row["conversion_check"] = deepcopy(conversion)
            if conversion.get("status") == "match":
                row["source_note"] = f"{row['source_note']}；Excel 校验一致（{conversion.get('label')}）"
            elif conversion.get("status") == "partial":
                row["source_note"] = f"{row['source_note']}；Excel 部分一致，{conversion.get('detail') or '需人工确认'}"
            else:
                row["source_note"] = f"{row['source_note']}；Excel 与配置不一致，{conversion.get('detail') or '需人工确认'}"
        elif review_note:
            row["source_note"] = f"{row['source_note']}；{review_note}"
        grade_options = building_compensation_grade_options(base_dir, county, str(row.get("key") or ""))
        if grade_options:
            row["grade_options"] = grade_options
        rows.append(row)
    if rows:
        _BUILDING_ROWS_CACHE[cache_key] = deepcopy(rows)
        return rows
    return _default_building_compensation_rows()


def _scope_table_path(base_dir: str | Path, city_name: str = "") -> Path:
    city = str(city_name or "").strip()
    if not city:
        return Path(base_dir) / "__missing_cost_scope_table__"
    filename = f"{city}农用地补偿区片划分表.docx"
    candidates = [
        COST_BASIS_ATTACHMENT_DIR / filename,
        Path("01_Source") / "03_attachment" / "农用地补偿区片划分表" / filename,
        Path("01_Source")
        / "03_attachment"
        / "湖南省自然资源厅关于调整征收农用地补偿区片空间范围的通知"
        / filename,
    ]
    if city in {"永州市", "永州"}:
        candidates.insert(0, COST_BASIS_ATTACHMENT_DIR / YONGZHOU_ZONE_TABLE)
    return _first_existing_path(base_dir, *candidates)


def _document_hash(path: Path) -> str:
    try:
        stat = path.stat()
    except OSError:
        return ""
    # Use abspath (cheap) instead of resolve()/getfinalpathname (a costly Windows
    # syscall) for the cache key so repeated hashing during one /calculate request
    # stays in-memory.
    cache_key = (os.path.normcase(os.path.abspath(path)), stat.st_mtime_ns, stat.st_size)
    cached = _DOCUMENT_HASH_CACHE.get(cache_key)
    if cached is not None:
        return cached
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    value = digest.hexdigest()
    _DOCUMENT_HASH_CACHE[cache_key] = value
    return value


def _county_base(name: Any) -> str:
    """Strip only the trailing county-level admin character so that
    '通道县' -> '通道', '道县' -> '道', '通道侗族自治县' -> '通道侗族自治'."""
    text = _normalize_location_key(name)
    return re.sub(r"[县市区州旗]$", "", text)


def _county_matches(input_county: str, table_county: str) -> bool:
    source = _normalize_location_key(input_county)
    target = _normalize_location_key(table_county)
    if not source or not target:
        return False
    if source == target:
        return True
    # A complete county-level token may lead a longer address. Check this
    # before comparing stripped base names so that "道县西洲街道" matches
    # "道县" without weakening the boundary between 市、县、区.
    if source.startswith(target) and re.search(r"[县市区州旗]$", target):
        return True
    if target.startswith(source) and re.search(r"[县市区州旗]$", source):
        return True
    source_level = source[-1] if re.search(r"[县市区州旗]$", source) else ""
    target_level = target[-1] if re.search(r"[县市区州旗]$", target) else ""
    if source_level and target_level and source_level != target_level:
        return False
    source_base = _county_base(input_county)
    target_base = _county_base(table_county)
    if not source_base or not target_base:
        return False
    if source_base == target_base:
        return True
    # Allow an autonomous-county long form to match its short form
    # (e.g. '通道' ↔ '通道侗族自治'), but never let a plain substring
    # like '道' inside '通道' produce a false cross-county match.
    longer, shorter = (
        (source_base, target_base)
        if len(source_base) >= len(target_base)
        else (target_base, source_base)
    )
    if longer.startswith(shorter) and re.search(r"族|自治", longer[len(shorter):]):
        return True
    # A county name appended after a province/city/prefecture boundary is valid
    # (e.g. '永州市道县' ↔ '道县'); the boundary char disambiguates it from the
    # '通道县' trap where '道县' follows an ordinary character.
    for container, county_token in ((source, target), (target, source)):
        if (
            re.search(r"[县市区州旗]$", county_token)
            and container != county_token
            and container.endswith(county_token)
            and len(container) > len(county_token)
            and container[len(container) - len(county_token) - 1] in "省市州盟"
        ):
            return True
    return False


_ATTACHMENT_STATUS_LABELS = {
    "available": "可用",
    "missing": "缺失",
    "scan_pending_structuring": "已扫描·待结构化",
    "verified": "已核验",
    "verified_partial": "部分核验",
    "pending": "待匹配本地政策",
    "structured": "已结构化",
    "needs_grade_selection": "需选择等别",
    "manual_input": "需人工录入",
    "reference_only": "说明依据",
    "pending_structuring": "待结构化",
}


def _attachment_status_label(status: Any) -> str:
    text = str(status or "").strip()
    return _ATTACHMENT_STATUS_LABELS.get(text, text or "可用")


def _inventory_key(source_path: Any, label: Any) -> str:
    raw = str(source_path or label or "").strip()
    if not raw:
        return ""
    return hashlib.md5(raw.encode("utf-8", errors="ignore")).hexdigest()[:12]


def _inventory_structured_label(status: Any) -> str:
    text = str(status or "").strip()
    return _ATTACHMENT_STATUS_LABELS.get(text, text or "待结构化")


def _apply_inventory_structured_fields(
    entry: Dict[str, Any],
    *,
    structured_status: str,
    structured_item_count: int = 0,
    price_fields: Iterable[str] | None = None,
    target_ref: str = "",
    next_action: str = "",
) -> Dict[str, Any]:
    entry["structured_status"] = structured_status
    entry["structured_status_label"] = _inventory_structured_label(structured_status)
    entry["structured_item_count"] = int(structured_item_count or 0)
    entry["price_fields"] = [str(item) for item in (price_fields or []) if str(item or "").strip()]
    entry["target_ref"] = target_ref or (entry["price_fields"][0] if entry["price_fields"] else "")
    entry["next_action"] = next_action
    entry["target_key"] = _inventory_key(entry.get("source_path"), entry.get("label"))
    return entry


def _dedup_counties_display(counties: Iterable[Any]) -> List[str]:
    """Collapse semantically duplicate county aliases for UI display
    (e.g. '通道县' and '通道侗族自治县' become a single '通道县')."""
    cleaned: List[str] = []
    for item in counties or []:
        name = str(item or "").strip()
        if not name or name == "*":
            continue
        if name not in cleaned:
            cleaned.append(name)
    groups: List[List[str]] = []
    for name in cleaned:
        placed = False
        for group in groups:
            if any(_county_matches(name, member) for member in group):
                group.append(name)
                placed = True
                break
        if not placed:
            groups.append([name])
    display: List[str] = []
    for group in groups:
        # Prefer the shortest representative as the canonical short form.
        display.append(min(group, key=len))
    return display


def _province_compensation_row(base_dir: str | Path, county_name: str) -> Dict[str, str]:
    path = _policy_path(base_dir)
    if not path.exists() or not county_name:
        return {}
    stat = path.stat()
    cache_key = (str(path.resolve()), stat.st_mtime_ns, stat.st_size, _normalize_location_key(county_name))
    cached = _PROVINCE_COMPENSATION_ROW_CACHE.get(cache_key)
    if cached is not None:
        return deepcopy(cached)
    document = Document(path)
    if not document.tables:
        return {}
    for row in document.tables[0].rows[2:]:
        values = [cell.text.strip().replace("\n", "") for cell in row.cells]
        if len(values) >= 8 and _county_matches(county_name, values[1]):
            result = {
                "city": values[0],
                "county": values[1],
                "zone_i": values[2],
                "zone_ii": values[3],
                "zone_iii": values[4],
                "dry_coefficient": values[5],
                "garden_coefficient": values[6],
                "forest_coefficient": values[7],
                "source_path": str(path),
                "source_hash": _document_hash(path),
            }
            _PROVINCE_COMPENSATION_ROW_CACHE[cache_key] = deepcopy(result)
            return result
    _PROVINCE_COMPENSATION_ROW_CACHE[cache_key] = {}
    return {}


_SCOPE_TABLE_CACHE: Dict[str, List[Dict[str, Any]]] = {}
_ZONE_ROMAN = {"I": "Ⅰ", "II": "Ⅱ", "III": "Ⅲ", "ⅰ": "Ⅰ", "ⅱ": "Ⅱ", "ⅲ": "Ⅲ"}
_ADMIN_SUFFIXES = (
    "街道办事处",
    "办事处",
    "街道",
    "镇",
    "瑶族乡",
    "乡",
    "社区居委会",
    "村委会",
    "居委会",
    "社区",
    "村",
)


def _normalize_zone_label(value: Any) -> str:
    text = str(value or "").strip().replace("区", "")
    return _ZONE_ROMAN.get(text, text)


def _normalize_location_key(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or ""))


def _admin_base_name(value: Any) -> str:
    text = _normalize_location_key(value)
    if not text:
        return ""
    changed = True
    while changed:
        changed = False
        for suffix in _ADMIN_SUFFIXES:
            if text.endswith(suffix) and len(text) > len(suffix):
                text = text[: -len(suffix)]
                changed = True
                break
    return text


def _split_scope_items(scope_text: str) -> List[str]:
    text = str(scope_text or "").strip().rstrip("、，,;；")
    if not text:
        return []
    return [item.strip() for item in re.split(r"[、,，;；]", text) if item.strip()]


def _parse_scope_clause(scope_text: str) -> Dict[str, Any]:
    text = str(scope_text or "").strip()
    if not text:
        return {"kind": "empty", "township": "", "places": [], "items": []}
    for separator in ("：", ":"):
        if separator in text:
            township, places_text = text.split(separator, 1)
            return {
                "kind": "township_places",
                "township": township.strip(),
                "places": _split_scope_items(places_text),
                "items": [],
            }
    items = _split_scope_items(text)
    if len(items) == 1:
        return {"kind": "single", "township": "", "places": [], "items": items}
    return {"kind": "list", "township": "", "places": [], "items": items}


def _load_scope_table_entries(path: Path) -> List[Dict[str, Any]]:
    cache_key = str(path.resolve())
    cached = _SCOPE_TABLE_CACHE.get(cache_key)
    if cached is not None:
        return cached
    if not path.exists():
        _SCOPE_TABLE_CACHE[cache_key] = []
        return []
    document = Document(path)
    if not document.tables:
        _SCOPE_TABLE_CACHE[cache_key] = []
        return []
    entries: List[Dict[str, Any]] = []
    for row_index, row in enumerate(document.tables[0].rows[1:], start=1):
        values = [cell.text.strip().replace("\n", "") for cell in row.cells]
        if len(values) < 3:
            continue
        county, zone, scope_text = values[0], _normalize_zone_label(values[1]), values[2]
        if not county or county == "县区名" or not zone:
            continue
        entries.append(
            {
                "row_index": row_index,
                "county": county,
                "zone": zone,
                "scope_text": scope_text,
                "clause": _parse_scope_clause(scope_text),
            }
        )
    _SCOPE_TABLE_CACHE[cache_key] = entries
    return entries


def _location_contains_name(location: str, name: str) -> bool:
    base = _admin_base_name(name)
    if not base:
        return False
    return base in _normalize_location_key(location)


def _scope_entry_match_score(location: str, entry: Dict[str, Any]) -> tuple[int, str]:
    clause = entry.get("clause") or {}
    scope_text = str(entry.get("scope_text") or "")
    kind = clause.get("kind")
    if kind == "township_places":
        township = str(clause.get("township") or "")
        if not _location_contains_name(location, township):
            return 0, ""
        places = list(clause.get("places") or [])
        if not places:
            return 60, f"{entry['county']}{township}"
        for place in places:
            if _location_contains_name(location, place):
                return 100, f"{entry['county']}{township}{place}"
        return 55, f"{entry['county']}{township}"
    items = list(clause.get("items") or [])
    if not items:
        return 0, ""
    matched = [item for item in items if _location_contains_name(location, item)]
    if not matched:
        return 0, ""
    if kind == "single":
        return 80, f"{entry['county']}{matched[0]}"
    if any(item.endswith(("镇", "乡")) or "瑶族乡" in item for item in matched):
        return 70, f"{entry['county']}{matched[0]}"
    return 65, f"{entry['county']}{'、'.join(matched)}"


def suggest_compensation_zone(
    location: str,
    base_dir: str | Path,
    *,
    county_name: str = "",
    city_name: str = "",
) -> Dict[str, Any]:
    location_text = str(location or "").strip()
    county = str(county_name or "").strip()
    city = str(city_name or "").strip()
    path = _scope_table_path(base_dir, city)
    source_hash = _document_hash(path)
    empty = {
        "compensation_zone": "",
        "matched": False,
        "confidence": "none",
        "match_detail": "",
        "source_path": str(path),
        "source_hash": source_hash,
        "scope_row_index": None,
        "scope_text": "",
        "warnings": [],
    }
    if not city:
        empty["warnings"].append("未识别所属地市，不得自动读取其他地市区片划分表，请人工确认征地区片。")
        return empty
    if not location_text:
        empty["warnings"].append("未提供位置全称，无法根据区片划分表建议征地区片。")
        return empty
    if not path.exists():
        empty["warnings"].append("未找到市级农用地补偿区片划分表，请人工确认征地区片。")
        return empty
    entries = _load_scope_table_entries(path)
    if not entries:
        empty["warnings"].append("区片划分表为空或无法解析，请人工确认征地区片。")
        return empty
    county_candidates = {county} if county else set()
    county_candidates.update(re.findall(r"[\u4e00-\u9fff]{2,6}(?:县|区|市)", location_text))
    if not county_candidates:
        empty["warnings"].append("未识别所属县市，不得跨县搜索区片划分表，请人工确认征地区片。")
        return empty
    filtered = [
        entry
        for entry in entries
        if any(_county_matches(candidate, entry["county"]) for candidate in county_candidates if candidate)
    ]
    if not filtered:
        empty["warnings"].append("区片划分表中未找到当前县市，不得回退搜索其他县市，请人工确认征地区片。")
        return empty
    best_score = 0
    best_entry: Dict[str, Any] | None = None
    best_detail = ""
    for entry in filtered:
        score, detail = _scope_entry_match_score(location_text, entry)
        if score > best_score:
            best_score = score
            best_entry = entry
            best_detail = detail
    if not best_entry or best_score <= 0:
        empty["warnings"].append("未在区片划分表中精确匹配到街道/乡镇/村社区，请估价师人工确认征地区片。")
        return empty
    confidence = {100: "village", 80: "township", 70: "town", 65: "street", 60: "township", 55: "township_partial"}.get(
        best_score,
        "approximate",
    )
    return {
        "compensation_zone": best_entry["zone"],
        "matched": True,
        "confidence": confidence,
        "match_detail": best_detail or best_entry["scope_text"],
        "source_path": str(path),
        "source_hash": source_hash,
        "scope_row_index": best_entry["row_index"],
        "scope_text": best_entry["scope_text"],
        "warnings": [],
    }


def _resolve_compensation_zone(
    data: Dict[str, Any],
    analysis: Dict[str, Any],
    suggestion: Dict[str, Any],
) -> tuple[str, bool]:
    current_zone = _normalize_zone_label(analysis.get("compensation_zone") or data.get("compensation_zone") or "")
    suggested_zone = _normalize_zone_label(suggestion.get("compensation_zone") or "")
    curr_location = _normalize_location_key(data.get("land_location_full") or data.get("land_location") or "")
    prev_location = _normalize_location_key(analysis.get("basis_land_location_full") or "")
    location_changed = bool(curr_location) and curr_location != prev_location
    override = _bool(analysis.get("compensation_zone_override"))
    if location_changed:
        override = False
    if suggested_zone and not override:
        if location_changed or not current_zone or current_zone == suggested_zone:
            return suggested_zone, False
        return current_zone or suggested_zone, True
    if override:
        return current_zone or suggested_zone or "Ⅰ", True
    return current_zone or suggested_zone or "Ⅰ", bool(current_zone and suggested_zone and current_zone != suggested_zone)


def applicable_cost_basis(
    data: Dict[str, Any],
    base_dir: str | Path,
    *,
    include_catalog_metadata: bool = True,
) -> Dict[str, Any]:
    county = str(data.get("county_name") or "").strip()
    analysis = data.get("cost_approx_analysis") or {}
    province_row = _province_compensation_row(base_dir, county)
    city_name = _expected_city(county, province_row)
    suggestion = suggest_compensation_zone(
        data.get("land_location_full") or data.get("land_location") or "",
        base_dir,
        county_name=county,
        city_name=city_name,
    )
    zone, zone_override = _resolve_compensation_zone(data, analysis, suggestion)
    land_class, land_subclass = _acquisition_land_classes(data)
    local_policy = _local_policy_defaults(county)
    row = province_row
    valuation_date = _date(data.get("valuation_date"))
    policy_applicable = valuation_date is None or valuation_date >= PROVINCE_POLICY_EFFECTIVE_DATE
    zone_key = {"I": "zone_i", "Ⅰ": "zone_i", "II": "zone_ii", "Ⅱ": "zone_ii", "III": "zone_iii", "Ⅲ": "zone_iii"}.get(zone, "zone_i")
    base_per_mu = row.get(zone_key, "") if policy_applicable else ""
    coefficient = _land_class_compensation_coefficient(land_class, land_subclass, row) if row else "1"
    if not policy_applicable:
        recommended_amount = "0.00"
    elif base_per_mu:
        amount = _decimal(base_per_mu) * _decimal(coefficient, D1) * MU_TO_SQM
        recommended_amount = _money(amount)
    else:
        recommended_amount = ""
    warnings: List[str] = list(suggestion.get("warnings") or [])
    if not row:
        warnings.append("未在湖南省征地补偿标准表中匹配到所属县市，请人工补充区片标准。")
    if not base_per_mu:
        warnings.append(f"当前县市未配置征地区片 {zone} 的补偿标准。")
    if not policy_applicable:
        warnings.append(
            f"估价期日早于{PROVINCE_POLICY_NO}生效日2024年1月31日，当前标准不得直接采用，请导入并确认估价期日有效的历史政策。"
        )
    result = {
        "county_name": county,
        "compensation_zone": zone,
        "compensation_zone_override": zone_override,
        "compensation_zone_suggestion": suggestion,
        "acquisition_land_class": land_class,
        "acquisition_land_subclass": land_subclass,
        "acquisition_land_class_confirmed": _bool(data.get("acquisition_land_class_confirmed", True)),
        "local_compensation_policy_name": str(data.get("local_compensation_policy_name") or local_policy["name"]),
        "local_compensation_policy_no": str(data.get("local_compensation_policy_no") or local_policy["document_no"]),
        "local_compensation_policy_date": str(data.get("local_compensation_policy_date") or local_policy["date"]),
        "green_seedling_standard_per_mu": _green_seedling_standard(land_subclass, county, base_dir),
        "green_seedling_standard_source": _green_seedling_standard_source(county, land_subclass, base_dir),
        "green_seedling_standard_review_status": green_seedling_standard_for_county(base_dir, county, land_subclass).get(
            "review_status", ""
        ),
        "compensation_zone_options": _compensation_zone_options(row),
        "paid_land_use_fee_grade_standards": deepcopy(HUNAN_PAID_LAND_USE_FEE_STANDARDS),
        "province_compensation": {
            **row,
            "policy_name": PROVINCE_POLICY_NAME,
            "policy_no": PROVINCE_POLICY_NO,
            "effective_date": PROVINCE_POLICY_EFFECTIVE_DATE.isoformat(),
            "applicable": policy_applicable,
            "base_per_mu": base_per_mu,
            "coefficient": coefficient,
            "recommended_amount_per_sqm": recommended_amount,
        },
        "warnings": warnings,
    }
    if include_catalog_metadata:
        result.update(
            {
                "cost_basis_attachment_inventory": cost_basis_attachment_inventory(base_dir, county),
                "building_compensation_add_catalog": building_compensation_add_catalog(base_dir, county),
                "building_compensation_policy_help": building_compensation_policy_help(base_dir, county),
            }
        )
    return result


def _item(
    key: str,
    label: str,
    category: str,
    amount: str,
    *,
    standard_value: str = "",
    standard_unit: str = "元/平方米",
    coefficient: str = "",
    formula: str = "",
    source: str = "policy_recommendation",
    source_note: str = "",
    policy_key: str = "",
    rule_key: str = "",
    rule_snapshot: Dict[str, Any] | None = None,
    grade_name: str = "",
    exclusion_reason: str = "",
    confirmed: bool = False,
    enabled: bool = True,
) -> Dict[str, Any]:
    return {
        "key": key,
        "label": label,
        "category": category,
        "standard_value": standard_value,
        "standard_unit": standard_unit,
        "quantity": None,
        "coefficient": coefficient,
        "amount_per_sqm": amount,
        "formula": formula,
        "source": source,
        "source_note": source_note,
        "policy_key": policy_key,
        "rule_key": rule_key,
        "rule_snapshot": deepcopy(rule_snapshot or {}),
        "grade_name": grade_name,
        "computed_amount_per_sqm": amount,
        "exclusion_reason": exclusion_reason,
        "confirmed": confirmed,
        "enabled": enabled,
    }


def _fill_missing_item_fields(item: Dict[str, Any], defaults: Dict[str, Any]) -> None:
    for key in (
        "label",
        "category",
        "standard_value",
        "standard_unit",
        "coefficient",
        "formula",
        "source",
        "source_note",
        "policy_key",
        "rule_key",
        "rule_snapshot",
        "grade_name",
        "computed_amount_per_sqm",
        "exclusion_reason",
    ):
        if item.get(key) in (None, "") and defaults.get(key) not in (None, ""):
            item[key] = defaults[key]
    if item.get("enabled") is None and defaults.get("enabled") is not None:
        item["enabled"] = defaults["enabled"]
    if item.get("amount_per_sqm") in (None, "") and defaults.get("amount_per_sqm") not in (None, ""):
        item["amount_per_sqm"] = defaults["amount_per_sqm"]


def _refresh_item_from_defaults(item: Dict[str, Any], defaults: Dict[str, Any], *, force: bool = False) -> bool:
    if not defaults:
        return False
    if item.get("source") in {"external_result", "manual_policy_replacement"}:
        return False
    changed = False
    for key in (
        "label",
        "category",
        "standard_value",
        "standard_unit",
        "coefficient",
        "amount_per_sqm",
        "formula",
        "source",
        "source_note",
        "policy_key",
        "rule_key",
        "rule_snapshot",
        "grade_name",
        "computed_amount_per_sqm",
        "enabled",
    ):
        if key not in defaults or defaults.get(key) is None:
            continue
        if not force and item.get(key) not in (None, ""):
            continue
        if item.get(key) != defaults.get(key):
            item[key] = deepcopy(defaults.get(key))
            changed = True
    if changed and force:
        item["confirmed"] = False
    return changed


def _same_text(left: Any, right: Any) -> bool:
    return str(left or "").strip() == str(right or "").strip()


def _formula_amount_for_item(item: Dict[str, Any]) -> str:
    key = item.get("key")
    standard = _decimal(item.get("standard_value"))
    coefficient = _decimal(item.get("coefficient"), D1)
    if key == "water_conservancy_fund" and item.get("standard_value") not in (None, "") and item.get("coefficient") not in (None, ""):
        return _money(standard * coefficient / Decimal("100"))
    if key == "land_compensation" and item.get("standard_value") not in (None, ""):
        return _money(standard * coefficient * MU_TO_SQM)
    if key == "farmland_occupation_tax" and item.get("standard_value") not in (None, ""):
        return _money(standard * coefficient) if item.get("coefficient") not in (None, "") else _money(standard)
    if key == "farmland_reclamation_fee" and item.get("standard_value") not in (None, ""):
        return _money(standard * Decimal("10000") * MU_TO_SQM)
    return ""


def _recalculate_formula_amount(item: Dict[str, Any], defaults: Dict[str, Any]) -> None:
    amount = _formula_amount_for_item(item)
    if not amount:
        return
    item["computed_amount_per_sqm"] = amount
    fields_match_defaults = (
        _same_text(item.get("standard_value"), defaults.get("standard_value"))
        and _same_text(item.get("coefficient"), defaults.get("coefficient"))
    )
    should_recalculate = not _bool(item.get("confirmed")) or not fields_match_defaults
    if should_recalculate:
        item["amount_per_sqm"] = amount


def _land_class_tax_coefficient(land_class: str, land_subclass: str) -> Decimal:
    detail = f"{land_class}{land_subclass}"
    if any(item in detail for item in ("园地", "林地", "草地", "乔木林地", "灌木林地")):
        return Decimal("0.8")
    return D1


def _land_class_compensation_coefficient(land_class: str, land_subclass: str, row: Dict[str, str]) -> str:
    detail = f"{land_class}{land_subclass}"
    if land_class == "林地" or any(token in detail for token in ("林", "竹", "灌")):
        return str(row.get("forest_coefficient") or "1").strip() or "1"
    if land_class == "园地" or "园" in detail:
        return str(row.get("garden_coefficient") or "1").strip() or "1"
    if land_subclass == "旱地" or land_class == "草地":
        return str(row.get("dry_coefficient") or "1").strip() or "1"
    return "1"


def _policy_rule_amount(
    rule: Dict[str, Any],
    *,
    dependency_amounts: Dict[str, Decimal],
    land_class: str,
    land_subclass: str,
) -> tuple[Decimal | None, str, str]:
    formula_key = str(rule.get("formula_key") or "")
    standard = _decimal(rule.get("standard_value"))
    if formula_key in {"direct", "manual_default"}:
        return standard, str(rule.get("standard_value") or ""), ""
    if formula_key == "paid_land_use_fee_times_rate":
        base_key = str(rule.get("depends_on") or "paid_land_use_fee")
        base = dependency_amounts.get(base_key)
        if base is None:
            return None, str(rule.get("standard_value") or ""), ""
        rate = standard / Decimal("100")
        return base * rate, _money(base), f"{_trim_decimal_text(standard)}%"
    if formula_key == "standard_times_land_coefficient":
        coefficient = _land_class_tax_coefficient(land_class, land_subclass)
        display_coefficient = "" if coefficient == D1 else _trim_decimal_text(coefficient)
        return standard * coefficient, str(rule.get("standard_value") or ""), display_coefficient
    if formula_key == "ten_thousand_per_mu_to_sqm":
        return standard * Decimal("10000") * MU_TO_SQM, str(rule.get("standard_value") or ""), ""
    return None, str(rule.get("standard_value") or ""), ""


def _configured_tax_items(
    config: Dict[str, Any],
    *,
    county_name: str,
    city_name: str,
    valuation_date: str,
    land_class: str,
    land_subclass: str,
    selected_grades: Dict[str, str] | None = None,
) -> tuple[List[Dict[str, Any]], List[str]]:
    warnings: List[str] = []
    selected_grades = selected_grades or {}
    is_forest = any(item in f"{land_class}{land_subclass}" for item in ("林地", "乔木林地", "灌木林地", "其他林地"))
    is_cultivated = land_class == "耕地" or any(item in land_subclass for item in ("水田", "水浇地", "旱地"))
    labels = {
        "water_conservancy_fund": "水利建设基金",
        "farmland_occupation_tax": "耕地占用税",
        "farmland_reclamation_fee": "耕地开垦费",
        "social_security_fund": "被征地农民社会保障费",
    }
    fee_order = [
        "water_conservancy_fund",
        "farmland_occupation_tax",
        "farmland_reclamation_fee",
        "social_security_fund",
    ]
    dependency_amounts: Dict[str, Decimal] = {}
    dependency_grades: Dict[str, str] = {}
    paid_grade = selected_grades.get("water_conservancy_fund", "")
    paid_rule = find_policy_rule(
        config,
        "paid_land_use_fee",
        county_name=county_name,
        city_name=city_name,
        valuation_date=valuation_date,
        land_subclass=land_subclass,
        grade_name=paid_grade,
    )
    if paid_rule:
        paid_amount, _, _ = _policy_rule_amount(
            paid_rule,
            dependency_amounts=dependency_amounts,
            land_class=land_class,
            land_subclass=land_subclass,
        )
        if paid_amount is not None:
            dependency_amounts["paid_land_use_fee"] = _qmoney(paid_amount)
            dependency_grades["paid_land_use_fee"] = str(paid_rule.get("grade_name") or "")
    elif paid_grade and paid_grade in HUNAN_PAID_LAND_USE_FEE_STANDARDS:
        paid_amount = _decimal(HUNAN_PAID_LAND_USE_FEE_STANDARDS[paid_grade])
        dependency_amounts["paid_land_use_fee"] = _qmoney(paid_amount)
        dependency_grades["paid_land_use_fee"] = paid_grade
    elif paid_grade:
        warnings.append("未匹配到所选等别的新增建设用地土地有偿使用费标准，水利建设基金无法自动折算，请估价师填写并确认。")
    items: List[Dict[str, Any]] = []
    for fee_key in fee_order:
        enabled = fee_key != "farmland_reclamation_fee" or is_cultivated
        rule = find_policy_rule(
            config,
            fee_key,
            county_name=county_name,
            city_name=city_name,
            valuation_date=valuation_date,
            land_subclass=land_subclass,
            grade_name=selected_grades.get(fee_key, "")
            if fee_key in ("farmland_reclamation_fee", "farmland_occupation_tax")
            else "",
        )
        if not rule:
            amount = "" if enabled else "0.00"
            if enabled:
                warnings.append(f"未匹配到{labels[fee_key]}有效政策规则，请估价师填写并确认。")
            items.append(
                _item(
                    fee_key,
                    labels[fee_key],
                    "tax",
                    amount,
                    standard_value="",
                    source="policy_config_missing" if enabled else "not_applicable",
                    source_note="无有效配置规则" if enabled else "当前征收地类不适用",
                    enabled=enabled,
                )
            )
            continue
        raw_amount, standard_value, coefficient = _policy_rule_amount(
            rule,
            dependency_amounts=dependency_amounts,
            land_class=land_class,
            land_subclass=land_subclass,
        )
        if raw_amount is None:
            amount = ""
            warnings.append(f"{labels[fee_key]}配置规则缺少依赖项或公式类型无法识别，请估价师校核。")
        else:
            amount = _money(raw_amount)
        if fee_key == "water_conservancy_fund":
            selected_paid_grade = selected_grades.get("water_conservancy_fund", "")
            if "paid_land_use_fee" in dependency_amounts:
                rule["dependency_amount_per_sqm"] = _money(dependency_amounts["paid_land_use_fee"])
                rule["dependency_grade_name"] = dependency_grades.get("paid_land_use_fee", "")
                standard_value = _money(dependency_amounts["paid_land_use_fee"])
                coefficient = str(rule.get("standard_value") or "10")
                rule["source_note"] = "；".join(
                    value
                    for value in (
                        _policy_workflow_note(paid_rule),
                        _policy_workflow_note(rule),
                    )
                    if value
                )
            elif selected_paid_grade:
                rule["dependency_grade_name"] = selected_paid_grade
        grade_name = (
            str(rule.get("dependency_grade_name") or selected_grades.get("water_conservancy_fund", ""))
            if fee_key == "water_conservancy_fund"
            else str(rule.get("grade_name") or "")
        )
        items.append(
            _item(
                fee_key,
                labels[fee_key],
                "tax",
                amount,
                standard_value=standard_value,
                coefficient=coefficient,
                formula=str(rule.get("formula_key") or ""),
                source="policy_config",
                source_note=_policy_workflow_note(rule),
                policy_key=str(rule.get("source_file") or ""),
                rule_key=str(rule.get("key") or ""),
                rule_snapshot=rule,
                grade_name=grade_name,
                enabled=enabled,
            )
        )
    if is_forest:
        warnings.append("森林植被恢复费暂未配置有效标准，请估价师按项目适用文件填写并确认。")
        items.append(
            _item(
                "forest_restoration_fee",
                "森林植被恢复费",
                "tax",
                "",
                source="policy_config_missing",
                source_note="当前地类涉及林地，需人工录入森林植被恢复费标准。",
                enabled=True,
            )
        )
    else:
        items.append(
            _item(
                "forest_restoration_fee",
                "森林植被恢复费",
                "tax",
                "0.00",
                standard_value="0.00",
                source="not_applicable",
                source_note="当前征收地类不涉及林地",
                enabled=False,
            )
        )
    return items, warnings


def _default_tax_items(land_class: str, land_subclass: str = "") -> List[Dict[str, Any]]:
    detail = land_subclass or land_class
    is_forest = "林" in detail
    is_cultivated = land_class == "耕地" or any(item in detail for item in ("水田", "水浇地", "旱"))
    return [
        _item("water_conservancy_fund", "水利建设基金", "tax", "1.60", standard_value="1.60", source_note="新增建设用地有偿使用费相关政策"),
        _item("farmland_occupation_tax", "耕地占用税", "tax", "20.00" if is_forest else "25.00", standard_value="25.00", coefficient="0.8" if is_forest else "", formula="适用税额", source_note="湖南省耕地占用税适用税额及地类调整政策"),
        _item("farmland_reclamation_fee", "耕地开垦费", "tax", "99.00" if is_cultivated else "0.00", standard_value="99.00" if is_cultivated else "0.00", source_note="湖南省耕地开垦费征收标准", enabled=is_cultivated),
        _item("social_security_fund", "被征地农民社会保障费", "tax", "30.00", standard_value="30.00", source_note="湖南省被征地农民社会保障政策"),
        _item("forest_restoration_fee", "森林植被恢复费", "tax", "20.00" if is_forest else "0.00", standard_value="20.00" if is_forest else "0.00", source_note="湖南省森林植被恢复费征收标准", enabled=is_forest),
    ]


def _development_keys(development_text: str) -> List[str]:
    text = development_text or ""
    if "七通" in text:
        return ["通路", "通电", "供水", "排水", "通讯", "供气", "供热", "场地平整"]
    if "三通" in text:
        return ["通路", "通电", "供水", "场地平整"]
    return ["通路", "供水", "排水", "通电", "通讯", "场地平整"]


def _default_development_survey_cases() -> List[Dict[str, Any]]:
    return [
        {
            "key": f"survey_{index + 1}",
            "name": "",
            "location": "",
            "survey_date": "",
            "source_type": "",
            "development_set": "",
            "total_per_sqm": "",
            "source_unit": "",
            "note": "",
            "confirmed": False,
        }
        for index in range(3)
    ]


def _development_survey_analysis(cases: Iterable[Dict[str, Any]]) -> Dict[str, str]:
    confirmed = [item for item in cases or [] if _bool(item.get("confirmed"))]
    values = [
        _decimal(item.get("total_per_sqm"))
        for item in confirmed
        if item.get("total_per_sqm") not in (None, "")
    ]
    values = [value for value in values if value > D0]
    average = _money(sum(values, D0) / Decimal(len(values))) if values else ""
    confirmed_count = str(len(confirmed))
    valid_count = str(len(values))
    if not values:
        status = "pending"
    elif len(values) < 3:
        status = "insufficient"
    else:
        status = "ready"
    return {
        "average_total_per_sqm": average,
        "case_count": str(len(list(cases or []))),
        "confirmed_count": confirmed_count,
        "valid_count": valid_count,
        "status": status,
    }


def _default_development_items(development_text: str) -> List[Dict[str, Any]]:
    recommended = {label: DEVELOPMENT_COST_RANGES.get(label, ("0", "0"))[1] for label in _development_keys(development_text)}
    return [
        {
            **_item(
                f"development_{index + 1}",
                label,
                "development",
                recommended[label],
                standard_value=recommended[label],
                source="project_survey_pending",
                source_note="参考区间上限，正式采用前请附≥3组同区域项目调查资料",
            ),
            "range_min": DEVELOPMENT_COST_RANGES.get(label, ("", ""))[0],
            "range_max": DEVELOPMENT_COST_RANGES.get(label, ("", ""))[1],
        }
        for index, label in enumerate(_development_keys(development_text))
    ]


def _usage_scenarios(data: Dict[str, Any], existing: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    existing_list = [dict(item) for item in existing or []]
    existing_by_key = {str(item.get("key") or ""): item for item in existing_list}
    text = " ".join(
        str(data.get(key) or "")
        for key in ("land_usage", "land_usage_key", "land_usage_short", "land_usage_full")
    )
    is_existing_grant_procedure = "办理" in str(data.get("transfer_purpose_mode") or data.get("transfer_purpose") or "")
    report_term = str(data.get("land_use_term_years") or data.get("land_use_term") or "").replace("年", "").strip()
    scenarios: List[Dict[str, Any]] = []
    candidates = [
        ("commercial", "商业服务业用地", ("商服", "商业"), "40", "40", "6.7"),
        ("residential", "居住用地", ("住宅", "居住"), "70", "35", "5.8"),
        ("industrial", "工矿用地", ("工业", "工矿"), "30", "5", "5.4"),
        ("warehouse", "仓储用地", ("仓储",), "50", "10", "5.4"),
        ("public", "公共管理与公共服务用地", ("公共", "公服"), "50", "10", "5.4"),
    ]
    for key, label, needles, term, value_added, reduction in candidates:
        if any(needle in text for needle in needles):
            if key == "residential" and is_existing_grant_procedure:
                reduction = "6.5"
            defaults = {
                    "key": key,
                    "label": label,
                    "use_term_years": report_term or term,
                    "profit_rate": "20" if key == "residential" and is_existing_grant_procedure else ("14" if key == "commercial" else ("15" if key == "residential" else "5")),
                    "value_added_rate": value_added,
                    "safe_rate": "",
                    "reduction_rate": reduction,
                    "location_correction_rate": "0",
                    "confirmed": False,
                }
            merged = {**defaults, **existing_by_key.get(key, {}), "key": key, "label": label}
            if report_term and not _bool(merged.get("confirmed")):
                merged["use_term_years"] = report_term
            scenarios.append(merged)
    if not scenarios:
        key = str(data.get("land_usage_key") or "other")
        label = str(data.get("land_usage_price_class") or data.get("land_usage") or "其他用地")
        defaults = {
                "key": key,
                "label": label,
                "use_term_years": report_term or "50",
                "profit_rate": "10",
                "value_added_rate": "10",
                "safe_rate": "",
                "reduction_rate": "5.4",
                "location_correction_rate": "0",
                "confirmed": False,
            }
        merged = {**defaults, **existing_by_key.get(key, {}), "key": key, "label": label}
        if report_term and not _bool(merged.get("confirmed")):
            merged["use_term_years"] = report_term
        scenarios.append(merged)
    return scenarios


def _default_building_compensation_rows() -> List[Dict[str, Any]]:
    return [
        {
            "key": "house_compensation",
            "label": "房屋补偿费",
            "standard": "1250",
            "standard_unit": "元/平方米",
            "calculation_basis": "按合法建筑面积",
            "quantity": "70",
            "divisor": "1",
            "amount": "87500",
            "note": "砖混二等：房屋补偿750元/㎡+装修补偿500元/㎡（附件1）",
            "confirmed": False,
        },
        {
            "key": "vacating_award",
            "label": "主动腾地奖",
            "standard": "100",
            "standard_unit": "元/平方米",
            "calculation_basis": "按合法建筑面积",
            "quantity": "70",
            "divisor": "1",
            "amount": "7000",
            "note": "按二层计算",
            "confirmed": False,
        },
        {
            "key": "moving_fee",
            "label": "搬家费",
            "standard": "2000",
            "standard_unit": "元/户",
            "calculation_basis": "按每户，4人/户",
            "quantity": "1",
            "divisor": "4",
            "amount": "500",
            "note": "按中间值计算",
            "confirmed": False,
        },
        {
            "key": "transition_fee",
            "label": "过渡费",
            "standard": "1000",
            "standard_unit": "元/户/月",
            "calculation_basis": "按每户，4人/户",
            "quantity": "1",
            "divisor": "4",
            "months": "12",
            "amount": "3000",
            "note": "过渡期12个月",
            "confirmed": False,
        },
    ]


def _default_resettlement_population_cases() -> List[Dict[str, Any]]:
    return [
        {"key": "case_a", "name": "", "location": "", "land_area_ha": "", "population": "", "population_per_ha": "", "confirmed": False},
        {"key": "case_b", "name": "", "location": "", "land_area_ha": "", "population": "", "population_per_ha": "", "confirmed": False},
        {"key": "case_c", "name": "", "location": "", "land_area_ha": "", "population": "", "population_per_ha": "", "confirmed": False},
    ]


def _default_risk_items(config: Dict[str, Any], scenarios: Iterable[Dict[str, Any]], scheme_key: str) -> List[Dict[str, Any]]:
    scheme = get_risk_scheme(config, scheme_key)
    if not scheme:
        return []
    usage_key = ""
    scenario_keys = [str(item.get("key") or "") for item in scenarios]
    if len(scenario_keys) == 1:
        usage_key = scenario_keys[0]
    assignment = scheme.get("assignment") or RISK_LEVEL_ADJUSTMENTS
    return [
        {
            "usage_key": usage_key,
            "group": group.get("label") or "",
            "label": factor.get("label") or "",
            "level": "",
            "group_weight": str(group.get("weight") or ""),
            "weight": str(factor.get("weight") or ""),
            "adjustment_rate": "",
            "level_options": [
                {"level": level, "adjustment_rate": str(rate)}
                for level, rate in assignment.items()
            ],
            "confirmed": False,
        }
        for group in scheme.get("groups") or []
        for factor in group.get("factors") or []
    ]


def _default_risk_groups(config: Dict[str, Any], scenarios: Iterable[Dict[str, Any]], scheme_key: str) -> List[Dict[str, Any]]:
    scheme = get_risk_scheme(config, scheme_key)
    if not scheme:
        return []
    usage_key = ""
    scenario_keys = [str(item.get("key") or "") for item in scenarios]
    if len(scenario_keys) == 1:
        usage_key = scenario_keys[0]
    return [
        {
            "usage_key": usage_key,
            "key": group.get("key") or "",
            "label": group.get("label") or "",
            "weight": str(group.get("weight") or ""),
            "computed_value": "",
            "override_enabled": False,
            "override_value": "",
            "override_reason": "",
            "effective_value": "",
            "confirmed": False,
        }
        for group in scheme.get("groups") or []
    ]


def _recalculate_attachment_analysis(analysis: Dict[str, Any]) -> Dict[str, str]:
    rows = analysis.get("building_compensation_rows") or []
    for row in rows:
        standard = _decimal(row.get("standard"))
        quantity = _decimal(row.get("quantity"))
        months = _decimal(row.get("months"), D1)
        divisor = _decimal(row.get("divisor"), D1)
        if divisor <= D0:
            divisor = D1
        row["amount"] = _money(standard * quantity * months / divisor)
    building_per_person = sum((_decimal(row.get("amount")) for row in rows), D0)
    cases = analysis.get("resettlement_population_cases") or []
    densities: List[Decimal] = []
    for item in cases:
        area = _decimal(item.get("land_area_ha"))
        population = _decimal(item.get("population"))
        if area > D0 and population >= D0:
            item["population_per_ha"] = _money(population / area)
            densities.append(population / area)
    average_density = sum(densities, D0) / Decimal(len(densities)) if densities else D0
    green_per_sqm = _decimal(analysis.get("green_seedling_standard_per_mu")) * MU_TO_SQM
    building_per_sqm = building_per_person * average_density / Decimal("10000")
    attachment_per_sqm = building_per_sqm + green_per_sqm
    return {
        "building_compensation_per_person": _money(building_per_person),
        "average_population_per_ha": _money(average_density),
        "building_compensation_per_sqm": _money(building_per_sqm),
        "green_seedling_per_sqm": _money(green_per_sqm),
        "attachment_compensation_per_sqm": _money(attachment_per_sqm),
    }


def _sum_items(items: Iterable[Dict[str, Any]]) -> Decimal:
    return sum(
        (_decimal(item.get("amount_per_sqm")) for item in items if item.get("enabled", True)),
        D0,
    )


def _risk_adjustment_value(items: Iterable[Dict[str, Any]]) -> Decimal:
    confirmed = [item for item in items if item.get("confirmed")]
    if not confirmed:
        return D0
    if not any(item.get("group_weight") not in (None, "") for item in confirmed):
        return sum((_decimal(item.get("weight")) * _decimal(item.get("adjustment_rate")) for item in confirmed), D0)
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for item in confirmed:
        grouped.setdefault(str(item.get("group") or "其他风险"), []).append(item)
    total = D0
    for group_items in grouped.values():
        weight_sum = sum((_decimal(item.get("weight")) for item in group_items), D0)
        group_adjustment = (
            sum((_decimal(item.get("weight")) * _decimal(item.get("adjustment_rate")) for item in group_items), D0)
            / weight_sum
            if weight_sum > D0
            else D0
        )
        total += _decimal(group_items[0].get("group_weight")) * group_adjustment
    return total


def _sync_risk_groups(
    risk_items: Iterable[Dict[str, Any]],
    risk_groups: Iterable[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    existing = {str(item.get("key") or item.get("label") or ""): dict(item) for item in risk_groups or []}
    labels = []
    for item in risk_items or []:
        label = str(item.get("group") or "")
        if label and label not in labels:
            labels.append(label)
    groups: List[Dict[str, Any]] = []
    for label in labels:
        group_items = [item for item in risk_items or [] if str(item.get("group") or "") == label]
        if not group_items:
            continue
        existing_group = existing.get(label) or next((group for group in existing.values() if group.get("label") == label), {})
        weight = str(existing_group.get("weight") or group_items[0].get("group_weight") or "")
        confirmed_items = [item for item in group_items if _bool(item.get("confirmed"))]
        weight_sum = sum((_decimal(item.get("weight")) for item in confirmed_items), D0)
        computed = (
            sum((_decimal(item.get("weight")) * _decimal(item.get("adjustment_rate")) for item in confirmed_items), D0) / weight_sum
            if confirmed_items and weight_sum > D0
            else D0
        )
        override_enabled = _bool(existing_group.get("override_enabled"))
        override_value = existing_group.get("override_value")
        legacy_override = existing_group.get("combined_rate") or existing_group.get("rate")
        if not override_enabled and legacy_override not in (None, ""):
            override_enabled = True
            override_value = legacy_override
        effective = _decimal(override_value) if override_enabled and override_value not in (None, "") else computed
        groups.append(
            {
                "usage_key": str(existing_group.get("usage_key") or group_items[0].get("usage_key") or ""),
                "key": str(existing_group.get("key") or label),
                "label": label,
                "weight": weight,
                "computed_value": _money(computed) if confirmed_items else "",
                "override_enabled": override_enabled,
                "override_value": str(override_value or ""),
                "override_reason": str(existing_group.get("override_reason") or ""),
                "effective_value": _money(effective) if (confirmed_items or override_enabled) else "",
                "confirmed": _bool(existing_group.get("confirmed")),
            }
        )
    return groups


def _risk_adjustment_value_from_groups(
    groups: Iterable[Dict[str, Any]],
    *,
    include_unconfirmed: bool = False,
) -> Decimal:
    applicable = [
        item
        for item in groups or []
        if (include_unconfirmed or _bool(item.get("confirmed"))) and item.get("effective_value") not in (None, "")
    ]
    if not applicable:
        return D0
    return sum((_decimal(item.get("weight")) * _decimal(item.get("effective_value")) for item in applicable), D0)


def _risk_adjustment_formula_from_groups(groups: Iterable[Dict[str, Any]]) -> str:
    confirmed = [item for item in groups or [] if _bool(item.get("confirmed")) and item.get("effective_value") not in (None, "")]
    if not confirmed:
        return "待校核"
    parts = [f"{item.get('weight') or '0'}×{item.get('effective_value') or '0'}%" for item in confirmed]
    return f"{'＋'.join(parts)}＝{_money(_risk_adjustment_value_from_groups(confirmed))}%"


def _risk_adjustment_formula(items: Iterable[Dict[str, Any]]) -> str:
    confirmed = [item for item in items if item.get("confirmed")]
    if not confirmed:
        return "待校核"
    if not any(item.get("group_weight") not in (None, "") for item in confirmed):
        parts = [f"{item.get('weight') or '0'}×{item.get('adjustment_rate') or '0'}%" for item in confirmed]
        return f"{'＋'.join(parts)}＝{_money(_risk_adjustment_value(confirmed))}%"
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for item in confirmed:
        grouped.setdefault(str(item.get("group") or "其他风险"), []).append(item)
    parts = []
    for group_items in grouped.values():
        weight_sum = sum((_decimal(item.get("weight")) for item in group_items), D0)
        group_adjustment = (
            sum((_decimal(item.get("weight")) * _decimal(item.get("adjustment_rate")) for item in group_items), D0)
            / weight_sum
            if weight_sum > D0
            else D0
        )
        parts.append(f"{group_items[0].get('group_weight') or '0'}×{_money(group_adjustment)}%")
    return f"{'＋'.join(parts)}＝{_money(_risk_adjustment_value(confirmed))}%"


def _all_enabled_confirmed(items: Iterable[Dict[str, Any]]) -> bool:
    enabled = [item for item in items if item.get("enabled", True)]
    return bool(enabled) and all(_bool(item.get("confirmed")) for item in enabled)


def _all_items_resolved(items: Iterable[Dict[str, Any]]) -> bool:
    items_list = list(items or [])
    if not items_list:
        return False
    for item in items_list:
        if item.get("enabled", True):
            if not _bool(item.get("confirmed")):
                return False
    return True


def _by_key(items: Iterable[Dict[str, Any]], key: str) -> Dict[str, Any]:
    return next((item for item in items if item.get("key") == key), {})


def _build_narratives(data: Dict[str, Any], analysis: Dict[str, Any]) -> Dict[str, str]:
    """Generate cost-approx report body text. Wording and section order must match the specimen report;
    policy_rules source_note/internal_note are workbench-only and must not be interpolated here."""
    acquisition_total = analysis["totals"]["acquisition_total"]
    tax_total = analysis["totals"]["tax_total"]
    development_total = analysis["totals"]["development_total"]
    location = data.get("land_location_full") or data.get("land_location") or "【请填写估价对象坐落】"
    land_class = analysis.get("acquisition_land_class") or "【请填写征收地类】"
    land_subclass = analysis.get("acquisition_land_subclass") or "【请填写征收地类细分类】"
    county = str(data.get("county_name") or "所属县市")
    county = str(data.get("county_name") or "所属县市")
    valuation_date = str(data.get("valuation_date") or "【请填写估价期日】")
    policy_basis = analysis.get("policy_basis") or {}
    zone = analysis.get("compensation_zone") or "【请填写区片】"
    province = policy_basis.get("province_compensation") or {}
    local_policy_name = analysis.get("local_compensation_policy_name") or "【请填写市县级征地补偿安置政策名称】"
    local_policy_no = analysis.get("local_compensation_policy_no") or "【请填写政策文号】"
    local_policy_date = analysis.get("local_compensation_policy_date") or "【请填写政策日期】"
    green_standard = analysis.get("green_seedling_standard_per_mu") or "【请填写青苗补偿标准】"
    green_sqm = _money(_decimal(green_standard) * MU_TO_SQM) if green_standard != "【请填写青苗补偿标准】" else "【请填写附着物和青苗补偿标准】"
    acquisition_items = analysis.get("acquisition_items") or []
    tax_items = analysis.get("tax_items") or []
    land_compensation = _by_key(acquisition_items, "land_compensation")
    building_compensation = _by_key(acquisition_items, "building_compensation")
    seedling_compensation = _by_key(acquisition_items, "seedling_compensation")
    attachment = _by_key(acquisition_items, "ground_attachment")
    base_per_mu = (
        land_compensation.get("standard_value")
        if land_compensation.get("standard_value") not in (None, "")
        else province.get("base_per_mu") or "【请填写征地补偿标准】"
    )
    water_fund = _by_key(tax_items, "water_conservancy_fund")
    occupation_tax = _by_key(tax_items, "farmland_occupation_tax")
    reclamation_fee = _by_key(tax_items, "farmland_reclamation_fee")
    social_fee = _by_key(tax_items, "social_security_fund")
    forest_fee = _by_key(tax_items, "forest_restoration_fee")
    water_rule = water_fund.get("rule_snapshot") or {}
    water_paid_standard = water_fund.get("standard_value") or water_rule.get("dependency_amount_per_sqm") or "16.00"
    water_rate = str(water_fund.get("coefficient") or water_rule.get("standard_value") or "10").replace("%", "")
    water_grade = water_fund.get("grade_name") or water_rule.get("dependency_grade_name") or "【请选择有偿使用费等别】"
    reclamation_rule = reclamation_fee.get("rule_snapshot") or {}
    occupation_grade = occupation_tax.get("grade_name") or "【请选择耕地占用税适用口径】"
    reclamation_grade = reclamation_fee.get("grade_name") or "【请选择耕地开垦费等别】"
    development_labels = "、".join(
        str(item.get("label")) for item in analysis.get("development_items") or [] if item.get("enabled", True)
    )
    result_by_key = {item.get("key"): item for item in analysis.get("usage_results") or []}
    scenario_by_key = {item.get("key"): item for item in analysis.get("usage_scenarios") or []}

    land_compensation_disabled = land_compensation.get("enabled") is False
    compensation_is_zero = land_compensation_disabled
    compensation_formula = f"{base_per_mu}×0.0015"
    compensation_basis_text = (
        ""
        if compensation_is_zero
        else (
            f"根据评估人员现场勘察和调查区域周边的土地利用状况，待估宗地所在区域土地利用类型主要为{land_class}。"
            f"根据《{PROVINCE_POLICY_NAME}》（{PROVINCE_POLICY_NO}，2024年1月31日）及《{PROVINCE_SCOPE_POLICY_NAME}》"
            f"（{PROVINCE_SCOPE_POLICY_NO}，2024年2月1日），估价对象属于{county}征地补偿区{zone}区范围，"
            f"{land_class}补偿区片价为{base_per_mu}元/亩，则："
        )
    )
    compensation_amount_text = land_compensation.get("amount_per_sqm") or "【请填写征地补偿标准】"
    compensation_result_text = (
        ""
        if compensation_is_zero
        else f"征地补偿费和安置补助费={compensation_formula}={compensation_amount_text}元/平方米。"
    )
    attachment_analysis = analysis.get("attachment_compensation_analysis") or {}
    
    occupation_formula = (
        f"{occupation_tax.get('standard_value') or '25'}×{occupation_tax.get('coefficient')}="
        f"{occupation_tax.get('amount_per_sqm')}元/平方米"
        if occupation_tax.get("coefficient") not in (None, "", "1")
        else f"{occupation_tax.get('amount_per_sqm') or '【请填写耕地占用税】'}元/平方米"
    )
    
    tax_parts = []
    idx = 1
    cn_nums = ["①", "②", "③", "④", "⑤"]

    if water_fund.get("enabled", True):
        tax_parts.extend(
            [
                f"{cn_nums[idx-1]}水利建设基金",
                (
                    "根据湖南省人民政府关于印发《湖南省水利建设基金筹集和使用管理办法》的通知（湘政发[2011]27号）及"
                    "《关于调整新增建设用地土地有偿使用费征收管理有关政策的通知》（湘财综[2018]42号），"
                    f"{county}新增建设用地土地有偿使用费征收等别为{water_grade}，征收标准为{water_paid_standard}元/平方米，"
                    f"按{water_rate}%征收水利建设基金，"
                    f"即{water_paid_standard}×{water_rate}%={water_fund.get('amount_per_sqm') or '【请填写水利建设基金】'}元/平方米。"
                ),
            ]
        )
        idx += 1

    if occupation_tax.get("enabled", True):
        tax_parts.extend(
            [
                f"{cn_nums[idx-1]}耕地占用税",
                (
                    f"根据《湖南省人民代表大会常务委员会关于湖南省耕地占用税适用税额的决定》，{county}耕地占用税适用税额为"
                    f"{occupation_tax.get('standard_value') or '25'}元/平方米，当前采用口径为{occupation_grade}。"
                    f"结合本次征收地类，耕地占用税为{occupation_formula}。"
                ),
            ]
        )
        idx += 1

    if reclamation_fee.get("enabled", True):
        tax_parts.extend(
            [
                f"{cn_nums[idx-1]}耕地开垦费",
                (
                    f"根据《湖南省耕地开垦费征收使用管理办法》，本次采用{reclamation_grade}口径，{land_subclass}耕地开垦费标准为"
                    f"{reclamation_rule.get('standard_value') or reclamation_fee.get('standard_value') or '【待校核】'}万元/亩，折合"
                    f"{reclamation_rule.get('standard_value') or reclamation_fee.get('standard_value') or '【待校核】'}×10000×0.0015="
                    f"{reclamation_fee.get('amount_per_sqm') or '【请填写耕地开垦费】'}元/平方米。"
                ),
            ]
        )
        idx += 1

    if social_fee.get("enabled", True):
        tax_parts.extend(
            [
                f"{cn_nums[idx-1]}社会保障资金",
                (
                    "根据湖南省人力资源社会保障厅《关于进一步做好被征地农民社会保障工作的通知》"
                    f"（湘人社规〔2023〕1号，2023年4月28日），被征地农民社会保障费为{social_fee.get('amount_per_sqm') or '30.00'}元/平方米。"
                ),
            ]
        )
        idx += 1

    if forest_fee.get("enabled", True):
        tax_parts.extend(
            [
                f"{cn_nums[idx-1]}森林植被恢复费",
                (
                    "根据《湖南省森林植被恢复费征收使用管理实施办法》（湘财税〔2024〕10号），结合林地类型及建设项目性质，"
                    f"本次森林植被恢复费为{forest_fee.get('amount_per_sqm') or '【请填写森林植被恢复费】'}元/平方米。"
                ),
            ]
        )
        idx += 1

    enabled_tax_amounts = []
    for item in tax_items:
        if not item.get("enabled", True):
            continue
        amt = item.get("amount_per_sqm")
        if amt in (None, ""):
            placeholder = f"【请填写{item.get('label')}】"
            enabled_tax_amounts.append(placeholder)
        else:
            enabled_tax_amounts.append(str(amt))

    tax_parts.append(f"则：各项税费合计为={'＋'.join(enabled_tax_amounts)}={tax_total}元/平方米。")
    tax_narrative = "\n".join(tax_parts)

    local_policy_details = "，".join(value for value in (local_policy_no, local_policy_date) if value)
    local_policy_reference = f"{local_policy_name}（{local_policy_details}）" if local_policy_details else local_policy_name
    seedling_disabled = seedling_compensation.get("enabled") is False
    building_disabled = building_compensation.get("enabled") is False
    green_seedling_paragraph = (
        ""
        if seedling_disabled
        else (
            f"根据{local_policy_reference}，{county}的{land_subclass}青苗补偿标准为"
            f"{analysis.get('green_seedling_standard_per_mu') or '【待校核】'}元/亩，"
            f"合{seedling_compensation.get('amount_per_sqm') or attachment_analysis.get('green_seedling_per_sqm') or green_sqm}元/平方米；"
        )
    )
    building_standard_intro = (
        ""
        if building_disabled
        else f"根据{local_policy_reference}，地上建筑物补偿标准计算如下："
    )
    attachment_population_narrative = "" if building_disabled else "\n".join(
        [
            (
                f"根据以上测算，区域地上建筑物平均补偿费用为"
                f"{attachment_analysis.get('building_compensation_per_person') or '【待校核】'}元/人。"
                "据湖南征地信息平台公布的数据，收集区域三个征地案例的安置农业人口数，"
                f"确定区域农业人口数约{attachment_analysis.get('average_population_per_ha') or '【待校核】'}人/公顷。"
                "具体补偿人口数量详见下表："
            ),
        ]
    )


    enabled_acquisition = [
        item for item in (land_compensation, building_compensation, seedling_compensation)
        if item and item.get("enabled", True) and _decimal(item.get("amount_per_sqm")) > D0
    ]
    acquisition_formula_parts = [str(item.get("amount_per_sqm") or "0.00") for item in enabled_acquisition]
    attachment_formula_parts = []
    if not building_disabled and _decimal(building_compensation.get("amount_per_sqm")) > D0:
        attachment_formula_parts.append(str(building_compensation.get("amount_per_sqm") or "0.00"))
    if not seedling_disabled and _decimal(seedling_compensation.get("amount_per_sqm")) > D0:
        attachment_formula_parts.append(str(seedling_compensation.get("amount_per_sqm") or "0.00"))
    attachment_formula = (
        "则地上附着物补偿标准为" + "＋".join(
            [
                label
                for enabled, label in (
                    (not building_disabled and _decimal(building_compensation.get("amount_per_sqm")) > D0, "地上建筑物费用"),
                    (not seedling_disabled and _decimal(seedling_compensation.get("amount_per_sqm")) > D0, "青苗补偿"),
                )
                if enabled
            ]
        )
        + f"\n={'＋'.join(attachment_formula_parts)}={_money(sum((_decimal(item) for item in attachment_formula_parts), D0))}元/平方米。\n"
        if attachment_formula_parts
        else ""
    )
    acquisition_solve_narrative = (
        attachment_formula
        + (f"故：土地取得费={'＋'.join(acquisition_formula_parts)}={acquisition_total}元/平方米。" if acquisition_formula_parts else "")
    )
    acquisition_parts = [
        f"根据{LAND_MANAGEMENT_LAW_REFERENCE}第四十八条规定，征收土地的补偿费用包括土地补偿费、安置补助费以及农村村民住宅、其他地上附着物和青苗等的补偿费用、被征地农民的社会保障费用。",
    ]
    if not compensation_is_zero:
        acquisition_parts.extend(["①征地补偿费和安置补助费", compensation_basis_text, compensation_result_text])
    if not building_disabled or not seedling_disabled:
        acquisition_parts.append("②地上附着物补偿")
        acquisition_parts.extend([green_seedling_paragraph, building_standard_intro, attachment_population_narrative])
    acquisition_parts.append(acquisition_solve_narrative)
    acquisition_narrative = "\n".join(part for part in acquisition_parts if str(part or "").strip())

    


    interest_narrative = "\n".join(
        [
            (
                f"根据待估宗地的开发程度和开发规模及实际调查分析，确定土地开发周期为{analysis.get('development_cycle_years')}年，"
                f"本次投资利息率取{analysis.get('interest_rate')}%。土地取得费及相关税费在征地时一次投入，土地开发费在开发期内均匀投入，单利计息，则："
            ),
            "投资利息=（土地取得费＋相关税费）×开发周期×利息率＋土地开发费×开发周期×1/2×利息率",
            (
                f"=（{acquisition_total}+{tax_total}）×{analysis.get('development_cycle_years')}×{analysis.get('interest_rate')}%"
                f"＋{development_total}×{analysis.get('development_cycle_years')}×1/2×{analysis.get('interest_rate')}%"
                f"={analysis['totals'].get('interest')}元/平方米。"
            ),
        ]
    )
    profit_lines = [
        "土地作为一种生产要素进入企业的生产过程，土地资产与其他生产要素相结合产生利润，土地投资应获得与其资产量相对应的回报。"
    ]
    value_added_lines = [
        "土地增值收益是待估土地因改变用途或进行土地开发，达到建设用地的某种利用条件而发生的价值增加。"
    ]
    term_lines = ["土地使用年期修正系数=1－1/（1＋r）ⁿ，其中r为土地还原率，n为设定土地使用年期。"]
    location_lines = []
    solve_lines = []
    for key, scenario in scenario_by_key.items():
        item = result_by_key.get(key) or {}
        profit_lines.extend(
            [
                f"{scenario.get('label')}投资利润率取{scenario.get('profit_rate')}%，则：",
                (
                    f"投资利润=（{acquisition_total}+{tax_total}+{development_total}）×{scenario.get('profit_rate')}%"
                    f"={item.get('profit') or '【计算求取投资利润】'}元/平方米。"
                ),
            ]
        )
        value_added_lines.extend(
            [
                f"{scenario.get('label')}土地增值收益率取{scenario.get('value_added_rate')}%，则：",
                (
                    f"土地增值收益={item.get('cost_price') or '【计算求取土地取得及开发成本】'}×{scenario.get('value_added_rate')}%"
                    f"={item.get('value_added_income') or '【计算求取土地增值收益】'}元/平方米。"
                ),
            ]
        )
        term_lines.append(
            f"{scenario.get('label')}年期修正系数=1－1/（1＋{scenario.get('reduction_rate')}%）{scenario.get('use_term_years')}="
            f"{item.get('term_correction_factor') or '【请核算年期修正系数】'}。"
        )
        location_lines.append(f"{scenario.get('label')}区位修正率合计为{item.get('location_correction_rate') or '0.00'}%。")
        solve_lines.extend(
            [
                "比准价格=（土地取得费＋相关税费＋土地开发费＋投资利息＋投资利润＋土地增值收益）×土地使用年期修正系数",
                "最终单价=比准价格×（1＋区位修正率）",
                f"{scenario.get('label')}：",
                (
                    f"P=（{acquisition_total}+{tax_total}+{development_total}+{item.get('interest')}+"
                    f"{item.get('profit')}+{item.get('value_added_income')}）×{item.get('term_correction_factor')}"
                    f"={item.get('comparable_price') or '【计算比准价格】'}元/平方米。"
                ),
                (
                    f"最终单价={item.get('comparable_price') or '【计算比准价格】'}×"
                    f"（1+{item.get('location_correction_rate')}%）={item.get('final_price')}元/平方米。"
                ),
            ]
        )
    policy_references = []
    for policy in analysis.get("policy_documents") or []:
        if not policy.get("enabled", True):
            continue
        detail = _policy_reference(policy)
        if detail:
            policy_references.append(detail)
    policy_text = "、".join(policy_references) or "【请确认成本逼近法政策依据】"
    basis_text = "\n".join(
        [
            (
                f"估价对象位于{location}，根据对估价对象周边区域用地调查，评估宗地以{land_class}为主，"
                f"此次采用成本逼近法评估待估宗地地价时，考虑估价期日被征用的土地为{land_class}。"
            ),
            (
                f"估价对象的估价期日为{valuation_date}，根据新土地管理法规定结合{policy_text}，"
                f"以及湖南省、{analysis.get('effective_local_city') or '所属市'}颁布其他各项文件为依据。"
                "根据成本逼近法测算地价的步骤，其各项费用如下："
            ),
        ]
    )
    return {
        "cost_approx_method_intro": str(data.get("cost_approx_method_intro") or (
            "成本逼近法是以开发土地所耗费的各项费用之和为主要依据，再加上一定的利润、利息、应缴纳的税金和土地增值收益来确定土地价格的估价方法。"
            "基本计算公式为：V＝Ea＋Ed＋T＋R1＋R2＋R3＝VE＋R3。"
        )),
        "cost_approx_basis_intro": basis_text,
        "cost_approx_acquisition_intro": "\n".join(
            part for part in [
                f"根据{LAND_MANAGEMENT_LAW_REFERENCE}第四十八条规定，征收土地的补偿费用包括土地补偿费、安置补助费以及农村村民住宅、其他地上附着物和青苗等的补偿费用、被征地农民的社会保障费用。",
                "①征地补偿费和安置补助费" if not compensation_is_zero else "",
                compensation_basis_text,
                compensation_result_text,
                "②地上附着物补偿" if (not building_disabled or not seedling_disabled) else "",
                green_seedling_paragraph,
                building_standard_intro,
            ] if str(part or "").strip()
        ),
        "cost_approx_attachment_population_narrative": attachment_population_narrative,
        "cost_approx_acquisition_solve": acquisition_solve_narrative,
        "cost_approx_acquisition_narrative": acquisition_narrative,
        "cost_approx_tax_narrative": tax_narrative,
        "cost_approx_development_narrative": (
            f"本次评估设定的土地开发程度涉及{development_labels}。根据估价人员市场调查资料，确定区域土地开发费为"
            f"{development_total}元/平方米。"
        ),
        "cost_approx_interest_narrative": interest_narrative,
        "cost_approx_profit_narrative": "\n".join(profit_lines),
        "cost_approx_value_added_narrative": "\n".join(value_added_lines),
        "cost_approx_interest_profit_narrative": "\n".join([interest_narrative, *profit_lines, *value_added_lines]),
        "cost_approx_term_narrative": "\n".join(term_lines),
        "cost_approx_location_narrative": (
            "据前分析，成本逼近法所计算的地价为估价对象所在区域的平均价格，估价对象地价还需要根据宗地所在区域内的位置和宗地自身条件进行个别因素修正。"
            + (
                "本次采用固定幅度等级修正法，估价师选择各因素优劣等级后，系统直接加总经确认的修正率。\n"
                if analysis.get("location_correction_mode") == "direct_sum"
                else "本次采用高级因素分析方式确定区位修正率。\n"
            )
            + "\n".join(location_lines)
        ),
        "cost_approx_solve_narrative": "\n".join(
            [
                *solve_lines,
                (
                    f"成本逼近法测算地价的结果为：估价对象于估价期日{valuation_date}，在设定土地用途、设定土地使用年限、"
                    "设定土地开发程度及设定利用条件下的国有土地使用权价格如上。"
                ),
            ]
        ) if solve_lines else "【请完成成本逼近法测算】",
    }


def _table_sections(analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Build the ten Baozhen-style cost-method tables shared by UI and Word."""
    def section(
        key: str,
        title: str,
        columns: List[str],
        rows: List[Dict[str, Any]],
        *,
        header_rows: List[List[Dict[str, Any]]] | None = None,
        group_columns: List[int] | None = None,
        source_target: str = "cost_items",
    ) -> Dict[str, Any]:
        return {
            "key": key,
            "title": title,
            "report_title": title,
            "columns": columns,
            "header_rows": header_rows or [[{"label": label} for label in columns]],
            "group_columns": group_columns or [],
            "rows": rows,
            "source_target": source_target,
        }

    def ref(path: str) -> str:
        return f"cost_approx_analysis.{path}"

    building_rows = []
    for index, item in enumerate(analysis.get("building_compensation_rows") or []):
        display_quantity = (
            f"{item.get('quantity') or '待校核'}人"
            if item.get("key") in {"moving_fee", "transition_fee"}
            else item.get("quantity") or "待校核"
        )
        building_rows.append(
            {
                "cells": [
                    item.get("label") or "待校核",
                    f"{item.get('standard') or '待校核'}{item.get('standard_unit') or item.get('unit') or ''}",
                    item.get("calculation_basis") or "待校核",
                    display_quantity,
                    item.get("amount") or "待校核",
                    item.get("note") or "",
                ],
                "cell_refs": [
                    ref(f"building_compensation_rows.{index}.label"),
                    ref(f"building_compensation_rows.{index}.standard"),
                    ref(f"building_compensation_rows.{index}.calculation_basis"),
                    ref(f"building_compensation_rows.{index}.quantity"),
                    ref(f"building_compensation_rows.{index}.amount"),
                    ref(f"building_compensation_rows.{index}.note"),
                ],
            }
        )

    population_rows = []
    for index, item in enumerate(analysis.get("resettlement_population_cases") or []):
        population_rows.append(
            {
                "cells": [
                    item.get("name") or f"案例{index + 1}",
                    item.get("location") or "待校核",
                    item.get("land_area_ha") or "待校核",
                    item.get("population") or "待校核",
                    item.get("population_per_ha") or "待校核",
                ],
                "cell_refs": [
                    ref(f"resettlement_population_cases.{index}.name"),
                    ref(f"resettlement_population_cases.{index}.location"),
                    ref(f"resettlement_population_cases.{index}.land_area_ha"),
                    ref(f"resettlement_population_cases.{index}.population"),
                    ref(f"resettlement_population_cases.{index}.population_per_ha"),
                ],
            }
        )
    population_rows.append(
        {
            "cells": ["平均值", "/", "/", "/", (analysis.get("attachment_compensation_analysis") or {}).get("average_population_per_ha") or "待校核"],
            "cell_refs": ["", "", "", "", ref("attachment_compensation_analysis.average_population_per_ha")],
        }
    )

    acquisition_total = analysis.get("totals", {}).get("acquisition_total") or "待校核"
    tax_total = analysis.get("totals", {}).get("tax_total") or "待校核"
    enabled_tax_items = [item for item in analysis.get("tax_items") or [] if item.get("enabled", True)]
    acquisition_tax_columns = ["土地取得费（元/平方米）"] + [
        f"{item.get('label') or '相关税费'}（元/平方米）" for item in enabled_tax_items
    ] + ["合计（元/平方米）"]
    acquisition_tax_rows = [
        {
            "cells": [acquisition_total]
            + [item.get("amount_per_sqm") or "待校核" for item in enabled_tax_items]
            + [_money(_decimal(acquisition_total) + _decimal(tax_total))],
            "cell_refs": [ref("totals.acquisition_total")]
            + [ref(f"tax_items.{item.get('key')}.amount_per_sqm") for item in enabled_tax_items]
            + [ref("totals.acquisition_and_tax_total")],
        }
    ]

    enabled_development_items = [item for item in analysis.get("development_items") or [] if item.get("enabled", True)]
    development_columns = ["土地开发项目"] + [item.get("label") or "待校核" for item in enabled_development_items] + ["合计"]
    development_ranges = [
        (
            f"{item.get('range_min')}-{item.get('range_max')}"
            if item.get("range_min") not in (None, "") and item.get("range_max") not in (None, "")
            else item.get("amount_per_sqm") or "待校核"
        )
        for item in enabled_development_items
    ]
    range_mins = sum((_decimal(item.get("range_min")) for item in enabled_development_items), D0)
    range_maxs = sum((_decimal(item.get("range_max")) for item in enabled_development_items), D0)
    development_rows = [
        {
            "cells": ["费用（元/平方米）"] + development_ranges + [f"{_money(range_mins)}-{_money(range_maxs)}"],
            "cell_refs": [""]
            + [ref(f"development_items.{item.get('key')}.range_min") for item in enabled_development_items]
            + [ref("totals.development_total")],
        }
    ]

    risk_items = analysis.get("risk_items") or []
    risk_groups = analysis.get("risk_groups") or []
    residential_scenario = next(
        (item for item in analysis.get("usage_scenarios") or [] if item.get("key") == "residential"),
        (analysis.get("usage_scenarios") or [{}])[0] if analysis.get("usage_scenarios") else {},
    )
    risk_usage_label = residential_scenario.get("label") or "居住用地"
    risk_impact_rows = [
        {
            "cells": [item.get("group") or "待校核", item.get("label") or "待校核", item.get("level") or "待校核"],
            "cell_refs": [
                ref(f"risk_items.{index}.group"),
                ref(f"risk_items.{index}.label"),
                ref(f"risk_items.{index}.level"),
            ],
        }
        for index, item in enumerate(risk_items)
    ]
    risk_weight_rows = [
        {
            "cells": [
                item.get("group") or "待校核",
                item.get("group_weight") or "待校核",
                item.get("label") or "待校核",
                item.get("weight") or "待校核",
            ],
            "cell_refs": [
                ref(f"risk_items.{index}.group"),
                ref(f"risk_items.{index}.group_weight"),
                ref(f"risk_items.{index}.label"),
                ref(f"risk_items.{index}.weight"),
            ],
        }
        for index, item in enumerate(risk_items)
    ]
    risk_assignment_rows = [
        {"cells": ["风险调整值", "0", "2%", "4%", "8%"], "cell_refs": ["", "", "", "", ""]},
    ]
    risk_result_rows = [
        {
            "cells": [
                "风险调整值",
                _risk_adjustment_formula_from_groups(risk_groups) if analysis.get("risk_mode") == "analysis" else _risk_adjustment_formula(risk_items),
            ],
            "cell_refs": ["", ref(f"usage_scenarios.{residential_scenario.get('key') or 'residential'}.reduction_rate")],
        }
    ]

    location_factor_pairs = [
        (index, item)
        for index, item in enumerate(analysis.get("location_factors") or [])
        if item.get("enabled") is not False
    ]
    location_factors = [item for _, item in location_factor_pairs]
    location_template = analysis.get("location_template_snapshot") or {}
    max_level_count = max([len(item.get("levels") or []) for item in location_factors] + [5])
    location_level_headers = [f"等级{index + 1}" for index in range(max_level_count)]
    location_rule_rows = [
        {
            "cells": (lambda levels: [
                item.get("label") or "待校核",
                *(levels + [""] * max_level_count)[:max_level_count],
                f"{item.get('grade_amplitude') or '待校核'}%",
            ])(list(item.get("levels") or RESIDENTIAL_LOCATION_GRADE_DESCRIPTIONS.get(item.get("label")) or ("优", "较优", "一般", "较劣", "劣"))),
            "cell_refs": [ref(f"location_factors.{index}.label"), *["" for _ in range(max_level_count)], ref(f"location_factors.{index}.grade_amplitude")],
        }
        for index, item in location_factor_pairs
    ]
    location_rows = [
        {
            "cells": [item.get("label") or "待校核", item.get("description") or "待校核", item.get("level") or "待校核", f"{item.get('correction_rate') or '待校核'}%"],
            "cell_refs": [
                ref(f"location_factors.{index}.label"),
                ref(f"location_factors.{index}.description"),
                ref(f"location_factors.{index}.level"),
                ref(f"location_factors.{index}.correction_rate"),
            ],
        }
        for index, item in location_factor_pairs
    ]
    location_rows.append(
        {
            "cells": ["合计", "/", "/", f"{sum((_decimal(item.get('correction_rate')) for item in location_factors if item.get('confirmed')), D0)}%"],
            "cell_refs": ["", "", "", ""],
        }
    )

    return [
        section("cost_building_compensation_rows", "表3-1 建筑物补偿标准", ["补偿项目名称", "附着物补偿标准（元/平方米）", "计算基数", "人均合法建筑面积（平方米）", "补偿金额（元）", "备注"], building_rows),
        section("cost_resettlement_population_rows", "表3-2 平均安置农业人口数确定表", ["项目名称", "位置", "用地面积（公顷）", "安置农业人口（人）", "安置人口数折合人/公顷"], population_rows),
        section("cost_acquisition_tax_rows", "表3-3 估价对象土地取得费及相关税费一览表", acquisition_tax_columns, acquisition_tax_rows),
        section("cost_development_rows", "表3-4道县城区土地开发费用分项一览表", development_columns, development_rows),
        section("cost_risk_impact_rows", "表3-5 风险因素因子对不同土地利用类型的影响程度分析", ["风险因素", "风险因子", risk_usage_label], risk_impact_rows, group_columns=[0], source_target="cost_adjustments"),
        section("cost_risk_weight_rows", "表3-6 风险因素、因子权重值", ["影响因素", "权重值", "影响因子", "权重值"], risk_weight_rows, group_columns=[0, 1], source_target="cost_adjustments"),
        section("cost_risk_assignment_rows", "表3-7 风险调整值赋值表", ["市场风险影响程度", "无影响（D）", "不严重（C）", "较严重（B）", "最严重（A）"], risk_assignment_rows, source_target="cost_adjustments"),
        section("cost_risk_result_rows", "表3-8 道县土地市场风险调整值结果", ["土地市场", risk_usage_label], risk_result_rows, source_target="cost_adjustments"),
        section(
            "cost_location_rule_rows",
            str(location_template.get("rule_table_title") or "表3-9 宗地地价区域个别因素修正系数表（%）"),
            ["因素", *location_level_headers, "等级修正幅度"],
            location_rule_rows,
            source_target="cost_adjustments",
        ),
        section(
            "cost_location_rows",
            str(location_template.get("result_table_title") or "表3-10 宗地地价影响因素条件说明、修正系数表"),
            ["因素", "因素说明", "因素优劣度", "因素修正"],
            location_rows,
            source_target="cost_adjustments",
        ),
    ]


def calculate_cost_approximation(
    data: Dict[str, Any],
    base_dir: str | Path,
    *,
    include_pricing_assistant: bool = False,
    include_catalog_metadata: bool = True,
    include_process_output: bool = True,
) -> Dict[str, Any]:
    result = deepcopy(data.get("cost_approx_analysis") or {})
    pricing_preview_mode = _bool(data.get("cost_pricing_preview_mode"))
    interactive_mode = _bool(data.get("cost_interactive_mode"))
    policy_config = load_cost_policy_config()
    result["policy_config_version"] = str(policy_config.get("version") or "")
    result["policy_config_hash"] = str(policy_config.get("source_hash") or "")
    policy = applicable_cost_basis(data, base_dir, include_catalog_metadata=include_catalog_metadata)
    land_class = str(policy.get("acquisition_land_class") or "耕地")
    land_subclass = str(policy.get("acquisition_land_subclass") or "水田")
    valuation_date_text = str(data.get("valuation_date") or "")
    result["acquisition_land_class"] = land_class
    result["acquisition_land_subclass"] = land_subclass
    result["acquisition_land_class_confirmed"] = _bool(data.get("acquisition_land_class_confirmed", True))
    result["policy_basis"] = policy
    result["local_compensation_policy_name"] = str(
        data.get("local_compensation_policy_name")
        or result.get("local_compensation_policy_name")
        or policy.get("local_compensation_policy_name")
        or ""
    )
    result["local_compensation_policy_no"] = str(
        data.get("local_compensation_policy_no")
        or result.get("local_compensation_policy_no")
        or policy.get("local_compensation_policy_no")
        or ""
    )
    result["local_compensation_policy_date"] = str(
        data.get("local_compensation_policy_date")
        or result.get("local_compensation_policy_date")
        or policy.get("local_compensation_policy_date")
        or ""
    )
    result["compensation_zone"] = str(policy.get("compensation_zone") or result.get("compensation_zone") or "Ⅰ")
    result["compensation_zone_options"] = list(policy.get("compensation_zone_options") or ["Ⅰ", "Ⅱ", "Ⅲ"])
    result["green_seedling_standard_source"] = str(policy.get("green_seedling_standard_source") or "")
    result["paid_land_use_fee_grade_standards"] = deepcopy(
        policy.get("paid_land_use_fee_grade_standards") or HUNAN_PAID_LAND_USE_FEE_STANDARDS
    )
    result["compensation_zone_override"] = _bool(policy.get("compensation_zone_override"))
    zone_suggestion = policy.get("compensation_zone_suggestion") or {}
    result["compensation_zone_suggestion"] = deepcopy(zone_suggestion)
    result["compensation_zone_source_path"] = str(zone_suggestion.get("source_path") or "")
    result["basis_land_location_full"] = str(data.get("land_location_full") or data.get("land_location") or "")
    current_context = {
        "basis_valuation_date": valuation_date_text,
        "basis_land_class": land_subclass,
        "basis_compensation_zone": str(result.get("compensation_zone") or ""),
        "basis_development_set": str(data.get("land_development_set") or "五通一平"),
    }
    confirmed_policy_or_cost = any(
        _bool(item.get("confirmed"))
        for key in ("policy_documents", "acquisition_items", "tax_items")
        for item in result.get(key) or []
    )
    confirmed_development = any(
        _bool(item.get("confirmed")) for item in result.get("development_items") or []
    )
    basis_context_mismatches: List[Dict[str, str]] = []
    basis_context_changes: List[Dict[str, str]] = []
    for key, label, target, is_frozen in (
        ("basis_valuation_date", "估价期日", "cost_policy", confirmed_policy_or_cost),
        ("basis_land_class", "征收地类细分类", "cost_items", confirmed_policy_or_cost),
        ("basis_compensation_zone", "征地区片", "cost_policy", confirmed_policy_or_cost),
        ("basis_development_set", "设定开发程度", "cost_items", confirmed_development),
    ):
        previous = str(result.get(key) or "")
        current = current_context[key]
        if previous and previous != current:
            basis_context_changes.append({"key": key, "label": label, "previous": previous, "current": current, "target": target})
            if is_frozen:
                basis_context_mismatches.append(
                    {
                        "level": "warning",
                        "message": f"{label}已由“{previous}”变更为“{current}”，系统已按新口径重新匹配标准值，请重新确认相关费用。",
                        "target": target,
                    }
                )
        result[key] = current
    policy_or_land_context_changed = any(
        item["key"] in {"basis_valuation_date", "basis_land_class", "basis_compensation_zone"}
        for item in basis_context_changes
    )
    tax_context_changed = any(
        item["key"] in {"basis_valuation_date", "basis_land_class"}
        for item in basis_context_changes
    )
    development_context_changed = any(item["key"] == "basis_development_set" for item in basis_context_changes)
    if policy_or_land_context_changed or not any(_bool(item.get("confirmed")) for item in result.get("acquisition_items") or []):
        result["green_seedling_standard_per_mu"] = str(policy.get("green_seedling_standard_per_mu") or "")
        result["green_seedling_standard_review_status"] = str(policy.get("green_seedling_standard_review_status") or "")
    else:
        result["green_seedling_standard_per_mu"] = str(
            result.get("green_seedling_standard_per_mu") or policy.get("green_seedling_standard_per_mu") or ""
        )
        result["green_seedling_standard_review_status"] = str(
            result.get("green_seedling_standard_review_status") or policy.get("green_seedling_standard_review_status") or ""
        )
    province = policy.get("province_compensation") or {}
    expected_city = _expected_city(data.get("county_name"), province)
    input_city = str(data.get("local_city") or "").strip()
    result["effective_local_city"] = expected_city or input_city
    selected_tax_grades = {
        str(item.get("key") or ""): str(item.get("grade_name") or "")
        for item in result.get("tax_items") or []
        if str(item.get("grade_name") or "").strip()
    }
    config_tax_items, config_tax_warnings = _configured_tax_items(
        policy_config,
        county_name=str(data.get("county_name") or ""),
        city_name=result["effective_local_city"],
        valuation_date=valuation_date_text,
        land_class=land_class,
        land_subclass=land_subclass,
        selected_grades=selected_tax_grades,
    )
    generated_policy = {
        "key": "hunan_compensation_2024",
        "name": PROVINCE_POLICY_NAME,
        "document_no": PROVINCE_POLICY_NO,
        "publish_date": "2024-01-31",
        "effective_date": "2024-01-31",
        "expiry_date": "",
        "region": "湖南省",
        "source_path": province.get("source_path") or "",
        "source_hash": province.get("source_hash") or "",
        "valuation_date": valuation_date_text,
        "applicable": bool(province.get("applicable", True)),
        "confirmed": False,
        "role": "province_compensation",
        "source_type": "system_recommendation",
        "reference_text": f"{PROVINCE_POLICY_NAME}（{PROVINCE_POLICY_NO}）",
        "enabled": True,
        "replaces_key": "",
        "note": "",
    }
    generated_local_policy = {
        "key": "local_compensation_policy",
        "name": result["local_compensation_policy_name"],
        "document_no": result["local_compensation_policy_no"],
        "publish_date": result["local_compensation_policy_date"],
        "effective_date": result["local_compensation_policy_date"],
        "expiry_date": "",
        "region": result["effective_local_city"],
        "source_path": str(
            _local_compensation_policy_xlsx_path(base_dir)
            if _local_compensation_policy_xlsx_path(base_dir).exists()
            else _local_compensation_policy_path(base_dir)
        )
        if (_local_compensation_policy_xlsx_path(base_dir).exists() or _local_compensation_policy_path(base_dir).exists())
        else "",
        "source_hash": _document_hash(
            _local_compensation_policy_xlsx_path(base_dir)
            if _local_compensation_policy_xlsx_path(base_dir).exists()
            else _local_compensation_policy_path(base_dir)
        )
        if (_local_compensation_policy_xlsx_path(base_dir).exists() or _local_compensation_policy_path(base_dir).exists())
        else "",
        "valuation_date": valuation_date_text,
        "applicable": True,
        "confirmed": False,
        "role": "local_compensation",
        "source_type": "system_recommendation",
        "reference_text": result["local_compensation_policy_name"],
        "enabled": True,
        "replaces_key": "",
        "note": (
            "；".join(
                validate_yongzhou_local_compensation_xlsx(
                    base_dir,
                    (_load_yongzhou_building_compensation_config(base_dir).get("rows") or []),
                ).get("warnings")
                or []
            )
            if _local_compensation_policy_xlsx_path(base_dir).exists() and "道县" in str(data.get("county_name") or "")
            else str((_load_yongzhou_building_compensation_config(base_dir).get("review_note") or "")).strip()
            if _local_compensation_policy_path(base_dir).exists() and "道县" in str(data.get("county_name") or "")
            else ""
        ),
    }
    scope_policy_path = _scope_policy_path(base_dir)
    scope_table_path = _scope_table_path(base_dir, _expected_city(data.get("county_name"), province))
    generated_scope_policy = {
        "key": "hunan_scope_2024",
        "name": PROVINCE_SCOPE_POLICY_NAME,
        "document_no": PROVINCE_SCOPE_POLICY_NO,
        "publish_date": "2024-02-01",
        "effective_date": "2024-02-01",
        "expiry_date": "",
        "region": "湖南省",
        "source_path": str(scope_table_path if scope_table_path.exists() else scope_policy_path),
        "source_hash": _document_hash(scope_table_path if scope_table_path.exists() else scope_policy_path),
        "valuation_date": valuation_date_text,
        "applicable": True,
        "confirmed": False,
        "role": "province_scope",
        "source_type": "system_recommendation",
        "reference_text": f"{PROVINCE_SCOPE_POLICY_NAME}（{PROVINCE_SCOPE_POLICY_NO}）",
        "enabled": True,
        "replaces_key": "",
        "note": "",
    }
    if not result.get("policy_documents"):
        result["policy_documents"] = [generated_policy, generated_scope_policy]
        if generated_local_policy["name"]:
            result["policy_documents"].append(generated_local_policy)
    else:
        result["policy_documents"] = [
            (
                {**item, **generated_policy, "confirmed": False}
                if item.get("key") == "hunan_compensation_2024" and not _bool(item.get("confirmed"))
                else item
            )
            for item in result.get("policy_documents") or []
        ]
        if not any(item.get("key") == "hunan_compensation_2024" for item in result["policy_documents"]):
            result["policy_documents"].append(generated_policy)
        if not any(item.get("key") == "hunan_scope_2024" for item in result["policy_documents"]):
            result["policy_documents"].append(generated_scope_policy)
        if generated_local_policy["name"] and not any(item.get("key") == "local_compensation_policy" for item in result["policy_documents"]):
            result["policy_documents"].append(generated_local_policy)
        result["policy_documents"] = [
            (
                {**item, **generated_local_policy, "confirmed": False}
                if item.get("key") == "local_compensation_policy" and not _bool(item.get("confirmed"))
                else item
            )
            for item in result["policy_documents"]
        ]
    for item in result.get("policy_documents") or []:
        item.setdefault("role", "supplemental")
        item.setdefault("source_type", "system_recommendation" if item.get("key") in {"hunan_compensation_2024", "hunan_scope_2024", "local_compensation_policy"} else "manual")
        item.setdefault("reference_text", item.get("name") or "")
        item.setdefault("enabled", True)
        item.setdefault("replaces_key", "")
        item.setdefault("note", "")
    replacement_keys = {
        str(item.get("replaces_key") or "")
        for item in result.get("policy_documents") or []
        if item.get("enabled", True) and str(item.get("replaces_key") or "").strip()
    }
    for item in result.get("policy_documents") or []:
        if item.get("key") in replacement_keys:
            item["enabled"] = False
    recommended = province.get("recommended_amount_per_sqm") or ""
    compensation_coefficient = province.get("coefficient") or "1"
    default_acquisition_items = [
        _item(
            "land_compensation",
            "征地补偿费和安置补助费",
            "acquisition",
            recommended,
            standard_value=province.get("base_per_mu") or "",
            standard_unit="元/亩",
            coefficient=compensation_coefficient,
            formula="区片标准×地类修正系数×0.0015",
            policy_key="hunan_compensation_2024",
            grade_name=f"{result['compensation_zone']}区",
            source_note=f"{data.get('county_name') or '所属县市'}征地区片{result['compensation_zone']}推荐值",
        ),
        _item(
            "building_compensation",
            "地上建筑物补偿",
            "acquisition",
            "0.00",
            formula="建筑物补偿标准×安置农业人口密度÷10000",
            source="structured_attachment_analysis",
            source_note="建筑物补偿标准与安置农业人口密度结构化测算",
        ),
        _item(
            "seedling_compensation",
            "青苗补偿",
            "acquisition",
            _money(_decimal(result.get("green_seedling_standard_per_mu")) * MU_TO_SQM),
            standard_value=result.get("green_seedling_standard_per_mu") or "",
            standard_unit="元/亩",
            formula="青苗标准×0.0015",
            source="structured_attachment_analysis",
            source_note="地方政策青苗补偿标准折算",
        ),
        _item(
            "ground_attachment",
            "地上附着物补偿",
            "acquisition",
            _money(_decimal(result.get("green_seedling_standard_per_mu")) * MU_TO_SQM),
            standard_value=result.get("green_seedling_standard_per_mu") or "",
            standard_unit="元/亩",
            formula="建筑物补偿＋青苗补偿",
            source="compat_derived",
            source_note="兼容字段：由地上建筑物补偿和青苗补偿派生，不参与新计算。",
            enabled=False,
        ),
    ]
    if not result.get("acquisition_items"):
        result["acquisition_items"] = default_acquisition_items
    else:
        default_by_key = {item["key"]: item for item in default_acquisition_items}
        existing_keys = {item.get("key") for item in result.get("acquisition_items") or []}
        for default_item in default_acquisition_items:
            if default_item["key"] not in existing_keys:
                result["acquisition_items"].append(deepcopy(default_item))
        for item in result.get("acquisition_items") or []:
            defaults = default_by_key.get(item.get("key"))
            if defaults:
                if policy_or_land_context_changed:
                    _refresh_item_from_defaults(item, defaults, force=True)
                else:
                    _fill_missing_item_fields(item, defaults)
                _recalculate_formula_amount(item, defaults)
            if item.get("key") == "land_compensation" and not _bool(item.get("confirmed")):
                item.update(
                    {
                        "standard_value": province.get("base_per_mu") or "",
                        "coefficient": compensation_coefficient,
                        "formula": "区片标准×地类修正系数×0.0015",
                        "grade_name": f"{result['compensation_zone']}区",
                        "source_note": f"{data.get('county_name') or '所属县市'}征地区片{result['compensation_zone']}推荐值",
                    }
                )
                _recalculate_formula_amount(item, defaults or {})
            if item.get("key") == "seedling_compensation" and not _bool(item.get("confirmed")):
                item.update(
                    {
                        "standard_value": result.get("green_seedling_standard_per_mu") or "",
                        "standard_unit": "元/亩",
                        "amount_per_sqm": _money(_decimal(result.get("green_seedling_standard_per_mu")) * MU_TO_SQM),
                        "computed_amount_per_sqm": _money(_decimal(result.get("green_seedling_standard_per_mu")) * MU_TO_SQM),
                        "formula": "青苗标准×0.0015",
                    }
                )
            if item.get("key") == "land_compensation":
                item["formula"] = "区片标准×地类修正系数×0.0015"
                _recalculate_formula_amount(item, defaults or {})
    for item in result.get("acquisition_items") or []:
        if item.get("key") == "ground_attachment":
            item["label"] = "地上附着物补偿"
            item["enabled"] = False
    if not result.get("tax_items"):
        result["tax_items"] = config_tax_items
    elif tax_context_changed:
        result["tax_items"] = config_tax_items
    else:
        default_tax_by_key = {item["key"]: item for item in config_tax_items}
        existing_tax_keys = {item.get("key") for item in result.get("tax_items") or []}
        for default_item in default_tax_by_key.values():
            if default_item["key"] not in existing_tax_keys:
                result["tax_items"].append(deepcopy(default_item))
        for item in result.get("tax_items") or []:
            defaults = default_tax_by_key.get(item.get("key"))
            if defaults:
                tax_grade_refresh_needed = item.get("key") in (
                    "farmland_reclamation_fee",
                    "water_conservancy_fund",
                    "farmland_occupation_tax",
                ) and (
                    not _same_text(item.get("rule_key"), defaults.get("rule_key"))
                    or not _same_text(item.get("grade_name"), defaults.get("grade_name"))
                    or not _same_text(item.get("standard_value"), defaults.get("standard_value"))
                    or not _same_text(item.get("amount_per_sqm"), defaults.get("amount_per_sqm"))
                )
                if tax_context_changed or tax_grade_refresh_needed:
                    _refresh_item_from_defaults(item, defaults, force=True)
                else:
                    _fill_missing_item_fields(item, defaults)
                _recalculate_formula_amount(item, defaults)
    if not result.get("development_items"):
        result["development_items"] = _default_development_items(str(data.get("land_development_set") or "五通一平"))
    elif development_context_changed:
        result["development_items"] = _default_development_items(str(data.get("land_development_set") or "五通一平"))
    default_development_by_label = {
        item["label"]: item
        for item in _default_development_items(str(data.get("land_development_set") or "五通一平"))
    }
    for item in result.get("development_items") or []:
        defaults = default_development_by_label.get(item.get("label")) or {}
        for key in ("range_min", "range_max"):
            if item.get(key) in (None, "") and defaults.get(key) not in (None, ""):
                item[key] = defaults[key]
    active_replacements = [
        item
        for item in result.get("policy_documents") or []
        if item.get("enabled", True) and str(item.get("replaces_key") or "").strip()
    ]
    replacement_item_keys = {
        "province_compensation": ("land_compensation",),
        "local_compensation": ("building_compensation", "seedling_compensation", "ground_attachment"),
        "tax_policy": tuple(item.get("key") for item in result.get("tax_items") or []),
    }
    for replacement in active_replacements:
        policy_key = str(replacement.get("key") or "")
        affected_keys = replacement_item_keys.get(str(replacement.get("role") or ""), ())
        for group_key in ("acquisition_items", "tax_items", "development_items"):
            for item in result.get(group_key) or []:
                if item.get("key") not in affected_keys or item.get("policy_key") == policy_key:
                    continue
                item["policy_key"] = policy_key
                item["confirmed"] = False
                item["source"] = "manual_policy_replacement"
                item["source_note"] = replacement.get("reference_text") or replacement.get("name") or ""
                item["standard_value"] = ""
                item["coefficient"] = ""
                item["amount_per_sqm"] = ""
    result.setdefault("external_results", [])
    result.setdefault("development_cycle_years", "1")
    result.setdefault("interest_rate", "3")
    result.setdefault("acquisition_investment_fraction", "1")
    result.setdefault("development_investment_fraction", "0.5")
    result.setdefault("location_correction_mode", "direct_sum")
    result.setdefault("risk_mode", "direct")
    result.setdefault("risk_scheme_key", "hunan_general")
    result["usage_scenarios"] = _usage_scenarios(data, result.get("usage_scenarios") or [])
    if not result.get("risk_items"):
        result["risk_items"] = _default_risk_items(policy_config, result["usage_scenarios"], str(result.get("risk_scheme_key") or "hunan_general"))
    if not result.get("risk_groups"):
        result["risk_groups"] = _default_risk_groups(policy_config, result["usage_scenarios"], str(result.get("risk_scheme_key") or "hunan_general"))
    result.setdefault("building_compensation_rows", building_compensation_rows_for_county(base_dir, str(data.get("county_name") or "")))
    # Catalog / 一览 / 清单 are static UI metadata (pure functions of county+policy
    # files). The pricing-assistant scenario recursion only needs the numeric
    # results, so skip rebuilding them there to keep interactive /calculate fast.
    if include_catalog_metadata:
        result["building_compensation_policy_help"] = building_compensation_policy_help(base_dir, str(data.get("county_name") or ""))
        result["building_compensation_add_catalog"] = building_compensation_add_catalog(base_dir, str(data.get("county_name") or ""))
        result["cost_basis_attachment_inventory"] = cost_basis_attachment_inventory(base_dir, str(data.get("county_name") or ""))
    else:
        result.pop("building_compensation_policy_help", None)
        result.pop("building_compensation_add_catalog", None)
        result.pop("cost_basis_attachment_inventory", None)
    result.setdefault("resettlement_population_cases", _default_resettlement_population_cases())
    if not result.get("development_survey_cases"):
        result["development_survey_cases"] = _default_development_survey_cases()
    result["development_survey_analysis"] = _development_survey_analysis(result.get("development_survey_cases") or [])
    default_building_by_key = {
        item["key"]: item
        for item in building_compensation_rows_for_county(base_dir, str(data.get("county_name") or ""))
    }
    building_policy_reference = "（".join(
        filter(
            None,
            [
                str(result.get("local_compensation_policy_name") or ""),
                str(result.get("local_compensation_policy_no") or ""),
            ],
        )
    )
    if building_policy_reference and "（" in building_policy_reference:
        building_policy_reference += "）"
    for row in result.get("building_compensation_rows") or []:
        defaults = default_building_by_key.get(row.get("key")) or {}
        for key, value in defaults.items():
            if row.get(key) in (None, ""):
                row[key] = value
        if defaults.get("source_path") and not row.get("source_path"):
            row["source_path"] = defaults.get("source_path")
        if defaults.get("source_hash") and not row.get("source_hash"):
            row["source_hash"] = defaults.get("source_hash")
        row.setdefault("review_status", defaults.get("review_status") or "")
        row.setdefault("pending_review", defaults.get("pending_review", False))
        row.setdefault("source_note", defaults.get("source_note") or building_policy_reference or "市县征收补偿安置政策，待校核")
        if defaults.get("grade_options") and not row.get("grade_options"):
            row["grade_options"] = deepcopy(defaults.get("grade_options"))
    result["attachment_compensation_analysis"] = _recalculate_attachment_analysis(result)
    default_template_key = default_location_template_key([item.get("key") for item in result["usage_scenarios"]])
    requested_template_key = str(result.get("location_template_key") or default_template_key)
    location_template = get_location_template(policy_config, requested_template_key) or get_location_template(policy_config, default_template_key)
    if not location_template:
        requested_template_key = "custom"
        location_template = get_location_template(policy_config, requested_template_key)
    previous_template_key = str((result.get("location_template_snapshot") or {}).get("key") or result.get("location_template_key") or "")
    location_template["key"] = requested_template_key
    result["location_template_key"] = requested_template_key
    result["location_template_snapshot"] = deepcopy(location_template)
    if not result.get("location_factors") or (previous_template_key and previous_template_key != requested_template_key):
        previous_by_key = {
            str(item.get("key") or item.get("label") or ""): dict(item)
            for item in result.get("location_factors") or []
        }
        generated_factors = _default_location_factors(location_template, result["usage_scenarios"])
        for factor in generated_factors:
            old = previous_by_key.get(str(factor.get("key") or factor.get("label") or ""))
            if old:
                for keep in ("description", "level", "correction_rate", "confirmed", "usage_key", "correction_rate_manual", "enabled"):
                    if old.get(keep) not in (None, ""):
                        factor[keep] = old.get(keep)
        result["location_factors"] = generated_factors
    if result.get("location_correction_mode") == "direct_sum":
        for item in result.get("location_factors") or []:
            if item.get("grade_amplitude") not in (None, "") and not _bool(item.get("correction_rate_manual")):
                item["correction_rate"] = _money(_grade_correction(item.get("level"), item.get("grade_amplitude"), item.get("levels") or []))
    risk_scheme = get_risk_scheme(policy_config, str(result.get("risk_scheme_key") or "hunan_general"))
    assignment = risk_scheme.get("assignment") or RISK_LEVEL_ADJUSTMENTS
    for item in result.get("risk_items") or []:
        if item.get("level") and item.get("adjustment_rate") in (None, ""):
            item["adjustment_rate"] = str(assignment.get(str(item.get("level")), ""))
        item.setdefault(
            "level_options",
            [{"level": level, "adjustment_rate": str(rate)} for level, rate in assignment.items()],
        )
    result["risk_groups"] = _sync_risk_groups(result.get("risk_items") or [], result.get("risk_groups") or [])

    attachment_analysis = result["attachment_compensation_analysis"]
    building_amount = attachment_analysis.get("building_compensation_per_sqm")
    seedling_amount = attachment_analysis.get("green_seedling_per_sqm")
    attachment_amount = attachment_analysis.get("attachment_compensation_per_sqm")
    for item in result["acquisition_items"]:
        if item.get("key") == "building_compensation":
            item["label"] = "地上建筑物补偿"
            item["standard_value"] = attachment_analysis.get("building_compensation_per_person")
            item["standard_unit"] = "元/人"
            item["amount_per_sqm"] = building_amount
            item["computed_amount_per_sqm"] = building_amount
            item["source"] = "structured_attachment_analysis"
            item["source_note"] = "建筑物补偿标准、安置农业人口密度结构化测算"
        if item.get("key") == "seedling_compensation" and item.get("source") != "manual_policy_replacement":
            item["label"] = "青苗补偿"
            item["amount_per_sqm"] = seedling_amount
            item["computed_amount_per_sqm"] = seedling_amount
            item["source"] = "structured_attachment_analysis"
            item["source_note"] = "地方政策青苗补偿标准折算"
        if item.get("key") == "ground_attachment":
            item["label"] = "地上附着物补偿"
            item["amount_per_sqm"] = attachment_amount
            item["computed_amount_per_sqm"] = attachment_amount
            item["enabled"] = False
            item["source"] = "compat_derived"
            item["source_note"] = "兼容字段：由地上建筑物补偿和青苗补偿派生，不参与新计算。"
    for item in result.get("acquisition_items") or []:
        if not item.get("enabled", True):
            item["computed_amount_per_sqm"] = item.get("computed_amount_per_sqm") or item.get("amount_per_sqm") or "0.00"
            item["amount_per_sqm"] = "0.00" if item.get("key") != "ground_attachment" else item.get("computed_amount_per_sqm")
    for group_key in ("tax_items", "development_items"):
        for item in result.get(group_key) or []:
            item["computed_amount_per_sqm"] = item.get("computed_amount_per_sqm") or item.get("amount_per_sqm") or "0.00"
            if not item.get("enabled", True):
                item["amount_per_sqm"] = "0.00"

    acquisition_total = _qmoney(_sum_items(result["acquisition_items"]))
    tax_total = _qmoney(_sum_items(result["tax_items"]))
    development_total = _qmoney(_sum_items(result["development_items"]))
    cycle = _decimal(result.get("development_cycle_years"), D1)
    interest_rate = _percent_decimal(result.get("interest_rate"))
    acquisition_fraction = _decimal(result.get("acquisition_investment_fraction"), D1)
    development_fraction = _decimal(result.get("development_investment_fraction"), Decimal("0.5"))
    interest = _qmoney(
        (acquisition_total + tax_total) * cycle * acquisition_fraction * interest_rate
        + development_total * cycle * development_fraction * interest_rate
    )

    usage_results = []
    rounding_trace: List[Dict[str, str]] = [
        {"key": "acquisition_total", "label": "土地取得费", "value": _money(acquisition_total)},
        {"key": "tax_total", "label": "相关税费", "value": _money(tax_total)},
        {"key": "development_total", "label": "土地开发费", "value": _money(development_total)},
        {"key": "interest", "label": "投资利息", "value": _money(interest)},
    ]
    for scenario in result["usage_scenarios"]:
        profit_rate = _percent_decimal(scenario.get("profit_rate"))
        value_added_rate = _percent_decimal(scenario.get("value_added_rate"))
        matching_risk_items = [
            item
            for item in result.get("risk_items") or []
            if item.get("confirmed") and (not item.get("usage_key") or item.get("usage_key") == scenario.get("key"))
        ]
        matching_risk_groups = [
            item
            for item in result.get("risk_groups") or []
            if (
                _bool(item.get("confirmed"))
                or (pricing_preview_mode and _bool(item.get("override_enabled")))
            )
            and item.get("effective_value") not in (None, "")
            and (not item.get("usage_key") or item.get("usage_key") == scenario.get("key"))
        ]
        if result.get("risk_mode") == "analysis" and matching_risk_groups and scenario.get("safe_rate") not in (None, ""):
            risk_adjustment = _risk_adjustment_value_from_groups(
                matching_risk_groups,
                include_unconfirmed=pricing_preview_mode,
            )
            scenario["reduction_rate"] = _money(_decimal(scenario.get("safe_rate")) + risk_adjustment)
        elif matching_risk_items and scenario.get("safe_rate") not in (None, ""):
            risk_adjustment = _risk_adjustment_value(matching_risk_items)
            scenario["reduction_rate"] = _money(_decimal(scenario.get("safe_rate")) + risk_adjustment)
        reduction_rate = _percent_decimal(scenario.get("reduction_rate"))
        term = _decimal(scenario.get("use_term_years"))
        matching_location_factors = [
            item
            for item in result.get("location_factors") or []
            if item.get("enabled") is not False
            and (
                _bool(item.get("confirmed"))
                or (
                    (pricing_preview_mode or interactive_mode)
                    and (
                        _bool(item.get("correction_rate_manual"))
                        or str(item.get("level") or "").strip()
                    )
                )
            )
            and (not item.get("usage_key") or item.get("usage_key") == scenario.get("key"))
        ]
        location_factor_total = sum(
            (_decimal(item.get("correction_rate")) for item in matching_location_factors),
            D0,
        )
        location_rate = _percent_decimal(
            location_factor_total if matching_location_factors else scenario.get("location_correction_rate")
        )
        if matching_location_factors:
            scenario["location_correction_rate"] = _money(location_factor_total)
        profit = _qmoney((acquisition_total + tax_total + development_total) * profit_rate)
        cost_price = _qmoney(acquisition_total + tax_total + development_total + interest + profit)
        value_added = _qmoney(cost_price * value_added_rate)
        unlimited_price = _qmoney(cost_price + value_added)
        term_factor = _qfactor(D1 - (D1 / ((D1 + reduction_rate) ** term))) if reduction_rate > D0 and term > D0 else D1
        comparable_price = _qmoney(unlimited_price * term_factor)
        final_price = _qprice(comparable_price * (D1 + location_rate))
        rounding_trace.extend(
            [
                {"key": f"{scenario.get('key')}.profit", "label": f"{scenario.get('label')}投资利润", "value": _money(profit)},
                {"key": f"{scenario.get('key')}.cost_price", "label": f"{scenario.get('label')}土地成本价格", "value": _money(cost_price)},
                {"key": f"{scenario.get('key')}.value_added_income", "label": f"{scenario.get('label')}土地增值收益", "value": _money(value_added)},
                {"key": f"{scenario.get('key')}.unlimited_price", "label": f"{scenario.get('label')}无限年期价格", "value": _money(unlimited_price)},
                {"key": f"{scenario.get('key')}.term_correction_factor", "label": f"{scenario.get('label')}年期修正系数", "value": _factor(term_factor)},
                {"key": f"{scenario.get('key')}.comparable_price", "label": f"{scenario.get('label')}比准价格", "value": _money(comparable_price)},
                {"key": f"{scenario.get('key')}.final_price", "label": f"{scenario.get('label')}最终单价", "value": _price(final_price)},
            ]
        )
        usage_results.append(
            {
                "key": scenario.get("key") or "other",
                "label": scenario.get("label") or "其他用地",
                "acquisition_total": _money(acquisition_total),
                "tax_total": _money(tax_total),
                "development_total": _money(development_total),
                "interest": _money(interest),
                "profit": _money(profit),
                "cost_price": _money(cost_price),
                "value_added_income": _money(value_added),
                "unlimited_price": _money(unlimited_price),
                "term_correction_factor": _factor(term_factor),
                "comparable_price": _money(comparable_price),
                "location_correction_rate": _money(location_rate * Decimal("100")),
                "final_price": _price(final_price),
            }
        )
    result["usage_results"] = usage_results
    result["totals"] = {
        "acquisition_total": _money(acquisition_total),
        "tax_total": _money(tax_total),
        "acquisition_and_tax_total": _money(acquisition_total + tax_total),
        "development_total": _money(development_total),
        "interest": _money(interest),
    }
    result["rounding_trace"] = rounding_trace

    warnings: List[Dict[str, str]] = [
        {"level": "warning", "message": message, "target": "cost_policy"}
        for message in policy.get("warnings") or []
    ]
    warnings.extend(
        {"level": "warning", "message": message, "target": "cost_items"}
        for message in config_tax_warnings
    )
    configured_grade_by_key = {
        str(item.get("key") or ""): str(item.get("grade_name") or "")
        for item in config_tax_items
    }
    for item in result.get("tax_items") or []:
        configured_grade = configured_grade_by_key.get(str(item.get("key") or ""), "")
        current_grade = str(item.get("grade_name") or "")
        if configured_grade and current_grade and configured_grade != current_grade:
            warnings.append(
                {
                    "level": "warning",
                    "message": f"{item.get('label') or '相关税费'}等别/口径已由系统匹配值“{configured_grade}”修改为“{current_grade}”，请同步核对标准值及依据文件。",
                    "target": "cost_items",
                }
            )
    water_fund_item = next(
        (item for item in result.get("tax_items") or [] if item.get("key") == "water_conservancy_fund"),
        None,
    )
    if water_fund_item:
        county_paid_rule = find_policy_rule(
            policy_config,
            "paid_land_use_fee",
            county_name=str(data.get("county_name") or ""),
            city_name=result["effective_local_city"],
            valuation_date=valuation_date_text,
            land_subclass=land_subclass,
        )
        county_default_grade = str(county_paid_rule.get("grade_name") or "") if county_paid_rule else ""
        selected_water_grade = str(water_fund_item.get("grade_name") or "")
        paid_standard = str(water_fund_item.get("standard_value") or "").strip()
        if (
            county_default_grade
            and selected_water_grade
            and county_default_grade != selected_water_grade
            and paid_standard
        ):
            county_label = str(data.get("county_name") or "当前县市")
            warnings.append(
                {
                    "level": "warning",
                    "message": (
                        f"水利建设基金等别已选「{selected_water_grade}」，"
                        f"依据湘财综〔2018〕42号附件1新增建设用地土地有偿使用费标准{paid_standard}元/平方米折算；"
                        f"{county_label}默认等别为「{county_default_grade}」，请估价师确认。"
                    ),
                    "target": "cost_items",
                }
            )
    warnings.extend(basis_context_mismatches)
    if input_city and expected_city and input_city != expected_city:
        warnings.append(
            {
                "level": "danger",
                "message": f"所属市级“{input_city}”与{data.get('county_name') or '当前县市'}对应的“{expected_city}”不一致，正文已按{expected_city}生成，请返回第二部分校正。",
                "target": "cost_policy",
            }
        )
    if not result.get("acquisition_land_class_confirmed"):
        warnings.append({"level": "warning", "message": "征收地类尚未确认。", "target": "cost_policy"})
    for group_key, group_label in (
        ("acquisition_items", "土地取得费"),
        ("tax_items", "相关税费"),
        ("development_items", "土地开发费"),
    ):
        pending = [item.get("label") for item in result.get(group_key) or [] if item.get("enabled", True) and not _bool(item.get("confirmed"))]
        if pending:
            warnings.append({"level": "warning", "message": f"{group_label}仍有未确认项目：{'、'.join(str(item) for item in pending)}。", "target": "cost_items"})
    pending_policies = [
        item.get("name") or item.get("document_no")
        for item in result.get("policy_documents") or []
        if item.get("enabled", True) and not _bool(item.get("confirmed"))
    ]
    if pending_policies:
        warnings.append({"level": "warning", "message": f"政策依据尚未确认：{'、'.join(str(item) for item in pending_policies)}。", "target": "cost_policy"})
    inapplicable_policies = [
        item.get("name") or item.get("document_no")
        for item in result.get("policy_documents") or []
        if item.get("enabled", True) and item.get("applicable") is False
    ]
    if inapplicable_policies:
        warnings.append({"level": "danger", "message": f"政策文件不适用于当前估价期日：{'、'.join(str(item) for item in inapplicable_policies)}。", "target": "cost_policy"})
    policy_date_mismatches = [
        item.get("name") or item.get("document_no")
        for item in result.get("policy_documents") or []
        if item.get("enabled", True)
        and _bool(item.get("confirmed"))
        and item.get("valuation_date")
        and _date(item.get("valuation_date")) != _date(valuation_date_text)
    ]
    if policy_date_mismatches:
        warnings.append(
            {
                "level": "danger",
                "message": f"估价期日已变化，冻结政策依据需重新校核：{'、'.join(str(item) for item in policy_date_mismatches)}。",
                "target": "cost_policy",
            }
        )
    structured_attachment_required = _by_key(result.get("acquisition_items") or [], "building_compensation").get("enabled", True)
    if structured_attachment_required and any(not _bool(item.get("confirmed")) for item in result.get("building_compensation_rows") or []):
        warnings.append({"level": "warning", "message": "建筑物补偿标准仍有未确认项目。", "target": "cost_items"})
    if structured_attachment_required and any(not _bool(item.get("confirmed")) for item in result.get("resettlement_population_cases") or []):
        warnings.append({"level": "warning", "message": "平均安置农业人口案例仍有未确认项目。", "target": "cost_items"})
    survey_analysis = result.get("development_survey_analysis") or {}
    survey_status = str(survey_analysis.get("status") or "pending")
    if survey_status == "pending":
        warnings.append(
            {
                "level": "warning",
                "message": "土地开发费尚未附项目调查资料；分项采用值仅作参考区间上限，正式报告前请录入并确认≥3组调查案例。",
                "target": "cost_development_survey",
            }
        )
    elif survey_status == "insufficient":
        warnings.append(
            {
                "level": "warning",
                "message": f"土地开发费调查案例仅确认 {survey_analysis.get('valid_count') or '0'} 组，建议至少 3 组同区域、同开发程度调查资料。",
                "target": "cost_development_survey",
            }
        )
    if any(not _bool(item.get("confirmed")) for item in result.get("development_survey_cases") or []):
        warnings.append({"level": "warning", "message": "土地开发费调查案例仍有未确认项目。", "target": "cost_development_survey"})
    pending_scenarios = [item.get("label") for item in result["usage_scenarios"] if not _bool(item.get("confirmed"))]
    if pending_scenarios:
        warnings.append({"level": "warning", "message": f"用途场景参数尚未确认：{'、'.join(str(item) for item in pending_scenarios)}。", "target": "cost_scenarios"})
    enabled_location_factors = _enabled_location_factors(result)
    if not enabled_location_factors:
        warnings.append({"level": "warning", "message": "尚未录入区位修正因素；当前按0%计算。", "target": "cost_location"})
    elif any(not _bool(item.get("confirmed")) for item in enabled_location_factors):
        warnings.append({"level": "warning", "message": "区位修正因素仍有未确认项目。", "target": "cost_location"})
    if result.get("risk_mode") == "analysis":
        if any(not _bool(item.get("confirmed")) for item in result.get("risk_items") or []):
            warnings.append({"level": "warning", "message": "土地还原率风险因素仍有未确认项目。", "target": "cost_scenarios"})
        if any(not _bool(item.get("confirmed")) for item in result.get("risk_groups") or []):
            warnings.append({"level": "warning", "message": "土地还原率风险组综合值仍未确认。", "target": "cost_scenarios"})
    green_seedling_match = green_seedling_standard_for_county(
        base_dir,
        str(data.get("county_name") or ""),
        str(result.get("acquisition_land_subclass") or land_subclass or ""),
    )
    seedling_item = _by_key(result.get("acquisition_items") or [], "seedling_compensation")
    manual_seedling_override = (
        seedling_item.get("source") == "manual_policy_replacement"
        and _bool(seedling_item.get("confirmed"))
        and str(seedling_item.get("amount_per_sqm") or "").strip() not in {"", "0", "0.00"}
    )
    pending_green_seedling = green_seedling_match.get("review_status") == "pending" and not manual_seedling_override
    if pending_green_seedling and str(green_seedling_match.get("standard_per_mu") or "").strip():
        warnings.append(
            {
                "level": "warning",
                "message": f"青苗补偿标准沿用默认口径（{green_seedling_match.get('standard_per_mu')}元/亩），待匹配县市本地政策。",
                "target": "cost_policy",
            }
        )
    result["warnings"] = warnings
    result["complete"] = (
        bool([item for item in result.get("policy_documents") or [] if item.get("enabled", True)])
        and all(_bool(item.get("confirmed")) for item in result.get("policy_documents") or [] if item.get("enabled", True))
        and all(item.get("applicable", True) for item in result.get("policy_documents") or [] if item.get("enabled", True))
        and not policy_date_mismatches
        and not basis_context_mismatches
        and (not input_city or not expected_city or input_city == expected_city)
        and result.get("acquisition_land_class_confirmed")
        and _all_items_resolved(result["acquisition_items"])
        and _all_items_resolved(result["tax_items"])
        and _all_items_resolved(result["development_items"])
        and (
            not structured_attachment_required
            or (
                all(_bool(item.get("confirmed")) for item in result.get("building_compensation_rows") or [])
                and all(_bool(item.get("confirmed")) for item in result.get("resettlement_population_cases") or [])
            )
        )
        and all(_bool(item.get("confirmed")) for item in result["usage_scenarios"])
        and bool(enabled_location_factors)
        and all(_bool(item.get("confirmed")) for item in enabled_location_factors)
        and (
            result.get("risk_mode") != "analysis"
            or (
                all(_bool(item.get("confirmed")) for item in result.get("risk_items") or [])
                and all(_bool(item.get("confirmed")) for item in result.get("risk_groups") or [])
            )
        )
        and not pending_green_seedling
    )
    if include_process_output:
        generated = _build_narratives(data, result)
        overrides = result.get("narrative_overrides") or {}
        result["generated_narratives"] = generated
        result["effective_narratives"] = {key: str(overrides.get(key) or value) for key, value in generated.items()}
        result["narrative_segment_sources"] = _cost_narrative_segment_sources(data, result)
        result["tables"] = _table_sections(result)
        content_blocks = [
            {"type": "narrative", "key": "cost_approx_method_intro"},
            {"type": "narrative", "key": "cost_approx_basis_intro"},
            {"type": "narrative", "key": "cost_approx_acquisition_intro"},
            {"type": "table", "key": "cost_building_compensation_rows"},
            {"type": "narrative", "key": "cost_approx_attachment_population_narrative"},
            {"type": "table", "key": "cost_resettlement_population_rows"},
            {"type": "narrative", "key": "cost_approx_acquisition_solve"},
            {"type": "narrative", "key": "cost_approx_tax_narrative"},
            {"type": "table", "key": "cost_acquisition_tax_rows"},
            {"type": "narrative", "key": "cost_approx_development_narrative"},
            {"type": "table", "key": "cost_development_rows"},
            {"type": "narrative", "key": "cost_approx_interest_narrative"},
            {"type": "narrative", "key": "cost_approx_profit_narrative"},
            {"type": "narrative", "key": "cost_approx_value_added_narrative"},
            {"type": "narrative", "key": "cost_approx_term_narrative"},
        ]
        if result.get("risk_mode") == "analysis":
            content_blocks.extend([
            {"type": "table", "key": "cost_risk_impact_rows"},
            {"type": "table", "key": "cost_risk_weight_rows"},
            {"type": "table", "key": "cost_risk_assignment_rows"},
            {"type": "table", "key": "cost_risk_result_rows"},
            ])
        content_blocks.extend([
            {"type": "narrative", "key": "cost_approx_location_narrative"},
            {"type": "table", "key": "cost_location_rule_rows"},
            {"type": "table", "key": "cost_location_rows"},
            {"type": "narrative", "key": "cost_approx_solve_narrative"},
        ])
        result["content_blocks"] = content_blocks
        result["results"] = [
            {
                "key": f"cost_approx_price_{item['key']}",
                "label": f"{item['label']}成本逼近法最终单价",
                "value": item["final_price"],
                "unit": "元/平方米",
                "formula": "（土地成本价格＋土地增值收益）×年期修正系数×（1＋区位修正率）",
            }
            for item in usage_results
        ]
    else:
        for key in (
            "generated_narratives",
            "effective_narratives",
            "narrative_segment_sources",
            "tables",
            "content_blocks",
            "results",
        ):
            result.pop(key, None)
    if len(usage_results) == 1:
        result["cost_approx_price"] = usage_results[0]["final_price"]
    address = data.get("land_location") or data.get("land_location_full") or ""
    other_info = data.get("other_info") or ""
    result["gas_station_location"] = extract_gas_station_location(address, other_info)
    result["gas_station_details"] = extract_gas_station_details(other_info)
    if include_pricing_assistant:
        result["pricing_assistant"] = build_cost_pricing_assistant(data, result, base_dir)
    else:
        result.pop("pricing_assistant", None)
    return result


def _pricing_assistant_formula_parts(analysis: Dict[str, Any]) -> List[Dict[str, str]]:
    totals = analysis.get("totals") or {}
    usage = (analysis.get("usage_results") or [{}])[0]
    parts = [
        {"key": "acquisition_total", "label": "土地取得费", "value": str(totals.get("acquisition_total") or "0.00")},
        {"key": "tax_total", "label": "相关税费", "value": str(totals.get("tax_total") or "0.00")},
        {"key": "development_total", "label": "土地开发费", "value": str(totals.get("development_total") or "0.00")},
        {"key": "interest", "label": "投资利息", "value": str(totals.get("interest") or "0.00")},
        {"key": "profit", "label": "投资利润", "value": str(usage.get("profit") or "0.00")},
        {"key": "value_added_income", "label": "土地增值收益", "value": str(usage.get("value_added_income") or "0.00")},
    ]
    return parts


def _apply_pricing_assistant_overrides(analysis: Dict[str, Any], overrides: Dict[str, str]) -> Dict[str, Any]:
    merged = deepcopy(analysis)
    tax_by_key = {str(item.get("key") or ""): item for item in merged.get("tax_items") or []}
    for fee_key, grade_name in overrides.items():
        item = tax_by_key.get(fee_key)
        if not item:
            continue
        item["grade_name"] = grade_name
        item["confirmed"] = False
    return merged


def _pricing_assistant_controls(analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
    controls: List[Dict[str, Any]] = []
    grade_options_by_key = {
        "water_conservancy_fund": list(HUNAN_PAID_LAND_USE_FEE_STANDARDS.keys()),
        "farmland_reclamation_fee": ["优等水田", "高等水田", "中等水田", "低等水田", "旱地"],
    }
    standard_only_tax_keys = {"farmland_occupation_tax", "social_security_fund", "forest_restoration_fee"}
    for item in analysis.get("tax_items") or []:
        fee_key = str(item.get("key") or "")
        if not item.get("enabled", True):
            continue
        options = grade_options_by_key.get(fee_key) or []
        if options:
            controls.append(
                {
                    "key": fee_key,
                    "label": str(item.get("label") or fee_key),
                    "type": "grade",
                    "value": str(item.get("grade_name") or ""),
                    "options": options,
                    "amount_per_sqm": str(item.get("amount_per_sqm") or ""),
                }
            )
        elif fee_key in standard_only_tax_keys:
            controls.append(
                {
                    "key": fee_key,
                    "label": str(item.get("label") or fee_key),
                    "type": "standard",
                    "value": str(item.get("standard_value") or item.get("amount_per_sqm") or ""),
                    "amount_per_sqm": str(item.get("amount_per_sqm") or ""),
                    "editable": True,
                    "note": "直接改数值，不设等别",
                }
            )
    land_item = _by_key(analysis.get("acquisition_items") or [], "land_compensation")
    if land_item.get("enabled", True):
        zone_options = list(analysis.get("compensation_zone_options") or ["Ⅰ", "Ⅱ", "Ⅲ"])
        controls.append(
            {
                "key": "land_compensation",
                "label": "征地区片",
                "type": "grade",
                "value": str(land_item.get("grade_name") or "").replace("区", "") or str(analysis.get("compensation_zone") or ""),
                "options": zone_options,
                "amount_per_sqm": str(land_item.get("amount_per_sqm") or ""),
            }
        )
    for row in analysis.get("building_compensation_rows") or []:
        row_key = str(row.get("key") or "")
        if not row_key:
            continue
        controls.append(
            {
                "key": f"building:{row_key}",
                "label": f"表3-1·{row.get('label') or row_key}",
                "type": "standard",
                "value": str(row.get("standard") or ""),
                "amount_per_sqm": str(row.get("amount") or ""),
                "editable": True,
                "note": "行补偿标准；改后随试算重算行金额与取得费",
            }
        )
    for item in analysis.get("acquisition_items") or []:
        fee_key = str(item.get("key") or "")
        if not item.get("enabled", True) or fee_key in {"land_compensation", "ground_attachment"}:
            continue
        controls.append(
            {
                "key": fee_key,
                "label": str(item.get("label") or fee_key),
                "type": "standard",
                "value": str(item.get("standard_value") or item.get("amount_per_sqm") or ""),
                "amount_per_sqm": str(item.get("amount_per_sqm") or ""),
                "editable": True,
            }
        )
    for item in analysis.get("development_items") or []:
        if not item.get("enabled", True):
            continue
        controls.append(
            {
                "key": f"development:{item.get('label')}",
                "label": f"开发费·{item.get('label')}",
                "type": "standard",
                "value": str(item.get("standard_value") or item.get("amount_per_sqm") or ""),
                "amount_per_sqm": str(item.get("amount_per_sqm") or ""),
                "editable": True,
            }
        )
    location_mode = str(analysis.get("location_correction_mode") or "direct_sum")
    for factor in analysis.get("location_factors") or []:
        if factor.get("enabled") is False:
            continue
        factor_key = str(factor.get("key") or factor.get("label") or "")
        if location_mode == "direct_sum":
            levels = [str(item) for item in (factor.get("levels") or ["优", "较优", "一般", "较劣", "劣"])]
            controls.append(
                {
                    "key": f"location:{factor_key}",
                    "label": f"区位·{factor.get('label') or factor_key}",
                    "type": "grade",
                    "value": str(factor.get("level") or ""),
                    "options": levels,
                    "amount_per_sqm": str(factor.get("correction_rate") or "0.00"),
                    "editable": True,
                    "note": "选择优劣度自动计算修正率（%）",
                }
            )
        else:
            controls.append(
                {
                    "key": f"location:{factor_key}",
                    "label": f"区位·{factor.get('label') or factor_key}",
                    "type": "standard",
                    "value": str(factor.get("correction_rate") or "0.00"),
                    "amount_per_sqm": str(factor.get("correction_rate") or "0.00"),
                    "editable": True,
                    "note": "修正率（%）",
                }
            )
    for group in analysis.get("risk_groups") or []:
        controls.append(
            {
                "key": f"risk:{group.get('key')}",
                "label": f"风险·{group.get('label') or group.get('key')}",
                "type": "standard",
                "value": str(
                    group.get("override_value")
                    or group.get("effective_value")
                    or group.get("combined_rate")
                    or group.get("rate")
                    or "0.00"
                ),
                "amount_per_sqm": str(group.get("effective_value") or group.get("override_value") or "0.00"),
                "editable": True,
                "note": "综合调整率（%）",
            }
        )
    return controls


def build_cost_pricing_assistant(data: Dict[str, Any], result: Dict[str, Any], base_dir: str | Path) -> Dict[str, Any]:
    baseline_usage = (result.get("usage_results") or [{}])[0]
    baseline_grades = {
        str(item.get("key") or ""): str(item.get("grade_name") or "")
        for item in result.get("tax_items") or []
        if str(item.get("grade_name") or "").strip()
    }
    tax_labels = {str(item.get("key") or ""): str(item.get("label") or "") for item in result.get("tax_items") or []}
    scenarios: List[Dict[str, Any]] = []
    for spec in COST_PRICING_SCENARIO_SPECS:
        variant_data = deepcopy(data)
        variant_analysis = _apply_pricing_assistant_overrides(result, spec.get("overrides") or {})
        variant_data["cost_approx_analysis"] = variant_analysis
        variant_result = calculate_cost_approximation(
            variant_data,
            base_dir,
            include_pricing_assistant=False,
            include_catalog_metadata=False,
            include_process_output=False,
        )
        variant_usage = (variant_result.get("usage_results") or [{}])[0]
        changed_items = [
            {
                "key": fee_key,
                "label": tax_labels.get(fee_key, fee_key),
                "from": baseline_grades.get(fee_key, "（未设定）"),
                "to": grade_name,
            }
            for fee_key, grade_name in (spec.get("overrides") or {}).items()
            if baseline_grades.get(fee_key, "") != grade_name
        ]
        scenarios.append(
            {
                "key": str(spec.get("key") or ""),
                "label": str(spec.get("label") or ""),
                "overrides": deepcopy(spec.get("overrides") or {}),
                "changed_items": changed_items,
                "final_price": str(variant_usage.get("final_price") or ""),
                "cost_price": str(variant_usage.get("cost_price") or ""),
                "tax_total": str((variant_result.get("totals") or {}).get("tax_total") or ""),
                "formula_parts": _pricing_assistant_formula_parts(variant_result),
            }
        )
    return {
        "baseline_final_price": str(baseline_usage.get("final_price") or ""),
        "baseline_cost_price": str(baseline_usage.get("cost_price") or ""),
        "formula_parts": _pricing_assistant_formula_parts(result),
        "controls": _pricing_assistant_controls(result),
        "scenarios": scenarios,
        "entry_points": [],
    }


def _cost_narrative_segment_sources(data: Dict[str, Any], analysis: Dict[str, Any]) -> Dict[str, List[Dict[str, str]]]:
    policy_basis = analysis.get("policy_basis") or {}
    province = policy_basis.get("province_compensation") or {}
    acquisition_items = analysis.get("acquisition_items") or []
    tax_items = analysis.get("tax_items") or []
    totals = analysis.get("totals") or {}
    usage_results = analysis.get("usage_results") or []
    scenario_by_key = {item.get("key"): item for item in analysis.get("usage_scenarios") or []}

    def pct(value: Any) -> str:
        text = str(value or "").strip()
        return f"{text}%" if text else ""

    def item(text: Any, field: str, prefix: str = "", suffix: str = "", priority: int = 0) -> Dict[str, str]:
        return {
            "text": str(text or "").strip(),
            "field": field,
            "prefix": prefix,
            "suffix": suffix,
            "priority": str(priority),
        }

    def inventory_field_for(*tokens: str) -> str:
        inventory = analysis.get("cost_basis_attachment_inventory") or []
        for index, entry in enumerate(inventory):
            haystack = f"{entry.get('label') or ''} {entry.get('source_path') or ''} {entry.get('category') or ''}"
            if any(token and token in haystack for token in tokens):
                return f"cost_approx_analysis.cost_basis_attachment_inventory.{entry.get('target_key') or index}"
        return "valuation_basis_docs_list"

    land_comp = next((i for i in acquisition_items if i.get("key") == "land_compensation"), {})
    attachment = next((i for i in acquisition_items if i.get("key") == "ground_attachment"), {})
    water_fund = next((i for i in tax_items if i.get("key") == "water_conservancy_fund"), {})
    occupation_tax = next((i for i in tax_items if i.get("key") == "farmland_occupation_tax"), {})
    reclamation_fee = next((i for i in tax_items if i.get("key") == "farmland_reclamation_fee"), {})
    social_fee = next((i for i in tax_items if i.get("key") == "social_security_fund"), {})
    forest_fee = next((i for i in tax_items if i.get("key") == "forest_restoration_fee"), {})
    building_comp = next((i for i in acquisition_items if i.get("key") == "building_compensation"), {})
    seedling_comp = next((i for i in acquisition_items if i.get("key") == "seedling_compensation"), {})

    sources = {
        "cost_approx_method_intro": [],
        "cost_approx_basis_intro": [
            item(data.get("land_location_full") or data.get("land_location"), "land_location_full", "估价对象位于", "，"),
            item(analysis.get("acquisition_land_class"), "acquisition_land_class", "评估宗地以", "为主"),
            item(analysis.get("acquisition_land_class"), "acquisition_land_class", "被征用的土地为", "。"),
            item(analysis.get("acquisition_land_subclass"), "acquisition_land_subclass", "征收地类细分类"),
            item(analysis.get("compensation_zone"), "cost_approx_analysis.compensation_zone", "征地补偿区", "区范围"),
            item(data.get("valuation_date"), "valuation_date", "估价期日"),
        ],
        "cost_approx_acquisition_narrative": [
            item(analysis.get("green_seedling_standard_per_mu"), "cost_approx_analysis.green_seedling_standard_per_mu", "标准为"),
            item(land_comp.get("standard_value"), "cost_approx_analysis.acquisition_items.land_compensation.standard_value", "区片价为", "元/亩"),
            item(land_comp.get("amount_per_sqm"), "cost_approx_analysis.acquisition_items.land_compensation.amount_per_sqm", "=", "元/平方米"),
            item(building_comp.get("amount_per_sqm"), "cost_approx_analysis.acquisition_items.building_compensation.amount_per_sqm", "地上建筑物费用", ""),
            item(seedling_comp.get("amount_per_sqm"), "cost_approx_analysis.acquisition_items.seedling_compensation.amount_per_sqm", "青苗补偿", ""),
            item(totals.get("acquisition_total"), "cost_approx_analysis.totals.acquisition_total", "=", "元/平方米"),
        ],
        "cost_approx_tax_narrative": [
            item(water_fund.get("grade_name"), "cost_approx_analysis.tax_items.water_conservancy_fund.grade_name", "征收等别为", "，"),
            item(water_fund.get("standard_value"), "cost_approx_analysis.tax_items.water_conservancy_fund.standard_value", "征收标准为", "元/平方米"),
            item(water_fund.get("coefficient"), "cost_approx_analysis.tax_items.water_conservancy_fund.coefficient", "按", "征收"),
            item(water_fund.get("amount_per_sqm"), "cost_approx_analysis.tax_items.water_conservancy_fund.amount_per_sqm", "=", "元/平方米"),
            item(occupation_tax.get("grade_name"), "cost_approx_analysis.tax_items.farmland_occupation_tax.grade_name", "适用口径为", "，"),
            item(occupation_tax.get("standard_value"), "cost_approx_analysis.tax_items.farmland_occupation_tax.standard_value", "适用税额为", "元/平方米"),
            item(occupation_tax.get("amount_per_sqm"), "cost_approx_analysis.tax_items.farmland_occupation_tax.amount_per_sqm", "耕地占用税为", "元/平方米"),
            item(reclamation_fee.get("grade_name"), "cost_approx_analysis.tax_items.farmland_reclamation_fee.grade_name", "开垦费等别为", "，"),
            item(reclamation_fee.get("standard_value"), "cost_approx_analysis.tax_items.farmland_reclamation_fee.standard_value", "开垦费标准为", "万元/亩"),
            item(reclamation_fee.get("amount_per_sqm"), "cost_approx_analysis.tax_items.farmland_reclamation_fee.amount_per_sqm", "开垦费为", "元/平方米"),
            item(social_fee.get("grade_name"), "cost_approx_analysis.tax_items.social_security_fund.grade_name", "社保费口径为", "，"),
            item(social_fee.get("amount_per_sqm"), "cost_approx_analysis.tax_items.social_security_fund.amount_per_sqm", "保障费为", "元/平方米"),
            item(forest_fee.get("amount_per_sqm"), "cost_approx_analysis.tax_items.forest_restoration_fee.amount_per_sqm", "恢复费为", "元/平方米"),
            item(totals.get("tax_total"), "cost_approx_analysis.totals.tax_total", "=", "元/平方米"),
        ],
        "cost_approx_development_narrative": [
            item(data.get("land_development_set"), "land_development_set", "土地开发程度涉及"),
            item(totals.get("development_total"), "cost_approx_analysis.totals.development_total", "区域土地开发费为", "元/平方米"),
        ],
        "cost_approx_interest_narrative": [
            item(analysis.get("development_cycle_years"), "cost_approx_analysis.development_cycle_years", "开发周期为", "年"),
            item(pct(analysis.get("interest_rate")), "cost_approx_analysis.interest_rate", "利息率取"),
            item(totals.get("acquisition_total"), "cost_approx_analysis.totals.acquisition_total", "（", "+"),
            item(totals.get("tax_total"), "cost_approx_analysis.totals.tax_total", "+", "）×"),
            item(totals.get("development_total"), "cost_approx_analysis.totals.development_total", "＋", "×"),
            item(totals.get("interest"), "cost_approx_analysis.totals.interest", "=", "元/平方米"),
        ],
        "cost_approx_profit_narrative": [],
        "cost_approx_value_added_narrative": [],
        "cost_approx_term_narrative": [],
        "cost_approx_location_narrative": [],
        "cost_approx_solve_narrative": [],
    }
    for policy in analysis.get("policy_documents") or []:
        if policy.get("enabled", True) and policy.get("name"):
            sources["cost_approx_basis_intro"].append(
                item(_policy_reference(policy), f"cost_approx_analysis.policy_documents.{policy.get('key')}")
            )
            sources["cost_approx_acquisition_narrative"].append(
                item(_policy_reference(policy), f"cost_approx_analysis.policy_documents.{policy.get('key')}")
            )
    sources["cost_approx_basis_intro"].append(
        item(analysis.get("effective_local_city"), "cost_approx_analysis.effective_local_city", "湖南省、", "颁布其他各项文件")
    )
    if analysis.get("local_compensation_policy_name"):
        sources["cost_approx_basis_intro"].append(
            item(analysis.get("local_compensation_policy_name"), "local_compensation_policy_name")
        )
    if analysis.get("local_compensation_policy_no"):
        sources["cost_approx_basis_intro"].append(
            item(analysis.get("local_compensation_policy_no"), "local_compensation_policy_no")
        )
    if analysis.get("local_compensation_policy_date"):
        sources["cost_approx_basis_intro"].append(
            item(analysis.get("local_compensation_policy_date"), "local_compensation_policy_date")
        )
    local_policy_reference = ""
    if analysis.get("local_compensation_policy_name"):
        local_policy_reference = str(analysis.get("local_compensation_policy_name") or "")
        local_policy_details = [
            str(analysis.get("local_compensation_policy_no") or "").strip(),
            str(analysis.get("local_compensation_policy_date") or "").strip(),
        ]
        local_policy_details = [part for part in local_policy_details if part]
        if local_policy_details:
            local_policy_reference = f"{local_policy_reference}（{'，'.join(local_policy_details)}）"
        for target_key in ("cost_approx_basis_intro", "cost_approx_acquisition_narrative"):
            sources[target_key].append(item(local_policy_reference, "local_compensation_policy_name", priority=20))
            sources[target_key].append(item(analysis.get("local_compensation_policy_no"), "local_compensation_policy_no"))
            sources[target_key].append(item(analysis.get("local_compensation_policy_date"), "local_compensation_policy_date"))
    sources["cost_approx_acquisition_narrative"].extend(
        [
            item(analysis.get("acquisition_land_class"), "acquisition_land_class", "土地利用类型主要为"),
            item(analysis.get("acquisition_land_subclass"), "acquisition_land_subclass", "的", "青苗补偿标准"),
            item(data.get("county_name"), "county_name", "属于", "征地补偿区"),
            item(data.get("county_name"), "county_name", "，", "的"),
        ]
    )
    static_tax_reference_targets = [
        ("水利建设基金", inventory_field_for("水利建设基金", "湘政发〔2011〕27号", "水利建设基金筹集")),
        ("新增建设用地土地有偿使用费", inventory_field_for("新增建设用地土地有偿使用费", "湘财综", "有偿使用费")),
        ("耕地占用税", inventory_field_for("耕地占用税")),
        ("耕地开垦费", inventory_field_for("耕地开垦费")),
        ("社会保障", inventory_field_for("社会保障", "社保")),
    ]
    for reference, (_, target_field) in zip(STATIC_TAX_POLICY_REFERENCES, static_tax_reference_targets):
        sources["cost_approx_tax_narrative"].append(item(reference, target_field, priority=25))
        sources["cost_approx_tax_narrative"].append(item(reference, "valuation_basis_docs_list", priority=5))
    sources["cost_approx_tax_narrative"].extend(
        [
            item(data.get("county_name"), "county_name", "，", "新增建设用地土地有偿使用费"),
            item(data.get("county_name"), "county_name", "，", "耕地占用税适用税额"),
        ]
    )
    sources["cost_approx_acquisition_intro"] = list(sources["cost_approx_acquisition_narrative"])
    sources["cost_approx_acquisition_solve"] = list(sources["cost_approx_acquisition_narrative"])
    sources["cost_approx_attachment_population_narrative"] = [
        item(
            (analysis.get("attachment_compensation_analysis") or {}).get("building_compensation_per_person"),
            "cost_approx_analysis.attachment_compensation_analysis.building_compensation_per_person",
            "金额为",
            "元/人",
        ),
        item(
            (analysis.get("attachment_compensation_analysis") or {}).get("average_population_per_ha"),
            "cost_approx_analysis.attachment_compensation_analysis.average_population_per_ha",
            "平均安置农业人口数为",
            "人/公顷",
        ),
        item(analysis.get("acquisition_land_subclass"), "acquisition_land_subclass"),
        item(
            (analysis.get("attachment_compensation_analysis") or {}).get("green_seedling_per_sqm"),
            "cost_approx_analysis.attachment_compensation_analysis.green_seedling_per_sqm",
            "折合",
            "元/平方米",
        ),
        item(
            attachment.get("amount_per_sqm"),
            "cost_approx_analysis.attachment_compensation_analysis.attachment_compensation_per_sqm",
            "=",
            "元/平方米",
        ),
    ]
    sources["cost_approx_acquisition_solve"].extend(sources["cost_approx_attachment_population_narrative"])

    for item_res in usage_results:
        key = item_res.get("key")
        scenario = scenario_by_key.get(key) or {}
        sources["cost_approx_profit_narrative"].extend([
            item(pct(scenario.get("profit_rate")), f"cost_approx_analysis.usage_scenarios.{key}.profit_rate", "投资利润率取"),
            item(item_res.get("profit"), f"cost_approx_analysis.usage_results.{key}.profit", "=", "元/平方米"),
        ])
        sources["cost_approx_value_added_narrative"].extend([
            item(pct(scenario.get("value_added_rate")), f"cost_approx_analysis.usage_scenarios.{key}.value_added_rate", "增值收益率取"),
            item(item_res.get("value_added_income"), f"cost_approx_analysis.usage_results.{key}.value_added_income", "=", "元/平方米"),
        ])
        sources["cost_approx_term_narrative"].extend([
            item(pct(scenario.get("reduction_rate")), f"cost_approx_analysis.usage_scenarios.{key}.reduction_rate", "1＋"),
            item(scenario.get("use_term_years"), f"cost_approx_analysis.usage_scenarios.{key}.use_term_years", "）"),
            item(item_res.get("term_correction_factor"), f"cost_approx_analysis.usage_results.{key}.term_correction_factor", "=", "。"),
        ])
        sources["cost_approx_location_narrative"].extend([
            item(
                item_res.get("location_correction_rate"),
                f"cost_approx_analysis.usage_results.{key}.location_correction_rate",
                f"{scenario.get('label')}区位修正率合计为",
                "%",
            ),
        ])
        sources["cost_approx_solve_narrative"].extend([
            item(item_res.get("cost_price"), f"cost_approx_analysis.usage_results.{key}.cost_price", "（"),
            item(item_res.get("interest"), "cost_approx_analysis.totals.interest", "+"),
            item(item_res.get("profit"), f"cost_approx_analysis.usage_results.{key}.profit", "+"),
            item(item_res.get("value_added_income"), f"cost_approx_analysis.usage_results.{key}.value_added_income", "+"),
            item(item_res.get("term_correction_factor"), f"cost_approx_analysis.usage_results.{key}.term_correction_factor", "）×"),
            item(item_res.get("comparable_price"), f"cost_approx_analysis.usage_results.{key}.comparable_price", "=", "元/平方米"),
            item(item_res.get("location_correction_rate"), f"cost_approx_analysis.usage_results.{key}.location_correction_rate", "（1+", "%）"),
            item(item_res.get("final_price"), f"cost_approx_analysis.usage_results.{key}.final_price", "=", "元/平方米"),
            item(data.get("valuation_date"), "valuation_date", "估价期日", "，"),
        ])

    valuation_date = data.get("valuation_date")
    if valuation_date:
        for source_items in sources.values():
            if not any(entry.get("field") == "valuation_date" for entry in source_items):
                source_items.append(item(valuation_date, "valuation_date", priority=30))

    return sources


def extract_gas_station_location(address: str, other_info: str = "") -> dict:
    if not address:
        address = ""
    if not other_info:
        other_info = ""

    city = ""
    district = ""
    detail = ""

    pattern = r'(?:(?P<province>[^省]+省))?(?:(?P<city>[^市]+市))?(?:(?P<district>[^区县]+[区县]))?(?P<detail>.*)'
    
    m = re.match(pattern, address)
    if m:
        city = m.group('city') or ""
        district = m.group('district') or ""
        detail = m.group('detail') or ""

    if not city or not district:
        m2 = re.search(r'([一-龥]+省)?([一-龥]+市)([一-龥]+[区县])([一-龥a-zA-Z0-9\-#]+)', other_info)
        if m2:
            if not city:
                city = m2.group(2) or ""
            if not district:
                district = m2.group(3) or ""
            if not detail and m2.group(4):
                detail = m2.group(4)

    if not city:
        m_city = re.search(r'([一-龥]+市)', address + " " + other_info)
        if m_city:
            city = m_city.group(1)
            
    if not district:
        m_dist = re.search(r'([一-龥]+[区县])', address + " " + other_info)
        if m_dist:
            district = m_dist.group(1)

    if not detail:
        detail = address

    return {
        "city": city.strip(),
        "district": district.strip(),
        "detail": detail.strip()
    }


def extract_gas_station_details(other_info: str) -> dict:
    if not other_info:
        other_info = ""

    details = {
        "canopy_area": None,
        "tank_area": None,
        "office_area": None,
        "store_area": None,
        "dispenser_count": None,
        "nozzle_count": None,
        "tanks": [],
        "loader_count": None
    }

    # 1. 罩棚面积
    m = re.search(r'(?:罩棚|罩棚折影)面积(?:约为)?\s*([\d\.]+)\s*(?:㎡|平方米)', other_info)
    if m:
        details["canopy_area"] = float(m.group(1))

    # 2. 罐区面积
    m = re.search(r'罐区面积(?:约为)?\s*([\d\.]+)\s*(?:㎡|平方米)', other_info)
    if m:
        details["tank_area"] = float(m.group(1))

    # 3. 站房/办公面积
    m = re.search(r'(?:站房|站房及辅助用房|办公室|站房及发配电房|办公)面积(?:约为)?\s*([\d\.]+)\s*(?:㎡|平方米)', other_info)
    if m:
        details["office_area"] = float(m.group(1))

    # 4. 便利店面积
    m = re.search(r'便利店面积(?:约为)?\s*([\d\.]+)\s*(?:㎡|平方米)', other_info)
    if m:
        details["store_area"] = float(m.group(1))

    # 5. 加油机数量
    m = re.search(r'(?:双枪|双枪及以上|四枪)?加油机\s*(\d+)\s*台', other_info)
    if m:
        details["dispenser_count"] = int(m.group(1))

    # 6. 加油枪数量
    m = re.search(r'加油枪\s*(\d+)\s*支', other_info)
    if m:
        details["nozzle_count"] = int(m.group(1))

    # 7. 加注机数量
    m = re.search(r'加注机\s*(\d+)\s*台', other_info)
    if m:
        details["loader_count"] = int(m.group(1))

    # 8. 储罐解析
    tanks_parsed = []
    # 正则1: 3个30立方米双层防爆汽油罐
    pattern1 = re.compile(r'(\d+)\s*个\s*([\d\.]+)\s*(?:m³|立方米|立方|立)(?:的)?\s*([^\d，。、；\s\(\)]*罐)')
    # 正则2: 双层防爆汽油罐3个，每个30立方米
    pattern2 = re.compile(r'([^\d，。、；\s\(\)]*罐)\s*(\d+)\s*个[，,]\s*(?:每|各)个\s*([\d\.]+)\s*(?:m³|立方米|立方|立)')
    # 正则3: 3个双层防爆汽油罐，每个30立方米
    pattern3 = re.compile(r'(\d+)\s*个\s*([^\d，。、；\s\(\)]*罐)[，,]\s*(?:每|各)个\s*([\d\.]+)\s*(?:m³|立方米|立方|立)')
    
    # 找正则2
    for m in pattern2.finditer(other_info):
        raw_type = m.group(1)
        count = int(m.group(2))
        volume = float(m.group(3))
        tanks_parsed.append((count, volume, raw_type))
        
    # 找正则3
    for m in pattern3.finditer(other_info):
        count = int(m.group(1))
        raw_type = m.group(2)
        volume = float(m.group(3))
        # 避免重复
        if not any(t[0] == count and t[1] == volume and t[2] == raw_type for t in tanks_parsed):
            tanks_parsed.append((count, volume, raw_type))

    # 找正则1
    for m in pattern1.finditer(other_info):
        count = int(m.group(1))
        volume = float(m.group(2))
        raw_type = m.group(3)
        if not any(t[0] == count and t[1] == volume and t[2] == raw_type for t in tanks_parsed):
            tanks_parsed.append((count, volume, raw_type))

    for count, volume, raw_type in tanks_parsed:
        is_double = False
        if any(keyword in raw_type.upper() for keyword in ["双层", "SF", "FF", "SS"]):
            is_double = True
            
        tank_type = "未知"
        if "汽油" in raw_type:
            tank_type = "汽油罐"
        elif "柴油" in raw_type:
            tank_type = "柴油罐"
        elif "CNG" in raw_type.upper():
            tank_type = "CNG"
        elif "LNG" in raw_type.upper():
            tank_type = "LNG"
        elif "LPG" in raw_type.upper():
            tank_type = "LPG"
        elif "阻隔防爆" in raw_type or "撬装" in raw_type:
            tank_type = "撬装罐"
        elif "埋地" in raw_type:
            tank_type = "埋地罐"

        for _ in range(count):
            details["tanks"].append({
                "volume": volume,
                "type": tank_type,
                "is_double_wall": is_double
            })

    return details


