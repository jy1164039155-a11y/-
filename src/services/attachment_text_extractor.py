# -*- coding: utf-8 -*-
from __future__ import annotations

import io
import tempfile
from dataclasses import dataclass
from importlib import metadata
from pathlib import Path
from typing import Any, Iterable


SUPPORTED_TEXT_EXTENSIONS = {".txt", ".text", ".md"}
SUPPORTED_DOCX_EXTENSIONS = {".docx"}
SUPPORTED_PDF_EXTENSIONS = {".pdf"}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
DEFAULT_TEXT_LAYER_MIN_CHARS = 80
PDF_OCR_DPI = 220
_PADDLE_OCR: Any = None
_PADDLE_INIT_ERROR: Exception | None = None


class AttachmentTextExtractionError(RuntimeError):
    pass


@dataclass
class ExtractedAttachmentText:
    text: str
    method: str
    page_count: int = 0
    warnings: list[str] | None = None

    def as_dict(self) -> dict:
        return {
            "text": self.text,
            "method": self.method,
            "page_count": self.page_count,
            "warnings": self.warnings or [],
        }


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf_with_pymupdf(data: bytes) -> ExtractedAttachmentText:
    import fitz  # type: ignore

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text") or ""
            if text.strip():
                parts.append(text)
        return ExtractedAttachmentText(
            text="\n".join(parts).strip(),
            method="pymupdf_text_layer",
            page_count=len(doc),
        )


def _extract_pdf_with_pdfplumber(data: bytes) -> ExtractedAttachmentText:
    import pdfplumber  # type: ignore

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
        return ExtractedAttachmentText(
            text="\n".join(parts).strip(),
            method="pdfplumber_text_layer",
            page_count=len(pdf.pages),
        )


def _extract_pdf_text_layer(data: bytes) -> ExtractedAttachmentText:
    errors: list[str] = []
    best_empty: ExtractedAttachmentText | None = None
    for extractor in (_extract_pdf_with_pymupdf, _extract_pdf_with_pdfplumber):
        try:
            extracted = extractor(data)
            if extracted.text.strip():
                return extracted
            if best_empty is None:
                best_empty = extracted
            errors.append(f"{extracted.method}: 未发现可复制文字层")
        except Exception as exc:
            errors.append(f"{extractor.__name__}: {exc}")
    if best_empty is not None:
        best_empty.warnings = errors
        return best_empty
    raise AttachmentTextExtractionError(f"PDF 文字层读取失败：{'; '.join(errors)}")


def _extract_docx_text(data: bytes) -> ExtractedAttachmentText:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return ExtractedAttachmentText(text="\n".join(parts).strip(), method="python_docx")


def _version_tuple(version: str) -> tuple[int, int, int]:
    parts: list[int] = []
    for part in version.replace("-", ".").split("."):
        digits = "".join(ch for ch in part if ch.isdigit())
        if digits == "":
            break
        parts.append(int(digits))
        if len(parts) == 3:
            break
    while len(parts) < 3:
        parts.append(0)
    return tuple(parts[:3])


def _installed_version(distribution_name: str, module: Any | None = None) -> str:
    if module is not None and getattr(module, "__version__", None):
        return str(getattr(module, "__version__"))
    try:
        return metadata.version(distribution_name)
    except metadata.PackageNotFoundError:
        return "0.0.0"


def _check_paddleocr_runtime_compatibility(paddle_module: Any, paddleocr_module: Any) -> None:
    paddle_version = _installed_version("paddlepaddle", paddle_module)
    paddleocr_version = _installed_version("paddleocr", paddleocr_module)
    if _version_tuple(paddle_version) >= (3, 3, 1) and _version_tuple(paddleocr_version) >= (3, 5, 0):
        raise AttachmentTextExtractionError(
            "检测到 PaddlePaddle/PaddleOCR 组合存在已知 PIR 兼容性风险："
            f"paddlepaddle={paddle_version}, paddleocr={paddleocr_version}。"
            "该组合可能触发 ConvertPirAttribute2RuntimeAttribute not support "
            "[pir::ArrayAttribute<pir::DoubleAttribute>]。"
            "请降级到 paddlepaddle==3.2.0 + paddleocr==3.3.3，或至少改用 paddlepaddle==3.2.2。"
        )


def _get_paddle_ocr():
    global _PADDLE_OCR, _PADDLE_INIT_ERROR
    if _PADDLE_OCR is not None:
        return _PADDLE_OCR
    if _PADDLE_INIT_ERROR is not None:
        raise AttachmentTextExtractionError(f"PaddleOCR 初始化失败：{_PADDLE_INIT_ERROR}") from _PADDLE_INIT_ERROR

    try:
        import paddle  # type: ignore
        import paddleocr as paddleocr_module  # type: ignore
        from paddleocr import PaddleOCR  # type: ignore
    except Exception as exc:
        _PADDLE_INIT_ERROR = exc
        raise AttachmentTextExtractionError(
            "当前环境未安装 PaddleOCR/PaddlePaddle，暂不能识别图片或扫描件。"
            "带文字层的 PDF/Word 仍会继续使用 PyMuPDF/python-docx 提取。"
        ) from exc

    _check_paddleocr_runtime_compatibility(paddle, paddleocr_module)

    init_attempts = [
        {"lang": "ch", "ocr_version": "PP-OCRv4", "use_textline_orientation": True},
        {"lang": "ch", "ocr_version": "PP-OCRv4"},
        {"lang": "ch", "use_textline_orientation": True},
        {"lang": "ch", "use_angle_cls": True, "show_log": False},
        {"lang": "ch"},
    ]
    for kwargs in init_attempts:
        try:
            _PADDLE_OCR = PaddleOCR(**kwargs)
            return _PADDLE_OCR
        except TypeError:
            continue
        except Exception as exc:
            _PADDLE_INIT_ERROR = exc
            raise AttachmentTextExtractionError(f"PaddleOCR 初始化失败：{exc}") from exc

    try:
        _PADDLE_OCR = PaddleOCR()
        return _PADDLE_OCR
    except Exception as exc:
        _PADDLE_INIT_ERROR = exc
        raise AttachmentTextExtractionError(f"PaddleOCR 初始化失败：{exc}") from exc


def _collect_paddle_texts(result: Any) -> list[str]:
    texts: list[str] = []
    if result is None:
        return texts
    if isinstance(result, tuple) and len(result) == 2 and isinstance(result[0], str) and isinstance(result[1], (int, float)):
        return [result[0]]
    if isinstance(result, dict):
        for key in ("rec_texts", "texts"):
            value = result.get(key)
            if isinstance(value, list):
                texts.extend(str(item) for item in value if str(item).strip())
        for key in ("text", "transcription"):
            value = result.get(key)
            if isinstance(value, str) and value.strip():
                texts.append(value)
        for value in result.values():
            texts.extend(_collect_paddle_texts(value))
        return texts
    if isinstance(result, (list, tuple)):
        for item in result:
            texts.extend(_collect_paddle_texts(item))
        return texts
    json_attr = getattr(result, "json", None)
    if json_attr is not None:
        try:
            json_value = json_attr() if callable(json_attr) else json_attr
            texts.extend(_collect_paddle_texts(json_value))
        except Exception:
            pass
    return texts


def _ocr_image_path_with_paddleocr(path: Path) -> str:
    ocr = _get_paddle_ocr()
    result = None
    if hasattr(ocr, "ocr"):
        try:
            result = ocr.ocr(str(path), cls=True)
        except TypeError:
            result = ocr.ocr(str(path))
    elif hasattr(ocr, "predict"):
        try:
            result = ocr.predict(input=str(path))
        except TypeError:
            result = ocr.predict(str(path))
    else:
        raise AttachmentTextExtractionError("PaddleOCR 实例缺少可调用的 ocr/predict 方法。")
    return "\n".join(_collect_paddle_texts(result)).strip()


def _extract_image_with_paddleocr(data: bytes, suffix: str) -> ExtractedAttachmentText:
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        text = _ocr_image_path_with_paddleocr(tmp_path)
        if not text:
            raise AttachmentTextExtractionError("PaddleOCR 未识别到有效文字。")
        return ExtractedAttachmentText(text=text, method="paddleocr_image")
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


def _ocr_pdf_pages_with_paddleocr(data: bytes) -> ExtractedAttachmentText:
    import fitz  # type: ignore

    parts: list[str] = []
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        with fitz.open(stream=data, filetype="pdf") as doc:
            page_count = len(doc)
            for page_index, page in enumerate(doc, 1):
                pix = page.get_pixmap(matrix=fitz.Matrix(PDF_OCR_DPI / 72, PDF_OCR_DPI / 72), alpha=False)
                image_path = tmp_path / f"page_{page_index}.png"
                pix.save(str(image_path))
                text = _ocr_image_path_with_paddleocr(image_path)
                if text:
                    parts.append(text)

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise AttachmentTextExtractionError("PaddleOCR 未从 PDF 页面图像中识别到有效文字。")
    return ExtractedAttachmentText(text=full_text, method="paddleocr_pdf_pages", page_count=page_count)


def supported_extensions() -> Iterable[str]:
    return sorted(SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_DOCX_EXTENSIONS | SUPPORTED_PDF_EXTENSIONS | SUPPORTED_IMAGE_EXTENSIONS)


def extract_text_from_attachment(
    filename: str,
    data: bytes,
    *,
    min_text_chars: int = DEFAULT_TEXT_LAYER_MIN_CHARS,
) -> ExtractedAttachmentText:
    suffix = Path(filename or "").suffix.lower()
    if not data:
        raise AttachmentTextExtractionError("附件内容为空，无法识别。")

    if suffix in SUPPORTED_PDF_EXTENSIONS:
        text_layer = _extract_pdf_text_layer(data)
        plain_len = len(text_layer.text.strip())
        if plain_len >= min_text_chars:
            return text_layer
        try:
            ocr_text = _ocr_pdf_pages_with_paddleocr(data)
            ocr_text.warnings = [
                f"PDF 文字层仅 {plain_len} 字，低于阈值 {min_text_chars}，已自动改用 PaddleOCR 页面识别。"
            ] + (text_layer.warnings or [])
            return ocr_text
        except AttachmentTextExtractionError as exc:
            if plain_len > 0:
                text_layer.warnings = [
                    f"PDF 文字层仅 {plain_len} 字，低于阈值 {min_text_chars}；PaddleOCR 后备失败：{exc}"
                ] + (text_layer.warnings or [])
                return text_layer
            raise AttachmentTextExtractionError(
                "该 PDF 未提取到足够文字层，可能是扫描件；尝试调用 PaddleOCR 失败。"
                f"原因：{exc}"
            ) from exc
    if suffix in SUPPORTED_DOCX_EXTENSIONS:
        extracted = _extract_docx_text(data)
        if extracted.text.strip():
            return extracted
        raise AttachmentTextExtractionError("该 Word 附件未提取到文本内容。")
    if suffix in SUPPORTED_TEXT_EXTENSIONS:
        text = _decode_text(data).strip()
        if text:
            return ExtractedAttachmentText(text=text, method="plain_text")
        raise AttachmentTextExtractionError("该文本附件为空。")
    if suffix in SUPPORTED_IMAGE_EXTENSIONS:
        return _extract_image_with_paddleocr(data, suffix)

    raise AttachmentTextExtractionError(
        f"暂不支持的附件格式：{suffix or '未知'}。当前支持：{', '.join(supported_extensions())}"
    )
