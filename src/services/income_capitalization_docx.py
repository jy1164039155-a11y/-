# -*- coding: utf-8 -*-
from __future__ import annotations

import base64
import io
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.services.income_capitalization import calculate_income_capitalization
from src.services.report_docx_styles import format_report_paragraph, format_report_table


CASE_SLOTS = ("A", "B", "C")
INCOME_BLOCKS = [
    ("narrative", "income_cap_method_intro"),
    ("narrative", "income_cap_gross_income_intro"),
    ("table", "income_rent_evidence_rows"),
    ("narrative", "income_cap_rent_factor_intro"),
    ("table", "income_rent_condition_rows"),
    ("narrative", "income_cap_rent_factor_basis"),
    ("table", "income_rent_index_rows"),
    ("table", "income_rent_correction_rows"),
    ("narrative", "income_cap_rent_solve_narrative"),
    ("narrative", "income_cap_annual_gross_narrative"),
    ("narrative", "income_cap_expense_narrative"),
    ("narrative", "income_cap_real_estate_net_narrative"),
    ("narrative", "income_cap_building_income_intro"),
    ("table", "income_cap_rate_rows"),
    ("narrative", "income_cap_building_income_solve"),
    ("narrative", "income_cap_land_income_narrative"),
    ("narrative", "income_cap_total_price_narrative"),
    ("narrative", "income_cap_unit_price_narrative"),
]


def _remove_node(node: Any) -> None:
    parent = node.getparent()
    if parent is not None:
        parent.remove(node)


def _marker(document: Document, text: str) -> Optional[Paragraph]:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(text):
            return paragraph
    return None


def _format_table(table: Table, *, font_size: float | None = None, header_rows: int = 1) -> None:
    column_count = max(1, len(table.columns))
    width_map = {
        2: (7.8, 7.8),
        3: (5.2, 5.2, 5.2),
        5: (2.4, 3.6, 3.3, 3.3, 3.3),
        6: (2.2, 3.2, 2.8, 2.8, 2.8, 2.8),
    }
    widths = width_map.get(column_count, tuple(16.1 / column_count for _ in range(column_count)))
    resolved_size = font_size if font_size is not None else (10.0 if column_count >= 6 else 10.5)
    format_report_table(table, widths_cm=widths, font_size_pt=resolved_size, header_rows=header_rows)


def _insert_paragraph_before(document: Document, stop: Paragraph, text: str, *, role: str = "body") -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    format_report_paragraph(paragraph, role=role)
    paragraph._p.getparent().remove(paragraph._p)
    stop._p.addprevious(paragraph._p)
    return paragraph


def _row_values(row: Dict[str, Any]) -> List[str]:
    if row.get("cells"):
        return [str(value or "") for value in row.get("cells") or []]
    return [
        str(row.get("group") or ""),
        str(row.get("subgroup") or ""),
        str(row.get("label") or ""),
        str(row.get("subject") or ""),
        str(row.get("a") or ""),
        str(row.get("b") or ""),
        str(row.get("c") or ""),
    ]


def _set_cell_text(cell: Any, value: Any) -> None:
    cell.text = "" if value == "空列" else str(value or "")


def _fill_header_rows(table: Table, section: Dict[str, Any], column_count: int) -> int:
    header_rows = section.get("header_rows") or [
        [{"label": label, "colspan": 1, "rowspan": 1} for label in section.get("columns") or []]
    ]
    for row_index, header_row in enumerate(header_rows):
        row = table.rows[row_index]
        col_index = 0
        for header_cell in header_row:
            if col_index >= column_count:
                break
            label = header_cell.get("label", "")
            colspan = max(1, int(header_cell.get("colspan") or 1))
            rowspan = max(1, int(header_cell.get("rowspan") or 1))
            end_row = min(row_index + rowspan - 1, len(header_rows) - 1)
            end_col = min(col_index + colspan - 1, column_count - 1)
            cell = table.cell(row_index, col_index)
            if end_row != row_index or end_col != col_index:
                cell = cell.merge(table.cell(end_row, end_col))
            _set_cell_text(cell, label)
            col_index += colspan
    return len(header_rows)


def _merge_group_columns(table: Table, header_count: int, group_columns: List[int]) -> None:
    body_count = len(table.rows) - header_count
    if body_count <= 1:
        return
    for column_index in group_columns:
        start = 0
        while start < body_count:
            absolute_start = header_count + start
            value = table.cell(absolute_start, column_index).text.strip()
            end = start
            while end + 1 < body_count and value and table.cell(header_count + end + 1, column_index).text.strip() == value:
                end += 1
            if value and end > start:
                for row_index in range(start + 1, end + 1):
                    _set_cell_text(table.cell(header_count + row_index, column_index), "")
                table.cell(absolute_start, column_index).merge(table.cell(header_count + end, column_index))
            start = end + 1


def _image_bytes(data_uri: str) -> io.BytesIO | None:
    if not data_uri:
        return None
    raw = data_uri
    if "," in raw and raw.strip().lower().startswith("data:"):
        raw = raw.split(",", 1)[1]
    try:
        return io.BytesIO(base64.b64decode(raw))
    except Exception:
        return None


def _clear_cell(cell: Any) -> None:
    for paragraph in cell.paragraphs:
        for child in list(paragraph._p):
            paragraph._p.remove(child)


def _fill_image_cell(cell: Any, label: str, data_uri: str) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image = _image_bytes(data_uri)
    if image is None:
        paragraph.add_run(label or "待上传")
        return
    try:
        blob = image.getvalue()
        source = Image.from_blob(blob)
        ratio = float(source.px_width or source.width) / float(source.px_height or source.height)
        max_width_cm = 7.0
        max_height_cm = 4.6
        if ratio >= max_width_cm / max_height_cm:
            width_cm = max_width_cm
            height_cm = max_width_cm / ratio
        else:
            height_cm = max_height_cm
            width_cm = max_height_cm * ratio
        paragraph.add_run().add_picture(io.BytesIO(blob), width=Cm(width_cm), height=Cm(height_cm))
        paragraph.add_run(f"\n{label}")
    except Exception:
        paragraph.add_run(label or "图片读取失败")


def _insert_rent_evidence_table(document: Document, stop: Paragraph, section: Dict[str, Any], analysis: Dict[str, Any]) -> None:
    _insert_paragraph_before(document, stop, str(section.get("report_title") or section.get("title") or ""), role="table_title")
    table = document.add_table(rows=3, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    by_slot = {item.get("slot"): item for item in analysis.get("rent_instances") or []}
    for row_index, slot in enumerate(CASE_SLOTS):
        item = by_slot.get(slot) or {}
        _fill_image_cell(table.cell(row_index, 0), f"实例{slot}照片", str(item.get("photo_data") or ""))
        _fill_image_cell(table.cell(row_index, 1), f"实例{slot}位置", str(item.get("location_image_data") or ""))
    table._tbl.getparent().remove(table._tbl)
    stop._p.addprevious(table._tbl)
    _format_table(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = None


def _insert_table_before(document: Document, stop: Paragraph, section: Dict[str, Any], analysis: Dict[str, Any]) -> None:
    if section.get("key") == "income_rent_evidence_rows":
        _insert_rent_evidence_table(document, stop, section, analysis)
        return
    _insert_paragraph_before(document, stop, str(section.get("report_title") or section.get("title") or ""), role="table_title")
    column_count = max(1, len(section.get("columns") or []))
    header_count = max(1, len(section.get("header_rows") or []))
    table = document.add_table(rows=header_count, cols=column_count)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _fill_header_rows(table, section, column_count)
    for row in section.get("rows") or []:
        cells = table.add_row().cells
        for cell, value in zip(cells, _row_values(row)):
            _set_cell_text(cell, value)
    _merge_group_columns(table, header_count, list(section.get("group_columns") or []))
    table._tbl.getparent().remove(table._tbl)
    stop._p.addprevious(table._tbl)
    _format_table(table, header_rows=header_count)


def _remove_existing_block(start: Paragraph, stop: Paragraph) -> None:
    node = start._p
    stop_node = stop._p
    while node is not None and node is not stop_node:
        next_node = node.getnext()
        _remove_node(node)
        node = next_node


def apply_income_capitalization_to_docx(docx_path: str, data: Dict[str, Any]) -> None:
    if not data.get("use_income_cap"):
        return
    document = Document(docx_path)
    start = _marker(document, "★收益还原法")
    stop = None
    if start is not None:
        seen = False
        for paragraph in document.paragraphs:
            if paragraph._p is start._p:
                seen = True
                continue
            text = paragraph.text.strip()
            if seen and (text.startswith("★") or text.startswith("三、地价的确定")):
                stop = paragraph
                break
    if start is None or stop is None:
        return

    analysis = calculate_income_capitalization(data)
    effective = analysis.get("effective_narratives") or analysis.get("generated_narratives") or {}
    table_map = {item.get("key"): item for item in analysis.get("tables") or []}
    _remove_existing_block(start, stop)
    _insert_paragraph_before(document, stop, "★收益还原法", role="method_heading")
    for block_type, key in INCOME_BLOCKS:
        if block_type == "narrative":
            text = str(effective.get(key) or data.get(key) or "").strip()
            for line in text.splitlines() or ["【待校核】"]:
                _insert_paragraph_before(document, stop, line)
        elif key in table_map:
            _insert_table_before(document, stop, table_map[key], analysis)
    document.save(docx_path)
