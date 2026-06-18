# -*- coding: utf-8 -*-
from __future__ import annotations

from typing import Any, Iterable, Optional

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


REPORT_BODY_FONT = "仿宋_GB2312"
REPORT_BODY_SIZE_PT = 14
REPORT_BODY_LINE_PT = 24
REPORT_FIRST_LINE_CM = 0.99
REPORT_TABLE_TITLE_BEFORE_PT = 7.85
REPORT_TABLE_TITLE_AFTER_PT = 0.5
REPORT_TABLE_LINE_PT = 13


def set_run_font(run: Any, *, font_name: str = REPORT_BODY_FONT, size_pt: float = REPORT_BODY_SIZE_PT, bold: Optional[bool] = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    for key in ("eastAsia", "ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), font_name)


def format_report_paragraph(paragraph: Paragraph, *, role: str = "body") -> Paragraph:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(REPORT_BODY_LINE_PT)
    pf.first_line_indent = Cm(REPORT_FIRST_LINE_CM)

    bold = role == "method_heading"
    if role == "table_title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        pf.space_before = Pt(REPORT_TABLE_TITLE_BEFORE_PT)
        pf.space_after = Pt(REPORT_TABLE_TITLE_AFTER_PT)
        pf.keep_with_next = True
        bold = False
    elif role == "method_heading":
        pf.keep_with_next = True
    elif role == "formula":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Pt(0)
        pf.left_indent = Cm(REPORT_FIRST_LINE_CM)
        pf.right_indent = Pt(0)
    elif role == "short_label":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(REPORT_FIRST_LINE_CM)
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)

    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, bold=bold)
    return paragraph


def set_repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "1")
        tr_pr.append(repeat)


def prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def format_report_table(
    table: Table,
    *,
    widths_cm: Optional[Iterable[float]] = None,
    font_size_pt: float = 10.5,
    header_rows: int = 1,
) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = tuple(widths_cm or ())
    for row_index, row in enumerate(table.rows):
        if row_index < header_rows:
            set_repeat_header(row)
        prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            if column_index < len(widths):
                cell.width = Cm(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                pf = paragraph.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)
                pf.right_indent = Pt(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = Pt(REPORT_TABLE_LINE_PT)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not paragraph.runs:
                    paragraph.add_run("")
                for run in paragraph.runs:
                    set_run_font(run, size_pt=font_size_pt, bold=False)
