# -*- coding: utf-8 -*-
from __future__ import annotations

import re
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.services.cost_approximation import calculate_cost_approximation
from src.services.report_docx_styles import format_report_paragraph, format_report_table


NARRATIVE_HEADINGS = {
    "cost_approx_basis_intro": "计算过程：",
    "cost_approx_acquisition_intro": "1、土地取得费",
    "cost_approx_tax_narrative": "2、相关税费（T）",
    "cost_approx_development_narrative": "3、土地开发费",
    "cost_approx_interest_narrative": "4、投资利息",
    "cost_approx_profit_narrative": "5、投资利润",
    "cost_approx_value_added_narrative": "6、土地增值收益",
    "cost_approx_term_narrative": "7、土地使用年期修正系数",
    "cost_approx_location_narrative": "8、区位修正系数",
    "cost_approx_solve_narrative": "9、成本逼近法价格计算",
}


def _remove_node(node: Any) -> None:
    parent = node.getparent()
    if parent is not None:
        parent.remove(node)


def _marker(document: Document, text: str) -> Optional[Paragraph]:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(text):
            return paragraph
    return None


TABLE_WIDTHS_CM = {
    "cost_building_compensation_rows": (2.3, 2.8, 2.7, 2.4, 2.2, 3.7),
    "cost_resettlement_population_rows": (3.0, 4.8, 2.4, 2.5, 3.4),
    "cost_acquisition_tax_rows": (2.5, 2.25, 2.25, 2.25, 2.25, 2.6),
    "cost_development_rows": (2.4, 1.9, 1.9, 1.9, 1.9, 1.9, 2.2, 2.0),
    "cost_risk_impact_rows": (3.8, 7.0, 5.3),
    "cost_risk_weight_rows": (4.2, 2.2, 7.4, 2.3),
    "cost_risk_assignment_rows": (4.5, 2.9, 2.9, 2.9, 2.9),
    "cost_risk_result_rows": (7.1, 9.0),
    "cost_location_rule_rows": (3.0, 2.05, 2.05, 2.05, 2.05, 2.05, 2.85),
    "cost_location_rows": (3.0, 7.2, 2.8, 3.1),
}


def _format_table(table: Table, section: Dict[str, Any]) -> None:
    column_count = max(1, len(table.columns))
    widths = TABLE_WIDTHS_CM.get(str(section.get("key") or ""))
    if not widths or len(widths) != column_count:
        widths = tuple(16.1 / column_count for _ in range(column_count))
    font_size = 9.0 if column_count >= 7 else 10.5
    format_report_table(
        table,
        widths_cm=widths,
        font_size_pt=font_size,
        header_rows=max(1, len(section.get("header_rows") or [])),
    )


def _insert_paragraph_before(document: Document, stop: Paragraph, text: str, *, role: str = "body") -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    format_report_paragraph(paragraph, role=role)
    paragraph._p.getparent().remove(paragraph._p)
    stop._p.addprevious(paragraph._p)
    return paragraph


def _narrative_line_role(text: str) -> str:
    stripped = str(text or "").strip()
    if re.match(r"^[A-Za-z][A-Za-z0-9]*\s*[=＝]", stripped):
        return "formula"
    if re.match(r"^\S+用地[：:]$", stripped):
        return "short_label"
    return "body"


def _row_values(row: Dict[str, Any], column_count: int) -> List[str]:
    if row.get("cells"):
        return [str(value or "") for value in row.get("cells") or []]
    values = [
        str(row.get("label") or ""),
        str(row.get("subject") or ""),
        str(row.get("a") or ""),
        str(row.get("b") or ""),
        str(row.get("c") or ""),
    ]
    if column_count == 4:
        values = [str(row.get("label") or ""), str(row.get("a") or ""), str(row.get("b") or ""), str(row.get("c") or "")]
    return values


def _set_cell_text(cell: Any, value: Any) -> None:
    cell.text = "" if value == "空列" else str(value or "")


def _fill_header_rows(table: Table, section: Dict[str, Any], column_count: int) -> int:
    header_rows = section.get("header_rows") or [
        [{"label": label, "colspan": 1, "rowspan": 1} for label in section.get("columns") or []]
    ]
    for row_index, header_row in enumerate(header_rows):
        col_index = 0
        for header_cell in header_row:
            if col_index >= column_count:
                break
            colspan = max(1, int(header_cell.get("colspan") or 1))
            rowspan = max(1, int(header_cell.get("rowspan") or 1))
            end_row = min(row_index + rowspan - 1, len(header_rows) - 1)
            end_col = min(col_index + colspan - 1, column_count - 1)
            cell = table.cell(row_index, col_index)
            if end_row != row_index or end_col != col_index:
                cell = cell.merge(table.cell(end_row, end_col))
            _set_cell_text(cell, header_cell.get("label", ""))
            col_index += colspan
    return len(header_rows)


def _merge_group_columns(table: Table, header_count: int, group_columns: List[int]) -> None:
    body_count = len(table.rows) - header_count
    for column_index in group_columns:
        start = 0
        while start < body_count:
            value = table.cell(header_count + start, column_index).text.strip()
            end = start
            while end + 1 < body_count and value and table.cell(header_count + end + 1, column_index).text.strip() == value:
                end += 1
            if value and end > start:
                for row_index in range(start + 1, end + 1):
                    _set_cell_text(table.cell(header_count + row_index, column_index), "")
                table.cell(header_count + start, column_index).merge(table.cell(header_count + end, column_index))
            start = end + 1


def _insert_table_before(document: Document, stop: Paragraph, section: Dict[str, Any]) -> None:
    _insert_paragraph_before(document, stop, str(section.get("report_title") or section.get("title") or ""), role="table_title")
    headers = list(section.get("columns") or [])
    header_count = max(1, len(section.get("header_rows") or []))
    table = document.add_table(rows=header_count, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _fill_header_rows(table, section, len(headers))
    for row in section.get("rows") or []:
        cells = table.add_row().cells
        for cell, value in zip(cells, _row_values(row, len(headers))):
            _set_cell_text(cell, value)
    _merge_group_columns(table, header_count, list(section.get("group_columns") or []))
    table._tbl.getparent().remove(table._tbl)
    stop._p.addprevious(table._tbl)
    _format_table(table, section)


def _remove_existing_block(start: Paragraph, stop: Paragraph) -> None:
    node = start._p
    stop_node = stop._p
    while node is not None and node is not stop_node:
        next_node = node.getnext()
        _remove_node(node)
        node = next_node


def apply_cost_approximation_to_docx(docx_path: str, data: Dict[str, Any], base_dir: str) -> None:
    if not data.get("use_cost_approx"):
        return
    document = Document(docx_path)
    start = _marker(document, "★成本逼近法")
    stop = None
    if start is not None:
        seen = False
        for paragraph in document.paragraphs:
            if paragraph._p is start._p:
                seen = True
                continue
            if seen and paragraph.text.strip().startswith("★"):
                stop = paragraph
                break
    if start is None or stop is None:
        return

    analysis = calculate_cost_approximation(data, base_dir)
    effective = analysis.get("effective_narratives") or analysis.get("generated_narratives") or {}
    table_map = {item.get("key"): item for item in analysis.get("tables") or []}
    _remove_existing_block(start, stop)
    _insert_paragraph_before(document, stop, "★成本逼近法", role="method_heading")
    for block in analysis.get("content_blocks") or []:
        block_type = str(block.get("type") or "")
        key = str(block.get("key") or "")
        heading = NARRATIVE_HEADINGS.get(key, "")
        if heading:
            _insert_paragraph_before(document, stop, heading)
        if block_type == "narrative":
            text = str(effective.get(key) or data.get(key) or "").strip()
            for line in text.splitlines() or ["【待校核】"]:
                _insert_paragraph_before(document, stop, line, role=_narrative_line_role(line))
        elif key in table_map:
            _insert_table_before(document, stop, table_map[key])
    document.save(docx_path)
