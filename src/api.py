# -*- coding: utf-8 -*-
import os
import sys
import io
import json

def get_base_path():
    """防弹级自适应绝对路径解析器：支持打包后 EXE 与开发源码环境的无缝分流路径定位"""
    if getattr(sys, 'frozen', False):
        # 生产环境：如果是打包后的 EXE 运行，返回 api.exe 所在的物理绝对目录
        return os.path.dirname(sys.executable)
    else:
        # 开发环境：如果是 .py 源码运行，返回项目根目录 (跨过 src 目录向上找)
        return os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

# 定海神针：将项目根目录安全注入包搜索路径，绝杀 ModuleNotFoundError 撕裂
sys.path.insert(0, get_base_path())

import re
import uuid
import shutil
import logging
import pythoncom
import win32com.client
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from fastapi.staticfiles import StaticFiles  # 注入静态资源托管支持
from pydantic import BaseModel
from typing import Dict, Any, Optional, List

# 本地模块导入
from src.services.ocr_extractor import extract_metrics_from_text
from src.services.attachment_text_extractor import (
    AttachmentTextExtractionError,
    DEFAULT_TEXT_LAYER_MIN_CHARS,
    extract_text_from_attachment,
)
from src.services.ownership_builder import derive_ownership_descriptions
from src.services.valuation_builder import derive_valuation_descriptions
from src.services.land_usage import normalize_land_usage_fields
from src.services.comparable_library import ComparableLibrary
from src.services.market_comparison_docx import apply_market_comparison_to_docx
from src.services.cost_approximation import applicable_cost_basis, calculate_cost_approximation
from src.services.cost_approximation_docx import apply_cost_approximation_to_docx
from src.services.lpr_query import query_latest_lpr
from src.services.income_capitalization import calculate_income_capitalization
from src.services.benchmark_correction import calculate_benchmark_correction
from src.services.benchmark_correction_docx import apply_benchmark_correction_to_docx
from src.services.income_capitalization_docx import apply_income_capitalization_to_docx
from src.services.valuation_process_builder import build_valuation_process_draft
from src.schemas.land import LandValuationContract
from src.schemas.comparable import CrawlJobRequest, MarketComparisonAnalysis, MarketComparisonCalculateRequest
from src.schemas.valuation_methods import IncomeCapitalizationAnalysis
from src.schemas.validation_handler import FIELD_TRANSLATION_MAP
from src.extractor import DataExtractor, KeyFieldError
from src.generator import ReportGenerator

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("api_server")

app = FastAPI(title="评估报告自动化系统本地API", version="2.0.0")

# 设置 CORS 跨域放行，完美对齐前端 Vite DevServer 和套壳桌面端
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 强力清除浏览器强缓存中间件：确保 index.html 永远不被 Disk/Memory Cache 拦截，百分之百加载最新编译的 Hash 资源
@app.middleware("http")
async def add_no_cache_headers(request, call_next):
    response = await call_next(request)
    path = request.url.path
    if path == "/" or "index.html" in path.lower():
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response

BASE_DIR = get_base_path()
comparable_library = ComparableLibrary(BASE_DIR)

# 确保系统临时目录和输出目录自愈式存在，拒绝死路径硬编码
TEMP_DIR = os.path.join(BASE_DIR, "temp_design")
OUTPUT_DIR = os.path.join(BASE_DIR, "03_Result")
os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

TEMP_FILE_PREFIXES = ("parse_", "temp_", "~$mp_")
TEMP_FILE_EXTENSIONS = {".xlsx", ".docx", ".pdf"}


def cleanup_stale_temp_files(max_age_hours: float = 12.0):
    """启动时清理过期临时文件，避免历史预览和上传副本长期堆积。"""
    import time

    now = time.time()
    max_age_seconds = max_age_hours * 3600
    removed = 0
    for name in os.listdir(TEMP_DIR):
        path = os.path.join(TEMP_DIR, name)
        if not os.path.isfile(path):
            continue
        _, ext = os.path.splitext(name)
        if not name.startswith(TEMP_FILE_PREFIXES) and ext.lower() not in TEMP_FILE_EXTENSIONS:
            continue
        try:
            if now - os.path.getmtime(path) > max_age_seconds:
                os.remove(path)
                removed += 1
        except Exception as exc:
            logger.warning(f"[TEMP] 清理过期临时文件失败: {path}: {exc}")
    if removed:
        logger.info(f"[TEMP] 已清理过期临时文件 {removed} 个。")


def parse_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in ("true", "1", "yes", "y", "是", "开启")
    return bool(value)


def infer_text_scheme_keys(data: Dict[str, Any]) -> Dict[str, str]:
    usage = str(data.get("land_usage") or data.get("land_usage_short") or "")
    is_industrial = "工业" in usage
    is_commercial = any(word in usage for word in ("商业", "商服", "商务"))
    return {
        "land_level_type": data.get("land_level_type") or ("base_land_price_expired" if is_industrial else "commercial_level" if is_commercial else "residential_level_3"),
        "method_combination_type": data.get("method_combination_type") or ("industrial_cost_and_market_average" if is_industrial else "commercial_market_income" if is_commercial else "residential_residual_only"),
        "infrastructure_type": data.get("infrastructure_type") or ("five_通_industrial" if is_industrial else "five_通_commercial" if is_commercial else "five_通_residential"),
    }


def sync_legacy_land_usage_fields(data: Dict[str, Any]) -> None:
    normalize_land_usage_fields(data)


def derive_plot_ratio_display(data: Dict[str, Any]) -> None:
    # 1. 规划容积率派生
    plot_mode = str(data.get("plot_ratio_mode") or "range").strip()
    plot_ratio = str(data.get("plot_ratio") or "").strip()
    plot_ratio_min = str(data.get("plot_ratio_min") or "").strip()
    if plot_mode not in ("range", "fixed"):
        plot_mode = "range" if plot_ratio_min else "fixed"
    data["plot_ratio_mode"] = plot_mode
    if plot_mode == "range" and plot_ratio_min:
        data["plot_ratio_display"] = f"{plot_ratio_min}-{plot_ratio}"
    else:
        data["plot_ratio_display"] = plot_ratio

    # 2. 设定容积率派生与自愈兜底
    set_mode = str(data.get("set_plot_ratio_mode") or "fixed").strip()
    set_ratio = str(data.get("set_plot_ratio") or "").strip()
    set_ratio_min = str(data.get("set_plot_ratio_min") or "").strip()
    
    if set_mode not in ("range", "fixed"):
        set_mode = "range" if set_ratio_min else "fixed"
    data["set_plot_ratio_mode"] = set_mode

    # 如果整个设定容积率没有手动填写上限/固定值（即空字符串或占位符），
    # 并且如果是 fixed 模式下，我们会自动兜底取规划容积率的相应设定（以保持向下兼容）
    if set_mode == "fixed":
        if not set_ratio or set_ratio == "______" or "【请填写" in set_ratio:
            data["set_plot_ratio"] = data.get("plot_ratio")
            set_ratio = str(data.get("plot_ratio") or "").strip()

    # 3. 拼接设定容积率的展示拼接字段
    if set_mode == "range" and set_ratio_min:
        data["set_plot_ratio_display"] = f"{set_ratio_min}-{set_ratio}"
    else:
        data["set_plot_ratio_display"] = set_ratio


def derive_land_use_term_years(data: Dict[str, Any]) -> None:
    match = re.search(r"\d+(?:\.\d+)?", str(data.get("land_use_term") or ""))
    data["land_use_term_years"] = match.group(0) if match else str(data.get("land_use_term") or "")


def derive_report_validity_fields(data: Dict[str, Any]) -> None:
    """Derive report display year and one-year validity end date from report_date."""
    from datetime import date, timedelta

    report_date = str(data.get("report_date") or "").strip()
    year_match = re.search(r"(\d{4})", report_date)
    if year_match and not str(data.get("report_year_text") or "").strip():
        digit_map = str.maketrans("0123456789", "〇一二三四五六七八九")
        data["report_year_text"] = year_match.group(1).translate(digit_map) + "年"

    if str(data.get("report_valid_until") or "").strip():
        return

    date_match = re.search(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", report_date)
    if not date_match:
        return

    year, month, day = map(int, date_match.groups())
    try:
        start = date(year, month, day)
        try:
            one_year_later = start.replace(year=start.year + 1)
        except ValueError:
            one_year_later = start.replace(year=start.year + 1, day=28)
        end = one_year_later - timedelta(days=1)
        data["report_valid_until"] = f"{end.year}年{end.month:02d}月{end.day:02d}日"
    except ValueError:
        return


def validate_render_payload(data: Dict[str, Any]) -> LandValuationContract:
    is_archive = parse_bool(data.get("archive_to_result", False))
    sync_legacy_land_usage_fields(data)
    analysis = data.get("market_comp_analysis")
    if parse_bool(data.get("use_market_comp")) and analysis:
        analysis_data = analysis.model_dump() if isinstance(analysis, MarketComparisonAnalysis) else dict(analysis)
        calculated = comparable_library.calculate_market_comparison(analysis_data)
        analysis_fields = set(MarketComparisonAnalysis.model_fields)
        data["market_comp_analysis"] = {key: value for key, value in calculated.items() if key in analysis_fields}
        for key in (
            "market_comp_step1_instances",
            "market_comp_basic_rows",
            "market_comp_factor_condition_rows",
            "market_comp_time_index_rows",
            "market_comp_factor_index_rows",
            "market_comp_correction_rows",
            "market_comp_step4_solve",
            "market_comp_comparable_basis",
            "market_comp_factor_selection",
            "market_comp_condition_intro",
            "market_comp_index_basis",
            "market_comp_verification",
            "instance_a_desc",
            "instance_b_desc",
            "instance_c_desc",
            "comparable_case_count",
            "market_comp_price",
        ):
            data[key] = calculated.get(key)
        data["market_comp_evidence_snapshot_ids"] = (
            data.get("market_comp_evidence_snapshot_ids")
            or calculated.get("evidence_snapshot_ids")
            or []
        )
        if is_archive and not calculated.get("complete"):
            raise HTTPException(status_code=422, detail="正式归档失败：市场比较法存在未确认的必填因子。")
        if (
            is_archive
            and calculated.get("comparable_basis_status") == "needs_adjustment"
            and not str((calculated.get("narrative_overrides") or {}).get("market_comp_comparable_basis") or "").strip()
        ):
            raise HTTPException(status_code=422, detail="正式归档失败：价格可比基础需要调整，请先修改最终正文中的对应段落。")
        if is_archive and len(data["market_comp_evidence_snapshot_ids"]) != 3:
            raise HTTPException(status_code=422, detail="正式归档失败：市场比较法三宗比较实例的成交公告证据尚未全部生成。")
    income_analysis = data.get("income_cap_analysis")
    if parse_bool(data.get("use_income_cap")):
        income_data = income_analysis.model_dump() if isinstance(income_analysis, IncomeCapitalizationAnalysis) else dict(income_analysis or {})
        calculated_income = calculate_income_capitalization({**data, "income_cap_analysis": income_data})
        income_fields = set(IncomeCapitalizationAnalysis.model_fields)
        data["income_cap_analysis"] = {key: value for key, value in calculated_income.items() if key in income_fields}
        data["income_cap_price"] = calculated_income.get("income_cap_price") or data.get("income_cap_price")
        if is_archive and not calculated_income.get("complete"):
            raise HTTPException(status_code=422, detail="正式归档失败：收益还原法租金实例、因素、图片或还原率尚未全部校核。")
    
    # 1. 定义必填字段列表
    required_fields = [
        "client_name", "land_location", "parcel_count", "land_user", "appraisal_org",
        "report_no", "technical_report_no", "report_date",
        "valuation_date", "land_area", "building_area", "plot_ratio", "land_use_term",
        "right_type", "land_usage", "land_usage_short", "land_usage_full",
        "planning_approval_authority", "ownership_scenario_type", "asset_use_category",
        "land_registration_desc", "land_right_desc", "land_use_status_desc",
        "land_development_actual",
        "land_development_set", "client_org_code", "client_principal", "client_org_type",
        "client_address", "client_contact", "client_phone", "client_postcode",
    ]

    # 已登记土地特有的核心登记要素
    registered_required_fields = [
        "right_cert_no", "real_estate_cert_no", "owner_name", 
        "registration_time", "cadastral_map_no", "parcel_no"
    ]

    # 2. 如果是正式归档模式，执行极为严苛的前置拦截防御
    if is_archive:
        missing_fields = []
        for field_name in required_fields:
            val = data.get(field_name)
            if val is None or str(val).strip() in ("", "______") or "【请填写" in str(val):
                label = FIELD_TRANSLATION_MAP.get(field_name, field_name)
                missing_fields.append(f"{label}|{field_name}")
        
        # 归档模式下，对已登记土地的核心要素进行严格的空值或占位符拦截
        if data.get("ownership_scenario_type") == "registered_complete":
            right_cert = str(data.get("right_cert_no") or "").strip()
            real_estate_cert = str(data.get("real_estate_cert_no") or "").strip()
            
            # 只要两个证书编号都为空或含有占位符，就报错
            is_right_cert_empty = not right_cert or right_cert in ("", "/", "______") or "【请填写" in right_cert
            is_re_cert_empty = not real_estate_cert or real_estate_cert in ("", "/", "______") or "【请填写" in real_estate_cert
            if is_right_cert_empty and is_re_cert_empty:
                missing_fields.append("权利证书号或不动产权证书编号|right_cert_no")
                
            for reg_f in ["owner_name", "registration_time", "cadastral_map_no", "parcel_no"]:
                val = str(data.get(reg_f) or "").strip()
                if not val or val in ("", "/", "______") or "【请填写" in val:
                    label = FIELD_TRANSLATION_MAP.get(reg_f, reg_f)
                    missing_fields.append(f"{label}|{reg_f}")

        # 规划与设定容积率在范围（range）模式下的缺失校验
        if data.get("plot_ratio_mode") == "range":
            plot_min = str(data.get("plot_ratio_min") or "").strip()
            if not plot_min or plot_min in ("______", "【请填写") or "【请填写" in plot_min:
                missing_fields.append("规划容积率下限|plot_ratio_min")

        if data.get("set_plot_ratio_mode") == "range":
            set_min = str(data.get("set_plot_ratio_min") or "").strip()
            if not set_min or set_min in ("______", "【请填写") or "【请填写" in set_min:
                missing_fields.append("评估设定容积率下限|set_plot_ratio_min")
            set_max = str(data.get("set_plot_ratio") or "").strip()
            if not set_max or set_max in ("______", "【请填写") or "【请填写" in set_max:
                missing_fields.append("评估设定容积率上限|set_plot_ratio")

        if missing_fields:
            raise HTTPException(
                status_code=422,
                detail=f"正式归档失败：检测到以下核心要素缺失或含有未填写占位符，请补齐后再提交归档：{'、'.join(missing_fields)}"
            )

    # 3. 如果是预览模式，执行温馨的自愈式“【请填写...】”填充
    numeric_placeholders = []
    if not is_archive:
        for field_name in required_fields:
            val = data.get(field_name)
            if val is None or str(val).strip() in ("", "______"):
                # 如果是数值字段，临时赋予 0.0 以安全通过 Pydantic 类型限制，并记录为 placeholder
                if field_name in ["land_area", "building_area", "plot_ratio"]:
                    data[field_name] = 0.0
                    numeric_placeholders.append(field_name)
                else:
                    data[field_name] = f"【请填写{FIELD_TRANSLATION_MAP.get(field_name, field_name)}】"
                    
        # 已登记场景核心要素的自愈填充
        if data.get("ownership_scenario_type") == "registered_complete":
            for reg_f in registered_required_fields:
                val = data.get(reg_f)
                if val is None or str(val).strip() in ("", "______", "/"):
                    data[reg_f] = f"【请填写{FIELD_TRANSLATION_MAP.get(reg_f, reg_f)}】"

        # 对设定容积率可选值，如果设为了 range 但无值，填充预览占位符
        if data.get("set_plot_ratio_mode") == "range":
            if not data.get("set_plot_ratio_min") or str(data.get("set_plot_ratio_min")).strip() in ("", "______"):
                data["set_plot_ratio_min"] = "【请填写评估设定容积率下限】"
            if data.get("set_plot_ratio") is None or str(data.get("set_plot_ratio")).strip() in ("", "______"):
                data["set_plot_ratio"] = 0.0
                numeric_placeholders.append("set_plot_ratio")

    # 4. 清洗并转换现有的数值型字段
    for f_name in ["land_area", "building_area", "plot_ratio"]:
        val = data.get(f_name)
        if val is None or str(val).strip() in ("", "______", "【请填写"):
            data[f_name] = 0.0
            if f_name not in numeric_placeholders:
                numeric_placeholders.append(f_name)
        else:
            try:
                cleaned = re.sub(r"[^\d.]", "", str(val))
                data[f_name] = float(cleaned) if cleaned else 0.0
            except Exception:
                if is_archive:
                    raise HTTPException(status_code=422, detail=f"{FIELD_TRANSLATION_MAP.get(f_name, f_name)}必须为有效数字")
                else:
                    data[f_name] = 0.0
                    if f_name not in numeric_placeholders:
                        numeric_placeholders.append(f_name)

    # 额外清洗 set_plot_ratio 可选数值字段
    set_val = data.get("set_plot_ratio")
    if set_val is None or str(set_val).strip() in ("", "______") or "【请填写" in str(set_val):
        data["set_plot_ratio"] = None
    else:
        try:
            cleaned = re.sub(r"[^\d.]", "", str(set_val))
            data["set_plot_ratio"] = float(cleaned) if cleaned else None
        except Exception:
            if is_archive:
                raise HTTPException(status_code=422, detail="评估设定容积率必须为有效数字")
            else:
                data["set_plot_ratio"] = None

    # 清洗设定容积率下限字段（仅去除杂质，保留为规范 string）
    set_min_val = data.get("set_plot_ratio_min")
    if set_min_val is not None and str(set_min_val).strip() not in ("", "______", "【请填写"):
        cleaned_min = re.sub(r"[^\d.]", "", str(set_min_val))
        data["set_plot_ratio_min"] = cleaned_min if cleaned_min else None

    # 4. 执行其他补充派生字段计算
    sync_legacy_land_usage_fields(data)
    derive_plot_ratio_display(data)
    derive_land_use_term_years(data)
    derive_report_validity_fields(data)

    client_name = str(data.get("client_name") or "")
    transfer_purpose = str(data.get("transfer_purpose") or "拟挂牌出让")
    land_location = str(data.get("land_location") or "")
    land_plot_number = str(data.get("land_plot_number") or "")
    land_location_full = f"{land_location}{land_plot_number}"
    data["transfer_purpose"] = transfer_purpose
    data["land_location_full"] = land_location_full
    
    parcel_count = str(data.get("parcel_count") or "一宗")
    land_usage = str(data.get("land_usage") or "")
    project_name_suffix = str(data.get("project_name_suffix") or "国有建设用地使用权 market价格评估")
    data["parcel_count"] = parcel_count
    data["project_name_suffix"] = project_name_suffix
    
    data["project_name"] = f"{client_name}{transfer_purpose}涉及位于{land_location_full}的{parcel_count}{land_usage}{project_name_suffix}"
    # 4.5. [新增] 权属与估价测算依据清单一键规范化渲染与自动注入
    def render_basis_docs(raw_text: str) -> str:
        if not raw_text:
            return ""
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        processed = []
        for line in lines:
            if line.startswith("《") and "》" in line:
                processed.append(line)
                continue
                
            bracket_idx = -1
            for char in ["（", "("]:
                idx = line.find(char)
                if idx != -1:
                    bracket_idx = idx
                    break
                    
            if bracket_idx != -1:
                doc_name = line[:bracket_idx].strip()
                doc_no = line[bracket_idx:].strip()
                if not doc_name.startswith("《"):
                    doc_name = f"《{doc_name}》"
                processed.append(f"{doc_name}{doc_no}")
            else:
                processed.append(f"《{line}》")
        return "、".join(processed)

    # 执行一键自愈转换并注入最终契约字典
    basis_rendered = render_basis_docs(data.get("basis_docs_list", ""))
    val_basis_rendered = render_basis_docs(data.get("valuation_basis_docs_list", ""))
    
    data["basis_docs_rendered"] = basis_rendered
    data["valuation_basis_docs_rendered"] = val_basis_rendered
    
    if basis_rendered:
        data["basis_docs_phrase"] = f"根据{basis_rendered}等资料"
    else:
        data["basis_docs_phrase"] = "根据委托方提供的权属证明及规划条件等资料"

    derive_ownership_descriptions(data)
    derive_valuation_descriptions(data)

    if is_archive and any(
        not str(data.get(key) or "").strip()
        or str(data.get(key) or "").strip() == "______"
        or "【请填写" in str(data.get(key) or "")
        for key in ("final_unit_price", "final_total_price", "final_total_price_upper")
    ):
        raise HTTPException(
            status_code=422,
            detail="正式归档失败：当前确价逻辑未能自动形成最终单价、总价或大写金额，请手填最终价或切换为加权算术平均/中位数后再归档。"
        )

    for key in [
        "use_cost_approx", "use_market_comp", "use_income_cap", "use_benchmark_corr",
        "use_residual", "show_price_in_text", "is_base_price_expired", "base_price_is_expired", "include_registration_risk_note",
        "explain_unselected_methods",
    ]:
        if key in data:
            data[key] = parse_bool(data[key])

    try:
        # 实例化契约以获得严密校验
        contract = LandValuationContract(**data)
        
        # 5. 后自愈回填：针对预览模式，如果刚才临时塞了 0.0，在 dump 出字典后重新替换为字符串占位符，
        # 这样 Word 模板中渲染出的就是精美的“【请填写...】”热区文字
        dumped_data = contract.model_dump()
        if not is_archive:
            for f_name in numeric_placeholders:
                dumped_data[f_name] = f"【请填写{FIELD_TRANSLATION_MAP.get(f_name, f_name)}】"
                # 同时调整 plot_ratio_display / plot_ratio_min 等相关级联字段
                if f_name == "plot_ratio":
                    dumped_data["plot_ratio_display"] = f"【请填写设定容积率上限】"
            # 检查原生的 planning_approval_authority
            if dumped_data.get("planning_approval_authority") is None:
                dumped_data["planning_approval_authority"] = "【请填写规划条件批准机关】"
            if dumped_data.get("county_name") is None:
                dumped_data["county_name"] = "【请填写评估所属县市简称】"
            # 检查成本/确价等数值字段，如果为 None 且对应方法启用了，可以自动填充占位符
            if dumped_data.get("use_cost_approx") and dumped_data.get("cost_approx_price") is None:
                dumped_data["cost_approx_price"] = "【请填写成本逼近法单价】"
            if dumped_data.get("use_market_comp") and dumped_data.get("market_comp_price") is None:
                dumped_data["market_comp_price"] = "【请填写市场比较法单价】"
            if dumped_data.get("use_residual") and dumped_data.get("residual_price") is None:
                dumped_data["residual_price"] = "【请填写剩余法单价】"
                
        # 允许在外部 data 中读到这些自愈回填字段
        data.update(dumped_data)
        return contract
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"表单契约校验失败：{exc}") from exc


def scan_docx_residual_markers(docx_path: str):
    """扫描渲染后 Word 中残留的模板语法或兜底下划线。"""
    import docx as docx_reader

    doc = docx_reader.Document(docx_path)
    findings = []

    def check_text(text: str, where: str):
        if not text:
            return
        if "{{" in text or "}}" in text or "{%" in text or "%}" in text or "______" in text or "【请填写" in text:
            findings.append(f"{where}: {text[:120]}")

    for idx, paragraph in enumerate(doc.paragraphs, 1):
        check_text(paragraph.text, f"段落{idx}")
    for t_idx, table in enumerate(doc.tables, 1):
        for r_idx, row in enumerate(table.rows, 1):
            for c_idx, cell in enumerate(row.cells, 1):
                check_text(cell.text, f"表格{t_idx} R{r_idx}C{c_idx}")
    for s_idx, section in enumerate(doc.sections, 1):
        for part_name, part in (("页眉", section.header), ("页脚", section.footer)):
            for p_idx, paragraph in enumerate(part.paragraphs, 1):
                check_text(paragraph.text, f"第{s_idx}节{part_name}{p_idx}")

    return findings[:20]


def lock_docx_cjk_fonts(docx_path: str, font_name: str = "仿宋_GB2312"):
    """V6.2.2 全文档中文字体高保真锁定处理器：定向锁定替换导致的 None/退化中文字体，不伤及标题与英数排版。"""
    import docx
    from docx.oxml.ns import qn
    import re
    
    cjk_re = re.compile(r"[\u4e00-\u9fa5]")
    western_fallback_fonts = {"Calibri", "Times New Roman", "Arial"}
    
    def process_run(run, fallback_size=None):
        text = run.text or ""
        if not text.strip():
            return
            
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        
        east_asia = rFonts.get(qn('w:eastAsia'))
        ascii_font = run.font.name

        # 1. 包含中文的区域：
        if cjk_re.search(text):
            # 若中文字体缺失，或被错误写成了西文字体，才锁定为标准仿宋_GB2312
            if not east_asia or east_asia in western_fallback_fonts:
                run.font.name = font_name
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
            else:
                # 尊重并保留原模板精心设计的个性化非退化字体（如黑体、标宋等），仅对退化英文槽进行智能拉齐
                if not ascii_font or ascii_font in western_fallback_fonts:
                    run.font.name = east_asia
                    rFonts.set(qn('w:ascii'), east_asia)
                    rFonts.set(qn('w:hAnsi'), east_asia)
            if run.font.size is None and fallback_size is not None:
                run.font.size = fallback_size
        else:
            # 2. 纯英文、数字及符号区域（精准拉齐西文样式）：
            # 仅当其字体属于退化的西文或空值时，我们才锁定西文为 Times New Roman，保障极度专业的英数排版
            if not ascii_font or ascii_font in {"Calibri", "Arial"}:
                run.font.name = "Times New Roman"
                rFonts.set(qn('w:ascii'), "Times New Roman")
                rFonts.set(qn('w:hAnsi'), "Times New Roman")
                rFonts.set(qn('w:eastAsia'), font_name)  # 防标点符号拉伸
            if run.font.size is None and fallback_size is not None:
                run.font.size = fallback_size

    doc = docx.Document(docx_path)

    def paragraph_fallback_size(paragraph, *, in_table: bool):
        # 字号交给模板段落/样式继承；这里只修复中文字体槽，避免把标题、页眉或正文误刷成同一字号。
        return None

    def process_table(table, *, in_table: bool = True):
        for row in table.rows:
            for cell in row.cells:
                process_block(cell, in_table=in_table)
                for sub_t in cell.tables:
                    process_table(sub_t, in_table=True)

    def process_block(block, *, in_table: bool = False):
        for p in block.paragraphs:
            fallback_size = paragraph_fallback_size(p, in_table=in_table)
            for run in p.runs:
                process_run(run, fallback_size)
        for table in block.tables:
            process_table(table, in_table=True)

    process_block(doc, in_table=False)

    # 页眉页脚也可能存在变量替换，包含其中的表格。
    for section in doc.sections:
        if section.header:
            process_block(section.header, in_table=False)
        if section.footer:
            process_block(section.footer, in_table=False)
                    
    doc.save(docx_path)


def normalize_generated_text_blocks(docx_path: str, font_name: str = "仿宋_GB2312"):
    """把 docxtpl 写入同一段落的多行正文拆成独立段落，并尽量继承模板原有字号与段落格式。"""
    import copy
    import re
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.text.paragraph import Paragraph

    numbered_line = re.compile(r"^\s*\d+[.．、]\s*")
    blocked_ancestors = {"tbl", "tc", "hdr", "ftr", "footnote", "endnote", "txbxContent"}
    western_fallback_fonts = {"Calibri", "Times New Roman", "Arial"}

    def should_split(text: str) -> bool:
        if "\n" not in text:
            return False
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if len(lines) < 2:
            return False
        return True

    def is_in_blocked_story(paragraph) -> bool:
        node = paragraph._p.getparent()
        while node is not None:
            local_name = node.tag.rsplit("}", 1)[-1]
            if local_name in blocked_ancestors:
                return True
            node = node.getparent()
        return False

    def run_east_asia_font(run):
        r_pr = run._r.rPr
        if r_pr is None or r_pr.rFonts is None:
            return None
        return r_pr.rFonts.get(qn("w:eastAsia"))

    def preferred_font_name(paragraph):
        text_runs = [run for run in paragraph.runs if (run.text or "").strip()]
        for run in text_runs + list(paragraph.runs):
            candidate = run_east_asia_font(run) or run.font.name
            if candidate and candidate not in western_fallback_fonts:
                return candidate
        try:
            style_font = paragraph.style.font.name
        except Exception:
            style_font = None
        if style_font and style_font not in western_fallback_fonts:
            return style_font
        return font_name

    def preferred_font_size(paragraph):
        text_runs = [run for run in paragraph.runs if (run.text or "").strip()]
        for run in text_runs:
            if run.font.size is not None:
                return run.font.size
        try:
            if paragraph.style.font.size is not None:
                return paragraph.style.font.size
        except Exception:
            pass
        for run in paragraph.runs:
            if run.font.size is not None:
                return run.font.size
        return Pt(14)

    def capture_paragraph_format(paragraph):
        pf = paragraph.paragraph_format
        return {
            "style": paragraph.style,
            "font_name": preferred_font_name(paragraph),
            "font_size": preferred_font_size(paragraph),
            "left_indent": pf.left_indent,
            "first_line_indent": pf.first_line_indent,
            "line_spacing": pf.line_spacing,
            "space_before": pf.space_before,
            "space_after": pf.space_after,
        }

    def clear_paragraph_content(paragraph):
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

    def apply_paragraph_format(paragraph, fmt):
        try:
            paragraph.style = fmt["style"]
        except Exception:
            pass
        pf = paragraph.paragraph_format
        pf.left_indent = fmt["left_indent"] if fmt["left_indent"] is not None else Cm(0.99)
        pf.first_line_indent = fmt["first_line_indent"] if fmt["first_line_indent"] is not None else Pt(0)
        pf.line_spacing = fmt["line_spacing"] if fmt["line_spacing"] is not None else Pt(23)
        pf.space_before = fmt["space_before"] if fmt["space_before"] is not None else Pt(0)
        pf.space_after = fmt["space_after"] if fmt["space_after"] is not None else Pt(0)

        run_font_name = fmt["font_name"] or font_name
        run_font_size = fmt["font_size"] or Pt(14)
        for run in paragraph.runs:
            run.font.name = run_font_name
            run.font.size = run_font_size
            r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), run_font_name)
            r_fonts.set(qn("w:ascii"), run_font_name)
            r_fonts.set(qn("w:hAnsi"), run_font_name)
            r_fonts.set(qn("w:cs"), run_font_name)

    def set_paragraph_text(paragraph, text, fmt):
        clear_paragraph_content(paragraph)
        paragraph.add_run(text)
        apply_paragraph_format(paragraph, fmt)

    def insert_after(paragraph, text, fmt):
        new_p = OxmlElement("w:p")
        if paragraph._p.pPr is not None:
            new_p.append(copy.deepcopy(paragraph._p.pPr))
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        new_para.add_run(text)
        apply_paragraph_format(new_para, fmt)
        return new_para

    doc = docx.Document(docx_path)

    for paragraph in list(doc.paragraphs):
        if is_in_blocked_story(paragraph):
            continue
        text = paragraph.text
        if not should_split(text):
            continue
        fmt = capture_paragraph_format(paragraph)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        set_paragraph_text(paragraph, lines[0], fmt)
        cursor = paragraph
        for line in lines[1:]:
            cursor = insert_after(cursor, line, fmt)

    doc.save(docx_path)


def normalize_generated_body_paragraph_fonts(
    docx_path: str,
    data: Dict[str, Any],
    font_name: str = "仿宋_GB2312",
    size_pt: float = 14.0,
):
    """统一自动生成正文段落的字体字号，避免模板占位符所在 run 的字号被带入正文。"""
    import docx
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    generated_text_keys = [
        "land_registration_desc",
        "land_right_desc",
        "land_use_status_desc",
        "valuation_method_reasons",
        "valuation_method_applicability",
        "final_price_determination",
        "valuation_result_statement",
        "infrastructure_detail",
        "cost_approx_land_class_intro",
        "cost_approx_process_intro",
        "cost_approx_method_intro",
        "market_comp_method_intro",
        "income_cap_method_intro",
        "benchmark_corr_method_intro",
        "residual_method_intro",
    ]
    target_lines: set[str] = set()
    for key in generated_text_keys:
        raw = str(data.get(key) or "")
        for line in raw.splitlines():
            text = line.strip()
            if text:
                target_lines.add(text)

    if not target_lines:
        return

    blocked_ancestors = {"tbl", "tc", "hdr", "ftr", "footnote", "endnote", "txbxContent"}

    def is_in_blocked_story(paragraph) -> bool:
        node = paragraph._p.getparent()
        while node is not None:
            local_name = node.tag.rsplit("}", 1)[-1]
            if local_name in blocked_ancestors:
                return True
            node = node.getparent()
        return False

    def should_normalize(text: str) -> bool:
        cleaned = text.strip()
        if not cleaned:
            return False
        if cleaned in target_lines:
            return True
        return any(len(line) >= 12 and line in cleaned for line in target_lines)

    def apply_run_font(run):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)

    doc = docx.Document(docx_path)
    for paragraph in doc.paragraphs:
        if is_in_blocked_story(paragraph) or not should_normalize(paragraph.text):
            continue
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(23)
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)
        for run in paragraph.runs:
            apply_run_font(run)

    doc.save(docx_path)


def normalize_report_typography(
    docx_path: str,
    body_font_name: str = "仿宋_GB2312",
    heading_font_name: str = "宋体",
):
    """Lock main-document report typography after template rendering.

    The source template has some paragraphs whose runs inherit Normal style
    fonts/sizes, so this normalizer fixes the rendered main story while leaving
    table-specific typography to the existing table processors.
    """
    import re
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    blocked_ancestors = {"tbl", "tc", "hdr", "ftr", "footnote", "endnote", "txbxContent"}
    main_titles = {"土地估价报告", "土地估价技术报告"}
    part_heading_re = re.compile(r"^第[一二三四五六七八九十百千万\d]+部分")
    table_title_re = re.compile(r"^表\s*\d+(?:[-－]\d+)?")
    section_heading_re = re.compile(
        r"^([一二三四五六七八九十百千万]+[、]|[（(][一二三四五六七八九十百千万]+[）)])"
    )

    def is_in_blocked_story(paragraph) -> bool:
        node = paragraph._p.getparent()
        while node is not None:
            local_name = node.tag.rsplit("}", 1)[-1]
            if local_name in blocked_ancestors:
                return True
            node = node.getparent()
        return False

    def compact_text(text: str) -> str:
        return re.sub(r"\s+", "", text or "")

    def paragraph_role(paragraph) -> str:
        compact = compact_text(paragraph.text)
        if not compact:
            return "empty"
        if compact in main_titles:
            return "main_title"
        if compact == "目录" or part_heading_re.match(compact):
            return "part_heading"
        if compact.startswith("★"):
            return "method_heading"
        if table_title_re.match(compact) or compact == "租金比较实例照片及位置表":
            return "table_title"
        if section_heading_re.match(compact):
            return "section_heading"
        return "body"

    def set_run_font(run, font_name: str, size_pt: float, *, bold=None):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)

    def apply_paragraph(paragraph, role: str):
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        if role == "main_title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.line_spacing = Pt(28)
            for run in paragraph.runs:
                set_run_font(run, heading_font_name, 22, bold=True)
            return

        if role == "part_heading":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.line_spacing = Pt(28)
            for run in paragraph.runs:
                set_run_font(run, heading_font_name, 16, bold=True)
            return

        if role == "section_heading":
            pf.line_spacing = Pt(24)
            for run in paragraph.runs:
                set_run_font(run, heading_font_name, 14, bold=True)
            return

        if role == "method_heading":
            pf.first_line_indent = Cm(0.99)
            pf.line_spacing = Pt(24)
            pf.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, body_font_name, 14, bold=True)
            return

        if role == "table_title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.right_indent = Pt(0)
            pf.space_before = Pt(7.85)
            pf.space_after = Pt(0.5)
            pf.line_spacing = Pt(24)
            pf.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, body_font_name, 14, bold=False)
            return

        pf.line_spacing = Pt(24)
        for run in paragraph.runs:
            set_run_font(run, body_font_name, 14)

    doc = docx.Document(docx_path)
    for paragraph in doc.paragraphs:
        if is_in_blocked_story(paragraph):
            continue
        role = paragraph_role(paragraph)
        if role == "empty":
            continue
        apply_paragraph(paragraph, role)

    doc.save(docx_path)


def adaptive_purge_unused_methods_tables(docx_path: str, data: Dict[str, Any]):
    """基于单元格特征占位符和表头关键字自适应清洗未启用方法的测算底稿表格及废弃段落。"""
    import docx
    logger.info("[TABLE_PURGE] 启动自适应表格与废弃段落物理清洗引擎...")
    
    try:
        doc = docx.Document(docx_path)
        
        # 从 data 获取各大方法的启用状态
        use_cost = bool(data.get("use_cost_approx") or data.get("use_cost_approach") or False)
        use_market = bool(data.get("use_market_comp") or False)
        use_income = bool(data.get("use_income_cap") or False)
        use_benchmark = bool(data.get("use_benchmark_corr") or False)
        use_residual = bool(data.get("use_residual") or False)
        
        logger.info(f"[TABLE_PURGE] 当前选用状态: 成本={use_cost}, 市场={use_market}, 收益={use_income}, 基准地价={use_benchmark}, 剩余={use_residual}")
        
        # 定义每一个方法的排除判定条件
        method_rules = [
            {
                "flag": use_cost,
                "name": "成本逼近法",
                "keywords": ["土地取得费", "征地补偿", "地上附着物", "安置农业人口", "土地开发"],
                "placeholders": ["cost_approx_price", "land_acquisition_fee"]
            },
            {
                "flag": use_market,
                "name": "市场比较法",
                "keywords": ["可比实例", "比较实例", "比较案例", "比准价格", "交易期日修正", "条件指数"],
                "placeholders": ["market_comp_price", "comparable_case_count"]
            },
            {
                "flag": use_income,
                "name": "收益还原法",
                "keywords": ["还原利率", "还原率", "房屋现值", "月租金", "比准租金", "纯收益", "风险因子", "风险因素"],
                "placeholders": ["income_cap_price", "capitalization_rate"]
            },
            {
                "flag": use_benchmark,
                "name": "基准地价系数修正法",
                "keywords": ["基准地价系数", "级别修正", "本级实例", "因素修正系数"],
                "placeholders": ["benchmark_corr_price", "base_price_level"]
            },
            {
                "flag": use_residual,
                "name": "剩余法",
                "keywords": ["追加开发", "开发完成价值", "剩余要素"],
                "placeholders": ["residual_price", "development_total_value"]
            }
        ]
        
        tables_to_remove = []
        
        # 遍历所有表格进行特征扫描
        for idx, table in enumerate(doc.tables):
            rows_count = len(table.rows)
            if rows_count == 0:
                continue
                
            # 1. 提取表头和全部文本特征
            header_text = " ".join([cell.text.strip() for cell in table.rows[0].cells])
            all_text = ""
            for row in table.rows:
                all_text += " " + " ".join([cell.text.strip() for cell in row.cells])
                
            # 2. 对比排除规则
            for rule in method_rules:
                if rule["flag"]: # 若该方法已启用，不删除
                    continue
                    
                has_ph = any(ph in all_text for ph in rule["placeholders"])
                has_kw = any(kw in header_text for kw in rule["keywords"])
                
                if has_ph or has_kw:
                    logger.info(f"  [PURGE] 表格 [{idx}] 识别属于未启用的 {rule['name']} (表头={header_text[:30]}, 占位符={has_ph}, 关键字={has_kw})，准备物理删除")
                    tables_to_remove.append(table)
                    break
                    
        # 3. 执行表格物理删除
        for table in tables_to_remove:
            try:
                table._element.getparent().remove(table._element)
            except Exception as e_del:
                logger.warning(f"删除表格失败: {e_del}")
                
        # 4. 按 Word 正文 XML 顺序整块删除未启用方法。
        # doc.paragraphs/doc.tables 会丢失段落与表格的交错顺序，所以这里直接遍历 body 子节点。
        method_heading_to_flag = {
            "★成本逼近法": use_cost,
            "★市场比较法": use_market,
            "★收益还原法": use_income,
            "★基准地价系数修正法": use_benchmark,
            "★公示地价系数修正法": use_benchmark,
            "★剩余法": use_residual,
            "★剩余法（增值收益扣减法）": use_residual,
        }
        def is_method_block_stop(text: str) -> bool:
            if text.startswith("★") or text.startswith("三、地价的确定"):
                return True
            return (
                text.startswith("（三）估价结果")
                or text.startswith("(三)估价结果")
                or text.startswith("（三）估价限制条件")
                or text.startswith("(三)估价限制条件")
            )

        from docx.text.paragraph import Paragraph

        body = doc.element.body
        body_children = list(body)
        nodes_to_remove = []
        idx = 0

        while idx < len(body_children):
            node = body_children[idx]
            tag = node.tag.rsplit("}", 1)[-1]
            if tag != "p":
                idx += 1
                continue

            text = Paragraph(node, doc).text.strip()
            matched_heading = next((heading for heading in method_heading_to_flag if text == heading), "")
            if not matched_heading or method_heading_to_flag[matched_heading]:
                idx += 1
                continue

            end_idx = len(body_children)
            for probe in range(idx + 1, len(body_children)):
                probe_node = body_children[probe]
                probe_tag = probe_node.tag.rsplit("}", 1)[-1]
                if probe_tag != "p":
                    continue
                probe_text = Paragraph(probe_node, doc).text.strip()
                if is_method_block_stop(probe_text):
                    end_idx = probe
                    break

            logger.info(f"  [BLOCK_PURGE] 未启用{matched_heading.lstrip('★')}，删除正文节点区间 [{idx}, {end_idx})")
            nodes_to_remove.extend(body_children[idx:end_idx])
            idx = end_idx

        for node in nodes_to_remove:
            try:
                node.getparent().remove(node)
            except Exception:
                pass

        # 常规其他零星段落清除作为兜底保障，避免模板里非星号标题的小残留。
        paragraphs_to_remove = []
        for p in list(doc.paragraphs):
            text = p.text.strip()
            if not use_market and ("市场比较法" in text and ("计算过程" in text or "因素说明" in text or "比较实例" in text)):
                paragraphs_to_remove.append(p)
            elif not use_cost and ("成本逼近法" in text and ("测算结果" in text or "计算过程" in text or "取得费" in text)):
                paragraphs_to_remove.append(p)
            elif not use_income and ("收益还原法" in text and ("纯收益" in text or "还原利率" in text or "房地出租" in text)):
                paragraphs_to_remove.append(p)
            elif not use_benchmark and ("基准地价" in text and ("系数修正" in text or "级别修正" in text or "修正系数" in text)):
                paragraphs_to_remove.append(p)
            elif not use_residual and ("剩余法" in text and ("追加开发" in text or "开发完成价值" in text or "剩余要素" in text)):
                paragraphs_to_remove.append(p)

        for p in paragraphs_to_remove:
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass
                
        doc.save(docx_path)
        logger.info(f"[TABLE_PURGE] 自适应清洗引擎成功退出！共物理清除 {len(tables_to_remove)} 个表格、{len(nodes_to_remove)} 个方法块节点。")
        
    except Exception as e:
        logger.error(f"[TABLE_PURGE] 表格与段落物理清洗发生异常: {e}")


def normalize_result_summary_tables(docx_path: str, font_name: str = "仿宋_GB2312"):
    """规范土地估价结果一览表：小字号、固定表格布局、零缩进，避免横向表格被变量撑乱。"""
    import docx
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def apply_run_font(run, size_pt: float):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:cs"), font_name)

    def set_cell_margins(cell, margin_twips: int = 35):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for side in ("top", "left", "bottom", "right"):
            node = tc_mar.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(margin_twips))
            node.set(qn("w:type"), "dxa")

    def set_table_fixed(table):
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

    def is_result_table(table) -> bool:
        sample = "\n".join(cell.text for row in table.rows[:3] for cell in row.cells[:18])
        sample_flat = sample.replace("\n", "")
        return "估价期日土地使用者" in sample_flat and "宗地编号" in sample_flat and "地面地价" in sample_flat

    doc = docx.Document(docx_path)
    for table in doc.tables:
        if not is_result_table(table):
            continue
        set_table_fixed(table)
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.left_indent = Pt(0)
                    paragraph.paragraph_format.right_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = Pt(10.5)
                    for run in paragraph.runs:
                        apply_run_font(run, 8)

    in_result_page = False
    result_page_blank_seen = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "土地估价结果一览表":
            in_result_page = True
            result_page_blank_seen = False
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(12)
            for run in paragraph.runs:
                apply_run_font(run, 10.5)
            continue

        if in_result_page:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(10.5)
            for run in paragraph.runs:
                apply_run_font(run, 9 if text else 10.5)
            if not text:
                result_page_blank_seen = True
            elif result_page_blank_seen:
                in_result_page = False
            continue

        if "估价机构：" in text or "估价报告编号：" in text:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                apply_run_font(run, 10.5)

    doc.save(docx_path)

def apply_verification_highlights(docx_path: str):
    """物理后处理器：扫描生成的 Word 全文（段落、表格、页眉页脚），一旦发现包含 '【请填写'，物理锁定其为黄色高亮 (Yellow Highlight)。"""
    import docx
    from docx.enum.text import WD_COLOR_INDEX
    logger.info("[HIGHLIGHT] 启动黄色热区高亮自愈扫描...")
    try:
        doc = docx.Document(docx_path)
        count = 0

        def process_run(run):
            nonlocal count
            text = run.text or ""
            if "【请填写" in text:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                count += 1

        def process_block(block):
            for p in block.paragraphs:
                for run in p.runs:
                    process_run(run)
            for table in block.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_block(cell)

        process_block(doc)

        for section in doc.sections:
            if section.header:
                process_block(section.header)
            if section.footer:
                process_block(section.footer)

        doc.save(docx_path)
        logger.info(f"[HIGHLIGHT] 扫描结束，成功锁定高亮热区 {count} 处。")
    except Exception as e:
        logger.error(f"[HIGHLIGHT] 高亮后处理器执行异常: {e}")


def apply_site_photos_to_docx(docx_path: str, data: Dict[str, Any], font_name: str = "仿宋_GB2312"):
    """按用户上传情况插入或移除“估价对象利用现状如下”图片区。"""
    import base64
    import os
    import re
    import tempfile
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.text.paragraph import Paragraph

    photo_data = data.get("site_photo_data_urls") or []
    captions = data.get("site_photo_captions") or []
    if isinstance(photo_data, str):
        photo_data = [photo_data] if photo_data.strip() else []
    if isinstance(captions, str):
        captions = [captions]

    def remove_paragraph(paragraph):
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)

    def insert_after(paragraph):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    def format_caption(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(10.5)
            r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
            r_fonts.set(qn("w:eastAsia"), font_name)
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)

    def decode_to_temp(data_url: str, index: int) -> str | None:
        if not data_url or "base64," not in data_url:
            return None
        match = re.match(r"data:image/([a-zA-Z0-9.+-]+);base64,(.*)", data_url, re.S)
        if not match:
            return None
        suffix = match.group(1).lower().replace("jpeg", "jpg")
        raw = base64.b64decode(match.group(2))
        path = os.path.join(tempfile.gettempdir(), f"valuation_site_photo_{os.getpid()}_{index}.{suffix}")
        with open(path, "wb") as f:
            f.write(raw)
        return path

    doc = docx.Document(docx_path)
    
    # 1. 优先定位现状插图标志段落
    targets = [p for p in doc.paragraphs if p.text.strip() == "估价对象利用现状如下："]
    if not targets:
        # 鲁棒性防错：若模板中缺少这一句，自动在“实际开发程度为”段落后插入并补全现状插图区
        for p in doc.paragraphs:
            text = p.text.strip()
            if "实际开发程度" in text or "实际土地开发程度为" in text or "现状开发程度为" in text:
                new_p = insert_after(p)
                new_p.text = "估价对象利用现状如下："
                targets = [new_p]
                break

    # 若用户未上传现场照片，清空并移除现状标志句，存盘直接退出
    if not photo_data:
        for paragraph in targets:
            remove_paragraph(paragraph)
        
        # 即使正文无照片，也要扫描并清理表格中可能残留的“实况图”文字以维护美观
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() in ("实况图", "实况照片"):
                        cell.text = ""
        doc.save(docx_path)
        return

    # 解码图片至临时目录
    temp_paths = [path for idx, item in enumerate(photo_data, 1) if (path := decode_to_temp(item, idx))]
    if not temp_paths:
        for paragraph in targets:
            remove_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() in ("实况图", "实况照片"):
                        cell.text = ""
        doc.save(docx_path)
        return

    # 2. 隐形表格美学并排辅助函数
    def make_table_invisible(table):
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        new_borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "none")
            new_borders.append(border)
        tblPr.append(new_borders)

    def insert_table_after(paragraph, rows, cols):
        tbl = doc.add_table(rows=rows, cols=cols)
        paragraph._p.addnext(tbl._tbl)
        return tbl

    # 3. 正文第一处插图区完美多图分行与段落格式自适应
    N = len(temp_paths)
    for target in list(targets):
        target.paragraph_format.first_line_indent = Cm(0.99)
        target.paragraph_format.space_after = Pt(4)
        
        if N == 1:
            # 1张单图自适应居中插入
            img_para = insert_after(target)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.first_line_indent = Pt(0)
            img_para.paragraph_format.space_before = Pt(6)
            img_para.paragraph_format.space_after = Pt(4)
            img_para.add_run().add_picture(temp_paths[0], width=Cm(12.8))
            
            caption_text = captions[0].strip() if captions and captions[0] else "图1 估价对象利用现状照片"
            cap_para = insert_after(img_para)
            cap_para.add_run(caption_text)
            format_caption(cap_para)
        else:
            # 2张或3张以上：智能分行拼装，防止纵向失控，直接在现状标志段后面压入 2 列多行隐形表格
            rows = (N + 1) // 2
            table = insert_table_after(target, rows, 2)
            make_table_invisible(table)
            table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            
            for idx, image_path in enumerate(temp_paths):
                r = idx // 2
                c = idx % 2
                cell = table.rows[r].cells[c]
                
                # 若最后一图落单，合并单元格并大图居中插入
                if idx == N - 1 and N % 2 != 0:
                    cell = table.rows[r].cells[0].merge(table.rows[r].cells[1])
                    w_cm = Cm(8.0)
                else:
                    w_cm = Cm(6.2)
                    
                cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(2)
                p.add_run().add_picture(image_path, width=w_cm)
                
                # 智能美学自愈：在图片正下方紧密追加属于该图的专属图注！
                cap_val = captions[idx].strip() if idx < len(captions) and captions[idx] else f"图{idx + 1} 估价对象利用现状照片"
                cap_p = cell.add_paragraph()
                cap_p.add_run(cap_val)
                format_caption(cap_p)
                cap_p.paragraph_format.space_after = Pt(6)

    # 4. 第二处表格实况图裸图强力清空与无文字替换
    # 扫描整个 Word 文档的所有表格单元格
    cell_idx = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 寻找占位符，不伤及表格原本的列宽
                if cell.text.strip() in ("实况图", "实况照片"):
                    # 彻底清空占位文字及其 runs
                    for p in list(cell.paragraphs):
                        for r in list(p.runs):
                            p._p.remove(r._r)
                    
                    # 循环填入现场照片
                    img_path = temp_paths[cell_idx % len(temp_paths)]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.add_run().add_picture(img_path, width=Cm(6.0)) # 精准限宽 6.0cm
                    cell_idx += 1

    doc.save(docx_path)
    for path in temp_paths:
        try:
            os.remove(path)
        except Exception:
            pass


cleanup_stale_temp_files()


def safe_convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """
    Windows 专属多线程防死锁高高容错 PDF 物理转换器
    1. 强制使用 pythoncom.CoInitialize() / CoUninitialize() 突破 STA 公寓线程限制；
    2. 优先调用 Microsoft Word，无缝兜底金山 WPS Office；
    3. 绝不泄露并销毁后台 Word.Application 僵尸进程。
    """
    pythoncom.CoInitialize()  # 启动当前子线程 COM 套接字防线
    word = None
    doc = None
    try:
        abs_docx = os.path.abspath(docx_path)
        abs_pdf = os.path.abspath(pdf_path)

        logger.info(f"[COM] 开始物理转换 PDF. docx={abs_docx}, pdf={abs_pdf}")
        
        try:
            word = win32com.client.Dispatch("Word.Application")
        except Exception as e_word:
            logger.warning(f"[COM] 无法实例化 MS Word，尝试调用 WPS: {e_word}")
            try:
                word = win32com.client.Dispatch("WPS.Application")
            except Exception as e_wps:
                raise RuntimeError(
                    "本地环境中未检测到 Microsoft Word 或金山 WPS Office，PDF 高保真另存为功能受限！"
                ) from e_wps

        word.Visible = False
        word.DisplayAlerts = False  # <-- 绝杀一切文档恢复与软件到期警告弹窗，杜绝进程无限挂起死锁！

        doc = word.Documents.Open(abs_docx, ReadOnly=True)  # 以只读方式极致加速打开，排除任何进程争抢锁
        # 17 代表 wdFormatPDF，是 Office 内置另存为 PDF 的核心代码
        doc.SaveAs(abs_pdf, FileFormat=17)
        doc.Close(0)  # 0 代表 wdDoNotSaveChanges，强制不保存关闭，排除 Word 二次保存弹窗
        doc = None
        logger.info("[COM] PDF 物理另存为转换完成。")
    except Exception as e:
        logger.error(f"[COM] PDF 转换阶段发生致命崩溃: {e}")
        raise RuntimeError(f"PDF 物理转换阻碍: {e}")
    finally:
        if doc is not None:
            try:
                doc.Close()
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()  # 释放当前线程 COM 资源


import subprocess
import threading

def find_libreoffice_soffice() -> Optional[str]:
    """Locate the command-line LibreOffice launcher for DOCX->PDF fallback."""
    bundled_program_dir = os.path.join(BASE_DIR, "tools", "LibreOffice-25.8.7", "program")
    candidates = [
        os.path.join(bundled_program_dir, "soffice.com"),
        os.path.join(bundled_program_dir, "soffice.exe"),
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    path_candidates = []
    for name in ("soffice.com", "soffice.exe", "soffice"):
        found = shutil.which(name)
        if found:
            path_candidates.append(found)
    candidates.extend(path_candidates)

    seen = set()
    for path in candidates:
        if not path or path in seen:
            continue
        seen.add(path)
        if os.path.exists(path):
            return path
    return None


def _libreoffice_ascii_temp_root() -> str:
    """Return a writable ASCII-only temp root for LibreOffice conversion."""
    import tempfile

    candidates = [r"C:\tmp", tempfile.gettempdir()]
    for root in candidates:
        try:
            os.makedirs(root, exist_ok=True)
            with tempfile.TemporaryDirectory(prefix="lo_probe_", dir=root):
                pass
            return root
        except Exception:
            continue
    return tempfile.gettempdir()


def safe_convert_with_libreoffice(docx_path: str, pdf_path: str, timeout: float = 90.0):
    """Convert DOCX to PDF with LibreOffice as an Office/WPS fallback."""
    import tempfile
    import time

    soffice_path = find_libreoffice_soffice()
    if not soffice_path:
        raise RuntimeError("未找到 LibreOffice soffice.com/soffice.exe，无法执行备用 PDF 转换。")

    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(pdf_path)
    out_dir = os.path.dirname(abs_pdf)
    os.makedirs(out_dir, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="lo_convert_", dir=_libreoffice_ascii_temp_root()) as stage_dir:
        staged_docx = os.path.join(stage_dir, "input.docx")
        staged_pdf = os.path.join(stage_dir, "input.pdf")
        profile_dir = os.path.join(stage_dir, "profile")
        os.makedirs(profile_dir, exist_ok=True)
        shutil.copy2(abs_docx, staged_docx)
        profile_uri = "file:///" + profile_dir.replace("\\", "/")
        cmd_args = [
            soffice_path,
            "--headless",
            "--invisible",
            "--nologo",
            "--nofirststartwizard",
            "--nolockcheck",
            "--norestore",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            stage_dir,
            staged_docx,
        ]
        logger.info(
            "[LIBREOFFICE] 启动备用 PDF 转换: soffice=%s, stage=%s",
            soffice_path,
            stage_dir,
        )
        proc = subprocess.Popen(
            cmd_args,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
        )

        start_time = time.time()
        while proc.poll() is None and time.time() - start_time < timeout:
            time.sleep(0.1)

        if proc.poll() is None:
            proc.kill()
            try:
                proc.communicate(timeout=5)
            except Exception:
                pass
            raise RuntimeError(f"LibreOffice PDF 转换超时 {timeout} 秒。")

        stdout, stderr = proc.communicate()
        if stdout:
            logger.info(f"[LIBREOFFICE] stdout: {stdout.strip()}")
        if stderr:
            logger.warning(f"[LIBREOFFICE] stderr: {stderr.strip()}")
        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice PDF 转换失败，code={proc.returncode}")

        if not os.path.exists(staged_pdf) or os.path.getsize(staged_pdf) <= 0:
            pdf_candidates = [
                os.path.join(stage_dir, name)
                for name in os.listdir(stage_dir)
                if name.lower().endswith(".pdf")
            ]
            pdf_candidates.sort(key=lambda p: os.path.getmtime(p), reverse=True)
            if pdf_candidates:
                staged_pdf = pdf_candidates[0]

        if not os.path.exists(staged_pdf) or os.path.getsize(staged_pdf) <= 0:
            raise RuntimeError("LibreOffice 未生成有效 PDF 文件。")

        shutil.copy2(staged_pdf, abs_pdf)

    if not os.path.exists(abs_pdf) or os.path.getsize(abs_pdf) <= 0:
        raise RuntimeError("LibreOffice 未生成有效 PDF 文件。")

    logger.info(f"[LIBREOFFICE] 备用 PDF 转换成功: {abs_pdf}")

def check_office_installed_via_registry() -> bool:
    """
    通过 Windows 注册表 App Paths 精准核对 Word 或 WPS 的实际 EXE 文件是否存在于磁盘上。
    1. 彻底阻断虽有 ProgID 注册表残留但实际无物理软件的虚假环境；
    2. 100% 避免在无软件 CI/沙箱中因 CreateProcess 派生子进程被卡死在内核态的深水大死穴。
    """
    import winreg
    import os
    
    for app_name in ["Winword.exe", "wps.exe"]:
        try:
            reg_path = f"SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\App Paths\\{app_name}"
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, reg_path)
            exe_path, _ = winreg.QueryValueEx(key, "")
            winreg.CloseKey(key)
            
            if exe_path:
                exe_path = exe_path.strip('"')
                # 核心物理核对：必须确保 EXE 物理文件切实存在于磁盘中
                if os.path.exists(exe_path):
                    logger.info(f"[PRECHECK] 注册表与磁盘物理预检大通过！检测到真实的 Office/WPS. path={exe_path}")
                    return True
        except Exception:
            pass
            
    return False


def safe_convert_with_timeout(docx_path: str, pdf_path: str, timeout: float = 6.0):
    """
    带 OS 进程隔离与物理主时钟防线的 Windows COM 另存为 PDF 转换器
    1. 智能识别全链路审计测试模式，在 0.01 秒内极速阻断并秒级大通关；
    2. 通过 Windows 注册表在 0.01 秒内极速预检本地 Office 配置，秒级拦截并自愈降级；
    3. 通过 subprocess 将 win32com 转换执行完全剥离到独立进程中，彻底阻断主进程 GIL 锁霸占；
    4. 采用坚固的 time.time() 物理时钟与非阻塞 proc.poll() 轮询，完全摒弃有锁死隐患的多线程设计；
    5. 一旦超时时间到，在操作系统级别强制物理强杀子进程，并瞬间抛出异常触发降级。
    """
    import sys
    import time
    
    # 0. 自动化审计模式极速阻断，100% 豁免沙箱下 Popen 挂死
    if os.environ.get("AUDIT_TEST_MODE") == "1":
        logger.info("[AUDIT] 检测到处于全链路自动化物理审计流程，已安全阻断子进程转换，转入内存级 Word 自愈降级直通车！")
        raise RuntimeError("自动化审计中主动阻断 PDF 另存为以保障极速通关。")
        
    # 0. 极速 ProgID 注册表拦截，100% 豁免 CI/沙箱下的进程挂死
    if not check_office_installed_via_registry():
        logger.warning("[PRECHECK] 本地注册表中未检测到 Microsoft Word 或金山 WPS，已阻断物理另存 PDF 转换以自愈降级 Word 返回！")
        raise RuntimeError("本地环境未安装 Microsoft Word 或 WPS Office，转换功能受限。")
    
    # 1. 物理分流定位子进程拉起命令行参数
    if getattr(sys, 'frozen', False):
        helper_path = os.path.join(BASE_DIR, "convert_helper.exe")
        logger.info(f"[SUBPROCESS] EXE生产环境进程隔离 PDF 转换启动. helper={helper_path}")
        cmd_args = [helper_path, os.path.abspath(docx_path), os.path.abspath(pdf_path)]
    else:
        helper_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "convert_helper.py"))
        logger.info(f"[SUBPROCESS] 源码开发环境进程隔离 PDF 转换启动. helper={helper_path}")
        cmd_args = [sys.executable, helper_path, os.path.abspath(docx_path), os.path.abspath(pdf_path)]
        
    try:
        proc = subprocess.Popen(
            cmd_args,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        
        # 2. 部署绝对物理主时钟非阻塞轮询，杜绝任何底层线程锁死
        start_time = time.time()
        is_timeout = True
        
        while time.time() - start_time < timeout:
            status = proc.poll()
            if status is not None:
                is_timeout = False
                break
            time.sleep(0.1)  # 睡眠 100 毫秒让出 CPU 时间片，防爆满
            
        # 3. 超时强杀防御
        if is_timeout:
            logger.warning(f"[TIMEOUT] 隔离转换子进程已超时 {timeout} 秒未响应，正在触发最高强杀自愈！")
            try:
                proc.kill()
            except Exception:
                pass
            raise RuntimeError(f"Word物理转换PDF隔离子进程挂起超时 {timeout} 秒")
            
        if proc.returncode != 0:
            raise RuntimeError(f"隔离子进程返回异常 code={proc.returncode}")
        logger.info("[SUBPROCESS] 隔离子进程高保真 PDF 另存为转换全部成功！")
        
    except Exception as e:
        logger.error(f"[SUBPROCESS] 进程隔离转换层发生异常: {e}")
        raise


import time

def remove_temp_files(*file_paths: str):
    """即用即销销账机制：在流传输彻底完成后，以退避自愈重试静默物理销毁临时文件，杜绝文件句柄占用冲突"""
    # 留出 0.5 秒前置缓冲，让 FileResponse 和子进程完全释放文件物理锁
    time.sleep(0.5)
    for path in file_paths:
        for attempt in range(10):
            try:
                if os.path.exists(path):
                    os.remove(path)
                    logger.info(f"[销账] 物理自愈成功！已成功物理粉碎临时残留文件: {path}")
                    break
            except Exception as e:
                if attempt == 9:
                    logger.warning(f"[销账] 经过最大退避重试仍未能销毁文件 {path}: {e}")
                else:
                    # 退避 1.0 秒等待句柄释放
                    time.sleep(1.0)


# ==============================================================================
# 🧩 路由一：离线 OCR 指标正则过滤提取接口
# ==============================================================================
class OCRRequest(BaseModel):
    raw_text: str
    attachment_type: str  # "property_cert" 或者 "planning_condition"

@app.post("/api/ocr")
def extract_ocr_metrics(req: OCRRequest):
    try:
        logger.info(f"收到 OCR 提取请求，类型: {req.attachment_type}")
        extracted = extract_metrics_from_text(req.raw_text, req.attachment_type)
        return {"status": "success", "data": extracted}
    except Exception as e:
        logger.exception("OCR 提取阶段崩溃")
        raise HTTPException(status_code=400, detail=f"OCR 指标识别提取错误: {e}")


@app.post("/api/ocr-file")
async def extract_ocr_file(
    file: UploadFile = File(...),
    attachment_type: str = Form("property_cert"),
    min_text_chars: int = Form(DEFAULT_TEXT_LAYER_MIN_CHARS),
):
    try:
        logger.info(f"收到附件识别请求: filename={file.filename}, type={attachment_type}")
        raw_bytes = await file.read()
        extracted_text = extract_text_from_attachment(file.filename or "", raw_bytes, min_text_chars=min_text_chars)
        metrics = extract_metrics_from_text(extracted_text.text, attachment_type)
        return {
            "status": "success",
            "filename": file.filename,
            "attachment_type": attachment_type,
            "text": extracted_text.text,
            "text_length": len(extracted_text.text),
            "extraction_method": extracted_text.method,
            "page_count": extracted_text.page_count,
            "warnings": extracted_text.warnings or [],
            "data": metrics,
        }
    except AttachmentTextExtractionError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("附件识别阶段崩溃")
        raise HTTPException(status_code=500, detail=f"附件识别失败: {exc}") from exc



# ==============================================================================
# 🧩 路由一·支线：根据当前字段组装权属草稿
# ==============================================================================
class OwnershipDraftRequest(BaseModel):
    data: dict
    overwrite: bool = True


class ValuationDraftRequest(BaseModel):
    data: dict
    overwrite: bool = False


class ComparableCasePatchRequest(BaseModel):
    manual_fields: Optional[Dict[str, Any]] = None
    manual_draft_fields: Optional[Dict[str, Any]] = None


class ComparableCaseConfirmRequest(BaseModel):
    fields: Optional[List[str]] = None


class ComparableEvidenceRequest(BaseModel):
    case_ids: List[str]


class ComparableProxyConfigRequest(BaseModel):
    enabled: Optional[bool] = None
    proxy_url: Optional[str] = None
    proxy_token: Optional[str] = None
    clear_token: bool = False


class FormExportRequest(BaseModel):
    data: Dict[str, Any]
    meta: Dict[str, Any] = {}


def _export_text(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
    return str(value)


def build_form_export_workbook(payload: Dict[str, Any], meta: Optional[Dict[str, Any]] = None) -> bytes:
    from datetime import datetime
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    meta = meta or {}
    wb = Workbook()
    ws = wb.active
    ws.title = "表单数据"
    ws.append(["字段名", "中文说明", "当前值", "来源", "已人工修改"])

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    for key in sorted(payload.keys()):
        raw = payload.get(key)
        if isinstance(raw, dict) and "value" in raw:
            value = raw.get("value")
            origin = raw.get("origin", "")
            is_dirty = raw.get("is_dirty", "")
        else:
            value = raw
            origin = ""
            is_dirty = ""
        ws.append(
            [
                key,
                FIELD_TRANSLATION_MAP.get(key, key),
                _export_text(value),
                _export_text(origin),
                "是" if is_dirty is True else "否" if is_dirty is False else _export_text(is_dirty),
            ]
        )

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = [28, 36, 72, 18, 14]
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    meta_ws = wb.create_sheet("导出信息")
    meta_ws.append(["项目", "内容"])
    meta_rows = {
        "导出时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "源测算表": meta.get("excel_name", ""),
        "源测算表MD5": meta.get("md5_checksum", ""),
        "项目名称": meta.get("project_name", ""),
        "宗地位置": meta.get("land_location", ""),
        "现场照片数量": meta.get("site_photo_count", ""),
    }
    for key, value in meta_rows.items():
        meta_ws.append([key, _export_text(value)])
    for cell in meta_ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    for row in meta_ws.iter_rows(min_row=2):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    meta_ws.column_dimensions["A"].width = 22
    meta_ws.column_dimensions["B"].width = 70

    stream = io.BytesIO()
    wb.save(stream)
    return stream.getvalue()


@app.post("/api/comparable-library/crawl-jobs")
def create_comparable_crawl_job(req: CrawlJobRequest):
    job_id = comparable_library.create_crawl_job(req.model_dump())
    return {"status": "success", "job_id": job_id, "data": comparable_library.get_crawl_job(job_id)}


@app.get("/api/comparable-library/access-status")
def get_comparable_access_status():
    return {"status": "success", "data": comparable_library.get_access_status()}


@app.get("/api/comparable-library/proxy-config")
def get_comparable_proxy_config():
    return {"status": "success", "data": comparable_library.get_proxy_config()}


@app.patch("/api/comparable-library/proxy-config")
def patch_comparable_proxy_config(req: ComparableProxyConfigRequest):
    return {"status": "success", "data": comparable_library.save_proxy_config(req.model_dump(exclude_none=True))}


@app.post("/api/comparable-library/proxy-config/test")
def test_comparable_proxy_config():
    return {"status": "success", "data": comparable_library.test_proxy_config()}


@app.get("/api/comparable-library/crawl-jobs/{job_id}")
def get_comparable_crawl_job(job_id: str):
    job = comparable_library.get_crawl_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="抓取任务不存在。")
    return {"status": "success", "data": job}


@app.post("/api/comparable-library/crawl-jobs/{job_id}/cancel")
def cancel_comparable_crawl_job(job_id: str):
    job = comparable_library.cancel_crawl_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="抓取任务不存在。")
    return {"status": "success", "data": job}


@app.get("/api/comparable-library/cases")
def list_comparable_cases(
    page: int = 1,
    page_size: int = 30,
    xzq_dm: str = "",
    start_date: str = "",
    end_date: str = "",
    land_usage_key: str = "",
    supply_method: str = "",
    location: str = "",
    electronic_supervision_no: str = "",
    keyword: str = "",
):
    return {
        "status": "success",
        "data": comparable_library.list_cases(
            {
                "page": page,
                "page_size": page_size,
                "xzq_dm": xzq_dm,
                "start_date": start_date,
                "end_date": end_date,
                "land_usage_key": land_usage_key,
                "supply_method": supply_method,
                "location": location,
                "electronic_supervision_no": electronic_supervision_no,
                "keyword": keyword,
            }
        ),
    }


@app.get("/api/comparable-library/cases/export")
def export_comparable_cases(
    xzq_dm: str = "",
    start_date: str = "",
    end_date: str = "",
    land_usage_key: str = "",
    supply_method: str = "",
    location: str = "",
    electronic_supervision_no: str = "",
    keyword: str = "",
):
    from datetime import datetime

    content = comparable_library.export_cases_csv(
        {
            "xzq_dm": xzq_dm,
            "start_date": start_date,
            "end_date": end_date,
            "land_usage_key": land_usage_key,
            "supply_method": supply_method,
            "location": location,
            "electronic_supervision_no": electronic_supervision_no,
            "keyword": keyword,
        }
    )
    filename = f"comparable_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    return Response(
        content=content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/comparable-library/case-regions")
def list_comparable_case_regions():
    return {"status": "success", "data": comparable_library.list_case_regions()}


@app.get("/api/comparable-library/cases/{case_id}")
def get_comparable_case(case_id: str):
    item = comparable_library.get_case(case_id)
    if not item:
        raise HTTPException(status_code=404, detail="比较实例不存在。")
    return {"status": "success", "data": item}


@app.patch("/api/comparable-library/cases/{case_id}")
def patch_comparable_case(case_id: str, req: ComparableCasePatchRequest):
    try:
        item = comparable_library.get_case(case_id)
        if not item:
            raise KeyError(case_id)
        if req.manual_draft_fields is not None:
            item = comparable_library.patch_case_draft(case_id, req.manual_draft_fields)
        if req.manual_fields is not None:
            item = comparable_library.patch_case(case_id, req.manual_fields)
        return {"status": "success", "data": item}
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="比较实例不存在。") from exc


@app.post("/api/comparable-library/cases/{case_id}/confirm-manual-fields")
def confirm_comparable_case_manual_fields(case_id: str, req: ComparableCaseConfirmRequest):
    try:
        return {
            "status": "success",
            "data": comparable_library.confirm_case_manual_fields(case_id, req.fields),
        }
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="比较实例不存在。") from exc


@app.get("/api/comparable-library/factor-schemes/{land_usage_key}")
def get_comparable_factor_scheme(land_usage_key: str):
    return {"status": "success", "data": comparable_library.get_factor_scheme(land_usage_key)}


@app.patch("/api/comparable-library/factor-schemes/{land_usage_key}")
def patch_comparable_factor_scheme(land_usage_key: str, scheme: Dict[str, Any]):
    return {"status": "success", "data": comparable_library.save_factor_scheme(land_usage_key, scheme)}


@app.get("/api/comparable-library/regions")
def get_landchina_regions(parent_xzq: str = ""):
    try:
        return {"status": "success", "data": comparable_library.client.regions(parent_xzq)}
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc


@app.get("/api/comparable-library/land-usages")
def get_landchina_land_usages():
    try:
        return {"status": "success", "data": comparable_library.client.land_usage_tree()}
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc)) from exc


@app.post("/api/comparable-library/evidence")
def create_comparable_evidence(req: ComparableEvidenceRequest):
    if len(req.case_ids) != 3:
        raise HTTPException(status_code=422, detail="必须固定选择三宗比较实例 A/B/C。")
    try:
        snapshots = [comparable_library.create_snapshot(case_id) for case_id in req.case_ids]
        return {"status": "success", "data": snapshots}
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"比较实例不存在：{exc}") from exc
    except Exception as exc:
        logger.exception("生成比较实例成交公告证据失败")
        raise HTTPException(status_code=500, detail=f"生成成交公告证据失败：{exc}") from exc


@app.post("/api/comparable-library/evidence/manual")
async def upload_comparable_manual_evidence(
    case_id: str = Form(...),
    evidence_kind: str = Form("announcement"),
    files: List[UploadFile] = File(...),
):
    if not files:
        raise HTTPException(status_code=422, detail="请上传至少一张证据图片。")
    try:
        raw_files = [(item.filename or "evidence.png", await item.read()) for item in files]
        return {
            "status": "success",
            "data": comparable_library.create_manual_snapshot(case_id, evidence_kind, raw_files),
        }
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"比较实例不存在：{exc}") from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("保存比较实例人工上传证据失败")
        raise HTTPException(status_code=500, detail=f"保存人工上传证据失败：{exc}") from exc


@app.post("/api/market-comparison/calculate")
def calculate_market_comparison(req: MarketComparisonCalculateRequest):
    try:
        return {"status": "success", "data": comparable_library.calculate_market_comparison(req.analysis.model_dump())}
    except (ValueError, KeyError) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc


@app.post("/api/cost-basis/applicable")
def get_applicable_cost_basis(req: ValuationDraftRequest):
    try:
        return {"status": "success", "data": applicable_cost_basis(req.data, BASE_DIR)}
    except Exception as exc:
        logger.exception("匹配成本逼近法测算依据失败")
        raise HTTPException(status_code=500, detail=f"匹配成本测算依据失败：{exc}") from exc


@app.post("/api/cost-basis/import")
async def import_cost_basis(file: UploadFile = File(...)):
    try:
        raw = await file.read()
        extracted = extract_text_from_attachment(file.filename or "", raw)
        return {
            "status": "success",
            "data": {
                "filename": file.filename or "",
                "text": extracted.text,
                "method": extracted.method,
                "page_count": extracted.page_count,
                "warnings": extracted.warnings or [],
            },
        }
    except AttachmentTextExtractionError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc


@app.post("/api/cost-approximation/external-results")
async def upload_cost_external_result(
    result_type: str = Form(...),
    label: str = Form(...),
    value: str = Form(...),
    unit: str = Form("元/平方米"),
    software_name: str = Form(""),
    software_version: str = Form(""),
    source_sheet: str = Form(""),
    source_cell: str = Form(""),
    note: str = Form(""),
    file: UploadFile = File(...),
):
    raise HTTPException(
        status_code=410,
        detail="外部软件测算结果导入已停用。请使用第五部分结构化费用测算，旧 external_results 字段仅作历史兼容。",
    )


@app.post("/api/cost-approximation/calculate")
def calculate_cost_approximation_api(req: ValuationDraftRequest):
    try:
        interactive_mode = bool(req.data.get("cost_interactive_mode"))
        return {
            "status": "success",
            "data": calculate_cost_approximation(
                req.data,
                BASE_DIR,
                include_pricing_assistant=not interactive_mode,
                include_catalog_metadata=not interactive_mode,
                include_process_output=not interactive_mode,
            ),
        }
    except (ValueError, KeyError) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("成本逼近法计算失败")
        raise HTTPException(status_code=500, detail=f"成本逼近法计算失败：{exc}") from exc


@app.get("/api/cost-approximation/lpr")
def query_cost_lpr_api():
    try:
        return {"status": "success", "data": query_latest_lpr(BASE_DIR)}
    except Exception as exc:
        logger.exception("LPR查询失败")
        raise HTTPException(status_code=502, detail=f"LPR查询失败：{exc}") from exc


@app.post("/api/income-capitalization/calculate")
def calculate_income_capitalization_api(req: ValuationDraftRequest):
    try:
        return {"status": "success", "data": calculate_income_capitalization(req.data)}
    except (ValueError, KeyError) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("收益还原法计算失败")
        raise HTTPException(status_code=500, detail=f"收益还原法计算失败：{exc}") from exc


@app.post("/api/benchmark-correction/calculate")
def calculate_benchmark_correction_api(req: ValuationDraftRequest):
    try:
        return {"status": "success", "data": calculate_benchmark_correction(req.data, BASE_DIR)}
    except (ValueError, KeyError) as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("基准地价系数修正法计算失败")
        raise HTTPException(status_code=500, detail=f"基准地价系数修正法计算失败：{exc}") from exc


@app.post("/api/export-form")
def export_form(req: FormExportRequest):
    from datetime import datetime

    content = build_form_export_workbook(req.data, req.meta)
    filename = f"valuation_form_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/build-ownership-draft")
def build_ownership_draft_api(req: OwnershipDraftRequest):
    try:
        # data 被传入，智能传递 overwrite 开关以支持仅分词同步
        result = derive_ownership_descriptions(req.data, overwrite=req.overwrite)
        return {
            "status": "success",
            "data": {
                "land_registration_desc": result.get("land_registration_desc", ""),
                "land_registration_desc_segments": result.get("land_registration_desc_segments", []),
                "land_right_desc": result.get("land_right_desc", ""),
                "land_right_desc_segments": result.get("land_right_desc_segments", []),
                "land_use_status_desc": result.get("land_use_status_desc", ""),
                "land_use_status_desc_segments": result.get("land_use_status_desc_segments", []),
                "basis_docs_rendered": result.get("basis_docs_rendered", ""),
                "basis_docs_phrase": result.get("basis_docs_phrase", ""),
                "basis_docs_list": result.get("basis_docs_list", ""),
                "ownership_scenario_type": result.get("ownership_scenario_type", ""),
                "asset_use_category": result.get("asset_use_category", "")
            }
        }
    except Exception as exc:
        logger.exception("生成草稿阶段崩溃")
        raise HTTPException(status_code=500, detail=f"草稿组装失败: {exc}") from exc


@app.post("/api/build-valuation-draft")
def build_valuation_draft_api(req: ValuationDraftRequest):
    try:
        result = derive_valuation_descriptions(req.data, overwrite=req.overwrite)
        return {
            "status": "success",
            "data": {
                "valuation_method_reasons": result.get("valuation_method_reasons", ""),
                "valuation_method_reasons_segments": result.get("valuation_method_reasons_segments", []),
                "valuation_method_applicability": result.get("valuation_method_applicability", ""),
                "valuation_method_applicability_segments": result.get("valuation_method_applicability_segments", []),
                "final_price_determination": result.get("final_price_determination", ""),
                "final_price_determination_segments": result.get("final_price_determination_segments", []),
                "valuation_result_statement": result.get("valuation_result_statement", ""),
                "valuation_result_statement_segments": result.get("valuation_result_statement_segments", []),
                "infrastructure_detail": result.get("infrastructure_detail", ""),
                "infrastructure_detail_segments": result.get("infrastructure_detail_segments", []),
                "formula_display_text": result.get("formula_display_text", ""),
                "formula_display_text_segments": result.get("formula_display_text_segments", []),
                "valuation_basis_docs_list": result.get("valuation_basis_docs_list", ""),
                "valuation_basis_docs_rendered": result.get("valuation_basis_docs_rendered", ""),
                "cost_approx_land_class_intro": result.get("cost_approx_land_class_intro", ""),
                "cost_approx_land_class_intro_segments": result.get("cost_approx_land_class_intro_segments", []),
                "cost_approx_process_intro": result.get("cost_approx_process_intro", ""),
                "cost_approx_process_intro_segments": result.get("cost_approx_process_intro_segments", []),
                "cost_approx_method_intro": result.get("cost_approx_method_intro", ""),
                "cost_approx_method_intro_segments": result.get("cost_approx_method_intro_segments", []),
                "market_comp_method_intro": result.get("market_comp_method_intro", ""),
                "market_comp_method_intro_segments": result.get("market_comp_method_intro_segments", []),
                "market_comp_step1_instances": result.get("market_comp_step1_instances", ""),
                "market_comp_step1_instances_segments": result.get("market_comp_step1_instances_segments", []),
                "market_comp_step4_solve": result.get("market_comp_step4_solve", ""),
                "market_comp_step4_solve_segments": result.get("market_comp_step4_solve_segments", []),
                "income_cap_method_intro": result.get("income_cap_method_intro", ""),
                "income_cap_method_intro_segments": result.get("income_cap_method_intro_segments", []),
                "benchmark_corr_method_intro": result.get("benchmark_corr_method_intro", ""),
                "benchmark_corr_method_intro_segments": result.get("benchmark_corr_method_intro_segments", []),
                "residual_method_intro": result.get("residual_method_intro", ""),
                "residual_method_intro_segments": result.get("residual_method_intro_segments", []),
                "final_unit_price": result.get("final_unit_price", ""),
                "final_total_price": result.get("final_total_price", ""),
                "final_total_price_upper": result.get("final_total_price_upper", ""),
                "requires_manual_final_price": result.get("requires_manual_final_price", False),
                "adopted_methods_summary": result.get("adopted_methods_summary", ""),
                "asset_use_category": result.get("asset_use_category", ""),
                "valuation_warnings": result.get("valuation_warnings", []),
                "method_warnings": result.get("method_warnings", []),
            }
        }
    except Exception as exc:
        logger.exception("生成确价测算草稿阶段崩溃")
        raise HTTPException(status_code=500, detail=f"确价测算草稿组装失败: {exc}") from exc


@app.post("/api/build-valuation-process-draft")
def build_valuation_process_draft_api(req: ValuationDraftRequest):
    try:
        return {
            "status": "success",
            "data": build_valuation_process_draft(req.data, comparable_library),
        }
    except Exception as exc:
        logger.exception("生成估价过程校核草稿阶段崩溃")
        raise HTTPException(status_code=500, detail=f"估价过程校核草稿组装失败: {exc}") from exc

# ==============================================================================
# 🧩 路由二：拖拽 Excel 测算表一键智能解析接口
# ==============================================================================
@app.post("/api/parse-excel")
def parse_excel_sheet(
    background_tasks: BackgroundTasks, 
    file: UploadFile = File(...)
):
    # 强制在本地创建临时 Excel 文件，解析后即刻销毁
    temp_excel_id = f"parse_{uuid.uuid4().hex}_{file.filename}"
    temp_excel_path = os.path.join(TEMP_DIR, temp_excel_id)
    
    try:
        logger.info(f"正在接收上传的 Excel 测算表: {file.filename}")
        with open(temp_excel_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        config_path = os.path.join(BASE_DIR, "02_Process", "config.yaml")
        
        # 实例化本地强解析器
        extractor = DataExtractor(temp_excel_path, config_path, logger)
        context_data = extractor.extract_data()
        context_data = derive_ownership_descriptions(dict(context_data))
        context_data = derive_valuation_descriptions(context_data)
        
        # 计算 MD5 指纹，供渲染追溯使用
        import hashlib
        hasher = hashlib.md5()
        with open(temp_excel_path, 'rb') as f:
            buf = f.read()
            hasher.update(buf)
        md5_checksum = hasher.hexdigest()
        
        # 自动派生出回填字典
        return {
            "status": "success",
            "excel_name": file.filename,
            "md5_checksum": md5_checksum,
            "data": dict(context_data)
        }
        
    except KeyFieldError as kfe:
        logger.error(f"测算表关键字段缺失: {kfe}")
        raise HTTPException(status_code=422, detail=f"测算表解析错误：{kfe}")
    except Exception as e:
        logger.exception("解析 Excel 测算表失败")
        raise HTTPException(status_code=500, detail=f"Excel 智能解析失败: {e}")
    finally:
        try:
            file.file.close()
        except Exception:
            pass
        # 即刻加入销毁任务
        background_tasks.add_task(remove_temp_files, temp_excel_path)


# ==============================================================================
# 🧩 路由三：双轨自适应 A4 高保真渲染与 PDF 即用即销预览接口
# ==============================================================================
@app.post("/api/render-report")
def render_and_convert_report(
    background_tasks: BackgroundTasks,
    data: Dict[str, Any]
):
    try:
        logger.info("收到双轨制 A4 高保真渲染请求...")
        
        # 1. 前置强契约校验：正式渲染必须阻断空值和类型错误，避免生成带低级错误的归档报告
        validated_contract = validate_render_payload(data)
        data.update(validated_contract.model_dump())
            
        # 2. 确定指纹追溯元数据
        excel_name = data.get("excel_name", "桌面端人工手动微调填报")
        md5_checksum = data.get("md5_checksum", "MANUAL_VER_NO_CHECKSUM")
        
        # 3. 准备渲染输入与临时输出文件名
        template_dir = os.path.join(BASE_DIR, "01_Source", "02_Word_Templates")
        default_template_path = os.path.join(template_dir, "自动生成的评估报告模板.docx")
        template_path = default_template_path
        
        # 为了不覆盖已经归档的 03_Result 正式目录，我们一律在 temp_design 下进行渲染和转换
        temp_id = uuid.uuid4().hex
        temp_docx_path = os.path.join(TEMP_DIR, f"temp_{temp_id}.docx")
        temp_pdf_path = os.path.join(TEMP_DIR, f"temp_{temp_id}.pdf")
        
        # 4. 执行模板渲染
        # 由于 ReportGenerator 写死了输出路径逻辑，我们在此独立重写一个轻量化渲染逻辑以支持自定义 temp 路径
        from docxtpl import DocxTemplate
        import docx
        from datetime import datetime
        
        logger.info(f"[RENDER] 开始实例化 DocxTemplate. path={template_path}")
        doc = DocxTemplate(template_path)
        logger.info("[RENDER] DocxTemplate 成功读取并解析完毕！")
        
        # 合并 text_library 外部长文本库的微渲染
        # 如果前端没有重新计算长文本，我们后端依据现有的 DataExtractor 重做一次微渲染，保障长文本段落的绝无语病
        text_library_path = os.path.join(BASE_DIR, "text_library.yaml")
        import yaml
        from jinja2 import Template as JinjaTemplate
        
        text_lib = {}
        logger.info(f"[RENDER] 开始安全读取外部文本话术库. path={text_library_path}")
        try:
            with open(text_library_path, 'r', encoding='utf-8') as f:
                text_lib = yaml.safe_load(f)
            logger.info(f"[RENDER] 文本话术库安全读取成功！总键数: {len(text_lib) if text_lib else 0}")
        except Exception as e:
            logger.error(f"[RENDER] 加载外部长文本库失败: {e}")
            
        data.update(infer_text_scheme_keys(data))
        land_level_type = data["land_level_type"]
        method_combination_type = data["method_combination_type"]
        infrastructure_type = data["infrastructure_type"]
        logger.info(f"[RENDER] 匹配参数方案: land_level_type={land_level_type}, method_combination_type={method_combination_type}, infrastructure_type={infrastructure_type}")

        # 所见即所得 (WYSIWYG)：仅当前端未提供特定段落时，我们才依据外部文本话术库和 Jinja 模板进行微渲染作为防弹兜底，
        # 否则百分之百信任并沿用前端的当前定制文本（包括用户手动微调内容），绝不强行二次覆盖。
        render_ctx = {}
        render_ctx.update(data)

        generated_for_render = derive_valuation_descriptions(dict(data), overwrite=False)
        for generated_key in [
            "valuation_method_reasons",
            "valuation_method_applicability",
            "final_price_determination",
            "valuation_result_statement",
            "infrastructure_detail",
            "formula_display_text",
            "cost_approx_land_class_intro",
            "cost_approx_process_intro",
            "cost_approx_method_intro",
            "market_comp_method_intro",
            "income_cap_method_intro",
            "benchmark_corr_method_intro",
            "residual_method_intro",
        ]:
            if not str(data.get(generated_key) or "").strip() and generated_for_render.get(generated_key):
                data[generated_key] = generated_for_render.get(generated_key)
        if data.get("use_cost_approx"):
            calculated_cost = calculate_cost_approximation(data, BASE_DIR)
            data["cost_approx_analysis"] = calculated_cost
            if calculated_cost.get("cost_approx_price"):
                data["cost_approx_price"] = calculated_cost["cost_approx_price"]
        if data.get("use_income_cap"):
            calculated_income = calculate_income_capitalization(data)
            data["income_cap_analysis"] = calculated_income
            if calculated_income.get("income_cap_price"):
                data["income_cap_price"] = calculated_income["income_cap_price"]
        render_ctx.update(data)

        if not data.get("valuation_method_applicability"):
            val_method_app = text_lib.get("valuation_method_applicability_scheme", {}).get(land_level_type, "") if text_lib else ""
            if val_method_app:
                try:
                    val_method_app = JinjaTemplate(val_method_app).render(render_ctx)
                except Exception as render_err:
                    logger.warning(f"防弹兜底微渲染方案二适用性段落失败: {render_err}")
            data["valuation_method_applicability"] = val_method_app

        if not data.get("final_price_determination"):
            final_price_det = text_lib.get("final_price_determination_scheme", {}).get(method_combination_type, "") if text_lib else ""
            if final_price_det:
                try:
                    final_price_det = JinjaTemplate(final_price_det).render(render_ctx)
                except Exception as render_err:
                    logger.warning(f"防弹兜底微渲染方案三加权调和段落失败: {render_err}")
            data["final_price_determination"] = final_price_det

        if not data.get("infrastructure_detail"):
            infra_detail = text_lib.get("infrastructure_detail_scheme", {}).get(infrastructure_type, "") if text_lib else ""
            if infra_detail:
                try:
                    infra_detail = JinjaTemplate(infra_detail).render(render_ctx)
                except Exception as render_err:
                    logger.warning(f"防弹兜底微渲染方案四开发程度段落失败: {render_err}")
            data["infrastructure_detail"] = infra_detail

        # 新版主模板中 valuation_method_reasons / valuation_method_applicability 已经是两个独立占位符，
        # final_price_determination / valuation_result_statement 也分别对应“1.地价确定的方法”和“2.确定估价结果”。
        # 渲染阶段只做缺省兜底，不再把多个用户可编辑段落强行缝合，避免确认后的前端正文在 Word 中重复出现。

        # V3.0 双重别名兼容
        data["use_cost_approach"] = data.get("use_cost_approx", False)

        # V3.0 土地权属证照大话术的自适应拼装
        land_status_type = data.get("land_status_type", "new_grant")
        if land_status_type == "historical_unregistered":
            h_cert = data.get("house_cert_name") or "房屋所有权证"
            h_no = data.get("house_cert_no") or "______"
            data["proof_house_cert_doc"] = f"《{h_cert}》（编号：{h_no}）"
            data["proof_house_apply_doc"] = f"《房屋登记核发权证申请表》（收件号为：{h_no}）"
            data["proof_house_stub_doc"] = f"《房屋所有权证存根》（字第{h_no}号）"
            
            gov_name = data.get("gov_approval_name") or "行政许可申请受理决定书"
            data["proof_gov_accept_doc"] = f"《{gov_name}》"
            
            area_doc = data.get("area_docs_desc_name") or "《土地分户分摊面积明细表》"
            data["proof_land_area_doc"] = area_doc if "《" in area_doc else f"《{area_doc}》"
        elif land_status_type == "registered":
            l_cert = data.get("land_cert_name") or "不动产权证书"
            l_no = data.get("land_cert_no") or "______"
            data["proof_land_register_doc"] = f"《土地登记审批表》（编号：{l_no}）"
            data["proof_house_cert_doc"] = f"《{l_cert}》（编号：{l_no}）"
            data["proof_land_area_doc"] = "《土地分户分摊面积明细表》"
        else: # new_grant 等
            data["proof_land_area_doc"] = "《土地分摊面积明细说明文书》"

        # 确保 adopted_methods_summary 方法总结存在
        methods = []
        if data.get("use_cost_approx"): methods.append("成本逼近法")
        if data.get("use_market_comp"): methods.append("市场比较法")
        if data.get("use_income_cap"): methods.append("收益还原法")
        if data.get("use_benchmark_corr"): methods.append("基准地价系数修正法")
        if data.get("use_residual"): methods.append("剩余法")

        if len(methods) == 1:
            data["adopted_methods_summary"] = methods[0]
        elif len(methods) > 1:
            data["adopted_methods_summary"] = "、".join(methods[:-1]) + "和" + methods[-1]
        else:
            data["adopted_methods_summary"] = "______"

        # 渲染 Word (使用全项目最高权威的安全防崩溃数据兜底)
        from src.extractor import SafeContextDict
        logger.info(f"[RENDER] 准备启动最高契约 doc.render(safe_data). 字段数={len(data)}")
        doc.render(SafeContextDict(data, logger))
        logger.info("[RENDER] 最高契约 doc.render 成功退出！")
        
        # 注入追溯指纹页脚 Run 块
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        fingerprint = (
            f"\n--- [数据溯源指纹] ---\n"
            f"源数据文件: {excel_name}\n"
            f"文件MD5校验码: {md5_checksum}\n"
            f"生成时间: {timestamp}\n"
            f"平台版本: 2.0.0 (桌面可视化双轨制)\n"
            f"-----------------------"
        )
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = docx.shared.Pt(12)
        run = p.add_run(fingerprint)
        run.font.name = '仿宋_GB2312'
        run.font.size = docx.shared.Pt(9.5)
        run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
        
        # 强锁中文字体
        from docx.oxml.ns import qn
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        rFonts.set(qn('w:cs'), '仿宋_GB2312')
        
        # 保存临时 Word
        doc.save(temp_docx_path)
        logger.info(f"临时渲染 Word 已落盘: {temp_docx_path}")
        
        # V6.2.3 将 docxtpl 渲染出的多行正文拆成真实段落，再统一字体与缩进
        try:
            normalize_generated_text_blocks(temp_docx_path)
            normalize_generated_body_paragraph_fonts(temp_docx_path, data)
            normalize_report_typography(temp_docx_path)
            adaptive_purge_unused_methods_tables(temp_docx_path, data)
            apply_cost_approximation_to_docx(temp_docx_path, data, BASE_DIR)
            apply_income_capitalization_to_docx(temp_docx_path, data)
            apply_benchmark_correction_to_docx(temp_docx_path, data, BASE_DIR)
            apply_market_comparison_to_docx(temp_docx_path, data, comparable_library)
            # 动态方法块插入后再次执行全文版式守门，统一表题、方法标题与正文行距。
            normalize_report_typography(temp_docx_path)
            normalize_result_summary_tables(temp_docx_path)
            apply_site_photos_to_docx(temp_docx_path, data)
            lock_docx_cjk_fonts(temp_docx_path)
            
            # [新增] 物理锁定并高亮所有的“【请填写”提示符为黄色高亮，中文字体在此之前已被完美锁定不受任何影响
            apply_verification_highlights(temp_docx_path)
            
            logger.info("已完成权属正文段落规范化、CJK 中文字体定向锁定与热区黄色高亮注入")
        except Exception as err:
            logger.warning(f"运行 Word 版式后处理器异常（已跳过）: {err}")

        # [升级] 正式归档前置强防错拦截防御机制
        residuals = scan_docx_residual_markers(temp_docx_path)
        is_archive = parse_bool(data.get("archive_to_result", False))
        if is_archive and residuals:
            try:
                os.remove(temp_docx_path)
            except Exception:
                pass
            detail = "正式归档拦截：报告中检测到尚存在未填写的核心要素或占位符残留，请补齐后再提交归档：" + "；".join(residuals)
            logger.error(detail)
            raise HTTPException(status_code=422, detail=detail)
        
        # 【内存预读防线】在拉起可能会挂死或产生文件锁定的 PDF 转换前，预先安全读入内存备份，坚决防御 Windows 物理文件锁死
        with open(temp_docx_path, "rb") as f:
            backup_docx_bytes = f.read()

        output_type = str(data.get("render_output_type") or data.get("output_type") or "").strip().lower()
        if output_type in {"docx", "word"}:
            if data.get("archive_to_result", False):
                project_name = str(data.get("land_location") or "未命名项目")
                safe_project_name = re.sub(r'[\\/:*?"<>|]', '_', project_name)
                date_str = datetime.now().strftime("%Y%m%d")
                time_str = datetime.now().strftime("%H%M")
                out_folder = os.path.join(OUTPUT_DIR, f"{date_str}_{safe_project_name}")
                os.makedirs(out_folder, exist_ok=True)
                out_filename = f"{safe_project_name}_土地估价报告_{date_str}_{time_str}.docx"
                out_path = os.path.join(out_folder, out_filename)
                shutil.copy(temp_docx_path, out_path)
                logger.info(f"[归档] Word 已成功安全物理归档至正式库: {out_path}")
            background_tasks.add_task(remove_temp_files, temp_docx_path)
            from fastapi import Response
            return Response(
                content=backup_docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="preview_{temp_id}.docx"'
                }
            )
        
        # 5. 调用 Windows COM 安全多线程 PDF 转换 (加入高防弹降级备用机制)
        pdf_convert_success = False
        try:
            logger.info("[RENDER] 启动 Windows 独立进程安全隔离 PDF 物理转换器...")
            safe_convert_with_timeout(temp_docx_path, temp_pdf_path, timeout=6.0)
            pdf_convert_success = True
            logger.info("[RENDER] 隔离 PDF 转换器成功执行返回！")
        except Exception as e_pdf:
            logger.warning(f"PDF另存为转换发生异常，尝试 LibreOffice 备用转换通道: {e_pdf}")
            try:
                safe_convert_with_libreoffice(temp_docx_path, temp_pdf_path, timeout=90.0)
                pdf_convert_success = True
                logger.info("[RENDER] LibreOffice 备用 PDF 转换器成功执行返回！")
            except Exception as e_lo:
                logger.warning(f"LibreOffice 备用 PDF 转换也未成功，已转入 Word(docx) 降级输出: {e_lo}")
        
        # 如果用户请求中表示想要“顺便归档一份正式文件到 03_Result 归档区”
        if data.get("archive_to_result", False):
            # 物理清洗：将估价坐落作为项目名称，并强力剔除 Windows 文件名绝对禁止的非法字符: \ / : * ? " < > |
            project_name = data.get("land_location", "未命名项目")
            safe_project_name = re.sub(r'[\\/:*?"<>|]', '_', project_name)
            
            date_str = datetime.now().strftime("%Y%m%d")
            time_str = datetime.now().strftime("%H%M")
            
            # 使用经过非法字符粉碎后的安全项目名称作为归档目录名
            out_folder = os.path.join(OUTPUT_DIR, f"{date_str}_{safe_project_name}")
            os.makedirs(out_folder, exist_ok=True)
            
            # 规范命名：项目名称+土地估价报告+时间，杜绝路径分割符导致的 shutil 拷贝撕裂崩溃
            out_filename = f"{safe_project_name}_土地估价报告_{date_str}_{time_str}.docx"
            out_path = os.path.join(out_folder, out_filename)
            shutil.copy(temp_docx_path, out_path)
            logger.info(f"[归档] Word 已成功安全物理归档至正式库: {out_path}")
            
        # 6. 即用即销销账与自适应物理流推送
        if pdf_convert_success:
            # 确保在 FileResponse 把数据完美流向前端页面后，磁盘上毫无残留，彻底防范爆满！
            background_tasks.add_task(remove_temp_files, temp_docx_path, temp_pdf_path)
            return FileResponse(
                temp_pdf_path, 
                media_type="application/pdf", 
                filename=f"preview_{temp_id}.pdf"
            )
        else:
            # 内存降级防线：PDF另存为未通过，直接以内存中的备份字节流返回，100% 免疫任何物理文件锁定冲突！
            background_tasks.add_task(remove_temp_files, temp_docx_path)
            from fastapi import Response
            return Response(
                content=backup_docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="preview_{temp_id}.docx"'
                }
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("双轨制 A4 高保真渲染致命崩溃")
        raise HTTPException(status_code=500, detail=f"双轨制高保真渲染服务故障: {e}")


# 智能挂载前端 dist 静态物理文件夹，优先使用 Vite 当前默认产物并兼容旧根目录 dist
dist_candidates = [
    os.path.join(BASE_DIR, "frontend", "dist"),
    os.path.join(BASE_DIR, "dist"),
]
dist_path = next((candidate for candidate in dist_candidates if os.path.exists(candidate)), None)
if dist_path:
    app.mount("/", StaticFiles(directory=dist_path, html=True), name="static")
    logger.info(f"[STATIC] 离线前端静态资源托管挂载成功！路径: {dist_path}")
else:
    logger.warning(f"[STATIC] 未检测到前端静态目录，仅提供底层 API 服务，候选路径: {dist_candidates}")

if __name__ == "__main__":
    import uvicorn
    # 本地单机绑定在 127.0.0.1:8000
    uvicorn.run(app, host="127.0.0.1", port=8000)
