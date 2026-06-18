# -*- coding: utf-8 -*-
from __future__ import annotations

import base64
import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

BASE_DIR = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(BASE_DIR))

from src.services.income_capitalization import calculate_income_capitalization
from src.services.income_capitalization_docx import apply_income_capitalization_to_docx
from tests.test_income_capitalization import _baozhen_income_payload


def image_data(label: str, color: tuple[int, int, int]) -> str:
    image = Image.new("RGB", (1200, 700), color)
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 58)
        small = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 34)
    except OSError:
        font = ImageFont.load_default()
        small = ImageFont.load_default()
    draw.rounded_rectangle((35, 35, 1165, 665), radius=28, outline="white", width=8)
    draw.text((75, 250), label, fill="white", font=font)
    draw.text((75, 335), "收益还原法视觉QA测试图片", fill="white", font=small)
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    return "data:image/png;base64," + base64.b64encode(stream.getvalue()).decode("ascii")


def build() -> Path:
    output = Path(__file__).resolve().parent / "income_cap_baozhen.docx"
    data = _baozhen_income_payload()
    colors = ((31, 111, 139), (45, 125, 82), (142, 82, 52))
    for item, color in zip(data["income_cap_analysis"]["rent_instances"], colors):
        item["photo_data"] = image_data(f"租金实例 {item['slot']} 现状照片", color)
        item["photo_name"] = f"rent_{item['slot']}.png"
        item["location_image_data"] = image_data(f"租金实例 {item['slot']} 位置图", tuple(min(channel + 35, 230) for channel in color))
        item["location_image_name"] = f"location_{item['slot']}.png"
    data["income_cap_analysis"] = calculate_income_capitalization(data)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    normal = document.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal.font.size = Pt(10.5)
    title = document.add_paragraph("收益还原法动态报告视觉 QA")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    document.add_paragraph("本页标题用于确认收益法动态块边界；下方正文、表格和图片均由运行后处理器生成。")
    document.add_paragraph("★收益还原法")
    document.add_paragraph("模板收益法旧内容，应被物理替换。")
    document.add_paragraph("★基准地价系数修正法")
    document.add_paragraph("下一方法边界保留成功。")
    document.save(output)

    apply_income_capitalization_to_docx(str(output), data)
    return output


if __name__ == "__main__":
    print(build())
