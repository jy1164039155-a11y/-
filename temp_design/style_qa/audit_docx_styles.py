from __future__ import annotations

import json
import re
import sys
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _pt(value: Any) -> float | None:
    return round(value.pt, 2) if value is not None else None


def _cm(value: Any) -> float | None:
    return round(value.cm, 3) if value is not None else None


def _enum(value: Any) -> str | None:
    return str(value) if value is not None else None


def _effective_style_value(style: Any, attr: str) -> Any:
    current = style
    while current is not None:
        value = getattr(current.paragraph_format, attr, None)
        if value is not None:
            return value
        current = current.base_style
    return None


def _effective_font_value(style: Any, attr: str) -> Any:
    current = style
    while current is not None:
        value = getattr(current.font, attr, None)
        if value is not None:
            return value
        current = current.base_style
    return None


def _paragraph_role(text: str) -> str:
    stripped = text.strip()
    if re.match(r"^表\s*\d", stripped):
        return "table_title"
    if stripped.startswith("★"):
        return "method_heading"
    if re.match(r"^\d+[、.．]", stripped):
        return "number_heading"
    if re.match(r"^[（(]\d+[）)]", stripped) or re.match(r"^[①②③④⑤⑥⑦⑧⑨⑩]", stripped):
        return "item"
    if stripped.startswith(("式中", "P＝", "P=", "A、", "B、", "C、", "D、", "E、", "F、", "G、")):
        return "formula"
    return "body"


def _paragraph_info(paragraph: Any) -> dict[str, Any]:
    pf = paragraph.paragraph_format
    style = paragraph.style
    run = next((item for item in paragraph.runs if item.text.strip()), paragraph.runs[0] if paragraph.runs else None)
    run_font = run.font if run is not None else None
    rpr = run._r.rPr if run is not None else None
    rfonts = rpr.rFonts if rpr is not None else None
    return {
        "text": paragraph.text.strip()[:120],
        "role": _paragraph_role(paragraph.text),
        "style": style.name if style is not None else None,
        "alignment_direct": _enum(paragraph.alignment),
        "alignment_effective_style": _enum(_effective_style_value(style, "alignment")) if style else None,
        "first_line_direct_cm": _cm(pf.first_line_indent),
        "first_line_effective_style_cm": _cm(_effective_style_value(style, "first_line_indent")) if style else None,
        "left_indent_direct_cm": _cm(pf.left_indent),
        "line_spacing_direct": str(pf.line_spacing) if pf.line_spacing is not None else None,
        "line_spacing_rule_direct": _enum(pf.line_spacing_rule),
        "space_before_direct_pt": _pt(pf.space_before),
        "space_after_direct_pt": _pt(pf.space_after),
        "font_name_direct": run_font.name if run_font is not None else None,
        "font_name_east_asia": rfonts.get(qn("w:eastAsia")) if rfonts is not None else None,
        "font_name_effective_style": _effective_font_value(style, "name") if style else None,
        "font_size_direct_pt": _pt(run_font.size) if run_font is not None else None,
        "font_size_effective_style_pt": _pt(_effective_font_value(style, "size")) if style else None,
        "bold_direct": run.bold if run is not None else None,
        "bold_effective_style": _effective_font_value(style, "bold") if style else None,
    }


def _cell_margins(cell: Any) -> dict[str, str | None]:
    tc_pr = cell._tc.tcPr
    tc_mar = tc_pr.first_child_found_in("w:tcMar") if tc_pr is not None else None
    result: dict[str, str | None] = {}
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}")) if tc_mar is not None else None
        result[edge] = node.get(qn("w:w")) if node is not None else None
    return result


def _table_info(table: Any, index: int) -> dict[str, Any]:
    first = table.cell(0, 0) if table.rows and table.columns else None
    body = table.cell(1, 0) if len(table.rows) > 1 and table.columns else first
    previous = table._tbl.getprevious()
    previous_text = ""
    if previous is not None and previous.tag.endswith("}p"):
        previous_text = "".join(previous.itertext()).strip()
    return {
        "index": index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "style": table.style.name if table.style is not None else None,
        "alignment": _enum(table.alignment),
        "autofit": table.autofit,
        "first_cell_vertical": _enum(first.vertical_alignment) if first is not None else None,
        "first_cell_margins_dxa": _cell_margins(first) if first is not None else None,
        "header_paragraph": _paragraph_info(first.paragraphs[0]) if first is not None else None,
        "body_paragraph": _paragraph_info(body.paragraphs[0]) if body is not None else None,
        "previous_text": previous_text[:120],
        "sample": " | ".join(cell.text.strip().replace("\n", "/")[:35] for cell in table.rows[0].cells) if table.rows else "",
    }


def audit(path: Path) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = [_paragraph_info(p) for p in doc.paragraphs if p.text.strip()]
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in paragraphs:
        grouped.setdefault(item["role"], []).append(item)
    samples = {key: values[:12] for key, values in grouped.items()}
    signatures = {}
    for role, values in grouped.items():
        keys = (
            "style",
            "alignment_direct",
            "first_line_direct_cm",
            "line_spacing_direct",
            "line_spacing_rule_direct",
            "space_before_direct_pt",
            "space_after_direct_pt",
            "font_name_east_asia",
            "font_size_direct_pt",
            "bold_direct",
        )
        signatures[role] = [
            {"count": count, **dict(zip(keys, signature))}
            for signature, count in Counter(tuple(item[key] for key in keys) for item in values).most_common(8)
        ]
    return {
        "path": str(path),
        "sections": [
            {
                "top_margin_cm": _cm(section.top_margin),
                "bottom_margin_cm": _cm(section.bottom_margin),
                "left_margin_cm": _cm(section.left_margin),
                "right_margin_cm": _cm(section.right_margin),
                "page_width_cm": _cm(section.page_width),
                "page_height_cm": _cm(section.page_height),
            }
            for section in doc.sections
        ],
        "paragraph_count": len(paragraphs),
        "paragraph_signatures": signatures,
        "paragraph_samples": samples,
        "tables": [_table_info(table, index) for index, table in enumerate(doc.tables)],
    }


if __name__ == "__main__":
    args = list(sys.argv[1:])
    output_path = None
    if "--out" in args:
        out_index = args.index("--out")
        output_path = Path(args[out_index + 1])
        del args[out_index : out_index + 2]
    result = [audit(Path(arg)) for arg in args]
    rendered = json.dumps(result, ensure_ascii=False, indent=2)
    if output_path is not None:
        output_path.write_text(rendered, encoding="utf-8")
    else:
        print(rendered)
